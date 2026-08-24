from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
from PIL import Image

from .capabilities.expression_rendering import render_expression_omml
from .capabilities.text_expression_rendering import build_text_expression_render_plans
from .document_contracts import (
    FOOTER_TEXT,
    HEADER_FOOTER_CONTRACT,
    HEADER_TEXT,
    PAGE_CONTRACT,
    TEXT_CONTRACT,
)
from .document_presentation import plan_ordered_answer_units, question_unit_rows
from .formula_audit import looks_like_formula, looks_like_symbolic_formula
from .omml_input import strip_structured_math_metadata
from .question_types import (
    explicit_question_type,
    infer_question_type,
    is_choice_question,
    is_short_answer_question,
    is_term_explanation_question,
    question_has_type,
)

MIN_WORD_FIGURE_HEIGHT_CM = 3.8
WIDE_WORD_FIGURE_ASPECT_RATIO = 2.4
MIN_WIDE_WORD_FIGURE_HEIGHT_CM = 4.4
INLINE_SCIENTIFIC_LIST_BREAK_RE = re.compile(
    r"(?<=[、，,;；])[^\S\r\n]*\n+[^\S\r\n]*(?=\([0-9A-Za-z+\-/,]{2,16}\))"
)
SUBSCRIPT_DIGITS = str.maketrans("₀₁₂₃₄₅₆₇₈₉", "0123456789")
GREEK_LATEX = {
    "α": r"\alpha",
    "β": r"\beta",
    "γ": r"\gamma",
    "δ": r"\delta",
    "θ": r"\theta",
    "λ": r"\lambda",
    "μ": r"\mu",
    "π": r"\pi",
    "σ": r"\sigma",
    "ν": r"\nu",
    "Δ": r"\Delta",
    "∆": r"\Delta",
}
GREEK_CHARS = "αβγδθλμπσνΔ∆"
SYMBOL_CHARS = rf"A-Za-z{GREEK_CHARS}"
SCRIPT_SUBSCRIPT_RE = rf"(?:\{{[A-Za-z0-9{GREEK_CHARS},]+\}}|[A-Za-z0-9{GREEK_CHARS}]+)"
SCRIPTED_ATOM_RE = (
    rf"(?:[Δ∆]_[rR][{SYMBOL_CHARS}]_{SCRIPT_SUBSCRIPT_RE}|"
    rf"[{SYMBOL_CHARS}]_{SCRIPT_SUBSCRIPT_RE})"
)
CHEM_SUBSCRIPT_RE = r"(?:[A-Za-z]{1,3}[₀₁₂₃₄₅₆₇₈₉]+[A-Za-z]?)"
REACTION_SUMMARY_RE = re.compile(
    rf"(?:{SCRIPTED_ATOM_RE}|{CHEM_SUBSCRIPT_RE})(?:\s*(?:[+→⇌])\s*(?:{SCRIPTED_ATOM_RE}|{CHEM_SUBSCRIPT_RE}))+"
)
SCRIPTED_SUMMARY_RE = re.compile(SCRIPTED_ATOM_RE)
UNICODE_LARGE_OPERATOR_RE = re.compile(
    rf"(?<![A-Za-z0-9])(?P<operator>[∑ΣΠ])(?:_(?:\{{(?P<braced>[A-Za-z0-9{GREEK_CHARS},]+)\}}|(?P<plain>[A-Za-z0-9{GREEK_CHARS}]+)))?"
)
CRYSTALLOGRAPHIC_LATEX_ATOM_RE = r"(?:[A-Za-z0-9]+|\\(?:bar|overline)\{[A-Za-z0-9+-]+\})"
CRYSTALLOGRAPHIC_LATEX_RE = re.compile(
    rf"(?:\(\{{?|\[\{{?|<\{{?|\{{)"
    rf"(?={CRYSTALLOGRAPHIC_LATEX_ATOM_RE}*\\(?:bar|overline)\{{)"
    rf"{CRYSTALLOGRAPHIC_LATEX_ATOM_RE}+"
    rf"(?:\}}?\)|\}}?\]|\}}?>|\}})"
)
SQRT_SUMMARY_RE = re.compile(r"√\s*(?:\(([^)]+)\)|([A-Za-z0-9]+))(?:\s*≈\s*[-+]?\d+(?:\.\d+)?)?")
FORMULA_TOKEN_RE = (
    rf"(?:"
    rf"(?:sin|cos|tan)\s*[{SYMBOL_CHARS}][A-Za-z0-9{GREEK_CHARS}]{{0,4}}"
    rf"|[-+]?\d+(?:\.\d+)?[{SYMBOL_CHARS}][A-Za-z0-9{GREEK_CHARS}]{{0,4}}"
    rf"|[{SYMBOL_CHARS}]?[dδΔ∆]?[{SYMBOL_CHARS}]{{1,4}}(?:_[A-Za-z0-9{GREEK_CHARS}]+)?"
    rf"|[-+]?\d+(?:\.\d+)?"
    rf")"
)
FORMULA_EXPR_RE = rf"{FORMULA_TOKEN_RE}(?:\s*(?:[+\-−·*])\s*{FORMULA_TOKEN_RE}|\s*/\s*{FORMULA_TOKEN_RE}|\s+{FORMULA_TOKEN_RE})*"
EQUATION_SUMMARY_RE = re.compile(
    r"(?<![A-Za-z0-9])((?:[Δ∆])?[A-Za-z][A-Za-z0-9]*)\s*=\s*"
    r"([-+]?\d+(?:\.\d+)?(?:×10\^?[-+]?\d+)?)(?:\s*([A-Za-z]+))?"
)
RELATION_SUMMARY_RE = re.compile(
    rf"(?<![A-Za-z0-9])({FORMULA_TOKEN_RE})\s*(≤|≥|<=|>=|<|>)\s*({FORMULA_EXPR_RE})(?![A-Za-z0-9])"
)
SYMBOLIC_EQUATION_SUMMARY_RE = re.compile(rf"(?<![A-Za-z0-9])({FORMULA_EXPR_RE})\s*=\s*({FORMULA_EXPR_RE})(?![A-Za-z0-9])")
DOLLAR_LATEX_SUMMARY_RE = re.compile(r"\$([^$\n]+)\$")
PARTIAL_DERIVATIVE_SUMMARY_RE = re.compile(
    r"[（(]\s*(?:∂|\\partial)\s*[（(]\s*"
    r"(?P<numerator>[A-Za-zΔ∆][A-Za-z0-9Δ∆]*(?:_[A-Za-z0-9]+)?)\s*[)）]\s*"
    r"/\s*(?:∂|\\partial)\s*(?P<denominator>[A-Za-z][A-Za-z0-9]*)\s*[)）]\s*"
    r"_?\s*(?P<condition>[A-Za-z][A-Za-z0-9]*)"
)
RATIO_EQUIVALENCE_SUMMARY_RE = re.compile(
    r"(?<![\d.])([-+]?\d+(?:\.\d+)?(?:\s*:\s*[-+]?\d+(?:\.\d+)?)+)"
    r"\s*(≈|=)\s*"
    r"([-+]?\d+(?:\.\d+)?(?:\s*:\s*[-+]?\d+(?:\.\d+)?)+)(?![\d.])"
)
ARROW_SUMMARY_RE = re.compile(r"[→⇌↔]")
AUDIT_PROMPT_FILL = "FFF2CC"
ANSWER_BODY_FIRST_LINE_INDENT_CM = TEXT_CONTRACT.answer_first_line_indent_cm
TOP_LEVEL_SUBQUESTION_RE = re.compile(
    r"^\s*(?:[（(]\s*(?:[1-9]\d*|[一二三四五六七八九十]+)\s*[）)]|第\s*(?:[1-9]\d*|[一二三四五六七八九十]+)\s*(?:小问|问))"
)
SUBQUESTION_MARKER_RE = re.compile(r"(?m)^(\s*)[（(]\s*((?:[1-9]|1\d|20))\s*[）)]\s*(?:[、.．]\s*)?")


def normalize_answer_hierarchy_markers(text: str) -> str:
    return SUBQUESTION_MARKER_RE.sub(r"\1(\2)", str(text or ""))


def figure_display_width_cm(image_path: Path) -> float:
    """Use more of the page for landscape figures so their labels remain readable."""
    try:
        with Image.open(image_path) as image:
            width, height = image.size
        ratio = width / max(height, 1)
    except Exception:
        return 10.5
    if ratio >= 1.7:
        return 13.8
    if ratio >= 1.25:
        return 12.2
    return 10.5


def _word_fit_report_path(image_path: Path) -> Path:
    if image_path.parent.name == "figures":
        return image_path.parent.parent / "figure_word_fit.json"
    return image_path.parent / "figure_word_fit.json"


def _record_word_fit(original: Path, fitted: Path, item: dict) -> None:
    report_path = _word_fit_report_path(original)
    try:
        data = json.loads(report_path.read_text(encoding="utf-8")) if report_path.exists() else {}
    except Exception:
        data = {}
    items = [entry for entry in data.get("items", []) if isinstance(entry, dict)] if isinstance(data, dict) else []
    key = str(original)
    items = [entry for entry in items if str(entry.get("original_path") or "") != key]
    items.append({"original_path": key, "fitted_path": str(fitted), **item})
    report_path.write_text(json.dumps({"schema_version": "answer_book.figure_word_fit.v1", "items": items}, ensure_ascii=False, indent=2), encoding="utf-8")


def prepare_figure_image_for_word(image_path: Path, display_width_cm: float) -> Path:
    """Pad very wide figures so Word insertion keeps a readable minimum height."""
    try:
        with Image.open(image_path) as image:
            source = image.convert("RGB")
            width, height = source.size
    except Exception:
        return image_path
    if width <= 0 or height <= 0 or display_width_cm <= 0:
        return image_path
    display_height_cm = display_width_cm * height / width
    display_aspect_ratio = display_width_cm / max(display_height_cm, 0.01)
    required_height_cm = MIN_WIDE_WORD_FIGURE_HEIGHT_CM if display_aspect_ratio >= WIDE_WORD_FIGURE_ASPECT_RATIO else MIN_WORD_FIGURE_HEIGHT_CM
    if display_height_cm >= required_height_cm:
        return image_path
    required_height = int((width * required_height_cm / display_width_cm) + 0.999)
    if required_height <= height:
        return image_path
    fit_path = image_path.with_name(f"{image_path.stem}_wordfit{image_path.suffix or '.png'}")
    top_pad = (required_height - height) // 2
    canvas = Image.new("RGB", (width, required_height), "white")
    canvas.paste(source, (0, top_pad))
    canvas.save(fit_path)
    _record_word_fit(
        image_path,
        fit_path,
        {
            "original_width_px": width,
            "original_height_px": height,
            "fitted_width_px": width,
            "fitted_height_px": required_height,
            "display_width_cm": round(display_width_cm, 2),
            "original_display_height_cm": round(display_height_cm, 2),
            "target_min_height_cm": required_height_cm,
            "top_pad_px": top_pad,
            "bottom_pad_px": required_height - height - top_pad,
            "method": "white_vertical_padding",
        },
    )
    return fit_path


def add_figure_picture(paragraph, image_path: Path) -> None:
    display_width_cm = figure_display_width_cm(image_path)
    fitted_path = prepare_figure_image_for_word(image_path, display_width_cm)
    paragraph.add_run().add_picture(str(fitted_path), width=Cm(display_width_cm))


def append_domain_text_runs(paragraph, text: str, *, highlight: bool = False) -> None:
    """Render deterministic academic notation while preserving surrounding prose."""

    source = str(text or "")
    cursor = 0
    for plan in (
        item
        for item in build_text_expression_render_plans(source)
        if item.expression_kind == "chemical_notation"
    ):
        if plan.start < cursor:
            continue
        if plan.start > cursor:
            run = paragraph.add_run(source[cursor:plan.start])
            set_run_font(run)
            if highlight:
                set_run_shading(run)
        if plan.preserve_parentheses:
            run = paragraph.add_run("（")
            set_run_font(run)
        try:
            paragraph._p.append(
                render_expression_omml(
                    plan.render_latex,
                    display=False,
                    location="answer_book_domain_notation",
                    expression_kind=plan.expression_kind,
                )
            )
        except Exception:
            run = paragraph.add_run(plan.raw)
            set_run_font(run)
        if plan.preserve_parentheses:
            run = paragraph.add_run("）")
            set_run_font(run)
        cursor = plan.end
    if cursor < len(source):
        run = paragraph.add_run(source[cursor:])
        set_run_font(run)
        if highlight:
            set_run_shading(run)


def set_run_font(
    run,
    east=TEXT_CONTRACT.east_asia_font,
    west=TEXT_CONTRACT.latin_font,
    size=TEXT_CONTRACT.body_size_pt,
    bold=False,
):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_run_shading(run, fill: str = AUDIT_PROMPT_FILL) -> None:
    rpr = run._element.get_or_add_rPr()
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_para(p, align=None):
    p.paragraph_format.line_spacing = TEXT_CONTRACT.line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align


def add_text_paragraph(doc: Document, text: str, bold: bool = False, size: float = 11, align=None, skip_formula_audit: bool = False):
    p = doc.add_paragraph()
    set_para(p, align)
    r = p.add_run(normalize_answer_hierarchy_markers(text))
    set_run_font(r, size=size, bold=bold)
    return p


def _latex_symbol(value: str) -> str:
    return "".join(GREEK_LATEX.get(ch, ch) for ch in str(value))


def _latex_variable(value: str) -> str:
    token = str(value)
    if token.startswith(("Δ", "∆")) and len(token) > 1:
        return rf"\Delta {_latex_symbol(token[1:])}"
    return _latex_symbol(token)


def _latex_atom(value: str) -> str:
    token = str(value).strip()
    composite = re.fullmatch(
        rf"([Δ∆])_([rR])([{SYMBOL_CHARS}])_(?:\{{([^{{}}]+)\}}|([A-Za-z0-9{GREEK_CHARS},]+))",
        token,
    )
    if composite:
        reaction_marker = r"\mathrm{r}" if composite.group(2).lower() == "r" else composite.group(2)
        subscript = composite.group(4) or composite.group(5) or ""
        return rf"\Delta_{{{reaction_marker}}} {_latex_symbol(composite.group(3))}_{{{_latex_symbol(subscript)}}}"
    if "_" in token:
        base, sub = token.split("_", 1)
        if sub.startswith("{") and sub.endswith("}"):
            sub = sub[1:-1]
        return f"{_latex_symbol(base)}_{{{_latex_symbol(sub)}}}"
    match = re.fullmatch(r"([A-Za-z]{1,3})([₀₁₂₃₄₅₆₇₈₉]+)([A-Za-z]?)", token)
    if match:
        head = match.group(1)
        digits = match.group(2).translate(SUBSCRIPT_DIGITS)
        tail = match.group(3)
        suffix = f"\\mathrm{{{tail}}}" if tail else ""
        return f"\\mathrm{{{head}}}_{{{digits}}}{suffix}"
    return _latex_symbol(token)


def _latex_reaction(value: str) -> str:
    parts = re.findall(rf"{SCRIPTED_ATOM_RE}|{CHEM_SUBSCRIPT_RE}|[+→⇌]", str(value))
    out: list[str] = []
    for part in parts:
        if part == "+":
            out.append("+")
        elif part == "→":
            out.append(r"\to")
        elif part == "⇌":
            out.append(r"\rightleftharpoons")
        else:
            out.append(_latex_atom(part))
    return "".join(out)


def _latex_sqrt(value: str) -> str:
    text = str(value).strip()
    approx = ""
    approx_match = re.search(r"≈\s*([-+]?\d+(?:\.\d+)?)", text)
    if approx_match:
        approx = rf"\approx {approx_match.group(1)}"
    inner_match = re.search(r"√\s*(?:\(([^)]+)\)|([A-Za-z0-9]+))", text)
    inner = (inner_match.group(1) or inner_match.group(2) or "").strip() if inner_match else ""
    if "/" in inner:
        numerator, denominator = [part.strip() for part in inner.split("/", 1)]
        body = rf"\frac{{{_latex_symbol(numerator)}}}{{{_latex_symbol(denominator)}}}"
    else:
        body = _latex_symbol(inner)
    return rf"\sqrt{{{body}}}{approx}"


def _latex_equation(match: re.Match[str]) -> str:
    variable = _latex_variable(match.group(1))
    value = match.group(2).replace("×", r"\times ")
    value = re.sub(r"10\^?([-+]?\d+)", r"10^{\1}", value)
    unit = match.group(3) or ""
    unit_text = rf"\ \mathrm{{{unit}}}" if unit else ""
    return f"{variable}={value}{unit_text}"


def _latex_relation_token(value: str) -> str:
    token = str(value).strip()
    if re.fullmatch(r"[-+]?\d+(?:\.\d+)?", token):
        return token
    trig = re.fullmatch(r"(sin|cos|tan)\s*(.+)", token)
    if trig:
        return rf"\{trig.group(1)} {_latex_symbol(trig.group(2))}"
    coefficient = re.fullmatch(r"([-+]?\d+(?:\.\d+)?)(.+)", token)
    if coefficient:
        return f"{coefficient.group(1)}{_latex_relation_token(coefficient.group(2))}"
    if len(token) >= 3 and token[0].isalpha() and token[1] in {"δ", "Δ", "∆"}:
        return f"{_latex_symbol(token[0])}{_latex_relation_token(token[1:])}"
    if token.startswith("d") and len(token) > 1:
        return rf"\mathrm{{d}}{_latex_symbol(token[1:])}"
    if token.startswith("δ") and len(token) > 1:
        return rf"\delta {_latex_symbol(token[1:])}"
    if token.startswith(("Δ", "∆")) and len(token) > 1:
        return rf"\Delta {_latex_symbol(token[1:])}"
    return _latex_atom(token) if "_" in token else _latex_variable(token)


def _latex_relation_operand(value: str) -> str:
    text = str(value).strip()
    if re.fullmatch(rf"{FORMULA_TOKEN_RE}\s*/\s*{FORMULA_TOKEN_RE}", text):
        numerator, denominator = [part.strip() for part in text.split("/", 1)]
        return rf"\frac{{{_latex_relation_token(numerator)}}}{{{_latex_relation_token(denominator)}}}"
    expr_parts = re.split(r"(\s*(?:[+\-−·*])\s*)", text)
    if len(expr_parts) > 1:
        converted: list[str] = []
        operators = {"−": "-", "·": r"\cdot", "*": r"\cdot"}
        for part in expr_parts:
            stripped = part.strip()
            if not stripped:
                continue
            if stripped in {"+", "-", "−", "·", "*"}:
                converted.append(operators.get(stripped, stripped))
            else:
                converted.append(_latex_relation_operand(stripped))
        return " ".join(converted)
    implicit_parts = [part for part in re.split(r"\s+", text) if part]
    if len(implicit_parts) > 1 and all(re.fullmatch(FORMULA_TOKEN_RE, part) for part in implicit_parts):
        return r" ".join(_latex_relation_token(part) for part in implicit_parts)
    return _latex_relation_token(text)


def _latex_relation(match: re.Match[str]) -> str:
    operators = {
        "≤": r"\le",
        ">=": r"\ge",
        "≥": r"\ge",
        "<=": r"\le",
        "<": "<",
        ">": ">",
    }
    left = _latex_relation_operand(match.group(1))
    operator = operators.get(match.group(2), match.group(2))
    right = _latex_relation_operand(match.group(3))
    return f"{left} {operator} {right}"


def _latex_symbolic_equation(match: re.Match[str]) -> str:
    left = _latex_relation_operand(match.group(1))
    right = _latex_relation_operand(match.group(2))
    return f"{left}={right}"


def _latex_ratio_equivalence(match: re.Match[str]) -> str:
    left = re.sub(r"\s*:\s*", ":", match.group(1))
    right = re.sub(r"\s*:\s*", ":", match.group(3))
    operator = r"\approx" if match.group(2) == "≈" else "="
    return f"{left}{operator}{right}"


def _latex_crystallographic_index(value: str) -> str:
    """Normalize a prose crystallographic index into renderable LaTeX."""

    raw = str(value or "")
    if raw.startswith("({") and raw.endswith("})"):
        return f"({raw[2:-2]})"
    if raw.startswith("[{") and raw.endswith("}]"):
        return f"[{raw[2:-2]}]"
    if raw.startswith("<{") and raw.endswith("}>"):
        return rf"\langle {raw[2:-2]} \rangle"
    if raw.startswith("(") and raw.endswith(")"):
        return raw
    if raw.startswith("[") and raw.endswith("]"):
        return raw
    if raw.startswith("<") and raw.endswith(">"):
        return rf"\langle {raw[1:-1]} \rangle"
    if raw.startswith("{") and raw.endswith("}"):
        return rf"\{{{raw[1:-1]}\}}"
    return raw


def _crystallographic_formula_candidates(text: str) -> list[tuple[int, int, str]]:
    return [
        (match.start(), match.end(), _latex_crystallographic_index(match.group(0)))
        for match in CRYSTALLOGRAPHIC_LATEX_RE.finditer(text)
    ]


def _answer_summary_formula_candidates(text: str) -> list[tuple[int, int, str]]:
    candidates: list[tuple[int, int, str]] = []
    for match in DOLLAR_LATEX_SUMMARY_RE.finditer(text):
        latex = match.group(1).strip()
        if latex:
            candidates.append((match.start(), match.end(), latex))
    for match in PARTIAL_DERIVATIVE_SUMMARY_RE.finditer(text):
        numerator = _latex_atom(match.group("numerator"))
        numerator = re.sub(r"^\\Delta(?=[A-Za-z])", r"\\Delta ", numerator)
        denominator = _latex_atom(match.group("denominator"))
        condition = _latex_atom(match.group("condition"))
        candidates.append(
            (
                match.start(),
                match.end(),
                rf"\left(\frac{{\partial {numerator}}}{{\partial {denominator}}}\right)_{{{condition}}}",
            )
        )
    for match in RATIO_EQUIVALENCE_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_ratio_equivalence(match)))
    for match in REACTION_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_reaction(match.group(0))))
    for match in SQRT_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_sqrt(match.group(0))))
    for match in RELATION_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_relation(match)))
    for match in SYMBOLIC_EQUATION_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_symbolic_equation(match)))
    for match in EQUATION_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_equation(match)))
    for match in UNICODE_LARGE_OPERATOR_RE.finditer(text):
        operator = r"\prod" if match.group("operator") == "Π" else r"\sum"
        subscript = match.group("braced") or match.group("plain") or ""
        latex = operator + (rf"_{{{_latex_symbol(subscript)}}}" if subscript else "")
        candidates.append((match.start(), match.end(), latex))
    for match in SCRIPTED_SUMMARY_RE.finditer(text):
        candidates.append((match.start(), match.end(), _latex_atom(match.group(0))))
    candidates.extend(_crystallographic_formula_candidates(text))
    candidates.extend(
        (plan.start, plan.end, plan.render_latex)
        for plan in build_text_expression_render_plans(text)
    )
    # Some domain transition chains contain labels such as G.P.区 or phase
    # names that intentionally remain prose.  Promote any arrow not already
    # covered by a richer reaction candidate so it is still a real Word math
    # object rather than a forbidden symbol in ordinary text.
    candidates.extend(
        (
            match.start(),
            match.end(),
            r"\to" if match.group(0) == "→" else r"\rightleftharpoons",
        )
        for match in ARROW_SUMMARY_RE.finditer(text)
    )
    candidates.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    selected: list[tuple[int, int, str]] = []
    cursor = -1
    for start, end, latex in candidates:
        if start < cursor:
            continue
        # A relation embedded in explanatory parentheses can be recognized as
        # ``(Delta U=Q+W`` while the closing parenthesis correctly remains
        # prose after the mathematical span.  Keep unmatched boundary
        # punctuation in prose; balanced mathematical groups such as
        # ``(TS)=...`` are unchanged.
        while latex.startswith("(") and latex.count("(") > latex.count(")"):
            latex = latex[1:].lstrip()
            start += 1
        while latex.endswith(")") and latex.count(")") > latex.count("("):
            latex = latex[:-1].rstrip()
            end -= 1
        if not latex:
            continue
        selected.append((start, end, latex))
        cursor = end
    return selected


def add_answer_summary_paragraph(doc: Document, answer_summary: str, size: float = 11, strict_formula_audit: bool = True):
    p = doc.add_paragraph()
    set_para(p)
    prefix = p.add_run("答：")
    set_run_font(prefix, size=size)
    text = normalize_answer_hierarchy_markers(strip_structured_math_metadata(answer_summary))
    cursor = 0
    for start, end, latex in _answer_summary_formula_candidates(text):
        plain = text[cursor:start]
        if plain:
            if strict_formula_audit and looks_like_symbolic_formula(plain):
                raise ValueError(f"Formula-like text remained in answer summary: {plain[:120]}")
            run = p.add_run(plain)
            set_run_font(run, size=size)
        p._p.append(render_expression_omml(latex, display=False, location="answer_summary"))
        cursor = end
    tail = text[cursor:]
    if tail:
        if strict_formula_audit and looks_like_symbolic_formula(tail):
            raise ValueError(f"Formula-like text remained in answer summary: {tail[:120]}")
        run = p.add_run(tail)
        set_run_font(run, size=size)
    return p


def _append_formula_text_runs(
    p,
    text: str,
    size: float = 11,
    strict_formula_audit: bool = True,
    audit_label: str = "text",
):
    value = normalize_answer_hierarchy_markers(strip_structured_math_metadata(text))
    cursor = 0
    for start, end, latex in _answer_summary_formula_candidates(value):
        plain = value[cursor:start]
        if plain:
            if strict_formula_audit and looks_like_symbolic_formula(plain):
                raise ValueError(f"Formula-like text remained in {audit_label}: {plain[:120]}")
            run = p.add_run(plain)
            set_run_font(run, size=size)
        p._p.append(render_expression_omml(latex, display=False, location=audit_label))
        cursor = end
    tail = value[cursor:]
    if tail:
        if strict_formula_audit and looks_like_symbolic_formula(tail):
            raise ValueError(f"Formula-like text remained in {audit_label}: {tail[:120]}")
        run = p.add_run(tail)
        set_run_font(run, size=size)
    return p


def add_labeled_formula_text_paragraph(
    doc: Document,
    label: str,
    text: str,
    size: float = 11,
    strict_formula_audit: bool = True,
):
    p = doc.add_paragraph()
    set_para(p)
    prefix = p.add_run(f"{label}：")
    set_run_font(prefix, size=size, bold=True)
    return _append_formula_text_runs(p, text, size=size, strict_formula_audit=strict_formula_audit, audit_label=label)


def add_formula_text_body_paragraph(
    doc: Document,
    text: str,
    size: float = 11,
    strict_formula_audit: bool = True,
):
    p = doc.add_paragraph()
    set_para(p)
    p.paragraph_format.first_line_indent = Cm(ANSWER_BODY_FIRST_LINE_INDENT_CM)
    return _append_formula_text_runs(p, text, size=size, strict_formula_audit=strict_formula_audit, audit_label="答案")


def add_mixed_paragraph(
    doc: Document,
    segments: list[dict],
    formulas: dict[str, dict],
    label: str = "",
    base_dir: Path | None = None,
    initial_paragraph=None,
    force_skip_formula_text_audit: bool = False,
):
    skip_formula_text_audit = str(label).strip() == "教材依据" or force_skip_formula_text_audit
    highlight_label = str(label).strip() == "待复核公式"
    p = initial_paragraph
    label_pending = bool(label) and initial_paragraph is None

    def ensure_text_paragraph():
        nonlocal p, label_pending
        if p is None:
            p = doc.add_paragraph()
            set_para(p)
            if label_pending and label:
                r = p.add_run(f"{label}：")
                set_run_font(r, bold=True)
                if highlight_label:
                    set_run_shading(r)
                label_pending = False
            if str(label).strip() == "易错点及注意事项":
                # Supporting notes use a slightly tighter rhythm so a short
                # tail does not create an almost-empty final page.
                p.paragraph_format.line_spacing = TEXT_CONTRACT.note_line_spacing
        return p

    def break_text_paragraph():
        nonlocal p
        p = None

    def add_text_run(text: str, highlight: bool = False):
        nonlocal p
        if not text:
            return
        normalized_text = normalize_answer_hierarchy_markers(strip_structured_math_metadata(text))
        # Models often line-wrap dense inline lists such as crystallographic
        # indices after every delimiter.  Treat those wraps as whitespace so
        # a short scientific list does not become one paragraph per item.
        normalized_text = INLINE_SCIENTIFIC_LIST_BREAK_RE.sub("", normalized_text)
        chunks = re.split(r"(\n+)", normalized_text)
        for chunk in chunks:
            if not chunk:
                continue
            if "\n" in chunk:
                break_text_paragraph()
                continue
            if p is not None and TOP_LEVEL_SUBQUESTION_RE.match(chunk):
                break_text_paragraph()
            paragraph = ensure_text_paragraph()
            cursor = 0
            for start, end, latex in _crystallographic_formula_candidates(chunk):
                if start > cursor:
                    append_domain_text_runs(paragraph, chunk[cursor:start], highlight=highlight)
                paragraph._p.append(render_expression_omml(latex, display=False, location="mixed_text"))
                cursor = end
            if cursor < len(chunk):
                append_domain_text_runs(paragraph, chunk[cursor:], highlight=highlight)

    def render_inline_formula(segment: dict, formula: dict) -> bool:
        if bool(segment.get("inline")):
            return True
        # Compatibility for fragments created before promoted sentence tokens
        # were explicitly marked inline. The provenance is deterministic and
        # distinguishes them from intentional standalone equations.
        source_note = str(formula.get("source_note") or "")
        return "Word 生成前从普通文本中识别出的公式片段" in source_note

    for seg in segments:
        typ = seg.get("type")
        if typ == "text":
            text = str(seg.get("text", ""))
            if not text:
                continue
            add_text_run(text, str(seg.get("highlight") or "") == "unconfirmed_evidence")
        elif typ == "formula_ref":
            fid = str(seg.get("formula_id", ""))
            formula = formulas.get(fid)
            if not formula:
                raise ValueError(f"Missing formula for formula_ref: {fid}")
            latex = str(formula.get("latex", ""))
            if render_inline_formula(seg, formula):
                ensure_text_paragraph()._p.append(render_expression_omml(latex, display=False, location="formula_ref"))
            elif bool(formula.get("display", True)) and not skip_formula_text_audit:
                break_text_paragraph()
                add_formula_paragraph(doc, latex)
            else:
                ensure_text_paragraph()._p.append(render_expression_omml(latex, display=False, location="formula_ref"))
        elif typ == "image_ref":
            image_path = Path(str(seg.get("path") or ""))
            if base_dir and not image_path.is_absolute():
                image_path = base_dir / image_path
            break_text_paragraph()
            if image_path.is_file():
                pic_p = doc.add_paragraph()
                set_para(pic_p, WD_ALIGN_PARAGRAPH.CENTER)
                add_figure_picture(pic_p, image_path)
            else:
                image_id = str(seg.get("image_id") or image_path)
                raise FileNotFoundError(f"Required Word image is missing: {image_id} ({image_path})")
        else:
            raise ValueError(f"Unsupported segment type: {typ}")
    if label_pending and label:
        ensure_text_paragraph()
    return p


def add_split_block(doc: Document, segments: list[dict], formulas: dict[str, dict], label: str = "", base_dir: Path | None = None):
    if label:
        add_text_paragraph(doc, f"{label}：", bold=True)
    skip_promoted_tail = False
    previous_was_formula = False
    for seg in segments:
        typ = seg.get("type")
        if typ == "text":
            text = str(seg.get("text", "")).strip()
            if previous_was_formula and re.fullmatch(r"[。．.，,；;：:、!?！？]+", text):
                skip_promoted_tail = False
                previous_was_formula = False
                continue
            if skip_promoted_tail and re.fullmatch(
                r"(?:(?:总|隔离|系统|环境|环|外|内)+\s*=\s*)?"
                r"[=<>{}\[\]()（）A-Za-z0-9⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻+\-−*/×·.,。，、\s]+",
                text,
            ):
                skip_promoted_tail = False
                previous_was_formula = False
                continue
            skip_promoted_tail = False
            if text:
                add_text_paragraph(doc, text)
            previous_was_formula = False
        elif typ == "formula_ref":
            # Inline promotion can leave a duplicate fragment immediately after
            # the authoritative display result (for example display ``b=a<100>``
            # followed by inline ``b=a`` + plain ``<100>``). Do not render that
            # duplicate or let the following step heading stick to it.
            if bool(seg.get("inline")):
                skip_promoted_tail = True
                previous_was_formula = True
                continue
            fid = str(seg.get("formula_id", ""))
            formula = formulas.get(fid)
            if not formula:
                raise ValueError(f"Missing formula for formula_ref: {fid}")
            source_note = str(formula.get("source_note") or formula.get("meaning") or "")
            if bool(formula.get("_program_mirrored_from_contract")) or "程序从计算结果账本中镜像" in source_note:
                # The ledger mirror is machine-audit evidence, not a second
                # student-facing result.  Older checkpoints only carry the
                # source_note marker, so keep that migration path deterministic.
                continue
            add_formula_paragraph(doc, str(formula.get("latex", "")))
            previous_was_formula = True
        elif typ == "image_ref":
            image_path = Path(str(seg.get("path") or ""))
            if base_dir and not image_path.is_absolute():
                image_path = base_dir / image_path
            if image_path.is_file():
                pic_p = doc.add_paragraph()
                set_para(pic_p, WD_ALIGN_PARAGRAPH.CENTER)
                add_figure_picture(pic_p, image_path)
            else:
                image_id = str(seg.get("image_id") or image_path)
                raise FileNotFoundError(f"Required Word image is missing: {image_id} ({image_path})")
            previous_was_formula = False
        else:
            raise ValueError(f"Unsupported segment type: {typ}")


def add_formula_paragraph(doc: Document, latex: str):
    p = doc.add_paragraph()
    set_para(p, WD_ALIGN_PARAGRAPH.CENTER)
    p._p.append(render_expression_omml(latex, display=True, location="formula_paragraph"))
    return p


def safe_warning_text(warnings: list) -> str:
    text = "；".join(str(x) for x in warnings)
    if looks_like_formula(text):
        return "该题存在需人工复核提示，详细内容请查看结构化答案审计信息。"
    return text


def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    # fldSimple must be a direct paragraph child. Nesting it in w:r is invalid
    # WordprocessingML and leaves the visible page number blank in Word.
    paragraph._p.append(fld)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def configure_compatible_font_table(doc: Document) -> None:
    """Declare macOS/compatible-reader alternates without changing Word fonts."""

    font_table_part = doc.part.part_related_by(RT.FONT_TABLE)
    font_table = parse_xml(font_table_part.blob)
    mappings = (
        (TEXT_CONTRACT.east_asia_font, TEXT_CONTRACT.east_asia_fallback_font, "roman"),
        (HEADER_FOOTER_CONTRACT.header_font, HEADER_FOOTER_CONTRACT.header_fallback_font, "swiss"),
        (HEADER_FOOTER_CONTRACT.footer_font, HEADER_FOOTER_CONTRACT.footer_fallback_font, "decorative"),
    )
    by_name = {node.get(qn("w:name")): node for node in font_table.findall(qn("w:font"))}
    for primary, fallback, family_name in mappings:
        font = by_name.get(primary)
        if font is None:
            font = OxmlElement("w:font")
            font.set(qn("w:name"), primary)
            font_table.append(font)
        for tag in ("w:altName", "w:charset", "w:family"):
            for child in font.findall(qn(tag)):
                font.remove(child)
        alt_name = OxmlElement("w:altName")
        alt_name.set(qn("w:val"), fallback)
        charset = OxmlElement("w:charset")
        charset.set(qn("w:val"), "86")
        family = OxmlElement("w:family")
        family.set(qn("w:val"), family_name)
        font.extend([alt_name, charset, family])
    font_table_part._blob = etree.tostring(font_table, xml_declaration=True, encoding="UTF-8", standalone=True)


def setup_document() -> Document:
    doc = Document()
    enable_field_updates(doc)
    configure_compatible_font_table(doc)
    section = doc.sections[0]
    section.page_width = Cm(PAGE_CONTRACT.width_cm)
    section.page_height = Cm(PAGE_CONTRACT.height_cm)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Cm(PAGE_CONTRACT.margin_cm))
    section.header_distance = Cm(PAGE_CONTRACT.header_distance_cm)
    section.footer_distance = Cm(PAGE_CONTRACT.footer_distance_cm)
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_p.add_run(HEADER_TEXT)
    set_run_font(
        r,
        east=HEADER_FOOTER_CONTRACT.header_font,
        size=HEADER_FOOTER_CONTRACT.header_size_pt,
        bold=True,
    )
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(FOOTER_TEXT)
    set_run_font(
        r,
        east=HEADER_FOOTER_CONTRACT.footer_font,
        size=HEADER_FOOTER_CONTRACT.footer_size_pt,
    )
    page_p = section.footer.add_paragraph()
    page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    start = page_p.add_run("第 ")
    set_run_font(start, size=10.5)
    add_page_field(page_p)
    end = page_p.add_run(" 页")
    set_run_font(end, size=10.5)
    return doc


def section_display_title(raw: str) -> str:
    return raw or "解析"


def _section_item(raw) -> dict:
    return raw if isinstance(raw, dict) else {"section": raw}


def is_choice_section(raw) -> bool:
    return is_choice_question(_section_item(raw))


def is_calculation_section(raw) -> bool:
    item = _section_item(raw)
    return (explicit_question_type(item) or infer_question_type(item)) == "计算题"


def is_short_answer_section(raw) -> bool:
    return is_short_answer_question(_section_item(raw))


def is_term_explanation_section(raw) -> bool:
    return is_term_explanation_question(_section_item(raw))


def is_graphic_section(raw) -> bool:
    item = _section_item(raw)
    # A calculation/short-answer parent may contain one drawing leaf. Keep the
    # parent renderer so its formulas and steps are not lost; figure blocks are
    # already rendered within that branch.
    return (explicit_question_type(item) or infer_question_type(item)) == "作图题"


def is_fill_section(raw) -> bool:
    return question_has_type(_section_item(raw), "填空题")


def should_hide_top_answer(section: str) -> bool:
    return is_calculation_section(section) or is_short_answer_section(section)


def should_show_answer_summary(section: str, answer: str, answer_summary: str) -> bool:
    if should_hide_top_answer(section):
        return False
    if is_choice_section(section):
        return False
    if not answer_summary:
        return False
    if answer_summary in {"待复核", "待补充", "见解析"}:
        return False
    return answer_summary != answer or answer == "见解析"


def display_block_label(section: str, label: str) -> str:
    raw = str(label or "")
    if raw == "待复核公式":
        # Keep the diagnostic label in fragments/review reports, but never
        # expose internal workflow language in the student-facing document.
        return "补充公式"
    if is_calculation_section(section) and raw == "解题步骤":
        return "答案"
    return raw


def should_split_block(section: str, label: str) -> bool:
    return is_calculation_section(section) and str(label or "") == "解题步骤"


def _answer_text(answer: str, answer_summary: str) -> str:
    summary = str(answer_summary or "").strip()
    if summary and summary not in {"待复核", "待补充", "见解析"}:
        return summary
    return str(answer or "").strip() or "待复核"


def _find_block(fragment: dict, label: str) -> dict | None:
    for block in fragment.get("blocks", []):
        if str(block.get("label", "")) == label:
            return block
    return None


def _formal_heading_marker(heading: str) -> str:
    text = str(heading or "").strip()
    circled = re.match(r"^([①-⑳]、)", text)
    if circled:
        return circled.group(1)
    match = re.match(r"^[（(][^）)]+[）)]", text)
    return match.group(0) if match else ""


def _formal_heading_key(value: str) -> str:
    return re.sub(r"[：:；;。？！?]+$", "", re.sub(r"\s+", "", str(value or "").strip()))


def _formal_block_segments(fragment: dict, block: dict) -> list[dict]:
    """Remove repeated source-question headings from the formal answer only."""

    projection = block.get("delivery_projection")
    if isinstance(projection, dict):
        allowed_types = {
            str(value)
            for value in projection.get("segment_types", []) or []
            if str(value).strip()
        }
        if allowed_types:
            return [
                segment
                for segment in block.get("segments", []) or []
                if isinstance(segment, dict) and str(segment.get("type") or "") in allowed_types
            ]
    elif str(block.get("label") or "") == "待复核公式":
        # Backward-compatible migration path for checkpoints created before
        # blocks carried explicit review/delivery audience metadata.
        return [
            segment
            for segment in block.get("segments", []) or []
            if isinstance(segment, dict) and segment.get("type") == "formula_ref"
        ]

    heading_to_marker: dict[str, str] = {}
    for row in question_unit_rows(fragment):
        heading = str(row.get("heading") or "").strip()
        parent = str(row.get("parent_heading") or "").strip()
        if heading:
            marker = _formal_heading_marker(heading)
            if marker:
                heading_to_marker[_formal_heading_key(heading)] = marker
                source_heading = str(row.get("source_heading") or "").strip()
                if source_heading:
                    heading_to_marker[_formal_heading_key(source_heading)] = marker
        if parent:
            marker = _formal_heading_marker(parent)
            if marker:
                heading_to_marker[_formal_heading_key(parent)] = marker
        source_parent = str(row.get("source_parent_heading") or "").strip()
        if source_parent:
            heading_to_marker[_formal_heading_key(source_parent)] = ""
    result: list[dict] = []
    for segment in block.get("segments", []) or []:
        if not isinstance(segment, dict) or segment.get("type") != "text":
            result.append(segment)
            continue
        text = str(segment.get("text") or "")
        key = _formal_heading_key(text)
        if key in heading_to_marker:
            marker = heading_to_marker[key]
            if marker:
                result.append({**segment, "text": marker})
            continue
        result.append(segment)
    return result


def _add_block_if_present(doc: Document, fragment: dict, formulas: dict[str, dict], label: str, base_dir: Path) -> None:
    block = _find_block(fragment, label)
    if block is not None:
        add_mixed_paragraph(doc, block.get("segments", []), formulas, label, base_dir)


def _add_number_and_evidence(doc: Document, number: str, fragment: dict, formulas: dict[str, dict], base_dir: Path) -> None:
    p = doc.add_paragraph()
    set_para(p)
    number_run = p.add_run(f"{number}、")
    set_run_font(number_run, size=11)
    block = _find_block(fragment, "教材依据")
    if block is None:
        return
    # Evidence gaps remain durable in evidence_selection.json and review
    # reports, but are process diagnostics rather than student-facing answer
    # content.  Render only confirmed citation segments in the deliverable.
    confirmed_segments = [
        segment
        for segment in block.get("segments", [])
        if isinstance(segment, dict)
        and str(segment.get("highlight") or "") != "unconfirmed_evidence"
        and str(segment.get("text") or "").strip("；; ")
    ]
    while confirmed_segments and str(confirmed_segments[-1].get("text") or "").strip() in {"；", ";"}:
        confirmed_segments.pop()
    if not confirmed_segments:
        return
    label_run = p.add_run("教材依据：")
    set_run_font(label_run, size=11, bold=True)
    add_mixed_paragraph(
        doc,
        confirmed_segments,
        formulas,
        "",
        base_dir,
        initial_paragraph=p,
        force_skip_formula_text_audit=True,
    )


def _add_ordered_answer_units(doc: Document, fragment: dict, formulas: dict[str, dict], base_dir: Path) -> bool:
    plan = plan_ordered_answer_units(fragment)
    if not plan.get("ok"):
        return False
    add_text_paragraph(doc, "解析：", bold=True, size=11)
    last_parent = ""
    for unit in plan.get("units", []):
        parent_number = str(unit.get("parent_number") or "")
        if parent_number and parent_number != last_parent:
            parent = add_text_paragraph(
                doc,
                _formal_heading_marker(str(unit.get("parent_heading") or "")),
                bold=True,
                size=11,
            )
            parent.paragraph_format.keep_with_next = True
            last_parent = parent_number
        heading = add_text_paragraph(
            doc,
            _formal_heading_marker(str(unit.get("heading") or "")),
            bold=not parent_number,
            size=11,
        )
        heading.paragraph_format.keep_with_next = True
        figure_segments = unit.get("figure_segments", [])
        if figure_segments:
            add_mixed_paragraph(doc, figure_segments, formulas, "图示", base_dir)
        analysis_segments = unit.get("analysis_segments", [])
        if analysis_segments:
            add_mixed_paragraph(doc, analysis_segments, formulas, "", base_dir)
        step_segments = unit.get("step_segments", [])
        if step_segments:
            add_split_block(doc, step_segments, formulas, "解题步骤", base_dir)
    return True


def _add_answer_text(doc: Document, answer_text: str, formulas: dict[str, dict], strict_formula_audit: bool, base_dir: Path) -> None:
    text = str(answer_text or "")
    if "$" in text or _answer_summary_formula_candidates(text):
        add_labeled_formula_text_paragraph(doc, "答案", text, size=11, strict_formula_audit=strict_formula_audit)
        return
    add_mixed_paragraph(
        doc,
        [{"type": "text", "text": text}],
        formulas,
        "答案",
        base_dir,
    )


def _add_indented_answer_text(doc: Document, answer_text: str, strict_formula_audit: bool) -> None:
    add_text_paragraph(doc, "答案：", bold=True, size=11)
    text = normalize_answer_hierarchy_markers(answer_text)
    if "$" in text or _answer_summary_formula_candidates(text):
        add_formula_text_body_paragraph(doc, text, size=11, strict_formula_audit=strict_formula_audit)
        return
    p = doc.add_paragraph()
    set_para(p)
    p.paragraph_format.first_line_indent = Cm(ANSWER_BODY_FIRST_LINE_INDENT_CM)
    append_domain_text_runs(p, text)


def build_docx_from_fragments(fragments_json: Path, output_docx: Path, *, strict_answer_summary_formula_audit: bool = True) -> Path:
    data = json.loads(fragments_json.read_text(encoding="utf-8"))
    doc = setup_document()
    add_text_paragraph(
        doc,
        "真题答案解析",
        bold=True,
        size=TEXT_CONTRACT.title_size_pt,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    current_section = ""
    for fragment in data.get("fragments", []):
        section = str(fragment.get("section", ""))
        if section and section != current_section:
            current_section = section
            add_text_paragraph(doc, section_display_title(section), bold=True, size=11)
        qid = str(fragment.get("question_id", "")).replace("_", "-")
        number = str(fragment.get("display_number") or fragment.get("number") or qid)
        formulas = {str(f.get("formula_id")): f for f in fragment.get("formulas", [])}
        answer = str(fragment.get("answer", ""))
        answer_summary = str(fragment.get("answer_summary", "")).strip()
        section_context = {
            "section": section,
            "question_type": fragment.get("question_type") or "",
            "subquestions": fragment.get("subquestions") or [],
        }
        if is_term_explanation_section(section_context):
            _add_number_and_evidence(doc, number, fragment, formulas, fragments_json.parent)
            _add_indented_answer_text(doc, _answer_text(answer, answer_summary), strict_answer_summary_formula_audit)
            continue
        if is_graphic_section(section_context):
            _add_number_and_evidence(doc, number, fragment, formulas, fragments_json.parent)
            _add_indented_answer_text(doc, _answer_text(answer, answer_summary), strict_answer_summary_formula_audit)
            # Preserve the original artifact before the generated answer figure
            # so users can compare the source condition with the solution.
            for label in ("原题图", "图示", "解析", "易错点及注意事项"):
                _add_block_if_present(doc, fragment, formulas, label, fragments_json.parent)
            continue
        if is_short_answer_section(section_context):
            _add_number_and_evidence(doc, number, fragment, formulas, fragments_json.parent)
            _add_indented_answer_text(doc, _answer_text(answer, answer_summary), strict_answer_summary_formula_audit)
            # A nominal short-answer question may contain calculation or
            # drawing leaves.  When every leaf can be mapped safely, preserve
            # the source unit order instead of emitting all prose first and a
            # later calculation block out of sequence.
            if _add_ordered_answer_units(doc, fragment, formulas, fragments_json.parent):
                for block in fragment.get("blocks", []):
                    label = str(block.get("label", ""))
                    if label in {"教材依据", "图示", "解析", "解题步骤", "待复核公式"}:
                        continue
                    add_mixed_paragraph(doc, block.get("segments", []), formulas, label, fragments_json.parent)
                continue
            for block in fragment.get("blocks", []):
                label = str(block.get("label", ""))
                if label in {"教材依据", "答案"}:
                    continue
                if label == "待复核公式":
                    segments = [
                        segment
                        for segment in block.get("segments", [])
                        if isinstance(segment, dict) and segment.get("type") == "formula_ref"
                    ]
                    if segments:
                        add_mixed_paragraph(doc, segments, formulas, "补充公式", fragments_json.parent)
                    continue
                if label == "解题步骤":
                    add_split_block(doc, _formal_block_segments(fragment, block), formulas, label, fragments_json.parent)
                    continue
                add_mixed_paragraph(doc, _formal_block_segments(fragment, block), formulas, label, fragments_json.parent)
            continue
        if is_calculation_section(section_context):
            _add_number_and_evidence(doc, number, fragment, formulas, fragments_json.parent)
            if _add_ordered_answer_units(doc, fragment, formulas, fragments_json.parent):
                for block in fragment.get("blocks", []):
                    label = str(block.get("label", ""))
                    if label in {"教材依据", "图示", "解析", "解题步骤"}:
                        continue
                    if label == "待复核公式":
                        continue
                    add_mixed_paragraph(doc, block.get("segments", []), formulas, label, fragments_json.parent)
                continue
            for block in fragment.get("blocks", []):
                label = str(block.get("label", ""))
                if label == "教材依据":
                    continue
                display_label = display_block_label(section_context, label)
                if should_split_block(section_context, label):
                    add_split_block(
                        doc,
                        _formal_block_segments(fragment, block),
                        formulas,
                        display_label,
                        fragments_json.parent,
                    )
                else:
                    add_mixed_paragraph(
                        doc,
                        _formal_block_segments(fragment, block),
                        formulas,
                        display_label,
                        fragments_json.parent,
                    )
            continue
        fill_answer_text = _answer_text(answer, answer_summary)
        if is_fill_section(section_context) and ("$" in fill_answer_text or _answer_summary_formula_candidates(fill_answer_text)):
            add_text_paragraph(doc, f"{number}、", size=11)
            _add_answer_text(doc, fill_answer_text, formulas, strict_answer_summary_formula_audit, fragments_json.parent)
            for block in fragment.get("blocks", []):
                add_mixed_paragraph(
                    doc,
                    _formal_block_segments(fragment, block),
                    formulas,
                    display_block_label(section_context, str(block.get("label", ""))),
                    fragments_json.parent,
                )
            continue
        if should_hide_top_answer(section_context):
            add_text_paragraph(doc, f"{number}、", size=11)
        else:
            add_text_paragraph(doc, f"{number}、{answer}", size=11)
        if should_show_answer_summary(section_context, answer, answer_summary):
            add_answer_summary_paragraph(doc, answer_summary, size=11, strict_formula_audit=strict_answer_summary_formula_audit)
        for block in fragment.get("blocks", []):
            label = str(block.get("label", ""))
            display_label = display_block_label(section_context, label)
            if should_split_block(section_context, label):
                add_split_block(
                    doc,
                    _formal_block_segments(fragment, block),
                    formulas,
                    display_label,
                    fragments_json.parent,
                )
            else:
                add_mixed_paragraph(
                    doc,
                    _formal_block_segments(fragment, block),
                    formulas,
                    display_label,
                    fragments_json.parent,
                )
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx
