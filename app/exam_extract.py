from __future__ import annotations

import json
import posixpath
import re
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

from .capabilities.catalog import capability_policy_contributions
from .formula_audit import looks_like_formula
from .omml_input import mixed_text_with_structured_math
from .question_requirements import answer_figure_required, source_image_required
from .text_utils import clean_text, cn_to_int

SECTION_TITLE_PREFIX = r"(?:选择题|判断题|正误题|填空题|名词解释题|名词解释|名解题|简答题|问答题|计算题|回答下列问题)"
SECTION_RE = re.compile(
    rf"^([一二三四五六七八九十]+)(?:\s*、\s*|\s*[.．]\s*(?={SECTION_TITLE_PREFIX})|\s+(?={SECTION_TITLE_PREFIX}))(.*)"
)
ITEM_RE = re.compile(r"^(\d{1,2})([、.．])\s*(.*)")
INLINE_ITEM_MARKER_RE = re.compile(
    r"(?<!\d)(?:[（(]\s*(?P<paren_number>\d{1,2})\s*[）)]|(?P<plain_number>\d{1,2})(?:[、．]|\.(?!\d)))\s*"
)
MULTIPART_CUE_RE = re.compile(
    r"(回答下列|回答以下|完成下列|完成以下|按要求|分别|逐项|各小问|小问|问题|如下|试求|计算下列|求下列|试述下列|说明下列|分析下列|作图|画出|绘制)"
)
PAREN_SUBQUESTION_RE = re.compile(r"^[（(]\s*(\d{1,2})\s*[）)]\s*(.*)")
ORDINAL_SUBQUESTION_RE = re.compile(r"^第\s*([一二三四五六七八九十\d]{1,3})\s*(?:小?问|题)\s*[:：、.．]?\s*(.*)")
UNNUMBERED_SECTION_RE = re.compile(
    r"^(选择题|判断题|正误题|填空题|名词解释题|名词解释|名解题|简答题|问答题|计算题|回答下列问题)\s*[（(].*(?:本题|每小题|共|分).*[）)]?\s*$"
)
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}
IMAGE_MARKER_PREFIX = "__ANSWER_BOOK_IMAGE__:"
TABLE_MARKER_PREFIX = "__ANSWER_BOOK_TABLE__:"
EXAM_GROUPING_POLICY_VERSION = "answer_book.exam_grouping.v8"
RESPONSE_CONSTRAINT_RE = re.compile(
    r"(?:"
    r"(?:计算|作答|答案|结果|数值)[^。；;？?]{0,30}(?:取|保留|精确到|写成|表示为)[^。；;？?]{0,24}(?:有效数字|小数|位|形式)?|"
    r"(?:取|保留)[一二三四五六七八九十\d]+位有效数字|"
    r"(?:单位|量纲)[^。；;？?]{0,16}(?:统一|采用|使用|写为)"
    r")"
)
ANSWER_ACTION_RE = re.compile(
    r"(?:求(?:出|得|解)?|计算|判断|说明|简述|论述|分析|比较|解释|写出|列出|证明|推导|画出|绘制|作图|标出)"
)


def int_to_cn(value: int) -> str:
    digits = "一二三四五六七八九"
    if 1 <= value <= 9:
        return digits[value - 1]
    if value == 10:
        return "十"
    if 11 <= value <= 19:
        return "十" + digits[value - 11]
    if 20 <= value <= 99:
        tens, ones = divmod(value, 10)
        return digits[tens - 1] + "十" + (digits[ones - 1] if ones else "")
    return str(value)


def _subject_title(text: str) -> str:
    clean = clean_text(text).strip("“”\"'")
    clean = re.sub(r"\s*部分\s*$", "", clean).strip()
    if "部分" in text and 2 <= len(clean) <= 20:
        return clean
    return ""


def _implicit_subject_title(text: str) -> str:
    """Recognize a subject heading only from its position before a section."""

    clean = clean_text(text).strip("“”\"'")
    if not 2 <= len(clean) <= 30:
        return ""
    if re.search(r"[。；;，,！？!?：:\d]", clean):
        return ""
    if ITEM_RE.match(clean) or PAREN_SUBQUESTION_RE.match(clean) or _unnumbered_section_title(clean):
        return ""
    return clean


def _unnumbered_section_title(text: str) -> bool:
    return bool(UNNUMBERED_SECTION_RE.match(clean_text(text)))


def _document_relationships(zf: ZipFile) -> dict[str, str]:
    try:
        root = etree.fromstring(zf.read("word/_rels/document.xml.rels"))
    except KeyError:
        return {}
    out: dict[str, str] = {}
    for rel in root.xpath(".//*[local-name()='Relationship']"):
        rid = str(rel.get("Id") or "").strip()
        target = str(rel.get("Target") or "").strip()
        if rid and target:
            out[rid] = target
    return out


def _media_zip_path(target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    if target.startswith("word/"):
        return posixpath.normpath(target)
    return posixpath.normpath(posixpath.join("word", target))


def _table_rows_from_xml(tbl) -> list[list[str]]:
    rows: list[list[str]] = []
    for tr in tbl.xpath("./w:tr", namespaces=NS):
        row: list[str] = []
        for tc in tr.xpath("./w:tc", namespaces=NS):
            row.append(clean_text(mixed_text_with_structured_math(tc).text))
        if any(cell for cell in row):
            rows.append(row)
    return rows


def _table_rows_from_docx_table(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _table_to_text(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    rendered_rows = [" | ".join(cell for cell in row if cell) for row in rows]
    return "表格：" + "；".join(row for row in rendered_rows if row)


def _table_marker(rows: list[list[str]]) -> str:
    return TABLE_MARKER_PREFIX + json.dumps({"rows": rows}, ensure_ascii=False, separators=(",", ":"))


def _table_from_marker(line: str) -> dict:
    if not str(line).startswith(TABLE_MARKER_PREFIX):
        return {}
    try:
        payload = json.loads(str(line)[len(TABLE_MARKER_PREFIX) :])
    except json.JSONDecodeError:
        return {}
    rows = payload.get("rows") if isinstance(payload, dict) else None
    if not isinstance(rows, list):
        return {}
    clean_rows = []
    for row in rows:
        if isinstance(row, list):
            clean_row = [clean_text(str(cell)) for cell in row]
            if any(clean_row):
                clean_rows.append(clean_row)
    if not clean_rows:
        return {}
    return {"rows": clean_rows, "text": _table_to_text(clean_rows)}


def _source_line_text(line: str) -> str:
    if str(line).startswith(TABLE_MARKER_PREFIX):
        return _table_from_marker(line).get("text", "")
    return str(line)


def _docx_paragraph_lines(path: Path, image_dir: Path | None = None) -> list[str]:
    lines: list[str] = []
    image_index = 0
    with ZipFile(path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
        relationships = _document_relationships(zf)
        for child in root.xpath("./w:body/*", namespaces=NS):
            local_name = etree.QName(child).localname
            if local_name == "p":
                text = clean_text(mixed_text_with_structured_math(child).text)
                if text:
                    lines.append(text)
                if image_dir is None:
                    continue
                for rid in child.xpath(".//a:blip/@r:embed", namespaces=NS):
                    target = relationships.get(str(rid))
                    if not target:
                        continue
                    media_path = _media_zip_path(target)
                    try:
                        image_bytes = zf.read(media_path)
                    except KeyError:
                        continue
                    image_index += 1
                    suffix = Path(media_path).suffix or ".png"
                    image_dir.mkdir(parents=True, exist_ok=True)
                    output = image_dir / f"source_image_{image_index:03d}{suffix}"
                    output.write_bytes(image_bytes)
                    lines.append(f"{IMAGE_MARKER_PREFIX}{output}")
            elif local_name == "tbl":
                rows = _table_rows_from_xml(child)
                if rows:
                    lines.append(_table_marker(rows))
    if lines:
        return lines
    doc = Document(path)
    fallback_lines = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    for table in doc.tables:
        rows = _table_rows_from_docx_table(table)
        if rows:
            fallback_lines.append(_table_marker(rows))
    return fallback_lines


def docx_paragraph_texts(path: Path) -> list[str]:
    out: list[str] = []
    for line in _docx_paragraph_lines(path):
        if line.startswith(IMAGE_MARKER_PREFIX):
            continue
        text = _source_line_text(line)
        if text:
            out.append(text)
    return out


def _snapshot_font(size: int):
    try:
        from PIL import ImageFont

        for path in (
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Helvetica.ttc",
        ):
            font_path = Path(path)
            if font_path.exists():
                return ImageFont.truetype(str(font_path), size)
        return ImageFont.load_default()
    except Exception:
        return None


def _wrap_snapshot_text(draw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in str(text or "").splitlines():
        paragraph = clean_text(raw)
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for char in paragraph:
            candidate = current + char
            try:
                width = draw.textlength(candidate, font=font)
            except Exception:
                width = len(candidate) * 14
            if current and width > max_width:
                lines.append(current)
                current = char
            else:
                current = candidate
        if current:
            lines.append(current)
    return lines


def _snapshot_table_text(item: dict) -> list[str]:
    out: list[str] = []
    for table in (item.get("attachments") or {}).get("tables") or []:
        if not isinstance(table, dict):
            continue
        text = clean_text(str(table.get("text") or ""))
        if text:
            out.append(text)
    return out


def _safe_snapshot_name(item: dict, index: int) -> str:
    qid = str(item.get("question_id") or f"question_{index:03d}")
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", qid).strip("._") or f"question_{index:03d}"
    return f"{safe}.png"


def _write_question_snapshot(item: dict, snapshot_dir: Path, index: int) -> str:
    """Render a readable whole-question preview for the structure-review UI."""
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return ""

    width = 1180
    padding = 44
    gap = 18
    max_content_width = width - padding * 2
    title_font = _snapshot_font(34)
    meta_font = _snapshot_font(24)
    body_font = _snapshot_font(28)
    small_font = _snapshot_font(22)
    if not all([title_font, meta_font, body_font, small_font]):
        return ""

    scratch = Image.new("RGB", (width, 200), "white")
    draw = ImageDraw.Draw(scratch)
    title = f"{item.get('number') or index}  {item.get('section') or item.get('section_raw') or ''}".strip()
    text_parts = [str(item.get("stem") or "").strip()]
    text_parts.extend(_snapshot_table_text(item))
    wrapped_text: list[tuple[str, object]] = []
    for part_index, part in enumerate(text_parts):
        if not part:
            continue
        if part_index:
            wrapped_text.append(("", body_font))
        for line in _wrap_snapshot_text(draw, part, body_font, max_content_width):
            wrapped_text.append((line, body_font))

    rendered_images: list[tuple[str, object]] = []
    for raw in item.get("image_refs") or []:
        path = Path(str(raw))
        if not path.exists() or not path.is_file():
            continue
        try:
            image = Image.open(path).convert("RGB")
            image.thumbnail((max_content_width, 620))
            rendered_images.append((path.name, image.copy()))
        except Exception:
            rendered_images.append((path.name, None))

    height = padding
    height += 44 + gap
    if item.get("question_id"):
        height += 30 + gap
    for line, _font in wrapped_text:
        height += (34 if line else 18) + 6
    if rendered_images:
        height += gap
    for _name, image in rendered_images:
        if image is not None:
            height += image.height + 58
        else:
            height += 42
    height += padding
    canvas = Image.new("RGB", (width, max(height, 360)), "#ffffff")
    draw = ImageDraw.Draw(canvas)

    y = padding
    draw.text((padding, y), title or f"题目 {index}", font=title_font, fill="#1f2937")
    y += 52
    if item.get("question_id"):
        draw.text((padding, y), str(item.get("question_id")), font=meta_font, fill="#64748b")
        y += 40
    for line, font in wrapped_text:
        if not line:
            y += 18
            continue
        draw.text((padding, y), line, font=font, fill="#1f2937")
        y += 40
    if rendered_images:
        y += gap
    for name, image in rendered_images:
        if image is None:
            draw.text((padding, y), f"[图片暂不能预览] {name}", font=small_font, fill="#94a3b8")
            y += 42
            continue
        x = padding + max(0, (max_content_width - image.width) // 2)
        draw.rounded_rectangle((x - 10, y - 10, x + image.width + 10, y + image.height + 10), radius=16, outline="#dbeafe", width=3, fill="#f8fafc")
        canvas.paste(image, (x, y))
        y += image.height + 24
        label = f"原题图片：{name}"
        label_width = draw.textlength(label, font=small_font)
        draw.text((padding + max(0, (max_content_width - label_width) // 2), y), label, font=small_font, fill="#64748b")
        y += 38

    snapshot_dir.mkdir(parents=True, exist_ok=True)
    output = snapshot_dir / _safe_snapshot_name(item, index)
    canvas.save(output)
    return str(output)


def _attach_question_snapshots(items: list[dict], output_json: Path) -> None:
    snapshot_dir = output_json.parent / "question_snapshots"
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        snapshot = _write_question_snapshot(item, snapshot_dir, index)
        if snapshot:
            item["question_snapshot_refs"] = [snapshot]


def section_kind(raw_title: str, body: list[str]) -> tuple[str, str]:
    if "回答下列问题" in raw_title or "回答问题" in raw_title:
        return "问答题", "qa"
    if "选择" in raw_title:
        return "选择题", "choice"
    if "判断" in raw_title or "正误" in raw_title:
        return "判断题", "judge"
    if "填空" in raw_title:
        return "填空题", "fill"
    if "名词解释" in raw_title or "名解" in raw_title:
        return "名词解释", "qa"
    if "计算" in raw_title:
        return "计算题", "calc"
    if "简答" in raw_title:
        return "简答题", "qa"
    text = raw_title + " " + " ".join(body[:5])
    if "判断" in text or "正误" in text:
        return "判断题", "judge"
    if "选择" in text:
        return "选择题", "choice"
    if "填空" in text:
        return "填空题", "fill"
    if "名词解释" in text or "名解" in text:
        return "名词解释", "qa"
    if "计算" in text:
        return "计算题", "calc"
    if "简答" in text:
        return "简答题", "qa"
    return "问答题", "qa"


def split_sections(paragraphs: list[str]) -> list[dict]:
    sections: list[dict] = []
    current: dict | None = None
    subject_index = 1
    subject_title = ""
    subject_seen = False
    last_major_no = 0
    for paragraph_index, para in enumerate(paragraphs):
        subject = _subject_title(para)
        next_paragraph = paragraphs[paragraph_index + 1] if paragraph_index + 1 < len(paragraphs) else ""
        if not subject and SECTION_RE.match(next_paragraph):
            subject = _implicit_subject_title(para)
        if subject and not SECTION_RE.match(para):
            if current:
                sections.append(current)
                current = None
            if subject_seen or last_major_no:
                subject_index += 1
            subject_seen = True
            subject_title = subject
            continue
        m = SECTION_RE.match(para)
        if m:
            if current:
                sections.append(current)
            major_no = cn_to_int(m.group(1)) or last_major_no + 1 or 1
            last_major_no = major_no
            current = {
                "cn": m.group(1),
                "major_no": major_no,
                "raw_title": para,
                "title_tail": clean_text(m.group(2)),
                "body": [],
                "subject_index": subject_index,
                "subject": subject_title,
            }
        elif _unnumbered_section_title(para):
            if current:
                sections.append(current)
            last_major_no += 1
            current = {
                "cn": int_to_cn(last_major_no),
                "major_no": last_major_no,
                "raw_title": para,
                "title_tail": clean_text(para),
                "body": [],
                "subject_index": subject_index,
                "subject": subject_title,
            }
        elif current:
            current["body"].append(para)
    if current:
        sections.append(current)
    if not sections:
        sections.append({"cn": "一", "major_no": 1, "raw_title": "一、问答题", "title_tail": "", "body": paragraphs, "subject_index": 1, "subject": ""})
    return sections


def has_image_hint(text: str) -> bool:
    generic_hint = bool(re.search(r"下图|如图|图中|画出|绘制|示意图", text))
    capability_hint = any(
        contribution.get("has_image_hint") is True
        for contribution in capability_policy_contributions(
            "exam_image_hint",
            {"text": text},
            text=text,
        )
        if isinstance(contribution, dict)
    )
    return generic_hint or capability_hint


def _has_multipart_cue(lines: list[str]) -> bool:
    text = clean_text("\n".join(lines[:3]))
    return bool(MULTIPART_CUE_RE.search(text) or text.rstrip().endswith(("：", ":")))


def _subquestion_entry(number: str, marker: str, text: str, raw: str) -> dict[str, str]:
    entry = {
        "number": str(number),
        "marker": marker,
        "stem": clean_text(text),
        "raw": clean_text(raw),
        "question_type": _requirement_type(text),
    }
    requirements = _infer_nested_requirements(entry["stem"], str(number))
    if requirements:
        entry["requirements"] = requirements
    constraints = _response_constraints(entry["stem"])
    if constraints:
        entry["response_constraints"] = constraints
    return entry


def _requirement_type(text: str) -> str:
    value = clean_text(text)
    if re.search(
        r"(画出|绘制|作图|画图|补全图|续画|绘出|"
        r"(?:画|作|绘制)(?:一幅|一个|出)?[^\n。；;]{0,20}示意图|"
        r"(?:请|需|要求|(?:在|于)[^\n。；;]{0,16}(?:图|坐标系)中)[^\n。；;]{0,6}标出|"
        r"标出.{0,20}(?:斑点|峰|相区|晶面|晶向|坐标|曲线))",
        value,
    ):
        return "作图题"
    if re.search(r"(计算|求出|求得|(?:^|[；;。])\s*(?:试)?求|质量比|质量分数|百分数|含量|比例|数值|多少)", value):
        return "计算题"
    if re.search(r"(判断|是否|可逆|正误)", value):
        return "判断题"
    return "简答题"


def _split_requirement_text(text: str) -> list[str]:
    value = re.sub(r"[（(]\s*\d+(?:\.\d+)?\s*分\s*[）)]", "", clean_text(text)).strip()
    if not value:
        return []
    chunks: list[str] = []
    # 先按明显问句边界拆，再把带“画出 A 和 B，计算 C”的复合句继续拆开。
    for piece in re.split(r"(?<=[？?])\s*", value):
        piece = piece.strip(" ；;。")
        if not piece:
            continue
        subpieces = re.split(
            r"[，,；;]\s*(?=(?:并|同时)?(?:示意)?(?:计算|求出|求得|判断|说明|写出|列出|画出|绘制|作图|标出))",
            piece,
        )
        chunks.extend(part.strip(" ；;。") for part in subpieces if part.strip(" ；;。"))
    out: list[str] = []
    for chunk in chunks:
        split = re.match(r"^(画出|绘制|作图|标出)(.+?)和(.+?)(?:，|,|；|;|$)", chunk)
        if split and any(keyword in split.group(2) + split.group(3) for keyword in ("曲线", "示意图", "组织", "图")):
            left = clean_text(split.group(1) + split.group(2))
            right = clean_text(split.group(1) + split.group(3))
            if left:
                out.append(left)
            if right:
                out.append(right)
            continue
        out.append(chunk)
    merged: list[str] = []
    index = 0
    while index < len(out):
        current = out[index].strip()
        context_prefix = bool(
            re.match(r"^(根据|依据|由)", current)
            and not re.search(r"(计算|求出|求得|判断|说明|写出|列出|画出|绘制|作图|标出|分析|比较|解释)", current)
        )
        if context_prefix and index + 1 < len(out):
            merged.append(clean_text(f"{current}，{out[index + 1]}"))
            index += 2
            continue
        if current and not re.fullmatch(r"[（(]?\s*\d+(?:\.\d+)?\s*分\s*[）)]?", current):
            merged.append(current)
        index += 1
    return merged


def _is_response_constraint(text: str) -> bool:
    value = clean_text(text).strip(" ；;。()（）")
    if not value or "?" in value or "？" in value:
        return False
    return bool(RESPONSE_CONSTRAINT_RE.search(value))


def _response_constraints(text: str) -> list[str]:
    return [part for part in _split_requirement_text(text) if _is_response_constraint(part)]


def _infer_nested_requirements(text: str, parent_number: str) -> list[dict[str, str]]:
    parts = _split_requirement_text(text)
    # Precision, unit and answer-format instructions govern how an answer is
    # written; they are not additional questions.  Preserve them separately on
    # the parent but never mint a synthetic circled child such as ``③``.
    meaningful = [part for part in parts if part and not _is_response_constraint(part)]
    if len(meaningful) < 2:
        return []
    types = [_requirement_type(part) for part in meaningful]
    # 多个问句或混合题型才拆；单纯一句话被误切时保持一级小问。
    if len(set(types)) < 2 and len(meaningful) <= 2 and not re.search(r"[？?].+[？?]", text):
        return []
    requirements: list[dict[str, str]] = []
    for index, (part, qtype) in enumerate(zip(meaningful, types), start=1):
        number = f"{parent_number}.{index}"
        requirements.append(
            {
                "number": number,
                "marker": number,
                "stem": clean_text(part),
                "raw": clean_text(part),
                "question_type": qtype,
            }
        )
    return requirements


def _extract_subquestion_from_line(line: str) -> dict[str, str] | None:
    text = clean_text(line)
    m = PAREN_SUBQUESTION_RE.match(text)
    if m:
        return _subquestion_entry(m.group(1), f"({m.group(1)})", m.group(2), text)
    m = ORDINAL_SUBQUESTION_RE.match(text)
    if m:
        raw_no = m.group(1)
        number = str(cn_to_int(raw_no) or raw_no)
        return _subquestion_entry(number, f"第{raw_no}问", m.group(2), text)
    m = ITEM_RE.match(text)
    if m:
        return _subquestion_entry(m.group(1), f"{m.group(1)}{m.group(2)}", m.group(3), text)
    return None


def _collect_subquestions(lines: list[str]) -> list[dict[str, str]]:
    subquestions: list[dict[str, str]] = []
    first_explicit_index = -1
    for line_index, line in enumerate(lines):
        entry = _extract_subquestion_from_line(line)
        if entry:
            if first_explicit_index < 0:
                first_explicit_index = line_index
            subquestions.append(entry)
    # Some source papers omit the visible ``(1)`` marker while retaining a
    # later explicit ``(2)``.  Recover exactly one missing predecessor only
    # when the immediately preceding tail contains a clear answer action.  The
    # setup/data paragraphs remain on the parent and are never turned into an
    # invented question.
    if subquestions and subquestions[0].get("number") == "2" and first_explicit_index > 0:
        action_lines: list[str] = []
        for raw in reversed(lines[:first_explicit_index]):
            value = clean_text(raw).strip()
            if not value:
                continue
            if ANSWER_ACTION_RE.search(value) and not _is_response_constraint(value):
                action_lines.insert(0, value)
                break
            if action_lines:
                break
        if action_lines:
            stem = clean_text("；".join(action_lines)).strip("；;。")
            inferred = _subquestion_entry("1", "(1)", stem, stem)
            inferred["inferred_missing_marker"] = True
            subquestions.insert(0, inferred)
    return subquestions


def _append_subquestion(current: dict, line: str, number: int, delimiter: str, text: str) -> None:
    current["stem_lines"].append(clean_text(line))
    current["_subquestion_started"] = True
    current["_last_subquestion_number"] = number
    current.setdefault("_subquestion_lines", []).append(clean_text(line))


def _looks_like_subquestion(current: dict, number: int, delimiter: str, text: str) -> bool:
    if not current or not _has_multipart_cue(current.get("stem_lines", [])):
        return False
    if delimiter in {".", "．"}:
        return True
    started = bool(current.get("_subquestion_started"))
    last_number = int(current.get("_last_subquestion_number") or 0)
    if not started:
        return number == 1
    return number == last_number + 1


def _split_inline_numbered_items(line: str) -> list[str]:
    """Split compact enumerations such as ``1、术语`` and ``(1) 术语``.

    This is intentionally used only in term-explanation sections; applying it
    to arbitrary calculation stems would confuse decimal values with item
    markers.
    """

    value = clean_text(line)
    matches = list(INLINE_ITEM_MARKER_RE.finditer(value))
    if len(matches) < 2 or matches[0].start() != 0:
        return [value]
    numbers = [int(match.group("paren_number") or match.group("plain_number")) for match in matches]
    if any(current != previous + 1 for previous, current in zip(numbers, numbers[1:])):
        return [value]
    rows: list[str] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(value)
        stem = clean_text(value[match.end() : end]).strip("；;，,。 ")
        if stem:
            number = match.group("paren_number") or match.group("plain_number")
            rows.append(f"{number}、 {stem}")
    return rows or [value]


def question_items(section: dict) -> list[dict]:
    major_no = int(section.get("major_no") or cn_to_int(section["cn"]) or 1)
    subject_index = int(section.get("subject_index") or 1)
    kind, prefix = section_kind(section["raw_title"], section["body"])
    final_section_title = f"{section['cn']}、{kind}"
    body = section["body"]
    if kind == "名词解释":
        body = [expanded for line in body for expanded in _split_inline_numbered_items(line)]
    items: list[dict] = []
    if prefix in {"judge", "choice", "fill"}:
        current: dict | None = None
        for line in body:
            m = ITEM_RE.match(line)
            if m:
                if current:
                    items.append(current)
                number = int(m.group(1))
                current = {
                    "question_id": f"{prefix}_s{subject_index:02d}_{major_no:02d}_{number:02d}",
                    "subject_index": subject_index,
                    "subject": section.get("subject", ""),
                    "major_number": str(major_no),
                    "section": final_section_title,
                    "section_raw": section["raw_title"],
                    "number": str(number),
                    "stem_lines": [clean_text(m.group(3))] if clean_text(m.group(3)) else [],
                }
            elif current:
                current["stem_lines"].append(line)
        if current:
            items.append(current)
    else:
        current = None
        heading: list[str] = []
        split_any = False
        for line in body:
            m = ITEM_RE.match(line)
            if m:
                number = int(m.group(1))
                delimiter = m.group(2)
                text = clean_text(m.group(3))
                if current and _looks_like_subquestion(current, number, delimiter, text):
                    _append_subquestion(current, line, number, delimiter, text)
                    continue
                split_any = True
                if current:
                    items.append(current)
                starts_multipart = bool(heading and _has_multipart_cue(heading))
                current = {
                    "question_id": f"{prefix}_s{subject_index:02d}_{major_no:02d}_{number:02d}",
                    "subject_index": subject_index,
                    "subject": section.get("subject", ""),
                    "major_number": str(major_no),
                    "section": final_section_title,
                    "section_raw": section["raw_title"],
                    "number": str(number),
                    "stem_lines": heading + ([clean_text(line) if starts_multipart else text] if text else []),
                    "_subquestion_started": starts_multipart,
                    "_last_subquestion_number": number if starts_multipart else 0,
                }
            elif current:
                if _extract_subquestion_from_line(line):
                    current["_subquestion_started"] = True
                current["stem_lines"].append(line)
            else:
                heading.append(line)
        if current:
            items.append(current)
        if not split_any:
            items.append(
                {
                    "question_id": f"{prefix}_s{subject_index:02d}_{major_no:02d}_01",
                    "subject_index": subject_index,
                    "subject": section.get("subject", ""),
                    "major_number": str(major_no),
                    "section": final_section_title,
                    "section_raw": section["raw_title"],
                    # The section itself is the question when it contains no
                    # numbered top-level item.  Preserve the source major
                    # number (for example “题九”) instead of inventing “1”.
                    "number": str(major_no),
                    "stem_lines": body,
                }
            )
    for item in items:
        image_refs: list[str] = []
        lines = []
        raw_stem_lines = item.pop("stem_lines", [])
        for raw in raw_stem_lines:
            text = clean_text(raw)
            if not text:
                continue
            if text.startswith(IMAGE_MARKER_PREFIX):
                image_refs.append(text[len(IMAGE_MARKER_PREFIX) :].strip())
                continue
            if text.startswith(TABLE_MARKER_PREFIX):
                table = _table_from_marker(text)
                if table:
                    item.setdefault("attachments", {}).setdefault("tables", []).append(table)
                continue
            lines.append(text)
        stem = "\n".join(lines)
        item["stem"] = stem
        subquestions = _collect_subquestions(lines)
        if subquestions:
            item["subquestions"] = subquestions
        else:
            # Composite instructions are not always explicitly numbered. When
            # a single sentence contains heterogeneous answer actions, retain
            # one visible parent and create typed leaf requirements so the
            # answer/coverage/figure stages cannot silently omit one modality.
            requirements = _infer_nested_requirements(stem, str(item.get("number") or "1"))
            if requirements:
                number = str(item.get("number") or "1")
                item["subquestions"] = [
                    {
                        "number": number,
                        "marker": f"({number})",
                        "stem": stem,
                        "raw": stem,
                        "question_type": "简答题",
                        # This parent exists only to retain independently typed
                        # answer actions.  It is not a visible nested level in
                        # the source question and must be flattened again when
                        # the formal answer is presented.
                        "synthetic_parent": True,
                        "requirements": requirements,
                    }
                ]
        item["image_refs"] = image_refs
        item["source_image_required"] = source_image_required(item)
        item["answer_figure_required"] = answer_figure_required(item)
        # Kept for schema compatibility; downstream stages now use the two
        # explicit requirements above instead of conflating source and answer.
        item["needs_figure"] = item["answer_figure_required"]
        item["formula_refs"] = [f"{item['question_id']}_source_formula"] if looks_like_formula(stem) else []
        item["status"] = "registered"
        for key in ("_subquestion_started", "_last_subquestion_number", "_subquestion_lines"):
            item.pop(key, None)
    combined = _combine_shared_composite_section(section, items)
    for item in combined:
        item["section_item_count"] = len(combined)
    return combined


def _combine_shared_composite_section(section: dict, items: list[dict]) -> list[dict]:
    """Represent one composite problem with a trailing shared figure as one group.

    Some exam DOCX files put a common diagram after the final numbered clause.
    Treating those clauses as independent questions assigns the diagram only to
    the last clause and deprives earlier clauses of required data.  A section
    explicitly labelled as one composite problem (``本题共``) is therefore kept
    as a parent question with numbered subquestions when it has a shared figure.
    This is subject-neutral and lets vision inspect the artifact exactly once.
    """

    title = clean_text(str(section.get("raw_title") or ""))
    if len(items) < 2 or "综合题" not in title or "本题共" not in title:
        return items
    image_refs = [
        str(raw).strip()
        for item in items
        for raw in item.get("image_refs", []) or []
        if str(raw).strip()
    ]
    if not image_refs:
        return items

    parent = dict(items[0])
    # Once a composite section is grouped into one parent, the visible parent
    # is the major question itself; its former item 1 remains subquestion (1).
    parent["number"] = str(section.get("major_no") or parent.get("major_number") or parent.get("number") or "1")
    parent["stem"] = "\n".join(
        f"({item.get('number') or index}) {clean_text(str(item.get('stem') or ''))}".strip()
        for index, item in enumerate(items, start=1)
    )
    parent["subquestions"] = []
    for index, item in enumerate(items, start=1):
        number = str(item.get("number") or index)
        stem = clean_text(str(item.get("stem") or ""))
        subquestion = _subquestion_entry(number, f"({number})", stem, stem)
        parent["subquestions"].append(subquestion)
    parent["image_refs"] = list(dict.fromkeys(image_refs))
    parent["source_image_required"] = source_image_required(parent)
    parent["answer_figure_required"] = answer_figure_required(parent)
    parent["needs_figure"] = parent["answer_figure_required"]
    parent["formula_refs"] = [f"{parent['question_id']}_source_formula"] if looks_like_formula(parent["stem"]) else []
    parent["attachment_scope"] = {
        "kind": "shared_composite_question",
        "source_item_numbers": [str(item.get("number") or index) for index, item in enumerate(items, start=1)],
        "reason": "composite_section_with_trailing_shared_figure",
    }
    return [parent]


def extract_exam_structure(exam_file: Path, output_json: Path) -> dict:
    image_dir = output_json.parent / "source_images"
    paragraphs = _docx_paragraph_lines(exam_file, image_dir=image_dir)
    items: list[dict] = []
    for section in split_sections(paragraphs):
        items.extend(question_items(section))
    _attach_question_snapshots(items, output_json)
    source_paragraphs = [line for line in paragraphs if not str(line).startswith(IMAGE_MARKER_PREFIX)]
    source_paragraphs = [text for line in source_paragraphs if (text := _source_line_text(str(line)))]
    data = {
        "schema_version": "structured_exam.v4.program_extracted",
        "grouping_policy_version": EXAM_GROUPING_POLICY_VERSION,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "exam_file": str(exam_file),
        "paragraph_count": len(source_paragraphs),
        "source_paragraphs": source_paragraphs,
        "items": items,
        "notes": ["Program extracted structure; review warnings before production."],
    }
    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_json.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return data
