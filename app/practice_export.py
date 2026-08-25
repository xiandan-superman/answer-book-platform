from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any, Callable
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from PIL import Image, UnidentifiedImageError

from .capabilities.expression_rendering import (
    build_expression_render_plan,
    preflight_expression_render,
    render_expression_omml,
)
from .capabilities.text_expression_rendering import (
    build_text_expression_render_plans,
    normalize_standard_state_latex,
    repair_json_escaped_latex,
)
from .practice_document_contracts import (
    PRACTICE_DOCUMENT_CONTRACT_VERSION,
    PRACTICE_PAGE_CONTRACT,
    PRACTICE_TEXT_CONTRACT,
)

BLUE = RGBColor(37, 99, 235)
DARK_BLUE = RGBColor(31, 77, 120)
ERROR_RED = RGBColor(185, 28, 28)
MUTED = RGBColor(100, 116, 139)
INK = RGBColor(0, 0, 0)
ASCII_FONT = PRACTICE_TEXT_CONTRACT.latin_font
CJK_FONT = PRACTICE_TEXT_CONTRACT.east_asia_font
CJK_FONT_FALLBACK = PRACTICE_TEXT_CONTRACT.east_asia_fallback_font
MATH_FONT = PRACTICE_TEXT_CONTRACT.math_font
CHART_FONT_PATHS = (
    Path("/System/Library/Fonts/PingFang.ttc"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
)


_SELF_CORRECTION_RE = re.compile(
    r"(?:自我纠错|纠错草稿|内部草稿|模型分析|思考过程|模型(?:刚才|此前)|"
    r"前文有误|刚才的答案|我刚才|更正如下|纠正(?:如下|为)|抱歉[，,]?)"
)

_LITERAL_SUBQUESTION_BREAK_RE = re.compile(
    r"\\n(?=\s*(?:[（(]\s*\d{1,2}\s*[）)]|[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]))"
)
_SOLVED_NUMERIC_FORMULA_RE = re.compile(
    r"=(?:\s*[-+]?\s*(?:\\frac\s*\{[^{}]+\}\s*\{[^{}]+\}|\d+(?:\.\d+)?))(?:\s*[;,，；]|\s*$)"
)
_ANSWER_FORMULA_ROLES = {"answer", "result", "solution", "答案", "结果", "解答"}


def _figure_image_path(spec: dict[str, Any]) -> Path | None:
    raw = str(spec.get("image_path") or spec.get("path") or "").strip()
    return Path(raw).expanduser() if raw else None


def _valid_figure_image(spec: dict[str, Any]) -> bool:
    path = _figure_image_path(spec)
    if path is None or not path.is_file():
        return False
    try:
        with Image.open(path) as image:
            image.verify()
        return True
    except (OSError, UnidentifiedImageError):
        return False


def _renderable_figure(spec: dict[str, Any]) -> bool:
    if _valid_figure_image(spec):
        return True
    series = spec.get("series") if isinstance(spec.get("series"), list) else []
    if any(isinstance(row, dict) and len(row.get("points") or []) >= 2 for row in series):
        return True
    nodes = spec.get("nodes") if isinstance(spec.get("nodes"), list) else []
    return len([node for node in nodes if isinstance(node, dict) and node.get("id")]) >= 2


def practice_export_exercise_id(item: dict[str, Any], index: int = 0) -> str:
    """Return the stable identity used to select questions for Word export."""
    for key in ("plan_item_id", "exercise_id", "question_id", "id", "number"):
        value = str(item.get(key) or "").strip()
        if value:
            return value
    return str(index + 1)


def practice_stem_answer_leak_reasons(item: dict[str, Any]) -> list[str]:
    """Detect structured stem assets that expose what the student must supply."""
    question_type = str(item.get("question_type") or "").strip()
    stem = str(item.get("stem") or "")
    has_blank = bool(
        re.search(
            r"_{2,}|[（(]\s*[）)]|\\(?:underline|underbrace)\s*\{",
            stem,
        )
    )
    reasons: list[str] = []
    for index, formula in enumerate(item.get("formulas") or [], start=1):
        if not isinstance(formula, dict) or not _matches_location(formula.get("location"), "stem"):
            continue
        latex = str(formula.get("latex") or "").strip()
        caption = str(formula.get("caption") or "").strip()
        role = str(formula.get("role") or "relation").strip().lower()
        if role in _ANSWER_FORMULA_ROLES or re.search(r"答案|最终结果|计算结果|标准结论", caption):
            reasons.append(f"第 {index} 个题干公式被标记为答案或结果")
            continue
        if question_type == "填空题" and has_blank and role != "given":
            reasons.append(f"第 {index} 个题干公式可能直接给出填空答案")
            continue
        if role != "given" and _SOLVED_NUMERIC_FORMULA_RE.search(latex):
            reasons.append(f"第 {index} 个题干公式包含已求得的数值结果")
    return reasons


def resolve_practice_export_payload(
    request_data: dict[str, Any],
    latest_data: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Resolve a Word request against the latest record and its exact selection."""
    request_data = request_data if isinstance(request_data, dict) else {}
    source = latest_data if isinstance(latest_data, dict) else request_data
    resolved = {**source}
    scope = str(request_data.get("export_scope") or "all").strip().lower()
    selected_ids = [
        str(value).strip()
        for value in request_data.get("selected_exercise_ids") or []
        if str(value).strip()
    ]
    if scope == "selected":
        exercises = source.get("exercises") if isinstance(source.get("exercises"), list) else []
        indexed = {
            practice_export_exercise_id(item, index): item
            for index, item in enumerate(exercises)
            if isinstance(item, dict)
        }
        resolved["exercises"] = [indexed[value] for value in dict.fromkeys(selected_ids) if value in indexed]
        # Whole-set quality belongs to the original collection. The selected
        # subset is validated item-by-item below and must not inherit stale
        # blockers from questions that are no longer part of the export.
        resolved["quality"] = {}
        resolved["requested_count"] = len(resolved["exercises"])
        resolved["export_scope"] = "selected"
        resolved["selected_exercise_ids"] = selected_ids
    return resolved


def validate_practice_export(data: dict[str, Any]) -> dict[str, Any]:
    """Block every deterministic defect before producing a formal Word file."""
    blocking_issues: list[str] = []
    warning_issues: list[str] = []
    quality = data.get("quality") if isinstance(data, dict) and isinstance(data.get("quality"), dict) else {}
    blocking_issues.extend(str(issue) for issue in quality.get("blocking_issues") or [] if str(issue).strip())
    if quality.get("release_level") == "review_candidate":
        quality_warnings = [str(issue).strip() for issue in quality.get("warnings") or [] if str(issue).strip()]
        warning_issues.extend(
            quality_warnings
            or ["网页质量状态已将本组题目标记为待复核；Word 同步以待复核候选版导出。"]
        )
    exercises = data.get("exercises") if isinstance(data, dict) else []
    if not exercises:
        blocking_issues.append("没有可导出的题目。")
    for index, item in enumerate(exercises or [], start=1):
        if not isinstance(item, dict):
            blocking_issues.append(f"第 {index} 题不是有效对象。")
            continue
        question_number = str(item.get("number") or index)
        joined = str(item.get("stem") or "")
        if not joined.strip():
            blocking_issues.append(f"第 {question_number} 题缺少题干。")
            continue
        # Visible text quality warnings are deliberately emitted only by
        # audit_practice_export_data().  This validator owns export blockers
        # and structural figure checks, avoiding two messages for one defect.
        if item.get("generation_status") == "failed":
            blocking_issues.append(f"第 {question_number} 题生成失败，不能进入正式题目卷。")
            continue
        for formula_index, formula in enumerate(item.get("formulas") or [], start=1):
            if not isinstance(formula, dict):
                blocking_issues.append(f"第 {question_number} 题第 {formula_index} 个公式不是有效对象。")
                continue
            latex = _normalize_standard_state_latex(_text(formula.get("latex"), 3000))
            plan = build_expression_render_plan(
                latex,
                question_id=str(item.get("exercise_id") or item.get("question_id") or question_number),
                location=f"formulas[{formula_index - 1}].latex",
                role=str(formula.get("role") or "relation"),
                display=bool(formula.get("display", True)),
            )
            if error := preflight_expression_render(plan):
                blocking_issues.append(f"第 {question_number} 题第 {formula_index} 个公式无法生成 Word 公式对象：{error}")
        for figure in item.get("figures") or []:
            if isinstance(figure, dict) and "figure_id" in figure and not str(figure.get("figure_id") or "").strip():
                blocking_issues.append(f"第 {question_number} 题存在缺少 ID 的题图，无法追踪题干与图片的绑定。")
            if not isinstance(figure, dict) or not _renderable_figure(figure):
                blocking_issues.append(f"第 {question_number} 题包含无法绘制的题图，不能用文字说明代替正式配图。")
    blocking_issues.extend(audit_practice_export_data(data))
    blocking_issues.extend(preflight_practice_inline_expressions(data))
    review = data.get("semantic_review") if isinstance(data.get("semantic_review"), dict) else {}
    review_items = {
        str(item.get("number") or "").strip(): item
        for item in review.get("items") or []
        if isinstance(item, dict) and str(item.get("number") or "").strip()
    }
    review_candidate_numbers: list[str] = []
    review_status = str(review.get("status") or "").strip().lower()
    if review and review_status not in {"disabled", "not_required"}:
        for index, item in enumerate(exercises or [], start=1):
            if not isinstance(item, dict) or item.get("generation_status") == "failed":
                continue
            number = str(item.get("number") or index).strip()
            item_review = review_items.get(number)
            status = str((item_review or {}).get("status") or "not_reviewed").strip().lower()
            risks = (item_review or {}).get("risks") if isinstance((item_review or {}).get("risks"), list) else []
            actionable = any(
                isinstance(risk, dict) and str(risk.get("severity") or "medium").strip().lower() in {"high", "medium"}
                for risk in risks
            )
            if status not in {"passed", "warning"} or actionable:
                review_candidate_numbers.append(number)
    if review_candidate_numbers:
        warning_issues.append(
            "第 " + "、".join(dict.fromkeys(review_candidate_numbers)) + " 题尚未完成学科复核；Word 可供查看和继续修改，但不应视为正式发布版。"
        )
    blocking_issues = list(dict.fromkeys(blocking_issues))
    warning_issues = list(dict.fromkeys(warning_issues))
    return {
        "ok": not blocking_issues,
        "release_level": "blocked" if blocking_issues else ("review_candidate" if warning_issues else "formal"),
        "issues": blocking_issues + warning_issues,
        "blocking_issues": blocking_issues,
        "warning_issues": warning_issues,
    }


def assert_practice_exportable(data: dict[str, Any]) -> None:
    report = validate_practice_export(data)
    if not report["ok"]:
        raise ValueError("正式导出被阻断：" + "；".join(report["blocking_issues"][:8]))


def _practice_twips(inches: float) -> str:
    return str(round(float(inches) * 1440))


def _practice_half_points(points: float) -> str:
    return str(round(float(points) * 2))


def _word_attr(node, name: str, namespace: str) -> str:
    if node is None:
        return ""
    return str(node.get(f"{{{namespace}}}{name}") or "")


def _practice_document_contract_issues(archive: ZipFile, root, namespaces: dict[str, str]) -> list[str]:
    """Check the independent practice-Word compatibility contract."""

    issues: list[str] = []
    word_ns = namespaces["w"]
    section_properties = root.xpath(".//w:sectPr", namespaces=namespaces)
    if not section_properties:
        issues.append("练习 Word 契约：缺少页面设置。")
    for index, section in enumerate(section_properties, 1):
        page_size = section.find("w:pgSz", namespaces=namespaces)
        page_margin = section.find("w:pgMar", namespaces=namespaces)
        if _word_attr(page_size, "w", word_ns) != _practice_twips(PRACTICE_PAGE_CONTRACT.width_inches):
            issues.append(f"练习 Word 契约：第 {index} 节页面宽度发生变化。")
        if _word_attr(page_size, "h", word_ns) != _practice_twips(PRACTICE_PAGE_CONTRACT.height_inches):
            issues.append(f"练习 Word 契约：第 {index} 节页面高度发生变化。")
        for side in ("top", "bottom"):
            if _word_attr(page_margin, side, word_ns) != _practice_twips(PRACTICE_PAGE_CONTRACT.top_bottom_margin_inches):
                issues.append(f"练习 Word 契约：第 {index} 节{side}页边距发生变化。")
        for side in ("left", "right"):
            if _word_attr(page_margin, side, word_ns) != _practice_twips(PRACTICE_PAGE_CONTRACT.left_right_margin_inches):
                issues.append(f"练习 Word 契约：第 {index} 节{side}页边距发生变化。")
        for edge in ("header", "footer"):
            if _word_attr(page_margin, edge, word_ns) != _practice_twips(PRACTICE_PAGE_CONTRACT.header_footer_distance_inches):
                issues.append(f"练习 Word 契约：第 {index} 节{edge}距离发生变化。")

    styles = etree.fromstring(archive.read("word/styles.xml"))

    def style(style_id: str):
        nodes = styles.xpath(f".//w:style[@w:styleId='{style_id}']", namespaces=namespaces)
        return nodes[0] if nodes else None

    normal = style("Normal")
    if normal is None:
        issues.append("练习 Word 契约：缺少 Normal 样式。")
    else:
        fonts = normal.find("w:rPr/w:rFonts", namespaces=namespaces)
        size = normal.find("w:rPr/w:sz", namespaces=namespaces)
        spacing = normal.find("w:pPr/w:spacing", namespaces=namespaces)
        indent = normal.find("w:pPr/w:ind", namespaces=namespaces)
        if _word_attr(fonts, "eastAsia", word_ns) != CJK_FONT:
            issues.append("练习 Word 契约：Normal 中文字体发生变化。")
        if _word_attr(fonts, "ascii", word_ns) != ASCII_FONT or _word_attr(fonts, "hAnsi", word_ns) != ASCII_FONT:
            issues.append("练习 Word 契约：Normal 西文字体发生变化。")
        if _word_attr(size, "val", word_ns) != _practice_half_points(PRACTICE_TEXT_CONTRACT.body_size_pt):
            issues.append("练习 Word 契约：Normal 正文字号发生变化。")
        if _word_attr(spacing, "line", word_ns) != str(round(PRACTICE_TEXT_CONTRACT.line_spacing * 240)):
            issues.append("练习 Word 契约：Normal 行距发生变化。")
        if _word_attr(spacing, "before", word_ns) != "0" or _word_attr(spacing, "after", word_ns) != "0":
            issues.append("练习 Word 契约：Normal 段前段后发生变化。")
        if _word_attr(indent, "firstLine", word_ns) != str(round(PRACTICE_TEXT_CONTRACT.first_line_indent_pt * 20)):
            issues.append("练习 Word 契约：Normal 首行缩进发生变化。")

    list_number = style("ListNumber")
    if list_number is None:
        issues.append("练习 Word 契约：缺少编号列表样式。")
    else:
        indent = list_number.find("w:pPr/w:ind", namespaces=namespaces)
        if _word_attr(indent, "left", word_ns) != str(round(PRACTICE_TEXT_CONTRACT.list_left_indent_pt * 20)):
            issues.append("练习 Word 契约：解析步骤列表左缩进发生变化。")
        if _word_attr(indent, "hanging", word_ns) != str(round(PRACTICE_TEXT_CONTRACT.list_hanging_indent_pt * 20)):
            issues.append("练习 Word 契约：解析步骤列表悬挂缩进发生变化。")

    list_paragraphs = root.xpath(".//w:p[w:pPr/w:pStyle[@w:val='ListNumber']]", namespaces=namespaces)
    explicit_list_ids = [
        _word_attr(paragraph.find("w:pPr/w:numPr/w:numId", namespaces=namespaces), "val", word_ns)
        for paragraph in list_paragraphs
    ]
    if list_paragraphs and any(not value for value in explicit_list_ids):
        issues.append("练习 Word 契约：解析步骤必须使用每题独立的显式编号，不能跨题连续编号。")
    if list_paragraphs:
        numbering = etree.fromstring(archive.read("word/numbering.xml"))
        restarting_ids = {
            _word_attr(node, "numId", word_ns)
            for node in numbering.xpath(
                ".//w:num[w:lvlOverride[@w:ilvl='0']/w:startOverride[@w:val='1']]",
                namespaces=namespaces,
            )
        }
        if any(value not in restarting_ids for value in explicit_list_ids if value):
            issues.append("练习 Word 契约：每道题的解析步骤必须从 1 重新编号。")

    font_table_text = archive.read("word/fontTable.xml").decode("utf-8", errors="ignore")
    if CJK_FONT not in font_table_text or CJK_FONT_FALLBACK not in font_table_text:
        issues.append("练习 Word 契约：中文字体或兼容字体声明缺失。")
    settings = etree.fromstring(archive.read("word/settings.xml"))
    math_fonts = settings.xpath(".//m:mathFont/@m:val", namespaces=namespaces)
    if PRACTICE_TEXT_CONTRACT.math_font not in math_fonts:
        issues.append("练习 Word 契约：数学字体发生变化。")

    titles = [
        "".join(node.xpath(".//w:t/text()", namespaces=namespaces)).strip()
        for node in root.xpath(".//w:p[w:pPr/w:pStyle[@w:val='Title']]", namespaces=namespaces)
    ]
    if titles not in [["专项练习题目卷"], ["专项练习"]]:
        issues.append(f"练习 Word 契约：总标题结构发生变化（{titles}）。")
    section_titles = [
        "".join(node.xpath(".//w:t/text()", namespaces=namespaces)).strip()
        for node in root.xpath(".//w:p[w:pPr/w:pStyle[@w:val='Heading1']]", namespaces=namespaces)
    ]
    valid_section_titles = [["练习题"], ["参考答案与解析"], ["练习题", "参考答案与解析"]]
    if section_titles not in valid_section_titles:
        issues.append(f"练习 Word 契约：题目/答案区块结构发生变化（{section_titles}）。")

    footer_parts = [name for name in archive.namelist() if re.fullmatch(r"word/footer\d+\.xml", name)]
    if not footer_parts:
        issues.append("练习 Word 契约：页脚或页码缺失。")
    else:
        footer = etree.fromstring(archive.read(sorted(footer_parts)[0]))
        page_fields = footer.xpath(".//w:p/w:fldSimple[contains(@w:instr, 'PAGE')]", namespaces=namespaces)
        footer_text = "".join(footer.xpath(".//w:t/text()", namespaces=namespaces))
        alignment = footer.find(".//w:p/w:pPr/w:jc", namespaces=namespaces)
        if len(page_fields) != 1 or "第 " not in footer_text or " 页" not in footer_text:
            issues.append("练习 Word 契约：页码必须保持“第 {PAGE} 页”。")
        if _word_attr(alignment, "val", word_ns) != "right":
            issues.append("练习 Word 契约：页码对齐方式发生变化。")
    return issues


def validate_docx_output(content: bytes, data: dict[str, Any]) -> dict[str, Any]:
    """Verify the generated Word package against its question contract."""
    exercises = [item for item in data.get("exercises") or [] if isinstance(item, dict)]
    expected_questions = len(exercises)
    expected_figures = sum(len(item.get("figures") or []) for item in exercises)
    expected_tables = sum(len(item.get("tables") or []) for item in exercises)
    expected_formulas = sum(len(item.get("formulas") or []) for item in exercises)
    issues: list[str] = []
    metrics = {
        "document_contract_version": PRACTICE_DOCUMENT_CONTRACT_VERSION,
        "expected_question_count": expected_questions,
        "question_heading_count": 0,
        "expected_figure_count": expected_figures,
        "media_count": 0,
        "drawing_count": 0,
        "expected_table_count": expected_tables,
        "table_count": 0,
        "expected_structured_formula_count": expected_formulas,
        "office_math_count": 0,
        "office_math_markup_leak_count": 0,
        "image_alt_text_count": 0,
        "page_number_field_count": 0,
    }
    required_parts = {
        "[Content_Types].xml",
        "_rels/.rels",
        "word/document.xml",
        "word/_rels/document.xml.rels",
        "word/styles.xml",
        "word/settings.xml",
    }
    try:
        with ZipFile(BytesIO(content)) as archive:
            names = set(archive.namelist())
            missing_parts = sorted(required_parts - names)
            if missing_parts:
                issues.append("DOCX 缺少必要结构：" + "、".join(missing_parts) + "。")
                return {"ok": False, "issues": issues, **metrics}
            if bad_member := archive.testzip():
                issues.append(f"DOCX 压缩包成员损坏：{bad_member}。")

            document_xml = archive.read("word/document.xml")
            root = etree.fromstring(document_xml)
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
            }
            issues.extend(_practice_document_contract_issues(archive, root, namespaces))
            heading_nodes = root.xpath(
                ".//w:p[w:pPr/w:pStyle[@w:val='Heading2']]",
                namespaces=namespaces,
            )
            question_headings = [
                "".join(node.xpath(".//w:t/text()", namespaces=namespaces)).strip()
                for node in heading_nodes
            ]
            expected_headings = [f"第 {index} 题" for index in range(1, expected_questions + 1)]
            metrics["question_heading_count"] = len(question_headings)
            if question_headings != expected_headings:
                issues.append(f"题号结构不完整：期望 {expected_headings}，实际 {question_headings}。")

            metrics["table_count"] = len(root.xpath(".//w:tbl", namespaces=namespaces))
            if metrics["table_count"] < expected_tables:
                issues.append(
                    f"题目包含 {expected_tables} 个表格，但 DOCX 仅写入 {metrics['table_count']} 个。"
                )
            metrics["office_math_count"] = len(root.xpath(".//m:oMath", namespaces=namespaces))
            if metrics["office_math_count"] < expected_formulas:
                issues.append(
                    f"题目包含 {expected_formulas} 个结构化公式，但 DOCX 仅写入 "
                    f"{metrics['office_math_count']} 个 Office 公式对象。"
                )
            office_math_texts = [
                "".join(node.xpath(".//m:t/text()", namespaces=namespaces))
                for node in root.xpath(".//m:oMath", namespaces=namespaces)
            ]
            metrics["office_math_markup_leak_count"] = sum(
                1 for text in office_math_texts if "$" in text or "\\" in text
            )
            if metrics["office_math_markup_leak_count"]:
                issues.append(
                    "DOCX Office 公式对象中仍含提供方的 LaTeX/Markdown 定界标记。"
                )
            metrics["drawing_count"] = len(root.xpath(".//w:drawing", namespaces=namespaces))
            doc_properties = root.xpath(".//wp:docPr", namespaces=namespaces)
            metrics["image_alt_text_count"] = sum(
                1 for node in doc_properties if str(node.get("descr") or "").strip()
            )
            if metrics["drawing_count"] < expected_figures:
                issues.append(
                    f"题目包含 {expected_figures} 幅图形，但 DOCX 仅写入 {metrics['drawing_count']} 个图片对象。"
                )
            if expected_figures and metrics["image_alt_text_count"] < expected_figures:
                issues.append(
                    f"DOCX 仅有 {metrics['image_alt_text_count']}/{expected_figures} 幅题图包含可追踪的替代文本。"
                )

            media_names = sorted(
                name for name in names if name.startswith("word/media/") and not name.endswith("/")
            )
            metrics["media_count"] = len(media_names)
            if metrics["media_count"] < expected_figures:
                issues.append(
                    f"题目包含 {expected_figures} 幅图形，但 DOCX 仅嵌入 {metrics['media_count']} 份媒体。"
                )
            for name in media_names:
                media = archive.read(name)
                if len(media) < 100 or not media.startswith((b"\x89PNG\r\n\x1a\n", b"\xff\xd8\xff")):
                    issues.append(f"DOCX 题图媒体无效或过小：{name}。")

            relationships = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
            relationship_targets = {
                str(node.get("Id") or ""): str(node.get("Target") or "")
                for node in relationships.xpath(".//pr:Relationship", namespaces=namespaces)
            }
            for embed_id in root.xpath(".//a:blip/@r:embed", namespaces=namespaces):
                target = relationship_targets.get(str(embed_id), "")
                package_target = "word/" + target.lstrip("/")
                if not target or package_target not in names:
                    issues.append(f"DOCX 图片关系 {embed_id} 未指向有效媒体文件。")

            footer_parts = [name for name in names if re.fullmatch(r"word/footer\d+\.xml", name)]
            metrics["page_number_field_count"] = sum(
                archive.read(name).count(b"PAGE") for name in footer_parts
            )
            if metrics["page_number_field_count"] < 1:
                issues.append("DOCX 页脚缺少页码字段。")

            visible_text = "".join(root.xpath(".//w:t/text()", namespaces=namespaces))
            if _SELF_CORRECTION_RE.search(visible_text):
                issues.append("DOCX 可见文本中仍含模型自我纠错或内部草稿表达。")
            if has_unrenderable_practice_markup(visible_text):
                issues.append("DOCX 可见文本中仍含未渲染的 LaTeX 标记。")
    except Exception as exc:
        issues.append(f"DOCX 无法解析：{exc}。")
        return {"ok": False, "issues": list(dict.fromkeys(issues)), **metrics}

    try:
        Document(BytesIO(content))
    except Exception as exc:
        issues.append(f"python-docx 无法打开文档：{exc}。")
    return {"ok": not issues, "issues": list(dict.fromkeys(issues)), **metrics}


def _text(value: Any, limit: int = 10000) -> str:
    return str(value or "").strip()[:limit]


def _set_run(
    run,
    *,
    size: float = PRACTICE_TEXT_CONTRACT.body_size_pt,
    bold: bool = False,
    color: RGBColor = INK,
) -> None:
    run.font.name = ASCII_FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    fonts.set(qn("w:cs"), ASCII_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    # Formal Word output uses black text throughout; semantic color is kept only
    # for the web presentation layer.
    run.font.color.rgb = INK


def _set_cell_free_style(style, *, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = ASCII_FONT
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    fonts.set(qn("w:cs"), ASCII_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("第 ")
    _set_run(prefix, size=PRACTICE_TEXT_CONTRACT.page_number_size_pt, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    suffix = paragraph.add_run(" 页")
    _set_run(suffix, size=PRACTICE_TEXT_CONTRACT.page_number_size_pt, color=MUTED)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PRACTICE_PAGE_CONTRACT.width_inches)
    section.page_height = Inches(PRACTICE_PAGE_CONTRACT.height_inches)
    section.top_margin = Inches(PRACTICE_PAGE_CONTRACT.top_bottom_margin_inches)
    section.bottom_margin = Inches(PRACTICE_PAGE_CONTRACT.top_bottom_margin_inches)
    section.left_margin = Inches(PRACTICE_PAGE_CONTRACT.left_right_margin_inches)
    section.right_margin = Inches(PRACTICE_PAGE_CONTRACT.left_right_margin_inches)
    section.header_distance = Inches(PRACTICE_PAGE_CONTRACT.header_footer_distance_inches)
    section.footer_distance = Inches(PRACTICE_PAGE_CONTRACT.header_footer_distance_inches)

    # Word uses the primary Windows font, while macOS/compatible readers can
    # resolve the declared alternate name instead of displaying tofu boxes.
    font_table_part = doc.part.part_related_by(RT.FONT_TABLE)
    font_table = parse_xml(font_table_part.blob)
    if not any(node.get(qn("w:name")) == CJK_FONT for node in font_table.findall(qn("w:font"))):
        font = OxmlElement("w:font")
        font.set(qn("w:name"), CJK_FONT)
        alt_name = OxmlElement("w:altName")
        alt_name.set(qn("w:val"), CJK_FONT_FALLBACK)
        charset = OxmlElement("w:charset")
        charset.set(qn("w:val"), "86")
        family = OxmlElement("w:family")
        family.set(qn("w:val"), "swiss")
        font.extend([alt_name, charset, family])
        font_table.append(font)
        font_table_part._blob = etree.tostring(font_table, xml_declaration=True, encoding="UTF-8", standalone=True)

    settings = doc.settings._element
    math_pr = settings.find(qn("m:mathPr"))
    if math_pr is None:
        math_pr = OxmlElement("m:mathPr")
        settings.append(math_pr)
    math_font = math_pr.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_pr.insert(0, math_font)
    math_font.set(qn("m:val"), MATH_FONT)
    normal = doc.styles["Normal"]
    _set_cell_free_style(normal, size=PRACTICE_TEXT_CONTRACT.body_size_pt, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = PRACTICE_TEXT_CONTRACT.line_spacing
    normal.paragraph_format.first_line_indent = Pt(PRACTICE_TEXT_CONTRACT.first_line_indent_pt)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, before, after in (
        ("Title", 0, 0),
        ("Subtitle", 0, 0),
        ("Heading 1", 0, 0),
        ("Heading 2", 0, 0),
        ("Heading 3", 0, 0),
    ):
        style = doc.styles[name]
        _set_cell_free_style(style, size=PRACTICE_TEXT_CONTRACT.body_size_pt, color=INK, bold=name != "Subtitle")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = PRACTICE_TEXT_CONTRACT.line_spacing
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.keep_with_next = True

    for name in ("List Number", "List Bullet"):
        style = doc.styles[name]
        _set_cell_free_style(style, size=PRACTICE_TEXT_CONTRACT.body_size_pt, color=INK)
        style.paragraph_format.left_indent = Pt(PRACTICE_TEXT_CONTRACT.list_left_indent_pt)
        style.paragraph_format.right_indent = Pt(0)
        style.paragraph_format.first_line_indent = Pt(-PRACTICE_TEXT_CONTRACT.list_hanging_indent_pt)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = PRACTICE_TEXT_CONTRACT.line_spacing
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _page_number(section.footer.paragraphs[0])


def _add_title_block(doc: Document, data: dict[str, Any], *, document_kind: str = "combined") -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titles = {
        "questions": "专项练习题目卷",
        "combined": "专项练习",
    }
    title.add_run(titles.get(document_kind, titles["combined"]))

    title.paragraph_format.first_line_indent = Pt(0)


_INLINE_MATH_RE = re.compile(
    r"(?<!\\)\$\$(.+?)(?<!\\)\$\$|\\\[(.+?)\\\]|(?<!\\)\$(.+?)(?<!\\)\$|\\\((.+?)\\\)",
    re.DOTALL,
)
_BARE_LATEX_RE = re.compile(
    r"\\(?:frac|sqrt|mathrm|mathbf|mathit|text|operatorname|Delta|sum|int|theta|alpha|beta|gamma|"
    r"partial|approx|times|cdot|left|right|begin|end)\b"
)
_BARE_LATEX_BRACED_COMMANDS = {
    "mathrm": 1,
    "mathbf": 1,
    "mathit": 1,
    "text": 1,
    "operatorname": 1,
    "sqrt": 1,
    "frac": 2,
}
_MATH_FILL_BLANK_RE = re.compile(r"(?<![\\_])_{2,}(?!_)")


def _normalize_inline_math_fill_blanks(match: re.Match[str]) -> str:
    """Keep answer blanks visible without treating underscores as subscripts."""

    original_latex = next((group for group in match.groups() if group is not None), "")
    latex = repair_json_escaped_latex(original_latex)

    if match.group(1) is not None:
        opening, closing = "$$", "$$"
    elif match.group(2) is not None:
        opening, closing = r"\[", r"\]"
    elif match.group(3) is not None:
        opening, closing = "$", "$"
    else:
        opening, closing = r"\(", r"\)"

    blanks = list(_MATH_FILL_BLANK_RE.finditer(latex))
    if not blanks:
        return match.group(0) if latex == original_latex else f"{opening}{latex}{closing}"

    trailing_blank = blanks[-1]
    prefix = latex[: trailing_blank.start()].rstrip()
    suffix = latex[trailing_blank.end() :]
    if not suffix.strip() and prefix.count("{") == prefix.count("}"):
        normalized_prefix = _MATH_FILL_BLANK_RE.sub(
            lambda _match: r"\underline{\hspace{2em}}", prefix
        )
        equation = f"{opening}{normalized_prefix}{closing}" if normalized_prefix else ""
        separator = " " if equation else ""
        return f"{equation}{separator}{trailing_blank.group(0)}{suffix}"

    normalized = _MATH_FILL_BLANK_RE.sub(
        lambda _match: r"\underline{\hspace{2em}}", latex
    )
    return f"{opening}{normalized}{closing}"


def preflight_practice_inline_expressions(data: dict[str, Any]) -> list[str]:
    """Fail before Word generation when visible inline math cannot become OMML."""

    issues: list[str] = []
    for index, item in enumerate(data.get("exercises") or [], start=1):
        if not isinstance(item, dict) or item.get("generation_status") == "failed":
            continue
        number = str(item.get("number") or index)
        fields: list[tuple[str, Any]] = [("题干", item.get("stem"))]
        fields.extend(
            (f"选项 {option_index}", option.get("text") if isinstance(option, dict) else option)
            for option_index, option in enumerate(item.get("options") or [], start=1)
        )
        for table_index, table in enumerate(item.get("tables") or [], start=1):
            if not isinstance(table, dict):
                continue
            fields.extend((f"表格 {table_index} 表头", value) for value in table.get("headers") or [])
            for row_index, row in enumerate(table.get("rows") or [], start=1):
                if isinstance(row, list):
                    fields.extend((f"表格 {table_index} 第 {row_index} 行", value) for value in row)
        for field, value in fields:
            text = normalize_practice_markup(value, limit=12000)
            for match in _INLINE_MATH_RE.finditer(text):
                latex = next((group for group in match.groups() if group is not None), "").strip()
                if not latex:
                    continue
                plan = build_expression_render_plan(
                    latex,
                    question_id=str(item.get("exercise_id") or item.get("question_id") or number),
                    location=f"practice_{field}",
                    role="relation",
                    display=False,
                )
                if error := preflight_expression_render(plan):
                    issues.append(f"第 {number} 题{field}包含无法生成 Word 公式对象的行内公式：{error}")
    return list(dict.fromkeys(issues))


def _repair_json_escaped_latex(value: str) -> str:
    return repair_json_escaped_latex(value)


def _normalize_standard_state_latex(value: str) -> str:
    return normalize_standard_state_latex(value)


def has_unrenderable_practice_markup(value: Any) -> bool:
    """Allow markup handled by the DOCX renderer; reject only leftovers."""
    text = _text(value, 12000)
    if any((ord(char) < 32 and char != "\n") or ord(char) == 127 for char in text):
        return True
    without_math = _INLINE_MATH_RE.sub("", text)
    if "$" in without_math or "\\(" in without_math or "\\)" in without_math:
        return True
    return bool(_BARE_LATEX_RE.search(without_math))


def _has_unrenderable_markup(value: Any) -> bool:
    """Backward-compatible internal alias for existing export callers."""
    return has_unrenderable_practice_markup(value)


def _read_latex_braced_argument(text: str, start: int) -> int | None:
    """Return the offset after one balanced LaTeX braced argument."""
    if start >= len(text) or text[start] != "{":
        return None
    depth = 0
    cursor = start
    while cursor < len(text):
        char = text[cursor]
        if char == "\\":
            cursor += 2
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return cursor + 1
        cursor += 1
    return None


def _repair_bare_latex_segment(text: str) -> str:
    """Wrap safe, standalone LaTeX commands that were emitted in prose.

    Models commonly emit chemical formulae such as ``\\mathrm{H_2O(l)}``
    directly in a sentence.  They are unambiguous but cannot be rendered by
    the browser or DOCX converters without math delimiters.  Only commands
    with a complete, known number of braced arguments are repaired here;
    anything ambiguous remains visible for the generation gate to reject.
    """
    output: list[str] = []
    cursor = 0
    command_pattern = re.compile(r"\\([A-Za-z]+)")
    while cursor < len(text):
        match = command_pattern.match(text, cursor)
        if not match:
            output.append(text[cursor])
            cursor += 1
            continue
        command = match.group(1)
        arity = _BARE_LATEX_BRACED_COMMANDS.get(command)
        if not arity:
            output.append(match.group(0))
            cursor = match.end()
            continue
        end = match.end()
        valid = True
        for _ in range(arity):
            while end < len(text) and text[end].isspace():
                end += 1
            argument_end = _read_latex_braced_argument(text, end)
            if argument_end is None:
                valid = False
                break
            end = argument_end
        if not valid:
            output.append(match.group(0))
            cursor = match.end()
            continue
        output.append(f"${text[cursor:end]}$")
        cursor = end
    return "".join(output)


def normalize_practice_markup(value: Any, *, limit: int = 12000) -> str:
    """Repair bare LaTeX and keep fill-in blanks out of invalid math syntax.

    Existing mathematical notation is otherwise preserved.  The same
    normalization is safe both for new generations and for historical exports.
    """
    text = _text(value, limit)
    parts: list[str] = []
    cursor = 0
    for match in _INLINE_MATH_RE.finditer(text):
        parts.append(_repair_bare_latex_segment(text[cursor : match.start()]))
        parts.append(_normalize_inline_math_fill_blanks(match))
        cursor = match.end()
    parts.append(_repair_bare_latex_segment(text[cursor:]))
    return "".join(parts)


_PRACTICE_QUESTION_TITLE_PREFIX_RE = re.compile(
    r"^\s*(?:#{1,6}\s*)?(?:第\s*\d+\s*题|题目\s*\d+)\s*(?:[：:.．、-]\s*)?"
)
_PRACTICE_SUBQUESTION_PREFIX_RE = re.compile(
    r"^\s*(?:[（(]\s*(?P<parenthesized>\d{1,2})\s*[）)]|(?P<plain>\d{1,2})\s*[.)）．、])\s*(?:[、.．:：-]\s*)?"
)
_PRACTICE_CIRCLED_ITEM_RE = re.compile(r"^\s*(?P<marker>[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳])\s*(?:[、.．:：-]\s*)?")
_PRACTICE_SECTION_LABEL_RE = re.compile(r"^\s*【(?:材料|已知|说明|注|提示)】")


def normalize_practice_question_text(value: Any, *, limit: int = 12000) -> str:
    """Own the generated question layout instead of trusting model line breaks.

    The generation schema stores a question as a text field, so this light
    normalizer gives every renderer the same hierarchy: the outer ``第 N 题``
    heading is supplied by the renderer, first-level subquestions use ASCII
    parentheses (``(1)``), and ordinary provider line wraps are joined into
    normal prose paragraphs.
    """
    source = normalize_practice_markup(value, limit=limit).replace("\r\n", "\n").replace("\r", "\n")
    # Some providers double-escape a requested paragraph break, leaving
    # visible ``\\n(1)`` text after JSON decoding. Convert only the safe,
    # unambiguous subquestion form so LaTeX commands remain untouched.
    source = _LITERAL_SUBQUESTION_BREAK_RE.sub("\n", source)
    blocks: list[str] = []
    paragraph: list[str] = []
    seen_content = False
    next_subquestion_number = 1

    def flush_paragraph() -> None:
        nonlocal paragraph
        text = " ".join(part for part in paragraph if part).strip()
        if text:
            blocks.append(text)
        paragraph = []

    for raw_line in source.split("\n"):
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        if not seen_content:
            line = _PRACTICE_QUESTION_TITLE_PREFIX_RE.sub("", line, count=1).strip()
            if not line:
                continue
        marker = _PRACTICE_SUBQUESTION_PREFIX_RE.match(line)
        circled = _PRACTICE_CIRCLED_ITEM_RE.match(line)
        if marker:
            flush_paragraph()
            # First-level item numbers are renderer-owned. Re-numbering here
            # prevents a malformed model response such as ``(1), (1)`` or
            # ``(2)`` from leaking into the final question sheet.
            blocks.append(f"({next_subquestion_number}) {line[marker.end():].strip()}".rstrip())
            next_subquestion_number += 1
            seen_content = True
            continue
        if circled:
            flush_paragraph()
            blocks.append(f"{circled.group('marker')} {line[circled.end():].strip()}".rstrip())
            seen_content = True
            continue
        if _PRACTICE_SECTION_LABEL_RE.match(line):
            flush_paragraph()
            blocks.append(line)
            seen_content = True
            continue
        paragraph.append(line)
        seen_content = True
    flush_paragraph()
    return "\n\n".join(blocks)


def audit_practice_export_data(data: dict[str, Any]) -> list[str]:
    """Report deterministic visible-text defects that forbid formal export."""
    issues: list[str] = []
    for index, item in enumerate(data.get("exercises") or [], start=1):
        if not isinstance(item, dict):
            issues.append(f"第{index}题不是有效题目对象")
            continue
        question_number = str(item.get("number") or index)
        fields: list[tuple[str, Any]] = [
            ("stem", item.get("stem")),
            ("target_skill", item.get("target_skill")),
        ]
        for option_index, option in enumerate(item.get("options") or [], start=1):
            option_text = option.get("text") if isinstance(option, dict) else option
            fields.append((f"选项 {option_index}", option_text))
        for table_index, table in enumerate(item.get("tables") or [], start=1):
            if not isinstance(table, dict):
                continue
            fields.append((f"表格 {table_index} 标题", table.get("title")))
            fields.extend((f"表格 {table_index} 表头", header) for header in table.get("headers") or [])
            for row_index, row in enumerate(table.get("rows") or [], start=1):
                if isinstance(row, list):
                    fields.extend((f"表格 {table_index} 第 {row_index} 行", cell) for cell in row)
        for field, raw_value in fields:
            value = _text(raw_value, 12000)
            if _has_unrenderable_markup(value):
                issues.append(f"第{question_number}题 {field} 含未渲染 Markdown/LaTeX 标记")
            if _SELF_CORRECTION_RE.search(value):
                issues.append(f"第{question_number}题 {field} 含自我纠错或模型草稿文本")
        for reason in practice_stem_answer_leak_reasons(item):
            issues.append(f"第{question_number}题存在题面答案泄漏：{reason}")
    return list(dict.fromkeys(issues))


def _split_export_paragraphs(value: Any, limit: int = 10000) -> list[str]:
    """Preserve author-intended blocks as Word paragraphs, never soft breaks."""
    text = _text(value, limit).replace("\r\n", "\n").replace("\r", "\n")
    return [line.strip() for line in text.split("\n") if line.strip()]


def _inline_export_text(value: Any, limit: int = 10000) -> str:
    return " ".join(_split_export_paragraphs(value, limit))


def _add_rich_text(paragraph, value: Any, *, limit: int = 10000) -> None:
    """Write text and inline LaTeX without leaking raw dollar delimiters."""
    text = normalize_practice_markup(_inline_export_text(value, limit), limit=limit)
    cursor = 0
    for match in _INLINE_MATH_RE.finditer(text):
        plain = text[cursor : match.start()]
        latex = _normalize_standard_state_latex(
            next((group for group in match.groups() if group is not None), "").strip()
        )
        # Generated scientific units are often written as mol$^{-1}$ or
        # cm$^3$. A script without its adjacent base becomes an empty-box
        # formula in Word, so move the complete unit into the Office formula.
        if latex.startswith(("^", "_")):
            base_match = re.search(r"([A-Za-z]+)$", plain)
            if base_match:
                base = base_match.group(1)
                plain = plain[:base_match.start()]
                latex = rf"\mathrm{{{base}}}{latex}"
        if plain:
            _add_plain_with_equations(paragraph, plain)
        try:
            paragraph._p.append(render_expression_omml(latex, display=False, location="practice_inline"))
        except Exception:
            run = paragraph.add_run(latex)
            _set_run(run, color=DARK_BLUE)
        cursor = match.end()
    if text[cursor:]:
        _add_plain_with_equations(paragraph, text[cursor:])


def _set_body_paragraph(paragraph, *, option: bool = False) -> None:
    paragraph.paragraph_format.left_indent = Pt(PRACTICE_TEXT_CONTRACT.list_left_indent_pt) if option else Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = (
        Pt(-PRACTICE_TEXT_CONTRACT.list_hanging_indent_pt)
        if option
        else Pt(PRACTICE_TEXT_CONTRACT.first_line_indent_pt)
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = PRACTICE_TEXT_CONTRACT.line_spacing
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _set_table_paragraph(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = PRACTICE_TEXT_CONTRACT.table_line_spacing
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_plain_with_equations(paragraph, text: str) -> None:
    """Convert un-delimited equations and domain notation without touching prose."""
    cursor = 0
    for plan in build_text_expression_render_plans(text):
        if plan.start > cursor:
            _add_markdown_runs(paragraph, text[cursor:plan.start])
        if plan.preserve_parentheses:
            _add_markdown_runs(paragraph, "（")
        try:
            paragraph._p.append(
                render_expression_omml(
                    plan.render_latex,
                    display=False,
                    location="practice_domain_notation",
                    expression_kind=plan.expression_kind,
                )
            )
        except Exception:
            _add_markdown_runs(paragraph, plan.raw)
        if plan.preserve_parentheses:
            _add_markdown_runs(paragraph, "）")
        cursor = plan.end
    if cursor < len(text):
        _add_markdown_runs(paragraph, text[cursor:])


def _add_markdown_runs(paragraph, text: str) -> None:
    """Render the small Markdown subset accepted from generated exercises."""
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            _set_run(run)
        run = paragraph.add_run(match.group(1))
        _set_run(run, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run(run)


def _option_text(value: Any) -> str:
    text = _text(value, 3000)
    text = re.sub(r"^\s*[A-Ha-h]\s*(?:[.．、:：]|[）)])\s*", "", text, count=1)
    text = re.sub(r"^\s*[（(]\s*[A-Ha-h]\s*[）)]\s*", "", text, count=1)
    text = text.replace("**", "")
    return re.sub(r"[。．.!！?？;；,，]+$", "", text).rstrip() + "。" if text else ""


def _add_question(doc: Document, item: dict[str, Any], index: int) -> None:
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run(f"第 {index} 题")
    stem = doc.add_paragraph()
    _set_body_paragraph(stem)
    if item.get("generation_status") == "failed":
        error = item.get("generation_error") if isinstance(item.get("generation_error"), dict) else {}
        message = _inline_export_text(error.get("message"), 500) or "上游模型未返回本题。"
        failed_run = stem.add_run(f"【生成失败】{message}已保留本题位置，请稍后在系统中重新生成本题。")
        _set_run(failed_run, size=11, bold=True, color=ERROR_RED)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)
        return
    stem_parts = _split_export_paragraphs(normalize_practice_question_text(item.get("stem"))) or [""]
    _add_rich_text(stem, stem_parts[0])
    for part in stem_parts[1:]:
        continuation = doc.add_paragraph()
        _set_body_paragraph(continuation)
        _add_rich_text(continuation, part)
    _add_structured_assets(doc, item, "stem")
    for option_index, option in enumerate(item.get("options") or []):
        if not isinstance(option, dict):
            continue
        option_parts = _split_export_paragraphs(_option_text(option.get("text")), 3000) or [""]
        paragraph = doc.add_paragraph()
        _set_body_paragraph(paragraph, option=True)
        label = paragraph.add_run(f"{chr(65 + option_index)}. ")
        _set_run(label, bold=False, color=INK)
        _add_rich_text(paragraph, option_parts[0], limit=3000)
        for part in option_parts[1:]:
            continuation = doc.add_paragraph()
            _set_body_paragraph(continuation, option=True)
            _add_rich_text(continuation, part, limit=3000)


def _add_answer(doc: Document, item: dict[str, Any], index: int) -> None:
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run(f"第 {index} 题")
    answer = doc.add_paragraph()
    _set_body_paragraph(answer)
    label = answer.add_run("参考答案：")
    _set_run(label, bold=True, color=BLUE)
    answer_parts = _split_export_paragraphs(item.get("answer")) or [""]
    _add_rich_text(answer, answer_parts[0])
    for part in answer_parts[1:]:
        continuation = doc.add_paragraph()
        _set_body_paragraph(continuation)
        _add_rich_text(continuation, part)
    _add_structured_assets(doc, item, "solution")
    steps = item.get("solution_steps") if isinstance(item.get("solution_steps"), list) else []
    if steps:
        subheading = doc.add_paragraph(style="Heading 3")
        subheading.add_run("解析")
        number_id = _new_solution_step_numbering(doc)
        for step in steps:
            paragraph = doc.add_paragraph(style="List Number")
            _assign_solution_step_numbering(paragraph, number_id)
            _add_rich_text(paragraph, step, limit=3000)
    points = [_text(value, 100) for value in (item.get("knowledge_points") or []) if _text(value, 100)]
    if points:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("涉及知识点：" + "、".join(points))
        _set_run(run, size=9.5, color=MUTED)


def _new_solution_step_numbering(doc: Document) -> int:
    """Create a fresh 1-based numbering instance for one question's steps."""

    style = doc.styles["List Number"]
    base_num_id = int(style._element.xpath("./w:pPr/w:numPr/w:numId/@w:val")[0])
    numbering = doc.part.numbering_part.element
    base_num = numbering.num_having_numId(base_num_id)
    number = numbering.add_num(int(base_num.abstractNumId.val))
    number.add_lvlOverride(0).add_startOverride(1)
    return int(number.numId)


def _assign_solution_step_numbering(paragraph, number_id: int) -> None:
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = int(number_id)


def _matches_location(value: Any, location: str) -> bool:
    return location in (_text(value, 30) or "stem")


def _add_formula(doc: Document, formula: dict[str, Any]) -> None:
    caption = _inline_export_text(formula.get("caption"), 300)
    if caption:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(caption)
        _set_run(run, size=9.5, color=MUTED)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        paragraph._p.append(
            render_expression_omml(
                _normalize_standard_state_latex(_text(formula.get("latex"), 3000)),
                display=True,
                location="practice_formula",
            )
        )
    except Exception:
        run = paragraph.add_run(_text(formula.get("latex"), 3000))
        _set_run(run, size=11, color=DARK_BLUE)


def _add_data_table(doc: Document, spec: dict[str, Any]) -> None:
    title = _inline_export_text(spec.get("title"), 300)
    if title:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(title)
        _set_run(run, size=9.5, bold=True, color=DARK_BLUE)
    headers = spec.get("headers") if isinstance(spec.get("headers"), list) else []
    rows = spec.get("rows") if isinstance(spec.get("rows"), list) else []
    columns = max(len(headers), max((len(row) for row in rows if isinstance(row, list)), default=0))
    if not columns:
        return
    table = doc.add_table(rows=1 if headers else 0, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    width = Inches(6.5 / columns)
    if headers:
        for index, value in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_table_paragraph(cell.paragraphs[0])
            _add_rich_text(cell.paragraphs[0], value, limit=500)
    for raw in rows:
        if not isinstance(raw, list):
            continue
        cells = table.add_row().cells
        for index in range(columns):
            cells[index].width = width
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_table_paragraph(cells[index].paragraphs[0])
            _add_rich_text(cells[index].paragraphs[0], raw[index] if index < len(raw) else "", limit=500)
    doc.add_paragraph()


def _chart_png(spec: dict[str, Any]) -> BytesIO | None:
    series = spec.get("series") if isinstance(spec.get("series"), list) else []
    has_points = any(isinstance(row, dict) and len(row.get("points") or []) >= 2 for row in series)
    nodes = [node for node in (spec.get("nodes") or []) if isinstance(node, dict) and node.get("id")]
    if not has_points and len(nodes) < 2:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.font_manager import FontProperties
        from matplotlib.patches import Circle, Ellipse, FancyBboxPatch, Polygon
    except Exception:
        return None
    chart_font_path = next((path for path in CHART_FONT_PATHS if path.exists()), None)
    chart_font = FontProperties(fname=str(chart_font_path)) if chart_font_path else None
    fig, ax = plt.subplots(figsize=(6.2, 3.2), dpi=160)
    if not has_points:
        ax.axis("off")
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        by_id = {str(node["id"]): node for node in nodes}
        plane_edges = [
            edge for edge in (spec.get("edges") or [])
            if isinstance(edge, dict)
            and re.search(r"晶面|阴影|平行\s*z", _text(edge.get("label"), 120), re.IGNORECASE)
            and str(edge.get("from")) in by_id
            and str(edge.get("to")) in by_id
        ]
        plane_ids: list[str] = []
        for edge in plane_edges:
            for endpoint in (str(edge.get("from")), str(edge.get("to"))):
                if endpoint not in plane_ids:
                    plane_ids.append(endpoint)
        if len(plane_ids) >= 3:
            polygon_points = [(float(by_id[node_id]["x"]), float(by_id[node_id]["y"])) for node_id in plane_ids]
            ax.add_patch(Polygon(
                polygon_points,
                closed=True,
                facecolor="#93c5fd",
                edgecolor="#2563eb",
                linewidth=1.1,
                alpha=0.28,
                hatch="//",
                zorder=0,
            ))
        for edge in spec.get("edges") or []:
            if not isinstance(edge, dict) or str(edge.get("from")) not in by_id or str(edge.get("to")) not in by_id:
                continue
            start, end = by_id[str(edge["from"])], by_id[str(edge["to"])]
            ax.annotate(
                "", xy=(float(end["x"]), float(end["y"])), xytext=(float(start["x"]), float(start["y"])),
                arrowprops={"arrowstyle": "->" if edge.get("directed") is not False else "-", "color": "#475569", "lw": 1.4},
                zorder=1,
            )
            if _text(edge.get("label"), 120):
                ax.text((float(start["x"]) + float(end["x"])) / 2, (float(start["y"]) + float(end["y"])) / 2,
                        _text(edge.get("label"), 120), ha="center", va="bottom", fontsize=8, fontproperties=chart_font)
        for node in nodes:
            x, y = float(node["x"]), float(node["y"])
            shape = _text(node.get("shape"), 20)
            label = _text(node.get("label"), 200)
            if not label:
                patch = Circle((x, y), 0.012, facecolor="#0f172a", edgecolor="#0f172a", lw=0.8, zorder=2)
            elif shape == "circle":
                patch = Circle((x, y), 0.065, facecolor="#eff6ff", edgecolor="#2563eb", lw=1.4, zorder=2)
            elif shape == "ellipse":
                patch = Ellipse((x, y), 0.2, 0.11, facecolor="#eff6ff", edgecolor="#2563eb", lw=1.4, zorder=2)
            else:
                patch = FancyBboxPatch((x - 0.09, y - 0.05), 0.18, 0.1, boxstyle="round,pad=0.01", facecolor="#eff6ff", edgecolor="#2563eb", lw=1.4, zorder=2)
            ax.add_patch(patch)
            if label:
                ax.text(x, y, label, ha="center", va="center", fontsize=8.5, fontproperties=chart_font, zorder=3)
        buffer = BytesIO()
        fig.savefig(buffer, format="png", bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buffer.seek(0)
        return buffer
    kind = _text(spec.get("figure_type"), 30)
    for index, row in enumerate(series):
        if not isinstance(row, dict):
            continue
        points = row.get("points") or []
        xs = [point[0] for point in points if isinstance(point, list) and len(point) >= 2]
        ys = [point[1] for point in points if isinstance(point, list) and len(point) >= 2]
        if not xs:
            continue
        label = _text(row.get("name"), 100) or f"Series {index + 1}"
        if kind == "bar":
            ax.bar(xs, ys, alpha=0.8, label=label)
        elif kind == "scatter":
            ax.scatter(xs, ys, s=22, label=label)
        else:
            ax.plot(xs, ys, marker="o", linewidth=1.8, markersize=3.5, label=label)
    for node in nodes:
        try:
            x, y = float(node["x"]), float(node["y"])
        except (TypeError, ValueError, KeyError):
            continue
        ax.scatter([x], [y], s=24, color="#0f172a", zorder=5)
        ax.annotate(
            _text(node.get("label"), 200),
            (x, y),
            xytext=(5, 6),
            textcoords="offset points",
            fontsize=8,
            fontproperties=chart_font,
            zorder=6,
        )
    ax.set_xlabel(_text(spec.get("x_label"), 100), fontproperties=chart_font)
    ax.set_ylabel(_text(spec.get("y_label"), 100), fontproperties=chart_font)
    ax.grid(alpha=0.2)
    if len(series) > 1 or any(_text(row.get("name"), 100) for row in series if isinstance(row, dict)):
        ax.legend(frameon=False, prop=chart_font)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _add_figure(doc: Document, spec: dict[str, Any]) -> None:
    title = _inline_export_text(spec.get("title"), 300)
    if title:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # A figure title at the bottom of one page with its image on the next
        # is visually misleading. Keep the title attached to the following
        # image paragraph while still allowing the description to flow.
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(title)
        _set_run(run, size=9.5, bold=True, color=DARK_BLUE)
    image_path = _figure_image_path(spec) if _valid_figure_image(spec) else None
    image = str(image_path) if image_path is not None else _chart_png(spec)
    if image is not None:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_run = paragraph.add_run()
        picture_run.add_picture(image, width=Inches(6.0))
        figure_id = _inline_export_text(spec.get("figure_id"), 120) or "practice_figure"
        alt_text = _inline_export_text(
            spec.get("description") or spec.get("title") or f"题图 {figure_id}",
            500,
        )
        for doc_property in picture_run._r.xpath(".//wp:docPr"):
            doc_property.set("name", figure_id)
            doc_property.set("title", title or figure_id)
            doc_property.set("descr", alt_text)
    description = _inline_export_text(spec.get("description"), 1500)
    if description:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(description)
        _set_run(run, size=9.5, color=MUTED)


def _add_structured_assets(doc: Document, item: dict[str, Any], location: str) -> None:
    for formula in item.get("formulas") or []:
        if isinstance(formula, dict) and _matches_location(formula.get("location"), location):
            _add_formula(doc, formula)
    for table in item.get("tables") or []:
        if isinstance(table, dict) and _matches_location(table.get("location"), location):
            _add_data_table(doc, table)
    for figure in item.get("figures") or []:
        if isinstance(figure, dict) and _matches_location(figure.get("location"), location):
            _add_figure(doc, figure)


def _apply_run_fonts(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:ascii"), ASCII_FONT)
            fonts.set(qn("w:hAnsi"), ASCII_FONT)
            fonts.set(qn("w:eastAsia"), CJK_FONT)
            fonts.set(qn("w:cs"), ASCII_FONT)


def _save_document(doc: Document) -> bytes:
    _apply_run_fonts(doc)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _new_practice_document(data: dict[str, Any], *, document_kind: str) -> Document:
    if not isinstance(data, dict) or not isinstance(data.get("exercises"), list) or not data["exercises"]:
        raise ValueError("没有可导出的专项练习。")
    doc = Document()
    _configure_document(doc)
    _add_title_block(doc, data, document_kind=document_kind)
    return doc


def build_practice_question_docx(
    data: dict[str, Any],
    *,
    progress_callback: Callable[[int, int], None] | None = None,
) -> bytes:
    doc = _new_practice_document(data, document_kind="questions")
    doc.add_paragraph("练习题", style="Heading 1")
    total = len(data["exercises"])
    for index, item in enumerate(data["exercises"], start=1):
        if isinstance(item, dict):
            _add_question(doc, item, index)
        if progress_callback is not None:
            progress_callback(index, total)
    return _save_document(doc)


def build_practice_solution_docx(data: dict[str, Any]) -> bytes:
    doc = _new_practice_document(data, document_kind="solutions")
    doc.add_paragraph("参考答案与解析", style="Heading 1")
    for index, item in enumerate(data["exercises"], start=1):
        if isinstance(item, dict):
            _add_answer(doc, item, index)
    return _save_document(doc)


def build_practice_docx(data: dict[str, Any]) -> bytes:
    """Backward-compatible combined export for older callers."""
    if not isinstance(data, dict) or not isinstance(data.get("exercises"), list) or not data["exercises"]:
        raise ValueError("没有可导出的专项练习。")
    doc = Document()
    _configure_document(doc)
    _add_title_block(doc, data)
    doc.add_paragraph("练习题", style="Heading 1")
    for index, item in enumerate(data["exercises"], start=1):
        if isinstance(item, dict):
            _add_question(doc, item, index)
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("参考答案与解析", style="Heading 1")
    for index, item in enumerate(data["exercises"], start=1):
        if isinstance(item, dict):
            _add_answer(doc, item, index)
    return _save_document(doc)
