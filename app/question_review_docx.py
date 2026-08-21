from __future__ import annotations

import json
import hashlib
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .docx_v4 import build_docx_from_fragments
from .final_acceptance import model_retry_summary
from .question_requirements import answer_figure_required
from .question_types import question_has_type, question_kind
from .render_word import export_docx_to_pdf, render_pdf_to_png
from .review_export import build_question_review

NON_REVIEW_QUALITY_WARNING_CODES = {
    "calculation_answer_missing_unit",
}


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def _qid(value: dict[str, Any]) -> str:
    return str(value.get("question_id", "")).strip()


def _add_note(notes: list[str], text: str) -> None:
    text = str(text or "").strip()
    if text and text not in notes:
        notes.append(text)


def _ignore_quality_note(filename: str, item: dict[str, Any], question: dict[str, Any] | None) -> bool:
    if filename != "content_quality_audit.json":
        return False
    code = str(item.get("code", ""))
    if code in NON_REVIEW_QUALITY_WARNING_CODES:
        return True
    kind = question_kind(question or {})
    if code.startswith("choice_") and kind != "choice":
        return True
    if code.startswith("calculation_") and kind != "calculation":
        return True
    return False


def _quality_notes_by_qid(stage_dir: Path) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {}
    structured_exam = _read_json(stage_dir / "structured_exam.json")
    questions_by_qid = {_qid(item): item for item in structured_exam.get("items", []) if _qid(item)}
    for filename in ("content_quality_audit.json", "answer_coverage_audit.json"):
        data = _read_json(stage_dir / filename)
        for item in data.get("issues", []) + data.get("warnings", []):
            if isinstance(item, dict):
                if _ignore_quality_note(filename, item, questions_by_qid.get(_qid(item))):
                    continue
                qid = _qid(item)
                message = str(item.get("message") or item.get("code") or "").strip()
            else:
                text = str(item)
                qid = text.split(":", 1)[0].strip()
                message = text
            if qid and message:
                out.setdefault(qid, [])
                _add_note(out[qid], message)
    allowed = _read_json(stage_dir / "user_allowed_audit_issues.json")
    for entry in allowed.get("entries", []) if isinstance(allowed, dict) else []:
        stage = str(entry.get("stage") or "审查")
        for item in entry.get("items", []) if isinstance(entry, dict) else []:
            if not isinstance(item, dict):
                continue
            qid = _qid(item)
            message = str(item.get("message") or "").strip()
            if qid and message:
                _add_note(out.setdefault(qid, []), f"用户已允许通过的{stage}问题：{message}")
    return out


def _fragment_text(block: dict[str, Any]) -> str:
    parts: list[str] = []
    for segment in block.get("segments", []):
        if not isinstance(segment, dict):
            continue
        if segment.get("type") == "text":
            parts.append(str(segment.get("text", "")).strip())
        elif segment.get("type") == "formula_ref":
            parts.append(f"[公式对象: {segment.get('formula_id', '')}]")
        elif segment.get("type") == "image_ref":
            parts.append(f"[图片: {segment.get('image_id') or segment.get('path') or ''}]")
    return "".join(part for part in parts if part)


def _resolve_image_path(stage_dir: Path, value: Any) -> Path | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    path = Path(raw)
    if path.exists():
        return path
    if not path.is_absolute():
        candidate = stage_dir / raw.replace("\\", "/")
        if candidate.exists():
            return candidate
    marker = "stage_outputs"
    normalized = raw.replace("\\", "/")
    if marker in normalized:
        suffix = normalized.split(marker, 1)[1].lstrip("/")
        candidate = stage_dir / suffix
        if candidate.exists():
            return candidate
    return None


def _figure_path_for_id(stage_dir: Path, figure_id: str) -> Path | None:
    figure_id = str(figure_id or "").strip()
    if not figure_id:
        return None
    candidate = stage_dir / "figures" / f"{figure_id}.png"
    return candidate if candidate.exists() else None


def _add_figure_review_record(
    records: dict[str, dict[str, Any]],
    *,
    question_id: Any,
    figure_id: Any = "",
    path: Path | None,
    stage: str,
    status: str = "",
    summary: str = "",
    role: str = "archive",
) -> None:
    qid = str(question_id or "").strip()
    if not qid or path is None or not path.exists():
        return
    try:
        content_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    except OSError:
        content_hash = str(path.resolve())
    figure_key = str(figure_id or path.stem).strip() or path.stem
    key = f"{qid}\0{figure_key}\0{content_hash}"
    role_priority = {"archive": 10, "candidate": 30, "selected_candidate": 60, "official": 100}
    item = records.setdefault(
        key,
        {
            "question_id": qid,
            "figure_id": figure_key,
            "path": path,
            "content_hash": content_hash,
            "role": role,
            "stages": [],
            "statuses": [],
            "summaries": [],
        },
    )
    if role_priority.get(role, 0) > role_priority.get(str(item.get("role") or ""), 0):
        item["role"] = role
        item["path"] = path
    label = str(stage or "").strip()
    if label and label not in item["stages"]:
        item["stages"].append(label)
    status_text = str(status or "").strip()
    if status_text and status_text not in item["statuses"]:
        item["statuses"].append(status_text)
    summary_text = str(summary or "").strip()
    if summary_text and summary_text not in item["summaries"]:
        item["summaries"].append(summary_text)


def _qa_status(item: dict[str, Any]) -> tuple[str, str]:
    qa = item.get("qa") if isinstance(item.get("qa"), dict) else {}
    if qa.get("ok") is True:
        status = "视觉QA通过"
    elif qa:
        status = "视觉QA未通过"
    else:
        status = ""
    summary = str(qa.get("summary") or qa.get("error") or "").strip()
    return status, summary


def collect_question_figure_review_items(stage_dir: Path) -> list[dict[str, Any]]:
    structured_exam = _read_json(stage_dir / "structured_exam.json")
    figure_qids = {
        _qid(question)
        for question in structured_exam.get("items", [])
        if isinstance(question, dict) and _qid(question) and answer_figure_required(question)
    }
    records: dict[str, dict[str, Any]] = {}

    fragments_data = _read_json(stage_dir / "answer_fragments.json")
    for fragment in fragments_data.get("fragments", []) if isinstance(fragments_data, dict) else []:
        if not isinstance(fragment, dict):
            continue
        qid = _qid(fragment)
        for block in fragment.get("blocks", []) or []:
            if not isinstance(block, dict):
                continue
            for segment in block.get("segments", []) or []:
                if not isinstance(segment, dict) or segment.get("type") != "image_ref":
                    continue
                image_id = str(segment.get("image_id") or "").strip()
                path = _resolve_image_path(stage_dir, segment.get("path")) or _figure_path_for_id(stage_dir, image_id)
                _add_figure_review_record(
                    records,
                    question_id=qid,
                    figure_id=image_id,
                    path=path,
                    stage="最终答案引用图",
                    status="正式采用",
                    role="official",
                )

    specs_data = _read_json(stage_dir / "figure_specs.json")
    for spec in specs_data.get("figures", []) if isinstance(specs_data, dict) else []:
        if not isinstance(spec, dict):
            continue
        qid = _qid(spec)
        figure_id = str(spec.get("figure_id") or "").strip()
        source = str(spec.get("source") or spec.get("kind") or "图规格").strip()
        path = _figure_path_for_id(stage_dir, figure_id)
        run_result = spec.get("run_result") if isinstance(spec.get("run_result"), dict) else {}
        status = "渲染通过" if run_result.get("ok") is True else ("渲染未通过" if run_result else "")
        _add_figure_review_record(
            records,
            question_id=qid,
            figure_id=figure_id,
            path=path,
            stage=f"图规格/{source}",
            status=status,
            role="official",
        )

    direct = _read_json(stage_dir / "direct_model_figures.json")
    for item in direct.get("generated", []) if isinstance(direct, dict) else []:
        if not isinstance(item, dict):
            continue
        path = _resolve_image_path(stage_dir, item.get("path"))
        _add_figure_review_record(
            records,
            question_id=item.get("question_id"),
            figure_id=item.get("figure_id"),
            path=path,
            stage=f"生图模型兜底/{item.get('model') or ''}".rstrip("/"),
            status="生成成功",
            role="official",
        )

    stage_manifest = _read_json(stage_dir / "figure_stage_images.json")
    for item in stage_manifest.get("items", []) if isinstance(stage_manifest, dict) else []:
        if not isinstance(item, dict):
            continue
        path = _resolve_image_path(stage_dir, item.get("path"))
        stage = str(item.get("stage") or "阶段归档图").strip()
        source = str(item.get("source") or item.get("kind") or "").strip()
        _add_figure_review_record(
            records,
            question_id=item.get("question_id"),
            figure_id=item.get("figure_id"),
            path=path,
            stage=f"阶段归档/{stage}" + (f"/{source}" if source else ""),
        )

    for report_name in ("figure_visual_qa.json",):
        report = _read_json(stage_dir / report_name)
        for item in report.get("items", []) if isinstance(report, dict) else []:
            if not isinstance(item, dict):
                continue
            status, summary = _qa_status(item)
            path = _resolve_image_path(stage_dir, item.get("path")) or _figure_path_for_id(stage_dir, item.get("figure_id"))
            _add_figure_review_record(
                records,
                question_id=item.get("question_id"),
                figure_id=item.get("figure_id"),
                path=path,
                stage="视觉QA正式审核",
                status=status,
                summary=summary,
                role="official",
            )

    # A task may retain reports from earlier repair phases.  Only the newest report
    # for each figure represents the current candidate decision; older reports are
    # audit evidence, not additional user-facing candidates.
    latest_targets: dict[tuple[str, str], tuple[float, Path, Any, dict[str, Any]]] = {}
    for repair_path in sorted(stage_dir.glob("figure_visual_qa_repair*.json")):
        report = _read_json(repair_path)
        for round_item in report.get("rounds", []) if isinstance(report, dict) else []:
            if not isinstance(round_item, dict):
                continue
            round_no = round_item.get("round") or ""
            for target in round_item.get("targets", []) or []:
                if not isinstance(target, dict):
                    continue
                qid = str(target.get("question_id") or "").strip()
                figure_id = str(target.get("figure_id") or "").strip()
                key = (qid, figure_id)
                modified = repair_path.stat().st_mtime
                previous = latest_targets.get(key)
                if previous is None or modified >= previous[0]:
                    latest_targets[key] = (modified, repair_path, round_no, target)

    for (qid, figure_id), (_, repair_path, round_no, target) in latest_targets.items():
        selected_strategy = str(target.get("selected_strategy") or "").strip()
        target_selected = target.get("selected") is True
        for candidate in target.get("candidates", []) or []:
            if not isinstance(candidate, dict):
                continue
            strategy = str(candidate.get("strategy") or "candidate").strip()
            path = _resolve_image_path(stage_dir, candidate.get("path"))
            selected = target_selected and bool(selected_strategy) and strategy == selected_strategy
            if selected:
                status = "已采用（回修候选，视觉QA通过）"
                role = "selected_candidate"
            elif candidate.get("passed") is True:
                status = "未采用（视觉QA通过但未选为最终图）"
                role = "candidate"
            else:
                status = "未采用（视觉QA未通过）"
                role = "candidate"
            notes = candidate.get("repair_notes") if isinstance(candidate.get("repair_notes"), list) else []
            readable_notes = [str(note).strip() for note in notes if str(note).strip()]
            summary = readable_notes[0][:260] + ("……" if readable_notes and len(readable_notes[0]) > 260 else "") if readable_notes else ""
            _add_figure_review_record(
                records,
                question_id=qid,
                figure_id=figure_id,
                path=path,
                stage=f"{repair_path.stem}/第{round_no}轮/{strategy}",
                status=status,
                summary=summary,
                role=role,
            )

    for image in sorted((stage_dir / "figures").glob("*.png")):
        figure_id = image.stem
        qid = ""
        for record in records.values():
            if str(record.get("figure_id") or "") == figure_id:
                qid = str(record.get("question_id") or "")
                break
        if not qid:
            qid = figure_id.rsplit("_", 1)[0] if "_" in figure_id else figure_id
        _add_figure_review_record(
            records,
            question_id=qid,
            figure_id=figure_id,
            path=image,
            stage="正式图目录",
            status="正式采用",
            role="official",
        )

    official_ids = {
        (str(item.get("question_id") or ""), str(item.get("figure_id") or ""))
        for item in records.values()
        if item.get("role") == "official"
    }
    official_hashes = {
        (
            str(item.get("question_id") or ""),
            str(item.get("figure_id") or ""),
            str(item.get("content_hash") or ""),
        )
        for item in records.values()
        if item.get("role") == "official"
    }
    for item in records.values():
        if item.get("role") != "selected_candidate":
            continue
        identity = (
            str(item.get("question_id") or ""),
            str(item.get("figure_id") or ""),
            str(item.get("content_hash") or ""),
        )
        if identity not in official_hashes:
            item["role"] = "candidate"
            item["statuses"] = [
                "阶段内曾采用，当前已被后续正式图替代"
                if value == "已采用（回修候选，视觉QA通过）"
                else value
                for value in item.get("statuses", [])
            ]
    records = {
        key: item
        for key, item in records.items()
        if item.get("role") != "archive"
        or (str(item.get("question_id") or ""), str(item.get("figure_id") or "")) not in official_ids
    }

    by_qid: dict[str, list[dict[str, Any]]] = {}
    for record in records.values():
        qid = str(record.get("question_id") or "").strip()
        if figure_qids and qid not in figure_qids and not any(qid.startswith(prefix) for prefix in figure_qids):
            continue
        by_qid.setdefault(qid, []).append(record)

    result: list[dict[str, Any]] = []
    order = [_qid(question) for question in structured_exam.get("items", []) if isinstance(question, dict) and _qid(question)]
    for qid in order + sorted(set(by_qid) - set(order)):
        figures = by_qid.get(qid, [])
        if not figures:
            continue
        role_order = {"official": 0, "selected_candidate": 1, "candidate": 2, "archive": 3}
        figures.sort(
            key=lambda item: (
                str(item.get("figure_id") or ""),
                role_order.get(str(item.get("role") or ""), 9),
                str(item.get("path") or ""),
            )
        )
        question = next((item for item in structured_exam.get("items", []) if isinstance(item, dict) and _qid(item) == qid), {})
        result.append(
            {
                "question_id": qid,
                "section": question.get("section", ""),
                "number": question.get("number", ""),
                "stem": question.get("stem", ""),
                "figures": figures,
            }
        )
    return result


def _snapshot_question(stage_dir: Path, output_dir: Path, fragment: dict[str, Any]) -> list[Path]:
    qid = _qid(fragment)
    if not qid:
        return []
    snapshot_dir = output_dir / "question_review_snapshots" / qid
    if snapshot_dir.exists():
        shutil.rmtree(snapshot_dir)
    snapshot_dir.mkdir(parents=True, exist_ok=True)
    snapshot_json = snapshot_dir / "fragment.json"
    snapshot_docx = snapshot_dir / f"{qid}.docx"
    snapshot_pdf = snapshot_dir / f"{qid}.pdf"
    snapshot_json.write_text(
        json.dumps(
            {
                "schema_version": "answer_book.answer_fragments.v4",
                "fragments": [fragment],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    build_docx_from_fragments(snapshot_json, snapshot_docx)
    export_docx_to_pdf(snapshot_docx, snapshot_pdf)
    return render_pdf_to_png(snapshot_pdf, snapshot_dir, prefix="snapshot")


def collect_question_review_items(stage_dir: Path) -> list[dict[str, Any]]:
    structured_exam = _read_json(stage_dir / "structured_exam.json")
    fragments_data = _read_json(stage_dir / "answer_fragments.json")
    fragments_by_qid = {_qid(fragment): fragment for fragment in fragments_data.get("fragments", []) if _qid(fragment)}
    quality_notes = _quality_notes_by_qid(stage_dir)
    review = build_question_review(stage_dir)
    notes_by_qid: dict[str, list[str]] = {}
    for row in review.get("review_rows", []):
        qid = _qid(row)
        if not qid:
            continue
        notes = notes_by_qid.setdefault(qid, [])
        for note in row.get("notes", []):
            _add_note(notes, note)
    for qid, notes in quality_notes.items():
        for note in notes:
            _add_note(notes_by_qid.setdefault(qid, []), note)
    for qid, fragment in fragments_by_qid.items():
        for flag in fragment.get("_review_flags", []):
            if isinstance(flag, dict):
                _add_note(notes_by_qid.setdefault(qid, []), str(flag.get("message") or flag.get("code") or ""))
        for warning in fragment.get("warnings", []):
            text = str(warning)
            if "存疑题目审查文档" in text or "特殊放行" in text or "程序自动绑定" in text:
                _add_note(notes_by_qid.setdefault(qid, []), text)

    items_by_qid = {_qid(item): item for item in structured_exam.get("items", []) if _qid(item)}
    out: list[dict[str, Any]] = []
    for qid in [str(item.get("question_id", "")).strip() for item in structured_exam.get("items", [])]:
        notes = notes_by_qid.get(qid, [])
        if not notes:
            continue
        fragment = fragments_by_qid.get(qid, {})
        question = items_by_qid.get(qid, {})
        out.append(
            {
                "question_id": qid,
                "section": question.get("section", ""),
                "number": question.get("number", ""),
                "stem": question.get("stem", ""),
                "answer": fragment.get("answer", ""),
                "notes": notes,
                "fragment": fragment,
            }
        )
    return out


def _set_run(run, size: float = 10.5, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_para(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _paragraph(doc: Document, text: str, bold: bool = False, color: str = "000000", size: float = 10.5):
    p = doc.add_paragraph()
    _set_para(p)
    r = p.add_run(str(text))
    _set_run(r, size=size, bold=bold, color=color)
    return p


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    _set_para(p, before=14 if level > 1 else 18, after=5)
    r = p.add_run(str(text))
    _set_run(r, size=16 if level == 1 else 13, bold=True, color="0B2545" if level == 1 else "1F4D78")
    return p


def _stage_label(stage: Any) -> str:
    labels = {
        "knowledge_planning": "考查内容判断",
        "evidence_selection": "教材依据确认",
        "answer_generation": "解析生成",
        "answer_coverage": "答案覆盖审查回修",
        "content_quality": "内容质量审查回修",
        "docx": "Word 生成审查回修",
    }
    text = str(stage or "").strip()
    return labels.get(text, text or "未知阶段")


def _add_model_retry_section(doc: Document, stage_dir: Path) -> None:
    summary = model_retry_summary(stage_dir)
    _heading(doc, "模型重试策略", 1)
    if not summary.get("applied"):
        _paragraph(doc, "本任务未触发模型重试策略。")
        return

    _paragraph(
        doc,
        f"本任务共有 {summary.get('retry_question_count', 0)} 道题触发模型重试，记录到 {summary.get('retry_event_count', 0)} 次重试事件。"
    )
    for index, row in enumerate(summary.get("items", [])[:80], start=1):
        question_id = str(row.get("question_id") or "未记录").strip()
        stage = _stage_label(row.get("stage"))
        attempt_count = row.get("attempt_count") or 0
        final_strategy = str(row.get("final_strategy") or "未记录").strip()
        final_model = str(row.get("final_model") or "未记录").strip()
        final_tokens = row.get("final_max_tokens") or "未记录"
        final_thinking = row.get("final_thinking") or "默认"
        _heading(doc, f"{index}. {question_id}", 2)
        _paragraph(doc, f"阶段：{stage}")
        _paragraph(doc, f"重试次数：{attempt_count}；最终策略：{final_strategy}；最终模型：{final_model}")
        _paragraph(doc, f"最终 max_tokens：{final_tokens}；Thinking：{final_thinking}")
        strategies = [str(item) for item in row.get("strategies", []) if str(item).strip()]
        if strategies:
            _paragraph(doc, "策略链：" + " -> ".join(strategies))
        errors = [str(item) for item in row.get("errors", []) if str(item).strip()]
        if errors:
            _paragraph(doc, "失败信息：" + "；".join(errors[:3]), color="8A1F11")


def _add_figure_review_section(doc: Document, stage_dir: Path) -> None:
    items = collect_question_figure_review_items(stage_dir)
    _heading(doc, "作图题全流程图片", 1)
    if not items:
        _paragraph(doc, "本任务未发现作图题图片产物。")
        return
    _paragraph(doc, "本节按图片内容去重：每张正式图只展示一次，仅保留当前有效回修报告中与正式图有实质差异的候选图，并明确标注是否采用。")
    for index, item in enumerate(items, start=1):
        title_parts = [f"{index}."]
        section = str(item.get("section") or "").strip()
        number = str(item.get("number") or "").strip()
        if section or number:
            title_parts.append(f"{section} 第 {number} 题")
        title_parts.append(f"（{item.get('question_id', '')}）")
        _heading(doc, " ".join(part for part in title_parts if part.strip()), 2)
        stem = str(item.get("stem") or "").strip()
        if stem:
            _paragraph(doc, f"题干：{stem[:500]}")
        for figure_index, figure in enumerate(item.get("figures", [])[:30], start=1):
            figure_id = str(figure.get("figure_id") or "").strip()
            path = figure.get("path")
            role = str(figure.get("role") or "archive")
            role_label = {
                "official": "正式采用",
                "selected_candidate": "已采用的回修候选",
                "candidate": "未采用的回修候选",
                "archive": "阶段归档兜底",
            }.get(role, "阶段图")
            _paragraph(doc, f"图 {figure_index}：{figure_id or Path(str(path)).stem}（{role_label}）", bold=True)
            stages = "；".join(str(value) for value in figure.get("stages", []) if str(value).strip())
            statuses = "；".join(str(value) for value in figure.get("statuses", []) if str(value).strip())
            summaries = "；".join(str(value) for value in figure.get("summaries", [])[:3] if str(value).strip())
            if stages:
                _paragraph(doc, f"阶段：{stages}")
            if statuses:
                _paragraph(doc, f"状态：{statuses}")
            if summaries:
                _paragraph(doc, f"说明：{summaries}")
            image = Path(path) if path else None
            if image and image.exists():
                p = doc.add_paragraph()
                _set_para(p)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(str(image), width=Cm(13.5))
                except Exception as exc:
                    _paragraph(doc, f"图片插入失败：{image}；{exc}", color="8A1F11")
            else:
                _paragraph(doc, f"图片文件不存在：{path}", color="8A1F11")


def build_figure_review_docx(stage_dir: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "作图题全流程图片.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title = doc.add_paragraph()
    _set_para(title, after=3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("作图题全流程图片")
    _set_run(run, size=22, bold=True)
    _paragraph(doc, "本文件单独汇总作图题在生成、正式引用、视觉审核和回修候选等阶段产生的图片。")
    _add_figure_review_section(doc, stage_dir)
    doc.save(docx_path)
    return docx_path


def build_question_review_docx(stage_dir: Path, output_dir: Path, render_snapshots: bool = True) -> Path:
    items = collect_question_review_items(stage_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "question_review.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title = doc.add_paragraph()
    _set_para(title, after=3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("存疑题目审查文档")
    _set_run(run, size=22, bold=True)
    _paragraph(doc, "本文件汇总平台自动生产过程中被特殊处理、产生警告或需要交付前关注的题目。")
    if not items:
        _heading(doc, "审查结果", 1)
        _paragraph(doc, "未发现需要单独审查的题目。")
        _add_model_retry_section(doc, stage_dir)
        doc.save(docx_path)
        return docx_path

    for index, item in enumerate(items, start=1):
        _heading(doc, f"{index}. {item.get('section', '')} 第 {item.get('number', '')} 题（{item.get('question_id', '')}）", 1)
        _heading(doc, "题目内容", 2)
        _paragraph(doc, item.get("stem", ""))
        _heading(doc, "存疑问题", 2)
        for note in item.get("notes", []):
            _paragraph(doc, f"- {note}")
        _heading(doc, "最终解析截图", 2)
        fragment = item.get("fragment") or {}
        rendered = []
        if render_snapshots and fragment:
            try:
                rendered = _snapshot_question(stage_dir, output_dir, fragment)
            except Exception as exc:
                _paragraph(doc, f"截图生成失败：{exc}")
        if rendered:
            for image in rendered[:3]:
                p = doc.add_paragraph()
                _set_para(p)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(image), width=Cm(10.0))
        else:
            _paragraph(doc, "未生成截图；以下为最终结构化解析文本。")
            for block in fragment.get("blocks", []):
                _paragraph(doc, f"{block.get('label', '')}：{_fragment_text(block)}")
        _heading(doc, "答案", 2)
        _paragraph(doc, item.get("answer", ""))
    _add_model_retry_section(doc, stage_dir)
    doc.save(docx_path)
    return docx_path
