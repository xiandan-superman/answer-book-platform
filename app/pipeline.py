from __future__ import annotations

import json
import shutil
import threading
import time
import traceback
from contextlib import contextmanager
from dataclasses import asdict, dataclass, replace
from pathlib import Path

from docx import Document

from .answer_coverage_audit import audit_answer_coverage
from .answer_generation import attach_program_evidence_block, bind_top_evidence, fallback_fragment, generate_answer_fragments, has_bound_evidence, write_demo_fragments
from .audit_model_repair import fill_missing_fragments_locally, repair_fragments_with_model_for_audit
from .audit_review_gate import auto_allow_audit_report, wait_for_user_review_decision
from .content_quality_audit import audit_content_quality
from .task_control import TaskCancelled, checkpoint, clear_task_control
from .docx_audit import audit_docx_v4
from .docx_model_repair import repair_fragments_with_model_for_docx
from .docx_v4 import build_docx_from_fragments
from .environment import check_environment
from .exam_audit import audit_exam_structure
from .exam_structure_review import wait_for_exam_structure_review
from .exam_extract import extract_exam_structure
from .evidence_selection import confirm_evidence_selection
from .evidence_audit import audit_retrieval_candidates
from .final_acceptance import build_final_acceptance_report
from .figures import audit_figures_with_vision, prepare_figures_for_fragments, repair_figures_with_model_for_visual_qa
from .figure_size_audit import audit_docx_figure_sizes
from .fragment_repair import repair_answer_fragments_for_docx
from .knowledge_planning import generate_knowledge_plans, load_knowledge_plans
from .model_usage_report import build_model_usage_report
from .paths import OUTPUTS_DIR, ensure_project_dirs
from .figure_schema_planning import attach_figure_schema_plans, plan_figure_schemas
from .question_review_docx import build_figure_review_docx, build_question_review_docx, collect_question_figure_review_items, collect_question_review_items
from .question_understanding import build_question_understandings
from .render_word import export_docx_to_pdf, render_pdf_to_png
from .render_audit import audit_rendered_pages
from .retrieval import build_candidates, candidates_for_question
from .review_notes import build_answer_review_notes
from .settings import get_provider, provider_supports_image_generation
from .task_store import append_event, load_task, task_dir, update_task
from .textbook_index_cache import install_textbook_index_cache
from .v4_schema import validate_v4_answer_fragment


CONTENT_QUALITY_MODEL_REPAIR_CODES = {
    "missing_required_figure",
    "calculation_missing_mistake_notes",
    "formula_absence_after_retry",
    "calculation_missing_steps",
    "calculation_missing_subquestion_steps",
    "calculation_invalid_subquestion_number",
    "missing_answer_unit_content",
    "missing_answer_unit_steps",
    "calculation_answer_missing_unit",
    "calculation_missing_substitution",
    "missing_analysis",
}
DOCX_MODEL_REPAIR_CODES = {"formula_like_normal_text"}


@dataclass
class PipelineOptions:
    use_model: bool = True
    allow_demo_without_key: bool = False
    render_with_word: bool = False
    reuse_fragments: bool = False
    require_preferred_formula_chain: bool = True


def stage_dir(task_id: str) -> Path:
    return task_dir(task_id) / "stage_outputs"


def output_dir(task_id: str) -> Path:
    return OUTPUTS_DIR / task_id


def write_json(path: Path, value) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def _provider_key_issue(label: str, provider) -> str:
    env_name = str(getattr(provider, "api_key_env", "") or "").strip()
    hint = f"（环境变量 {env_name}）" if env_name else "（providers.local.json 或 .env）"
    return f"{label} {getattr(provider, 'name', 'unknown')} 未配置 API Key {hint}"


def _validate_required_provider_keys(
    *,
    use_model: bool,
    allow_demo_without_key: bool,
    provider,
    reasoning_provider,
    answer_provider,
    vision_provider,
    image_provider,
    image_model: str,
) -> list[str]:
    if not use_model or allow_demo_without_key:
        return []
    checks = [
        ("基础/作图规则模型", provider),
        ("知识点与教材依据模型", reasoning_provider),
        ("答案生成模型", answer_provider),
        ("读图模型", vision_provider),
    ]
    if str(image_model or "").strip():
        checks.append(("作图生图模型", image_provider))
    issues: list[str] = []
    seen: set[tuple[str, str]] = set()
    for label, cfg in checks:
        name = str(getattr(cfg, "name", "") or "")
        key = (label, name)
        if key in seen:
            continue
        seen.add(key)
        if not str(getattr(cfg, "api_key", "") or "").strip():
            issues.append(_provider_key_issue(label, cfg))
    return issues


class FigureProgressTracker:
    """Persist figure sub-stage events and a heartbeat while remote calls are pending."""

    def __init__(self, output_path: Path, heartbeat_seconds: int = 15):
        self.output_path = output_path
        self.heartbeat_seconds = heartbeat_seconds
        self.started_at = time.monotonic()
        self._lock = threading.Lock()
        self._state: dict = {"stage": "figures", "status": "running", "recent_events": []}

    def emit(self, event: str, detail: dict | None = None) -> None:
        detail = dict(detail or {})
        now = time.strftime("%H:%M:%S")
        with self._lock:
            record = {"time": now, "event": event, **detail}
            events = list(self._state.get("recent_events") or [])
            events.append(record)
            self._state.update(detail)
            self._state.update(
                {
                    "stage": "figures",
                    "status": "running",
                    "active_event": event,
                    "updated_at": now,
                    "elapsed_seconds": max(0, int(time.monotonic() - self.started_at)),
                    "recent_events": events[-20:],
                }
            )
            write_json(self.output_path, self._state)

    @contextmanager
    def operation(self, name: str, **detail):
        self.emit("operation_started", {"operation": name, **detail})
        stop = threading.Event()

        def heartbeat() -> None:
            while not stop.wait(self.heartbeat_seconds):
                self.emit("heartbeat", {"operation": name, **detail})

        thread = threading.Thread(target=heartbeat, daemon=True)
        thread.start()
        try:
            yield
        except Exception as exc:
            self.emit("operation_failed", {"operation": name, **detail, "error": str(exc)[:300]})
            raise
        else:
            self.emit("operation_completed", {"operation": name, **detail})
        finally:
            stop.set()
            thread.join(timeout=0.2)


def attach_figure_generation_audit(content_quality: dict, sdir: Path) -> dict:
    audit_path = sdir / "figure_generation_audit.json"
    if audit_path.exists():
        try:
            content_quality["figure_generation_audit"] = json.loads(audit_path.read_text(encoding="utf-8"))
        except Exception as exc:
            content_quality["figure_generation_audit"] = {"error": str(exc)[:300]}
    visual_qa_path = sdir / "figure_visual_qa.json"
    if visual_qa_path.exists():
        try:
            content_quality["figure_visual_qa"] = json.loads(visual_qa_path.read_text(encoding="utf-8"))
        except Exception as exc:
            content_quality["figure_visual_qa"] = {"error": str(exc)[:300]}
    if audit_path.exists() or visual_qa_path.exists():
        write_json(sdir / "content_quality_audit.json", content_quality)
    return content_quality


def figure_visual_qa_issue_count(report: dict) -> int:
    if not isinstance(report, dict) or not report.get("enabled"):
        return 0
    count = len(report.get("failed", []) if isinstance(report.get("failed"), list) else [])
    item_figure_ids: set[str] = set()
    for item in report.get("items", []):
        if not isinstance(item, dict):
            continue
        figure_id = str(item.get("figure_id") or "").strip()
        if figure_id:
            item_figure_ids.add(figure_id)
        qa = item.get("qa") if isinstance(item.get("qa"), dict) else {}
        if qa.get("ok") is not True:
            count += 1
    for item in report.get("skipped", []):
        figure_id = str(item.get("figure_id") or "").strip() if isinstance(item, dict) else ""
        if isinstance(item, dict) and str(item.get("reason") or "") == "figure image missing" and figure_id not in item_figure_ids:
            count += 1
    return count


def _docx_issue_code(issue: str) -> str:
    text = str(issue or "")
    if (
        "Formula-like text must not be written" in text
        or "Formula-like text leaked" in text
        or "formula-like normal text" in text
    ):
        return "formula_like_normal_text"
    if "OMML formula count" in text:
        return "omml_formula_count_below_expected"
    if "raw radical" in text or "√" in text:
        return "raw_radical_normal_text"
    if "raw subscript" in text:
        return "raw_subscript_normal_text"
    if "raw latex" in text or "leftharpoons" in text or "\\" in text:
        return "raw_latex_marker"
    return "docx_audit_issue"


def _filter_audit_report_for_model_repair(report: dict, allowed_codes: set[str]) -> dict:
    filtered = dict(report)
    for key in ("issues", "warnings"):
        items = []
        for item in report.get(key, []) if isinstance(report, dict) else []:
            code = str(item.get("code") or "").strip() if isinstance(item, dict) else ""
            if code in allowed_codes:
                items.append(item)
        filtered[key] = items
    filtered["issue_count"] = len(filtered.get("issues", []))
    filtered["warning_count"] = len(filtered.get("warnings", []))
    filtered["ok"] = not filtered["issue_count"]
    return filtered


def _filter_docx_issues_for_model_repair(issues: list[str]) -> list[str]:
    return [issue for issue in issues if _docx_issue_code(issue) in DOCX_MODEL_REPAIR_CODES]


def build_user_allowed_docx_placeholder(fragments_json: Path, docx_path: Path, stage_dir: Path, reason: str) -> dict:
    data = json.loads(fragments_json.read_text(encoding="utf-8")) if fragments_json.exists() else {"fragments": []}
    fragments = list(data.get("fragments", [])) if isinstance(data.get("fragments"), list) else []
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("真题答案解析（待复核版）", level=0)
    doc.add_paragraph("本任务存在用户已允许继续的 DOCX 审计问题。系统已生成可打开的待复核总文档，正式排查请结合 question_review.docx 与审查记录。")
    doc.add_paragraph(f"降级原因：{reason}")
    for fragment in fragments:
        if not isinstance(fragment, dict):
            continue
        section = str(fragment.get("section") or "").strip()
        number = str(fragment.get("number") or "").strip()
        title = " ".join(x for x in [section, f"{number}、" if number else ""] if x).strip() or str(fragment.get("question_id") or "题目")
        doc.add_heading(title, level=1)
        answer = str(fragment.get("answer") or "待复核").strip() or "待复核"
        doc.add_paragraph(f"答案：{answer}")
        warnings = [str(item).strip() for item in fragment.get("warnings", []) if str(item).strip()]
        flags = [item for item in fragment.get("_review_flags", []) if isinstance(item, dict)]
        if warnings or flags:
            doc.add_paragraph("复核提示：")
            for warning in warnings[:5]:
                doc.add_paragraph(warning, style="List Bullet")
            for flag in flags[:5]:
                message = str(flag.get("message") or flag.get("code") or "").strip()
                if message:
                    doc.add_paragraph(message, style="List Bullet")
        else:
            doc.add_paragraph("复核提示：该题随本次待复核总文档一并保留。")

    doc.save(docx_path)
    report = {
        "ok": docx_path.exists(),
        "docx": str(docx_path),
        "fragment_count": len(fragments),
        "reason": reason,
        "mode": "user_allowed_placeholder",
    }
    write_json(stage_dir / "docx_user_allowed_placeholder.json", report)
    return report


def _build_user_allowed_candidate_payload(fragments_json: Path, stage_dir: Path) -> tuple[Path, int]:
    data = json.loads(fragments_json.read_text(encoding="utf-8")) if fragments_json.exists() else {"fragments": []}
    fragments = []
    candidate_count = 0
    for fragment in data.get("fragments", []) if isinstance(data.get("fragments"), list) else []:
        candidate = fragment.get("_review_candidate_fragment") if isinstance(fragment, dict) else None
        if isinstance(candidate, dict):
            replacement = dict(candidate)
            warnings = [str(item) for item in replacement.get("warnings", []) if str(item).strip()]
            notice = "用户确认允许使用待复核前的模型解析候选版本。"
            if notice not in warnings:
                warnings.append(notice)
            replacement["warnings"] = warnings
            meta = dict(replacement.get("_meta") or {})
            meta["review_candidate_used_by_user"] = True
            replacement["_meta"] = meta
            fragments.append(replacement)
            candidate_count += 1
        else:
            fragments.append(fragment)
    if candidate_count == 0:
        return fragments_json, 0
    candidate_json = stage_dir / "answer_fragments.user_allowed_candidate.json"
    write_json(candidate_json, {**data, "fragments": fragments, "review_candidate_count": candidate_count})
    return candidate_json, candidate_count


def build_user_allowed_docx_candidate(fragments_json: Path, docx_path: Path, stage_dir: Path, reason: str) -> dict:
    report = {
        "ok": False,
        "docx": str(docx_path),
        "reason": reason,
        "mode": "user_allowed_candidate",
    }
    try:
        source_json, candidate_count = _build_user_allowed_candidate_payload(fragments_json, stage_dir)
        report["candidate_count"] = candidate_count
        report["source_json"] = str(source_json)
        report["source_mode"] = "review_candidate_fragment" if candidate_count else "current_fragments"
        build_docx_from_fragments(source_json, docx_path, strict_answer_summary_formula_audit=False)
        report["ok"] = docx_path.exists()
    except Exception as exc:
        report["error"] = str(exc)
        report["traceback"] = traceback.format_exc()
    write_json(stage_dir / "docx_user_allowed_candidate.json", report)
    return report


def build_and_audit_docx_with_repair(
    task_id: str,
    fragments_json: Path,
    docx_path: Path,
    sdir: Path,
    mark,
    *,
    structured_exam: dict | None = None,
    candidates: list | None = None,
    selection_data: dict | None = None,
    provider=None,
    model: str = "",
    use_model: bool = False,
) -> dict:
    attempts: list[dict] = []

    def attempt(label: str) -> list[str]:
        try:
            build_docx_from_fragments(fragments_json, docx_path)
            fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
            expected_formula_count = sum(len(x.get("formulas", [])) for x in fragments_data.get("fragments", []))
            issues = audit_docx_v4(docx_path, min_formulas=expected_formula_count)
            attempts.append({"attempt": label, "ok": not issues, "issues": issues[:30]})
            return issues
        except Exception as exc:
            attempts.append({"attempt": label, "ok": False, "error": str(exc), "traceback": traceback.format_exc()})
            return [str(exc)]

    issues = attempt("initial")
    if not issues:
        write_json(sdir / "docx_repair.json", {"needed": False, "attempts": attempts})
        return {"ok": True, "issues": [], "repair": {"needed": False, "attempts": attempts}}

    repair_payload = {"needed": True, "attempts": attempts}
    model_repair_issues = _filter_docx_issues_for_model_repair(issues)
    if model_repair_issues and use_model and provider is not None and getattr(provider, "api_key", "") and structured_exam is not None:
        mark(
            "docx_model_repair",
            "started",
            {
                "issues": model_repair_issues[:10],
                "skipped_issues": [issue for issue in issues if issue not in model_repair_issues][:10],
                "reason": "DOCX 审计失败，先把具体题目和失败原因交给模型单题回修。",
            },
        )
        model_repair_report = repair_fragments_with_model_for_docx(
            fragments_json,
            structured_exam,
            candidates or [],
            selection_data=selection_data or {},
            provider=provider,
            model=model,
            docx_issues=model_repair_issues,
            backup_path=sdir / "answer_fragments.before_docx_model_repair.json",
        )
        repair_payload["model_repair"] = model_repair_report
        mark("docx_model_repair", "applied" if model_repair_report.get("changed") else "skipped", model_repair_report)
        if model_repair_report.get("changed"):
            issues = attempt("after_model_repair")
            repair_payload["attempts"] = attempts
            repair_payload["ok"] = not issues
            if not issues:
                write_json(sdir / "docx_repair.json", repair_payload)
                return {"ok": True, "issues": [], "repair": repair_payload}
    elif issues:
        mark(
            "docx_model_repair",
            "skipped",
            {
                "reason": "DOCX 审计问题不在模型回修白名单内，跳过模型回修，进入程序自修。",
                "issues": issues[:30],
            },
        )
        repair_payload["model_repair"] = {
            "ok": False,
            "changed": False,
            "repaired_count": 0,
            "repaired_question_ids": [],
            "issues": ["DOCX 审计问题不在模型回修白名单内，跳过模型回修，进入程序自修。"],
            "skipped_issues": issues[:30],
        }

    mark("docx_repair", "started", {"initial_error": attempts[-1].get("error"), "initial_issues": attempts[-1].get("issues", [])[:10]})
    repair_report = repair_answer_fragments_for_docx(fragments_json, sdir / "answer_fragments.before_docx_repair.json")
    repair_payload["repair"] = repair_report
    if repair_report.get("ok") and repair_report.get("changed"):
        mark("docx_repair", "applied", repair_report)
        issues = attempt("after_repair")
    else:
        mark("docx_repair", "skipped", repair_report)
    repair_payload["attempts"] = attempts
    repair_payload["ok"] = not issues
    write_json(sdir / "docx_repair.json", repair_payload)
    return {"ok": not issues, "issues": issues, "repair": repair_payload}


def run_pipeline(task_id: str, options: PipelineOptions | None = None) -> dict:
    options = options or PipelineOptions()
    ensure_project_dirs()
    record = load_task(task_id)
    sdir = stage_dir(task_id)
    odir = output_dir(task_id)
    sdir.mkdir(parents=True, exist_ok=True)
    odir.mkdir(parents=True, exist_ok=True)
    summary: dict = {"task_id": task_id, "stages": []}

    def mark(stage: str, status: str, detail=None):
        row = {"stage": stage, "status": status, "detail": detail or {}}
        summary["stages"].append(row)
        append_event(task_id, f"{stage}_{status}", row)
        write_json(sdir / "pipeline_status.json", summary)

    try:
        clear_task_control(task_id)
        checkpoint(task_id)
        update_task(task_id, status="running", current_stage="environment", error="")
        env = check_environment()
        write_json(sdir / "environment_check.json", env)
        if options.require_preferred_formula_chain and not env.get("formula_conversion", {}).get("preferred_chain_ready"):
            mark("environment", "failed", {"formula_conversion": env.get("formula_conversion", {})})
            raise RuntimeError("Preferred formula conversion chain is not ready")
        mark(
            "environment",
            "passed",
            {
                "word_mac": env["microsoft_word"]["mac"],
                "word_windows": env["microsoft_word"]["windows"],
                "formula_conversion": env.get("formula_conversion", {}),
            },
        )

        exam_path = Path(record.exam_path).expanduser()
        textbooks_dir = Path(record.textbooks_dir).expanduser()
        if not exam_path.exists():
            raise FileNotFoundError(f"Exam file not found: {exam_path}")
        if not textbooks_dir.exists():
            raise FileNotFoundError(f"Textbooks dir not found: {textbooks_dir}")
        thinking_mode = getattr(record, "model_thinking", "auto") or "auto"
        provider = replace(get_provider(record.provider), thinking_mode=thinking_mode)
        reasoning_provider_name = getattr(record, "reasoning_provider", "") or record.provider
        reasoning_provider = replace(get_provider(reasoning_provider_name), thinking_mode=thinking_mode)
        reasoning_model = str(getattr(record, "reasoning_model", "") or (record.model if reasoning_provider_name == record.provider else reasoning_provider.default_model) or record.model).strip()
        answer_provider_name = getattr(record, "answer_provider", "") or record.provider
        answer_provider = replace(get_provider(answer_provider_name), thinking_mode=thinking_mode)
        answer_model = str(getattr(record, "answer_model", "") or (record.model if answer_provider_name == record.provider else answer_provider.default_model) or record.model).strip()
        vision_provider = get_provider(getattr(record, "vision_provider", "") or record.provider)
        vision_model = str(getattr(record, "vision_model", "") or getattr(vision_provider, "vision_model", "") or record.model).strip()
        image_provider = get_provider(getattr(record, "image_provider", "") or record.provider)
        image_model = str(getattr(record, "image_model", "") or getattr(image_provider, "image_model", "") or "").strip()
        if not provider_supports_image_generation(image_provider):
            image_model = ""
        if image_model:
            image_provider = replace(image_provider, image_model=image_model)
        key_issues = _validate_required_provider_keys(
            use_model=options.use_model,
            allow_demo_without_key=options.allow_demo_without_key,
            provider=provider,
            reasoning_provider=reasoning_provider,
            answer_provider=answer_provider,
            vision_provider=vision_provider,
            image_provider=image_provider,
            image_model=image_model,
        )
        if key_issues:
            mark("provider_config", "failed", {"issues": key_issues})
            raise RuntimeError("；".join(key_issues))

        checkpoint(task_id)
        update_task(task_id, current_stage="extract_exam")
        structured_exam = extract_exam_structure(exam_path, sdir / "structured_exam.json")
        exam_issues = audit_exam_structure(structured_exam, sdir / "exam_structure_audit.json")
        if exam_issues:
            mark("extract_exam", "failed", {"issues": exam_issues[:30]})
            raise RuntimeError("Exam structure audit failed")
        mark("extract_exam", "passed", {"question_count": len(structured_exam.get("items", []))})

        checkpoint(task_id)
        update_task(task_id, current_stage="exam_structure_review")
        mark("exam_structure_review", "started", {"question_count": len(structured_exam.get("items", []))})
        structured_exam = wait_for_exam_structure_review(task_id, structured_exam, sdir, sdir / "structured_exam.json")
        mark(
            "exam_structure_review",
            "passed",
            {
                "question_count": len(structured_exam.get("items", [])),
                "reviewed": True,
            },
        )

        checkpoint(task_id)
        update_task(task_id, current_stage="question_understanding")
        mark("question_understanding", "started", {"question_count": len(structured_exam.get("items", []))})
        understanding_report = build_question_understandings(
            structured_exam,
            sdir / "question_understanding.json",
            provider=vision_provider if options.use_model and vision_provider.api_key else None,
            model=vision_model,
            progress_json=sdir / "question_understanding_progress.json",
        )
        attached_understanding_count = sum(
            1
            for item in structured_exam.get("items", [])
            if isinstance(item, dict) and isinstance(item.get("question_understanding"), dict)
        )
        write_json(sdir / "structured_exam.json", structured_exam)
        mark(
            "question_understanding",
            "passed",
            {
                "question_count": understanding_report.get("question_count", 0),
                "vision_required_count": understanding_report.get("vision_required_count", 0),
                "vision_used_count": understanding_report.get("vision_used_count", 0),
                "attached_understanding_count": attached_understanding_count,
                "downstream_reuse": [
                    "knowledge_planning",
                    "evidence_selection",
                    "answer_generation",
                ],
            },
        )

        checkpoint(task_id)
        update_task(task_id, current_stage="figure_schema_planning")
        figure_schema_plan = plan_figure_schemas(
            structured_exam,
            sdir / "figure_schema_plan.json",
            provider=provider if options.use_model and provider.api_key else None,
            model=record.model,
        )
        structured_exam = attach_figure_schema_plans(structured_exam, figure_schema_plan)
        write_json(sdir / "structured_exam.json", structured_exam)
        mark(
            "figure_schema_planning",
            "passed",
            {
                "planned_count": figure_schema_plan.get("planned_count", 0),
                "schema_found_count": sum(1 for item in figure_schema_plan.get("items", []) if (item.get("schema_resolution") or {}).get("status") == "schema_found"),
                "schema_proposed_count": sum(1 for item in figure_schema_plan.get("items", []) if (item.get("schema_resolution") or {}).get("status") == "schema_proposed"),
            },
        )

        checkpoint(task_id)
        update_task(task_id, current_stage="textbook_index")
        if not record.selected_textbooks:
            raise RuntimeError("当前任务没有绑定已索引教材。请先在教材管理页选择教材并建立索引，再创建解析任务。")
        index_detail = install_textbook_index_cache(
            record.selected_textbooks,
            sdir,
            record.textbook_display_names or {},
        )
        mark("textbook_index", "passed" if index_detail.get("page_map_ok", True) else "failed", index_detail)
        if not index_detail.get("page_map_ok", True):
            issues = index_detail.get("page_map_issues") or []
            messages = [str(issue.get("message", issue)) for issue in issues if isinstance(issue, dict)]
            raise RuntimeError("教材页码读取失败：" + "；".join(messages[:5]))

        checkpoint(task_id)
        update_task(task_id, current_stage="knowledge_planning")
        plans_json = sdir / "knowledge_plans.json"
        if options.use_model and reasoning_provider.api_key:
            plan_result = generate_knowledge_plans(
                structured_exam,
                reasoning_provider,
                reasoning_model,
                plans_json,
                use_model=True,
                progress_json=sdir / "knowledge_planning_progress.json",
            )
        elif options.allow_demo_without_key:
            plan_result = generate_knowledge_plans(
                structured_exam,
                reasoning_provider,
                reasoning_model,
                plans_json,
                use_model=False,
                progress_json=sdir / "knowledge_planning_progress.json",
            )
        else:
            raise RuntimeError(f"API key not configured for reasoning provider: {reasoning_provider.name}")
        plan_detail = asdict(plan_result)
        mark("knowledge_planning", "passed" if plan_result.ok else "failed", plan_detail)
        if not plan_result.ok:
            raise RuntimeError("Knowledge planning failed")
        knowledge_plans = load_knowledge_plans(plans_json)

        checkpoint(task_id)
        update_task(task_id, current_stage="retrieval")
        candidates = build_candidates(
            structured_exam,
            sdir / "textbook_blocks.csv",
            sdir / "textbook_page_map.csv",
            sdir / "retrieval_candidates.csv",
            knowledge_plans=knowledge_plans,
        )
        retrieval_issues = audit_retrieval_candidates(structured_exam, sdir / "retrieval_candidates.csv", sdir / "retrieval_audit.json")
        if retrieval_issues:
            mark("retrieval", "failed", {"issues": retrieval_issues[:30]})
            raise RuntimeError("Retrieval audit failed")
        mark("retrieval", "passed", {"candidate_count": len(candidates)})

        checkpoint(task_id)
        update_task(task_id, current_stage="evidence_selection")
        selection_result, confirmed_candidates = confirm_evidence_selection(
            structured_exam,
            knowledge_plans,
            candidates,
            reasoning_provider,
            reasoning_model,
            sdir / "evidence_selection.json",
            sdir / "textbook_blocks.csv",
            sdir / "textbook_page_map.csv",
            progress_json=sdir / "evidence_selection_progress.json",
            use_model=options.use_model,
        )
        selection_detail = asdict(selection_result)
        mark("evidence_selection", "passed" if selection_result.ok else "failed", selection_detail)
        if not selection_result.ok:
            raise RuntimeError("Evidence selection failed")
        selection_data = json.loads((sdir / "evidence_selection.json").read_text(encoding="utf-8"))
        evidence_selections = {
            str(selection.get("question_id", "")).strip(): selection
            for selection in selection_data.get("selections", [])
            if str(selection.get("question_id", "")).strip()
        }
        candidates = confirmed_candidates

        checkpoint(task_id)
        update_task(task_id, current_stage="answer_generation")
        fragments_json = sdir / "answer_fragments.json"
        if options.reuse_fragments:
            if not fragments_json.exists():
                raise RuntimeError("Cannot reuse fragments: answer_fragments.json not found")
            fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
            fragments = list(fragments_data.get("fragments", []))
            existing_ids = {str(fragment.get("question_id", "")).strip() for fragment in fragments}
            questions_by_id = {str(question.get("question_id", "")).strip(): question for question in structured_exam.get("items", [])}
            recovered_missing = []
            recovered_evidence = []
            migrated_evidence_reason = []
            for question in structured_exam.get("items", []):
                qid = str(question.get("question_id", "")).strip()
                if not qid or qid in existing_ids:
                    continue
                evidence = candidates_for_question(candidates, qid)
                fragment = fallback_fragment(
                    question,
                    evidence,
                    "previous model run did not produce valid structured JSON; auto-filled for review",
                )
                fragments.append(attach_program_evidence_block(fragment, evidence))
                existing_ids.add(qid)
                recovered_missing.append(qid)
            for fragment in fragments:
                qid = str(fragment.get("question_id", "")).strip()
                if not qid or qid not in questions_by_id or has_bound_evidence(fragment):
                    continue
                evidence = candidates_for_question(candidates, qid)
                if not evidence:
                    continue
                bind_top_evidence(
                    fragment,
                    evidence,
                    reason="历史结构化结果未提供不绑定候选证据的原因，程序按检索排序补充最相关教材证据。",
                )
                recovered_evidence.append(qid)
            for fragment in fragments:
                qid = str(fragment.get("question_id", "")).strip()
                meta = dict(fragment.get("_meta") or {})
                if not qid or not has_bound_evidence(fragment) or meta.get("evidence_binding"):
                    continue
                warnings = [str(x) for x in fragment.get("warnings", [])]
                if not any("程序自动绑定" in warning for warning in warnings):
                    continue
                evidence_ids = [str(x) for x in fragment.get("evidence_ids", []) if str(x).strip()]
                reason = "历史结果已由程序补充候选证据；原模型未提供不绑定候选证据的原因。"
                meta["evidence_binding"] = {
                    "strategy": "program_top_evidence",
                    "reason": reason,
                    "bound_evidence_ids": evidence_ids,
                }
                fragment["_meta"] = meta
                migrated_evidence_reason.append(qid)
            if recovered_missing:
                fragments_data["fragments"] = fragments
                fragments_data.setdefault("issues", []).append(
                    {
                        "question_id": ",".join(recovered_missing),
                        "issues": ["missing fragments auto-filled for review"],
                    }
                )
                fragments_data["recovered_count"] = int(fragments_data.get("recovered_count", 0)) + len(recovered_missing)
                fragments_data["fallback_count"] = int(fragments_data.get("fallback_count", 0)) + len(recovered_missing)
            if recovered_evidence:
                fragments_data["fragments"] = fragments
                fragments_data.setdefault("recovery_events", []).extend(
                    {"question_id": qid, "strategy": "program_evidence_binding", "reason": "历史结构化结果未提供不绑定候选证据的原因"} for qid in recovered_evidence
                )
                fragments_data["recovered_count"] = int(fragments_data.get("recovered_count", 0)) + len(recovered_evidence)
            if migrated_evidence_reason:
                fragments_data["fragments"] = fragments
                fragments_data.setdefault("recovery_events", []).extend(
                    {"question_id": qid, "strategy": "program_evidence_binding_metadata_migration"} for qid in migrated_evidence_reason
                )
            if recovered_missing or recovered_evidence or migrated_evidence_reason:
                write_json(fragments_json, fragments_data)
            fragment_issues = []
            for idx, fragment in enumerate(fragments_data.get("fragments", []), start=1):
                for issue in validate_v4_answer_fragment(fragment):
                    fragment_issues.append(f"fragment {idx}: {issue}")
            if fragment_issues:
                mark("answer_generation", "failed", {"reused": True, "issues": fragment_issues[:30]})
                raise RuntimeError("Reused answer fragments failed v4 validation")
            generation_detail = {
                "ok": True,
                "reused": True,
                "fragment_count": len(fragments_data.get("fragments", [])),
                "recovered_count": len(recovered_missing) + len(recovered_evidence),
                "fallback_count": len(recovered_missing),
                "evidence_bound_count": len(recovered_evidence),
                "evidence_binding_metadata_migrated_count": len(migrated_evidence_reason),
                "output": str(fragments_json),
                "issues": [],
            }
        elif options.use_model and answer_provider.api_key:
            generation = generate_answer_fragments(
                structured_exam,
                candidates,
                answer_provider,
                answer_model,
                fragments_json,
                progress_json=sdir / "answer_generation_progress.json",
                evidence_selections=evidence_selections,
            )
            generation_detail = asdict(generation)
        elif options.allow_demo_without_key:
            generation = write_demo_fragments(structured_exam, candidates, fragments_json)
            generation_detail = asdict(generation)
        else:
            raise RuntimeError(f"API key not configured for answer provider: {answer_provider.name}")
        mark("answer_generation", "passed" if generation_detail["ok"] else "failed", generation_detail)
        if not generation_detail["ok"]:
            raise RuntimeError("Answer generation failed v4 validation")

        checkpoint(task_id)
        update_task(task_id, current_stage="answer_coverage")
        fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
        review_notes = build_answer_review_notes(fragments_data, sdir / "answer_review_notes.json")
        coverage = audit_answer_coverage(structured_exam, fragments_data, sdir / "answer_coverage_audit.json")
        if not coverage["ok"] and options.use_model and answer_provider.api_key:
            mark("answer_coverage_model_repair", "started", {"issues": coverage["issues"][:30], "warnings": coverage["warnings"][:30]})
            model_repair = repair_fragments_with_model_for_audit(
                fragments_json,
                structured_exam,
                candidates,
                selection_data=selection_data,
                provider=answer_provider,
                model=answer_model,
                audit_stage="answer_coverage",
                audit_report=coverage,
                backup_path=sdir / "answer_fragments.before_answer_coverage_model_repair.json",
            )
            mark("answer_coverage_model_repair", "applied" if model_repair.get("changed") else "skipped", model_repair)
            fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
            review_notes = build_answer_review_notes(fragments_data, sdir / "answer_review_notes.json")
            coverage = audit_answer_coverage(structured_exam, fragments_data, sdir / "answer_coverage_audit.json")
        if not coverage["ok"]:
            mark("answer_coverage_local_repair", "started", {"issues": coverage["issues"][:30]})
            local_repair = fill_missing_fragments_locally(
                fragments_json,
                structured_exam,
                candidates,
                "覆盖检查失败后，本地生成待复核占位解析，等待用户确认。",
            )
            mark("answer_coverage_local_repair", "applied" if local_repair.get("changed") else "skipped", local_repair)
            fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
            review_notes = build_answer_review_notes(fragments_data, sdir / "answer_review_notes.json")
            coverage = audit_answer_coverage(structured_exam, fragments_data, sdir / "answer_coverage_audit.json")
        if not coverage["ok"]:
            coverage = wait_for_user_review_decision(
                task_id,
                "answer_coverage",
                coverage,
                sdir,
                title="覆盖检查仍未通过",
                output_json=sdir / "answer_coverage_audit.json",
            )
        mark(
            "answer_coverage",
            "passed" if coverage["ok"] else "failed",
            {
                "question_count": coverage["question_count"],
                "fragment_count": coverage["fragment_count"],
                "covered_count": coverage["covered_count"],
                "issue_count": coverage["issue_count"],
                "warning_count": coverage["warning_count"],
                "review_note_count": review_notes["note_count"],
                "issues": coverage["issues"][:30],
                "warnings": coverage["warnings"][:30],
            },
        )
        if not coverage["ok"]:
            raise RuntimeError("Answer coverage audit failed")

        checkpoint(task_id)
        update_task(task_id, current_stage="figures")
        figure_specs = sdir / "figure_specs.json"
        figure_progress = FigureProgressTracker(sdir / "figure_progress.json")
        figure_progress.emit("stage_started", {"question_count": len(structured_exam.get("items", []))})
        with figure_progress.operation("prepare_figures"):
            generated_figures = prepare_figures_for_fragments(
                structured_exam,
                fragments_json,
                figure_specs,
                sdir / "figures",
                provider=image_provider,
                model=image_model or record.model,
                code_provider=answer_provider,
                code_model=answer_model,
                progress_callback=figure_progress.emit,
            )
        with figure_progress.operation("visual_qa", figure_count=len(generated_figures), model=vision_model):
            figure_qa = audit_figures_with_vision(
                structured_exam,
                figure_specs,
                sdir / "figures",
                sdir / "figure_visual_qa.json",
                provider=vision_provider,
                model=vision_model,
                progress_callback=figure_progress.emit,
            )
        if figure_visual_qa_issue_count(figure_qa) and options.use_model and answer_provider.api_key:
            with figure_progress.operation("visual_qa_repair", model=answer_model):
                figure_repair = repair_figures_with_model_for_visual_qa(
                    structured_exam,
                    fragments_json,
                    figure_specs,
                    sdir / "figures",
                    sdir / "figure_visual_qa.json",
                    sdir / "figure_visual_qa_repair.json",
                    qa_report=figure_qa,
                    provider=answer_provider,
                    model=answer_model,
                    vision_provider=vision_provider,
                    vision_model=vision_model,
                    max_rounds=1,
                    progress_callback=figure_progress.emit,
                )
            figure_qa = figure_repair.get("latest_visual_qa") if isinstance(figure_repair.get("latest_visual_qa"), dict) else figure_qa
            generated_figures = sorted((sdir / "figures").glob("*.png"))
            mark(
                "figure_visual_qa_model_repair",
                "applied" if figure_repair.get("changed") else "skipped",
                {
                    "changed": figure_repair.get("changed"),
                    "rounds": figure_repair.get("rounds", [])[:5],
                    "visual_qa_issue_count": figure_visual_qa_issue_count(figure_qa),
                },
            )
        fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
        figure_qa_issue_count = figure_visual_qa_issue_count(figure_qa)
        mark(
            "figures",
            "failed" if figure_qa_issue_count else "passed",
            {
                "generated_count": len(generated_figures),
                "visual_qa_enabled": figure_qa.get("enabled"),
                "visual_qa_count": len(figure_qa.get("items", [])),
                "visual_qa_issue_count": figure_qa_issue_count,
            },
        )
        figure_progress.emit(
            "stage_completed",
            {
                "generated_count": len(generated_figures),
                "visual_qa_issue_count": figure_qa_issue_count,
                "status": "failed" if figure_qa_issue_count else "passed",
            },
        )

        checkpoint(task_id)
        update_task(task_id, current_stage="content_quality")
        mark("content_quality", "started", {"message": "开始进行内容质量审查。"})
        drafts_path = sdir / "answer_drafts.json"
        selection_path = sdir / "evidence_selection.json"
        drafts_data = json.loads(drafts_path.read_text(encoding="utf-8")) if drafts_path.exists() else {"drafts": []}
        selection_data = json.loads(selection_path.read_text(encoding="utf-8")) if selection_path.exists() else {"selections": []}
        content_quality = audit_content_quality(
            structured_exam,
            fragments_data,
            drafts_data,
            selection_data,
            sdir / "content_quality_audit.json",
        )
        model_repair_quality = _filter_audit_report_for_model_repair(content_quality, CONTENT_QUALITY_MODEL_REPAIR_CODES)
        model_repair_has_targets = bool(model_repair_quality.get("issues") or model_repair_quality.get("warnings"))
        if model_repair_has_targets and options.use_model and answer_provider.api_key:
            skipped_quality_issues = [
                item
                for item in list(content_quality.get("issues", [])) + list(content_quality.get("warnings", []))
                if str(item.get("code") or "").strip() not in CONTENT_QUALITY_MODEL_REPAIR_CODES
            ]
            mark(
                "content_quality_model_repair",
                "started",
                {
                    "issues": model_repair_quality["issues"][:30],
                    "warnings": model_repair_quality["warnings"][:30],
                    "skipped_issues": skipped_quality_issues[:30],
                },
            )
            model_repair = repair_fragments_with_model_for_audit(
                fragments_json,
                structured_exam,
                candidates,
                selection_data=selection_data,
                provider=answer_provider,
                model=answer_model,
                audit_stage="content_quality",
                audit_report=model_repair_quality,
                backup_path=sdir / "answer_fragments.before_content_quality_model_repair.json",
            )
            mark("content_quality_model_repair", "applied" if model_repair.get("changed") else "skipped", model_repair)
            if model_repair.get("changed"):
                repaired_figures = prepare_figures_for_fragments(
                    structured_exam,
                    fragments_json,
                    figure_specs,
                    sdir / "figures",
                    provider=image_provider,
                    model=image_model or record.model,
                    code_provider=answer_provider,
                    code_model=answer_model,
                )
                repaired_figure_qa = audit_figures_with_vision(
                    structured_exam,
                    figure_specs,
                    sdir / "figures",
                    sdir / "figure_visual_qa.json",
                    provider=vision_provider,
                    model=vision_model,
                )
                if figure_visual_qa_issue_count(repaired_figure_qa) and options.use_model and answer_provider.api_key:
                    figure_repair = repair_figures_with_model_for_visual_qa(
                        structured_exam,
                        fragments_json,
                        figure_specs,
                        sdir / "figures",
                        sdir / "figure_visual_qa.json",
                        sdir / "figure_visual_qa_repair.after_content_quality.json",
                        qa_report=repaired_figure_qa,
                        provider=answer_provider,
                        model=answer_model,
                        vision_provider=vision_provider,
                        vision_model=vision_model,
                        max_rounds=1,
                        progress_callback=figure_progress.emit,
                    )
                    repaired_figure_qa = figure_repair.get("latest_visual_qa") if isinstance(figure_repair.get("latest_visual_qa"), dict) else repaired_figure_qa
                    repaired_figures = sorted((sdir / "figures").glob("*.png"))
                    mark(
                        "figure_visual_qa_model_repair_after_content_quality",
                        "applied" if figure_repair.get("changed") else "skipped",
                        {
                            "changed": figure_repair.get("changed"),
                            "rounds": figure_repair.get("rounds", [])[:5],
                            "visual_qa_issue_count": figure_visual_qa_issue_count(repaired_figure_qa),
                        },
                    )
                mark(
                    "figures_after_content_quality_model_repair",
                    "failed" if figure_visual_qa_issue_count(repaired_figure_qa) else "passed",
                    {
                        "generated_count": len(repaired_figures),
                        "visual_qa_enabled": repaired_figure_qa.get("enabled"),
                        "visual_qa_count": len(repaired_figure_qa.get("items", [])),
                        "visual_qa_issue_count": figure_visual_qa_issue_count(repaired_figure_qa),
                        "paths": [str(path) for path in repaired_figures[:20]],
                    },
                )
            fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
            review_notes = build_answer_review_notes(fragments_data, sdir / "answer_review_notes.json")
            content_quality = audit_content_quality(
                structured_exam,
                fragments_data,
                drafts_data,
                selection_data,
                sdir / "content_quality_audit.json",
            )
        elif not content_quality["ok"]:
            mark(
                "content_quality_model_repair",
                "skipped",
                {
                    "reason": "质量审查问题不在模型回修白名单内，跳过模型回修，进入程序自修。",
                    "issues": content_quality["issues"][:30],
                    "warnings": content_quality["warnings"][:30],
                },
            )
        if not content_quality["ok"]:
            mark("content_quality_local_repair", "started", {"issues": content_quality["issues"][:30]})
            local_repair = repair_answer_fragments_for_docx(fragments_json, sdir / "answer_fragments.before_content_quality_local_repair.json")
            mark("content_quality_local_repair", "applied" if local_repair.get("changed") else "skipped", local_repair)
            fragments_data = json.loads(fragments_json.read_text(encoding="utf-8"))
            content_quality = audit_content_quality(
                structured_exam,
                fragments_data,
                drafts_data,
                selection_data,
                sdir / "content_quality_audit.json",
            )
        if not content_quality["ok"]:
            content_quality = auto_allow_audit_report(
                sdir,
                "content_quality",
                content_quality,
                title="质量审查仍有问题",
                output_json=sdir / "content_quality_audit.json",
            )
        content_quality = attach_figure_generation_audit(content_quality, sdir)
        mark(
            "content_quality",
            "passed",
            {
                "question_count": content_quality["question_count"],
                "checked_count": content_quality["checked_count"],
                "issue_count": content_quality["issue_count"],
                "warning_count": content_quality["warning_count"],
                "review_required": False,
                "auto_allowed": bool(content_quality.get("auto_allowed") or content_quality.get("user_allowed")),
                "issues": content_quality["issues"][:30],
                "warnings": content_quality["warnings"][:30],
            },
        )

        checkpoint(task_id)
        update_task(task_id, current_stage="docx")
        docx_path = odir / "answer_book.docx"
        mark("docx", "started", {"docx": str(docx_path), "message": "开始生成并审查最终 Word 文档。"})
        docx_result = build_and_audit_docx_with_repair(
            task_id,
            fragments_json,
            docx_path,
            sdir,
            mark,
            structured_exam=structured_exam,
            candidates=candidates,
            selection_data=selection_data,
            provider=answer_provider,
            model=answer_model,
            use_model=options.use_model,
        )
        issues = docx_result["issues"]
        docx_audit_report = {"ok": not issues, "issues": issues, "warnings": []}
        write_json(sdir / "docx_audit.json", docx_audit_report)
        if issues:
            docx_audit_report = auto_allow_audit_report(
                sdir,
                "docx",
                docx_audit_report,
                title="DOCX 审计仍未通过",
                output_json=sdir / "docx_audit.json",
            )
            issues = docx_audit_report.get("issues", [])
        placeholder_report = None
        if not docx_path.exists() and (docx_audit_report.get("auto_allowed") or docx_audit_report.get("user_allowed")):
            candidate_report = build_user_allowed_docx_candidate(
                fragments_json,
                docx_path,
                sdir,
                "系统按规则默认允许 DOCX 审计问题继续，按当前完整解析内容重新生成正式总版 Word。",
            )
            mark("docx_user_allowed_candidate", "applied" if candidate_report.get("ok") else "failed", candidate_report)
            if not candidate_report.get("ok"):
                placeholder_report = build_user_allowed_docx_placeholder(
                    fragments_json,
                    docx_path,
                    sdir,
                    "系统按规则默认允许 DOCX 审计问题继续，但正式总版 Word 重新生成仍失败。",
                )
                mark("docx_placeholder", "applied", placeholder_report)
        if not docx_path.exists():
            mark("docx", "failed", {"issues": ["answer_book.docx was not generated"], "repair": docx_result.get("repair", {})})
            raise RuntimeError("DOCX was marked allowed but answer_book.docx was not generated")
        figure_size_audit = audit_docx_figure_sizes(docx_path)
        write_json(sdir / "figure_size_audit.json", figure_size_audit)
        mark(
            "figure_size_audit",
            "passed" if figure_size_audit["ok"] else "failed",
            {
                "figure_count": len(figure_size_audit.get("figures", [])),
                "issues": figure_size_audit.get("issues", [])[:20],
                "warnings": figure_size_audit.get("warnings", [])[:20],
            },
        )
        if not figure_size_audit["ok"]:
            raise RuntimeError("Embedded figure size audit failed")
        mark("docx", "passed", {"docx": str(docx_path), "repair": docx_result.get("repair", {}), "placeholder": placeholder_report})

        checkpoint(task_id)
        update_task(task_id, current_stage="question_review")
        mark("question_review", "started", {"message": "开始生成存疑审查文档。"})
        review_items = collect_question_review_items(sdir)
        review_docx = build_question_review_docx(sdir, odir, render_snapshots=options.render_with_word)
        figure_review_items = collect_question_figure_review_items(sdir)
        figure_review_docx = build_figure_review_docx(sdir, odir)
        write_json(
            sdir / "question_review_docx.json",
            {
                "ok": review_docx.exists(),
                "review_question_count": len(review_items),
                "docx": str(review_docx),
                "figure_review_question_count": len(figure_review_items),
                "figure_review_docx": str(figure_review_docx),
            },
        )
        write_json(
            sdir / "figure_review_docx.json",
            {
                "ok": figure_review_docx.exists(),
                "review_question_count": len(figure_review_items),
                "docx": str(figure_review_docx),
            },
        )
        mark(
            "question_review",
            "passed",
            {
                "review_question_count": len(review_items),
                "docx": str(review_docx),
                "figure_review_question_count": len(figure_review_items),
                "figure_review_docx": str(figure_review_docx),
            },
        )

        if options.render_with_word:
            checkpoint(task_id)
            update_task(task_id, current_stage="render")
            mark("render", "started", {"message": "开始生成 PDF/PNG 并进行渲染复核。"})
            rendered = odir / "word_rendered"
            pdf = rendered / "answer_book.pdf"
            export_docx_to_pdf(docx_path, pdf)
            pngs = render_pdf_to_png(pdf, rendered)
            render_issues = audit_rendered_pages(rendered, min_pages=1)
            write_json(sdir / "render_audit.json", {"ok": not render_issues, "issues": render_issues})
            if render_issues:
                mark("render", "failed", {"issues": render_issues[:30]})
                raise RuntimeError("Rendered page audit failed")
            mark("render", "passed", {"pdf": str(pdf), "png_count": len(pngs)})

        report = {
            "task_id": task_id,
            "status": "passed",
            "docx": str(docx_path),
            "pipeline_status": str(sdir / "pipeline_status.json"),
            "rendered": options.render_with_word,
            "content_quality_review_required": not content_quality.get("ok", False),
        }
        checkpoint(task_id)
        update_task(task_id, current_stage="acceptance")
        mark("acceptance", "started", {"message": "开始整理验收结果。"})
        write_json(sdir / "acceptance_report.json", report)
        mark("acceptance", "passed", report)
        checkpoint(task_id)
        update_task(task_id, current_stage="model_usage_report")
        mark("model_usage_report", "started", {"message": "开始生成模型调用汇总文档。"})
        model_usage_report = build_model_usage_report(sdir, odir, task_id)
        mark("model_usage_report", "passed", {"report": str(model_usage_report)})
        checkpoint(task_id)
        update_task(task_id, current_stage="final_acceptance")
        mark("final_acceptance", "started", {"require_render": options.render_with_word})
        final_report = build_final_acceptance_report(sdir, odir, require_render=options.render_with_word)
        if not final_report["ok"]:
            mark("final_acceptance", "failed", {"issues": final_report["issues"][:30]})
            raise RuntimeError("Final acceptance audit failed")
        update_task(task_id, status="completed", current_stage="completed", error="")
        mark("final_acceptance", "passed", {"warning_count": final_report["warning_count"], "report": str(sdir / "final_acceptance_report.json")})
        return report
    except TaskCancelled as exc:
        write_json(sdir / "pipeline_error.json", {"error": str(exc), "cancelled": True})
        mark("pipeline", "cancelled", {"error": str(exc)})
        update_task(task_id, status="cancelled", current_stage="cancelled", error=str(exc))
        raise
    except Exception as exc:
        tb = traceback.format_exc()
        write_json(sdir / "pipeline_error.json", {"error": str(exc), "traceback": tb})
        try:
            build_model_usage_report(sdir, odir, task_id)
        except Exception:
            pass
        update_task(task_id, status="failed", error=str(exc))
        mark("pipeline", "failed", {"error": str(exc)})
        raise
