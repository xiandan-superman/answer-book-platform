from __future__ import annotations

import unicodedata
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from lxml import etree
from PIL import Image, ImageStat


def _page_number(path: Path) -> int:
    try:
        return int(path.stem.rsplit("-", 1)[1])
    except (IndexError, ValueError):
        return 0


def _dark_row_profile(gray: Image.Image, *, top_fraction: float = 0.18) -> list[int]:
    width, height = gray.size
    top = max(1, min(height, int(height * top_fraction)))
    pixels = gray.load()
    return [sum(1 for x in range(width) if pixels[x, y] < 180) for y in range(top)]


def _ink_runs(profile: list[int], *, minimum_pixels: int) -> list[tuple[int, int, int, int]]:
    runs: list[tuple[int, int, int, int]] = []
    start = -1
    peak = 0
    total = 0
    for row, count in enumerate(profile + [0]):
        if count >= minimum_pixels:
            if start < 0:
                start = row
                peak = 0
                total = 0
            peak = max(peak, count)
            total += count
        elif start >= 0:
            runs.append((start, row - 1, peak, total))
            start = -1
    return runs


def _long_horizontal_rule_rows(
    gray: Image.Image,
    *,
    top_fraction: float = 0.18,
    row_profile: list[int] | None = None,
) -> list[int]:
    width, height = gray.size
    top = max(1, min(height, int(height * top_fraction)))
    minimum_run = int(width * 0.35)
    pixels = gray.load()
    rows: list[int] = []
    profile = row_profile if row_profile is not None else _dark_row_profile(gray, top_fraction=top_fraction)
    candidate_rows = [y for y, count in enumerate(profile[:top]) if count >= minimum_run]
    for y in candidate_rows:
        run = 0
        longest = 0
        for x in range(width):
            if pixels[x, y] < 180:
                run += 1
                longest = max(longest, run)
            else:
                run = 0
        if longest >= minimum_run:
            rows.append(y)
    return rows


def inspect_header_clipping(page_paths: list[Path]) -> dict[str, Any]:
    """Detect conservative, renderer-visible header clipping risks.

    The gate intentionally avoids OCR.  It catches ink cut by the physical top
    page edge, text colliding with a long header rule, and a repeated header
    band whose rendered height collapses on only some pages.
    """

    issues: list[str] = []
    page_details: list[dict[str, object]] = []
    header_bands: list[tuple[Path, int, int, int]] = []
    for page_path in page_paths:
        try:
            with Image.open(page_path) as image:
                gray = image.convert("L")
                width, height = gray.size
                profile = _dark_row_profile(gray)
                minimum_pixels = max(2, int(width * 0.0015))
                runs = _ink_runs(profile, minimum_pixels=minimum_pixels)
                edge_rows = max(2, int(height * 0.003))
                edge_peak = max(profile[:edge_rows], default=0)
                global_peak = max(profile, default=0)
                edge_clipped = edge_peak >= max(6, int(global_peak * 0.08))

                text_runs = [run for run in runs if run[2] < int(width * 0.35) and run[3] >= minimum_pixels * 8]
                first_band = text_runs[0] if text_runs else None
                long_rule_rows = _long_horizontal_rule_rows(gray, row_profile=profile)
                long_rule_row_set = set(long_rule_rows)
                long_rule_groups = _ink_runs(
                    [1 if row in long_rule_row_set else 0 for row in range(len(profile))],
                    minimum_pixels=1,
                )
                long_rules = [start for start, end, _, _ in long_rule_groups if end - start + 1 <= 3]
                if first_band is not None:
                    # A chart/table border shortly below the repeated header is
                    # ordinary body content, not a collision.  Only a rule
                    # crossing the text band or its immediate anti-aliased
                    # edge can actually clip the header glyphs.
                    rule_window_end = first_band[1] + max(3, int(height * 0.003))
                    long_rules = [row for row in long_rules if row <= rule_window_end]
                rule_collision = False
                for row in long_rules:
                    above = max(profile[max(0, row - 2) : row], default=0)
                    below = max(profile[row + 1 : row + 3], default=0)
                    if above >= minimum_pixels * 4 and below >= minimum_pixels * 4:
                        rule_collision = True
                        break

                if first_band is not None and first_band[0] < int(height * 0.12):
                    header_bands.append((page_path, first_band[0], first_band[1], first_band[3]))
                if edge_clipped:
                    issues.append(f"{page_path.name} header ink touches the physical top edge and may be clipped")
                if rule_collision:
                    issues.append(f"{page_path.name} header text collides with a long horizontal rule")
                page_details.append(
                    {
                        "page": page_path.name,
                        "edge_clipped": edge_clipped,
                        "rule_collision": rule_collision,
                        "first_header_band": list(first_band[:2]) if first_band else [],
                    }
                )
        except Exception as exc:
            issues.append(f"{page_path.name} header clipping inspection failed: {exc}")

    heights = sorted(end - start + 1 for _, start, end, _ in header_bands)
    if len(heights) >= 3:
        median_height = heights[len(heights) // 2]
        for page_path, start, end, ink_total in header_bands:
            height = end - start + 1
            if median_height >= 8 and height <= median_height * 0.6 and ink_total >= 20:
                issue = (
                    f"{page_path.name} repeated header band height {height}px is below 60% "
                    f"of the {median_height}px median and may be vertically clipped"
                )
                if issue not in issues:
                    issues.append(issue)
    return {"ok": not issues, "issues": issues, "pages": page_details}


def audit_rendered_pages_report(rendered_dir: Path, min_pages: int = 1) -> dict[str, Any]:
    issues: list[str] = []
    pages = sorted(rendered_dir.glob("page-*.png"), key=_page_number)
    if len(pages) < min_pages:
        issues.append(f"rendered PNG page count {len(pages)} below expected minimum {min_pages}")
    for page in pages:
        try:
            with Image.open(page) as img:
                width, height = img.size
                if width < 500 or height < 700:
                    issues.append(f"{page.name} too small: {width}x{height}")
                gray = img.convert("L")
                stat = ImageStat.Stat(gray)
                extrema = gray.getextrema()
                histogram = gray.histogram()
                very_dark_ratio = sum(histogram[:120]) / max(width * height, 1)
                if extrema[1] - extrema[0] < 8:
                    issues.append(f"{page.name} appears blank or nearly uniform")
                if stat.mean and stat.mean[0] > 254 and very_dark_ratio < 0.001:
                    issues.append(f"{page.name} is almost entirely white")
        except Exception as exc:
            issues.append(f"{page.name} could not be inspected: {exc}")
    header_clipping = inspect_header_clipping(pages)
    issues.extend(header_clipping["issues"])
    return {
        "schema_version": "answer_book.rendered_page_audit.v2",
        "ok": not issues,
        "page_count": len(pages),
        "minimum_page_count": min_pages,
        "issues": issues,
        "header_clipping": header_clipping,
    }


def audit_rendered_pages(rendered_dir: Path, min_pages: int = 1) -> list[str]:
    return list(audit_rendered_pages_report(rendered_dir, min_pages=min_pages)["issues"])


def _normalized_delivery_text(value: str) -> str:
    return "".join(char.lower() for char in str(value or "") if char.isalnum() or "\u4e00" <= char <= "\u9fff")


_W_TEXT = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
_W_DRAWING = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
_M_TEXT = "{http://schemas.openxmlformats.org/officeDocument/2006/math}t"
_M_MATH = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"
_M_MATH_PARAGRAPH = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara"
_FORMULA_COMPARISON_SYMBOLS = frozenset("=+-−×÷*/^_()[]{}<>≤≥≈≠∂√∞±·")


def _normalized_formula_text(value: str) -> str:
    return "".join(
        char.lower()
        for char in unicodedata.normalize("NFKC", str(value or ""))
        if char.isalnum() or "\u4e00" <= char <= "\u9fff" or char in _FORMULA_COMPARISON_SYMBOLS
    )


def _paragraph_comparison_spans(paragraph: Any) -> tuple[list[str], list[str]]:
    """Return prose spans split at formulas/drawings plus standalone formulas.

    A Word paragraph can interleave ordinary runs and OMML.  Treating the
    presence of any formula as a reason to discard the whole paragraph makes a
    long explanation untestable.  Splitting at non-prose objects keeps stable
    prose anchors without inventing adjacency across an equation or figure.
    """

    prose_spans: list[str] = []
    formulas: list[str] = []
    prose_parts: list[str] = []

    def flush_prose() -> None:
        value = "".join(prose_parts)
        if value.strip():
            prose_spans.append(value)
        prose_parts.clear()

    def visit(node: Any) -> None:
        if node.tag in {_M_MATH, _M_MATH_PARAGRAPH}:
            flush_prose()
            formula = "".join(str(item.text or "") for item in node.iter() if item.tag == _M_TEXT)
            if formula.strip():
                formulas.append(formula)
            return
        if node.tag == _W_DRAWING:
            flush_prose()
            return
        if node.tag == _W_TEXT:
            prose_parts.append(str(node.text or ""))
            return
        for child in node:
            visit(child)

    visit(paragraph)
    flush_prose()
    return prose_spans, formulas


def _comparison_anchors(document: Document) -> list[tuple[str, str]]:
    """Build conservative text/formula anchors from body and table paragraphs."""

    anchors: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()

    def add(value: str, kind: str) -> None:
        normalized = _normalized_formula_text(value) if kind == "formula" else _normalized_delivery_text(value)
        minimum_length = 3 if kind == "formula" else 6
        if len(normalized) < minimum_length:
            return
        if len(normalized) >= 28:
            positions = (0, max(0, len(normalized) // 2 - 12), max(0, len(normalized) - 24))
            candidates = [normalized[position : position + 24] for position in positions]
        else:
            candidates = [normalized]
        for candidate in candidates:
            key = (kind, candidate)
            if candidate and key not in seen:
                seen.add(key)
                anchors.append(key)

    # document.paragraphs excludes paragraphs inside tables.  Walking body XML
    # keeps ordinary prose and table-cell content under the same rules.
    for paragraph in document.element.body.xpath(".//w:p"):
        prose_spans, formulas = _paragraph_comparison_spans(paragraph)
        for span in prose_spans:
            add(span, "text")
        for formula in formulas:
            add(formula, "formula")
    return anchors


def _pdf_text_and_image_count(pdf_path: Path) -> tuple[str, int]:
    import pypdfium2 as pdfium
    from pypdfium2.raw import FPDF_PAGEOBJ_IMAGE

    texts: list[str] = []
    image_count = 0
    with pdfium.PdfDocument(str(pdf_path)) as document:
        for page in document:
            text_page = page.get_textpage()
            texts.append(text_page.get_text_range())
            image_count += sum(1 for _ in page.get_objects(filter=[FPDF_PAGEOBJ_IMAGE]))
            text_page.close()
            page.close()
    return "\n".join(texts), image_count


def audit_docx_pdf_consistency(docx_path: Path, pdf_path: Path) -> dict:
    """Block stale Word exports and silently omitted embedded figures."""

    issues: list[str] = []
    if not docx_path.is_file() or not pdf_path.is_file():
        return {
            "ok": False,
            "issues": ["DOCX or rendered PDF is missing for delivery consistency audit"],
            "anchor_count": 0,
            "matched_anchor_count": 0,
            "anchor_match_ratio": 0.0,
            "text_anchor_count": 0,
            "formula_anchor_count": 0,
            "docx_drawing_count": 0,
            "pdf_image_count": 0,
            "compared": False,
            "comparison_status": "not_comparable",
            "text_comparison_status": "not_comparable",
            "formula_comparison_status": "not_comparable",
            "image_comparison_status": "not_comparable",
            "not_comparable_reasons": ["DOCX or rendered PDF is missing"],
        }
    document = Document(docx_path)
    anchors = _comparison_anchors(document)
    pdf_text, pdf_image_count = _pdf_text_and_image_count(pdf_path)
    normalized_pdf = _normalized_delivery_text(pdf_text)
    normalized_formula_pdf = _normalized_formula_text(pdf_text)
    matched = sum(
        1
        for kind, anchor in anchors
        if anchor and anchor in (normalized_formula_pdf if kind == "formula" else normalized_pdf)
    )
    match_ratio = matched / max(len(anchors), 1)
    if anchors and match_ratio < 0.6:
        issues.append(
            f"rendered PDF content does not match current DOCX: anchor match ratio {match_ratio:.3f} below 0.600"
        )
    with ZipFile(docx_path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    docx_drawing_count = len(root.xpath(".//w:drawing", namespaces=namespaces))
    if pdf_image_count < docx_drawing_count:
        issues.append(
            f"rendered PDF image count {pdf_image_count} below DOCX drawing count {docx_drawing_count}"
        )
    text_anchor_count = sum(1 for kind, _ in anchors if kind == "text")
    formula_anchor_count = sum(1 for kind, _ in anchors if kind == "formula")
    text_compared = text_anchor_count > 0
    formula_compared = formula_anchor_count > 0
    image_compared = docx_drawing_count > 0
    compared = text_compared or formula_compared or image_compared
    return {
        "ok": not issues,
        "issues": issues,
        "anchor_count": len(anchors),
        "matched_anchor_count": matched,
        "anchor_match_ratio": round(match_ratio, 4),
        "text_anchor_count": text_anchor_count,
        "formula_anchor_count": formula_anchor_count,
        "docx_drawing_count": docx_drawing_count,
        "pdf_image_count": pdf_image_count,
        "compared": compared,
        "comparison_status": "compared" if compared else "not_comparable",
        "text_comparison_status": "compared" if text_compared else "not_comparable",
        "formula_comparison_status": "compared" if formula_compared else "not_comparable",
        "image_comparison_status": "compared" if image_compared else "not_comparable",
        "not_comparable_reasons": [] if compared else ["no comparable DOCX text, formula, or drawing anchors"],
    }
