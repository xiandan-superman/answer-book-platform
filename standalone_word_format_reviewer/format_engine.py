from __future__ import annotations

import re
from dataclasses import asdict, dataclass, replace
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

PROFILE_LABELS = {"answer": "真题答案", "lecture": "讲义"}
FRONT_TITLE = "复习规划及本书使用指南"
FRONT_SECTION_RE = re.compile(r"^[一二三四五六七八九十]+[、.]\s*\S+")
ALLOWED_CHINESE_FONTS = ("宋体", "仿宋", "黑体", "微软雅黑", "华文新魏")
ALLOWED_LATIN_FONTS = ("Times New Roman", "Arial", "Calibri", "Cambria")
ALLOWED_ALIGNMENTS = ("left", "center", "justify")
ALIGNMENT_VALUES = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
ALIGNMENT_NAMES = {value: key for key, value in ALIGNMENT_VALUES.items()}


@dataclass(frozen=True)
class ParagraphSpec:
    chinese_font: str = "宋体"
    latin_font: str = "Times New Roman"
    size_pt: float = 11
    bold: bool | None = None
    alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY
    line_spacing: float = 1.5
    first_line_indent_pt: float = 0
    space_before_pt: float = 0
    space_after_pt: float = 0
    keep_with_next: bool = False
    color: str = "000000"


@dataclass(frozen=True)
class Profile:
    key: str
    label: str
    body_indent_pt: float
    style_specs: dict[str, ParagraphSpec]
    front_boundary_styles: tuple[str, ...]
    footer_font: str
    footer_size_pt: float
    footer_color: str


@dataclass(frozen=True)
class AuditIssue:
    code: str
    location: str
    item: str
    current: str
    expected: str
    suggestion: str
    severity: str = "error"
    fixable: bool = True
    affected: int = 1

    def to_dict(self) -> dict:
        return asdict(self)


BODY_ANSWER = ParagraphSpec(first_line_indent_pt=10)
BODY_LECTURE = ParagraphSpec(first_line_indent_pt=22.1)
FRONT_TITLE_SPEC = ParagraphSpec(size_pt=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_pt=0)
FRONT_SECTION_SPEC = ParagraphSpec(size_pt=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=0)
FIGURE_SPEC = ParagraphSpec(alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_pt=0)

PROFILES: dict[str, Profile] = {
    "answer": Profile(
        key="answer",
        label="真题答案",
        body_indent_pt=10,
        front_boundary_styles=("北航", "Heading 1"),
        footer_font="华文新魏",
        footer_size_pt=9,
        footer_color="000000",
        style_specs={
            "Normal": BODY_ANSWER,
            "北航": ParagraphSpec(chinese_font="黑体", size_pt=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.25),
            "Heading 1": ParagraphSpec(chinese_font="黑体", size_pt=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.25, keep_with_next=True),
            "911代码": ParagraphSpec(chinese_font="黑体", size_pt=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.25),
            "说明": ParagraphSpec(chinese_font="黑体", size_pt=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.25),
            "一二三": ParagraphSpec(size_pt=11.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=0),
            "部分": ParagraphSpec(chinese_font="黑体", size_pt=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_pt=0),
            "图": FIGURE_SPEC,
            "表": ParagraphSpec(alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, first_line_indent_pt=0),
        },
    ),
    "lecture": Profile(
        key="lecture",
        label="讲义",
        body_indent_pt=22.1,
        front_boundary_styles=("一级标题",),
        footer_font="华文新魏",
        footer_size_pt=10.5,
        footer_color="404040",
        style_specs={
            "Normal": BODY_LECTURE,
            "一级标题": ParagraphSpec(size_pt=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_pt=0, keep_with_next=True),
            "二级标题": ParagraphSpec(size_pt=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_pt=0, keep_with_next=True),
            "三级标题": ParagraphSpec(size_pt=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_pt=0, keep_with_next=True),
            "注": ParagraphSpec(size_pt=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_pt=5),
            "图": FIGURE_SPEC,
        },
    ),
}


def _spec_option(spec: ParagraphSpec) -> dict:
    return {
        "chinese_font": spec.chinese_font,
        "latin_font": spec.latin_font,
        "size_pt": float(spec.size_pt),
        "bold": spec.bold,
        "alignment": ALIGNMENT_NAMES.get(spec.alignment, "justify"),
        "line_spacing": float(spec.line_spacing),
        "first_line_indent_pt": float(spec.first_line_indent_pt),
        "space_before_pt": float(spec.space_before_pt),
        "space_after_pt": float(spec.space_after_pt),
        "keep_with_next": bool(spec.keep_with_next),
        "color": spec.color,
    }


def default_task_options(profile_key: str) -> dict:
    try:
        base = PROFILES[profile_key]
    except KeyError as exc:
        raise ValueError("文档类型必须是 answer 或 lecture") from exc
    table_spec = base.style_specs.get("表", base.style_specs["Normal"])
    return {
        "page": {"size": "b5", "margin_cm": 1.7, "header_distance_cm": 1.3, "footer_distance_cm": 1.3},
        "styles": {name: _spec_option(spec) for name, spec in base.style_specs.items()},
        "special": {"front_title": _spec_option(FRONT_TITLE_SPEC), "front_section": _spec_option(FRONT_SECTION_SPEC)},
        "header": {
            **_spec_option(ParagraphSpec(chinese_font="黑体", size_pt=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, first_line_indent_pt=0)),
            "content": "",
            "bottom_border": True,
        },
        "footer": _spec_option(ParagraphSpec(chinese_font=base.footer_font, size_pt=base.footer_size_pt, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, first_line_indent_pt=0, color=base.footer_color)),
        "page_number": _spec_option(ParagraphSpec(size_pt=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, first_line_indent_pt=0)),
        "table": {**_spec_option(table_spec), "alignment": "center" if profile_key == "answer" else "preserve"},
        "image": {"alignment": "center", "max_width_cm": 14.8, "max_height_cm": 22.3},
    }


def _merge_known(target: dict, incoming: dict) -> None:
    for key, value in incoming.items():
        if key not in target:
            continue
        if isinstance(target[key], dict) and isinstance(value, dict):
            _merge_known(target[key], value)
        else:
            target[key] = value


def _validate_spec_option(value: dict, *, allow_preserve: bool = False) -> None:
    if str(value.get("chinese_font")) not in ALLOWED_CHINESE_FONTS:
        raise ValueError("不支持所选中文字体")
    if str(value.get("latin_font")) not in ALLOWED_LATIN_FONTS:
        raise ValueError("不支持所选英文数字字体")
    size = float(value.get("size_pt"))
    if not 8 <= size <= 36:
        raise ValueError("字号必须在8至36 pt之间")
    alignment = str(value.get("alignment"))
    valid_alignments = set(ALLOWED_ALIGNMENTS) | ({"preserve"} if allow_preserve else set())
    if alignment not in valid_alignments:
        raise ValueError("段落对齐方式无效")
    if float(value.get("line_spacing")) not in (1.0, 1.25, 1.5, 2.0):
        raise ValueError("行距必须为1、1.25、1.5或2倍")
    for key in ("first_line_indent_pt", "space_before_pt", "space_after_pt"):
        number = float(value.get(key))
        if not 0 <= number <= 48:
            raise ValueError("段落间距或缩进超出允许范围")
    if value.get("bold") not in (True, False, None):
        raise ValueError("加粗设置无效")
    color = str(value.get("color") or "")
    if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
        raise ValueError("文字颜色必须为6位十六进制色值")


def normalize_task_options(profile_key: str, task_options: dict | None = None) -> dict:
    options = default_task_options(profile_key)
    raw = task_options if isinstance(task_options, dict) else {}
    legacy = {key: raw.get(key) for key in ("chinese_font", "latin_font", "body_size_pt", "line_spacing") if key in raw}
    _merge_known(options, raw)
    if legacy:
        body = options["styles"]["Normal"]
        body["chinese_font"] = legacy.get("chinese_font", body["chinese_font"])
        body["latin_font"] = legacy.get("latin_font", body["latin_font"])
        body["size_pt"] = legacy.get("body_size_pt", body["size_pt"])
        body["line_spacing"] = legacy.get("line_spacing", body["line_spacing"])
        for role in ("表", "注"):
            if role in options["styles"]:
                for key in ("chinese_font", "latin_font", "size_pt"):
                    options["styles"][role][key] = body[key]
                if role == "注":
                    options["styles"][role]["line_spacing"] = body["line_spacing"]
        for key in ("chinese_font", "latin_font", "size_pt"):
            options["table"][key] = body[key]

    page = options["page"]
    if page["size"] not in ("b5", "a4"):
        raise ValueError("页面尺寸仅支持JIS B5或A4")
    for key in ("margin_cm", "header_distance_cm", "footer_distance_cm"):
        page[key] = float(page[key])
        if not 0.8 <= page[key] <= 3.5:
            raise ValueError("页面距离必须在0.8至3.5 cm之间")
    for value in options["styles"].values():
        _validate_spec_option(value)
    for value in options["special"].values():
        _validate_spec_option(value)
    _validate_spec_option(options["header"])
    _validate_spec_option(options["footer"])
    _validate_spec_option(options["page_number"])
    _validate_spec_option(options["table"], allow_preserve=True)
    if profile_key == "answer":
        for key in _spec_option(PROFILES["answer"].style_specs["表"]):
            options["styles"]["表"][key] = options["table"][key]
    options["header"]["content"] = str(options["header"].get("content") or "").strip()[:200]
    options["header"]["bottom_border"] = bool(options["header"].get("bottom_border", True))
    if options["image"]["alignment"] not in ("left", "center"):
        raise ValueError("图片对齐方式无效")
    for key in ("max_width_cm", "max_height_cm"):
        options["image"][key] = float(options["image"][key])
        if not 3 <= options["image"][key] <= 30:
            raise ValueError("图片最大尺寸超出允许范围")
    return options


def _option_spec(value: dict, fallback: ParagraphSpec) -> ParagraphSpec:
    return replace(
        fallback,
        chinese_font=str(value["chinese_font"]),
        latin_font=str(value["latin_font"]),
        size_pt=float(value["size_pt"]),
        bold=value["bold"],
        alignment=ALIGNMENT_VALUES[str(value["alignment"])],
        line_spacing=float(value["line_spacing"]),
        first_line_indent_pt=float(value["first_line_indent_pt"]),
        space_before_pt=float(value["space_before_pt"]),
        space_after_pt=float(value["space_after_pt"]),
        keep_with_next=bool(value["keep_with_next"]),
        color=str(value["color"]).upper(),
    )


def _profile(profile_key: str, task_options: dict | None = None) -> Profile:
    try:
        base = PROFILES[profile_key]
    except KeyError as exc:
        raise ValueError("文档类型必须是 answer 或 lecture") from exc
    options = normalize_task_options(profile_key, task_options)
    style_specs = dict(base.style_specs)
    for style_name, spec in style_specs.items():
        style_specs[style_name] = _option_spec(options["styles"][style_name], spec)
    style_specs["图"] = replace(style_specs["图"], alignment=ALIGNMENT_VALUES[options["image"]["alignment"]])
    return replace(base, style_specs=style_specs)


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _set_run_fonts(run, spec: ParagraphSpec, *, set_bold: bool = True) -> None:
    run.font.name = spec.latin_font
    run.font.size = Pt(spec.size_pt)
    run.font.color.rgb = _rgb(spec.color)
    if set_bold and spec.bold is not None:
        run.font.bold = spec.bold
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), spec.latin_font)
    r_fonts.set(qn("w:hAnsi"), spec.latin_font)
    r_fonts.set(qn("w:eastAsia"), spec.chinese_font)
    r_fonts.set(qn("w:cs"), spec.latin_font)


def _set_style(style, spec: ParagraphSpec) -> None:
    style.font.name = spec.latin_font
    style.font.size = Pt(spec.size_pt)
    style.font.color.rgb = _rgb(spec.color)
    if spec.bold is not None:
        style.font.bold = spec.bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key, value in (("ascii", spec.latin_font), ("hAnsi", spec.latin_font), ("eastAsia", spec.chinese_font), ("cs", spec.latin_font)):
        r_fonts.set(qn(f"w:{key}"), value)
    fmt = style.paragraph_format
    fmt.alignment = spec.alignment
    fmt.line_spacing = spec.line_spacing
    fmt.first_line_indent = Pt(spec.first_line_indent_pt)
    fmt.space_before = Pt(spec.space_before_pt)
    fmt.space_after = Pt(spec.space_after_pt)
    fmt.keep_with_next = spec.keep_with_next


def _format_paragraph(paragraph: Paragraph, spec: ParagraphSpec, *, force_body_bold: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.alignment = spec.alignment
    fmt.line_spacing = spec.line_spacing
    fmt.first_line_indent = Pt(spec.first_line_indent_pt)
    fmt.space_before = Pt(spec.space_before_pt)
    fmt.space_after = Pt(spec.space_after_pt)
    fmt.keep_with_next = spec.keep_with_next
    for run in paragraph.runs:
        if run._r.xpath(".//w:drawing|.//w:object|.//m:oMath"):
            continue
        _set_run_fonts(run, spec, set_bold=force_body_bold or spec.bold is not None)


def _all_body_paragraphs(document: DocumentObject) -> Iterable[Paragraph]:
    yield from document.paragraphs
    seen_cells: set[object] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                yield from cell.paragraphs


def _style_name(paragraph: Paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def _is_image_only(paragraph: Paragraph) -> bool:
    return not paragraph.text.strip() and bool(paragraph._p.xpath(".//w:drawing|.//w:pict"))


def _front_section_indexes(document: DocumentObject, profile: Profile) -> set[int]:
    boundary = len(document.paragraphs)
    for index, paragraph in enumerate(document.paragraphs):
        if _style_name(paragraph) in profile.front_boundary_styles:
            boundary = index
            break
    return {
        index
        for index, paragraph in enumerate(document.paragraphs[:boundary])
        if FRONT_SECTION_RE.match(paragraph.text.strip())
    }


def _style_signature(style) -> tuple:
    r_pr = style.element.rPr
    east_asia = ""
    ascii_font = ""
    if r_pr is not None:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            east_asia = r_fonts.get(qn("w:eastAsia"), "")
            ascii_font = r_fonts.get(qn("w:ascii"), "") or r_fonts.get(qn("w:hAnsi"), "")
    color = str(style.font.color.rgb or "")
    fmt = style.paragraph_format
    return (
        east_asia,
        ascii_font or style.font.name or "",
        round(style.font.size.pt, 2) if style.font.size else None,
        style.font.bold,
        fmt.alignment,
        round(float(fmt.line_spacing), 2) if isinstance(fmt.line_spacing, float) else fmt.line_spacing,
        round(fmt.first_line_indent.pt, 2) if fmt.first_line_indent else 0,
        round(fmt.space_before.pt, 2) if fmt.space_before else 0,
        round(fmt.space_after.pt, 2) if fmt.space_after else 0,
        bool(fmt.keep_with_next),
        color,
    )


def _expected_signature(spec: ParagraphSpec) -> tuple:
    return (
        spec.chinese_font,
        spec.latin_font,
        spec.size_pt,
        spec.bold,
        spec.alignment,
        spec.line_spacing,
        spec.first_line_indent_pt,
        spec.space_before_pt,
        spec.space_after_pt,
        spec.keep_with_next,
        spec.color,
    )


def _style_matches_spec(style, spec: ParagraphSpec) -> bool:
    """Compare required style fields while treating optional bold as a wildcard."""

    actual = _style_signature(style)
    expected = _expected_signature(spec)
    return all(
        current == required
        for index, (current, required) in enumerate(zip(actual, expected))
        if not (index == 3 and spec.bold is None)
    )


def _format_spec_text(spec: ParagraphSpec) -> str:
    align = {
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }.get(spec.alignment, "指定对齐")
    bold = "、加粗" if spec.bold else ""
    return f"中文{spec.chinese_font}、英文数字{spec.latin_font}、{spec.size_pt:g} pt{bold}、{align}、{spec.line_spacing:g}倍行距、首行缩进{spec.first_line_indent_pt:g} pt"


def _describe_style(style) -> str:
    signature = _style_signature(style)
    return f"中文字体{signature[0] or '未指定'}，英文字体{signature[1] or '未指定'}，字号{signature[2] or '未指定'} pt"


def _almost(value: float | None, target: float, tolerance: float = 0.06) -> bool:
    return value is not None and abs(value - target) <= tolerance


def _has_bottom_border(paragraph: Paragraph) -> bool:
    p_bdr = paragraph._p.pPr.find(qn("w:pBdr")) if paragraph._p.pPr is not None else None
    return p_bdr is not None and p_bdr.find(qn("w:bottom")) is not None


def _set_bottom_border(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")


def _remove_bottom_border(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    p_bdr = p_pr.find(qn("w:pBdr")) if p_pr is not None else None
    bottom = p_bdr.find(qn("w:bottom")) if p_bdr is not None else None
    if p_bdr is not None and bottom is not None:
        p_bdr.remove(bottom)


def _unique_headers(document: DocumentObject):
    seen: set[int] = set()
    for section in document.sections:
        header = section.header
        marker = id(header.part)
        if marker not in seen:
            seen.add(marker)
            yield header


def _unique_footers(document: DocumentObject):
    seen: set[int] = set()
    for section in document.sections:
        footer = section.footer
        marker = id(footer.part)
        if marker not in seen:
            seen.add(marker)
            yield footer


def _has_page_field(paragraph: Paragraph) -> bool:
    xml = paragraph._p.xml.upper()
    return "PAGE" in xml and ("INSTRTEXT" in xml or "FLDSIMPLE" in xml)


def _append_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def _set_math_font(document: DocumentObject) -> None:
    settings = document.settings.element
    math_pr = settings.find(qn("m:mathPr"))
    if math_pr is None:
        math_pr = OxmlElement("m:mathPr")
        settings.append(math_pr)
    math_font = math_pr.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_pr.append(math_font)
    math_font.set(qn("m:val"), "Cambria Math")


def _math_font(document: DocumentObject) -> str:
    settings = document.settings.element
    math_font = settings.find(".//" + qn("m:mathFont"))
    return math_font.get(qn("m:val"), "") if math_font is not None else ""


def _formula_count(document: DocumentObject) -> int:
    return len(document.element.body.xpath(".//m:oMath"))


def _conflicting_direct_format(paragraph: Paragraph, spec: ParagraphSpec) -> bool:
    fmt = paragraph.paragraph_format
    checks = (
        (fmt.alignment, spec.alignment),
        (float(fmt.line_spacing) if isinstance(fmt.line_spacing, float) else fmt.line_spacing, spec.line_spacing),
        (round(fmt.first_line_indent.pt, 2) if fmt.first_line_indent else None, spec.first_line_indent_pt),
        (round(fmt.space_before.pt, 2) if fmt.space_before else None, spec.space_before_pt),
        (round(fmt.space_after.pt, 2) if fmt.space_after else None, spec.space_after_pt),
    )
    for current, expected in checks:
        if current is not None and current != expected:
            return True
    for run in paragraph.runs:
        r_pr = run._r.rPr
        if r_pr is None:
            continue
        if run.font.size is not None and not _almost(run.font.size.pt, spec.size_pt):
            return True
        if spec.bold is not None and run.font.bold is not None and run.font.bold != spec.bold:
            return True
        if run.font.color.rgb is not None and str(run.font.color.rgb) != spec.color:
            return True
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            east = r_fonts.get(qn("w:eastAsia"))
            latin = r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi"))
            if east and east != spec.chinese_font:
                return True
            if latin and latin != spec.latin_font:
                return True
    return False


def _explicit_format_matches(paragraph: Paragraph, spec: ParagraphSpec) -> bool:
    fmt = paragraph.paragraph_format
    if fmt.alignment != spec.alignment:
        return False
    if not isinstance(fmt.line_spacing, float) or round(fmt.line_spacing, 2) != spec.line_spacing:
        return False
    indent = round(fmt.first_line_indent.pt, 2) if fmt.first_line_indent else 0
    if indent != spec.first_line_indent_pt:
        return False
    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        return False
    for run in text_runs:
        if run.font.size is None or not _almost(run.font.size.pt, spec.size_pt):
            return False
        if spec.bold is not None and run.font.bold != spec.bold:
            return False
        r_pr = run._r.rPr
        r_fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
        if r_fonts is None:
            return False
        if r_fonts.get(qn("w:eastAsia")) != spec.chinese_font:
            return False
        if (r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi"))) != spec.latin_font:
            return False
    return True


def _exact_height_rows(document: DocumentObject) -> int:
    count = 0
    for table in document.tables:
        for row in table.rows:
            tr_pr = row._tr.trPr
            if tr_pr is None:
                continue
            for height in tr_pr.findall(qn("w:trHeight")):
                if height.get(qn("w:hRule")) == "exact":
                    count += 1
    return count


def _remove_exact_row_heights(document: DocumentObject) -> None:
    for table in document.tables:
        for row in table.rows:
            tr_pr = row._tr.trPr
            if tr_pr is None:
                continue
            for height in list(tr_pr.findall(qn("w:trHeight"))):
                if height.get(qn("w:hRule")) == "exact":
                    tr_pr.remove(height)


def _resolved_options(profile_key: str, header_text: str, task_options: dict | None) -> dict:
    options = normalize_task_options(profile_key, task_options)
    raw_header = task_options.get("header") if isinstance(task_options, dict) else None
    if not isinstance(raw_header, dict) or "content" not in raw_header:
        options["header"]["content"] = str(header_text or "").strip()[:200]
    return options


def _page_dimensions(options: dict) -> tuple[float, float]:
    return (21.0, 29.7) if options["page"]["size"] == "a4" else (18.2, 25.7)


def _table_spec(options: dict, fallback: ParagraphSpec) -> ParagraphSpec:
    value = dict(options["table"])
    if value["alignment"] == "preserve":
        value["alignment"] = ALIGNMENT_NAMES.get(fallback.alignment, "justify")
    return _option_spec(value, fallback)


def audit_docx(path: str | Path, profile_key: str, header_text: str, task_options: dict | None = None) -> dict:
    normalized_options = _resolved_options(profile_key, header_text, task_options)
    profile = _profile(profile_key, normalized_options)
    front_title_spec = _option_spec(normalized_options["special"]["front_title"], FRONT_TITLE_SPEC)
    front_section_spec = _option_spec(normalized_options["special"]["front_section"], FRONT_SECTION_SPEC)
    header_spec = _option_spec(normalized_options["header"], ParagraphSpec())
    footer_text_spec = _option_spec(normalized_options["footer"], ParagraphSpec())
    page_spec = _option_spec(normalized_options["page_number"], ParagraphSpec())
    page_width_cm, page_height_cm = _page_dimensions(normalized_options)
    document = Document(str(path))
    issues: list[AuditIssue] = []
    front_sections = _front_section_indexes(document, profile)
    special_paragraphs = {
        paragraph._p
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == FRONT_TITLE or index in front_sections
    }

    for index, section in enumerate(document.sections, start=1):
        dimensions_ok = _almost(section.page_width.cm, page_width_cm) and _almost(section.page_height.cm, page_height_cm)
        if not dimensions_ok:
            size_label = "A4" if normalized_options["page"]["size"] == "a4" else "JIS B5"
            issues.append(AuditIssue("page_size", f"第{index}节", "页面尺寸", f"{section.page_width.cm:.2f} × {section.page_height.cm:.2f} cm", f"{size_label}：{page_width_cm:g} × {page_height_cm:g} cm", f"调整为{size_label}纵向页面"))
        margins = (section.top_margin.cm, section.bottom_margin.cm, section.left_margin.cm, section.right_margin.cm)
        margin_cm = normalized_options["page"]["margin_cm"]
        if not all(_almost(value, margin_cm) for value in margins):
            issues.append(AuditIssue("page_margin", f"第{index}节", "页边距", " / ".join(f"{value:.2f}" for value in margins) + " cm", f"上下左右均为{margin_cm:.2f} cm", "统一四边页边距"))
        distances = (section.header_distance.cm, section.footer_distance.cm)
        expected_distances = (normalized_options["page"]["header_distance_cm"], normalized_options["page"]["footer_distance_cm"])
        if not all(_almost(value, expected) for value, expected in zip(distances, expected_distances)):
            issues.append(AuditIssue("header_footer_distance", f"第{index}节", "页眉页脚距离", f"{distances[0]:.2f} / {distances[1]:.2f} cm", f"页眉{expected_distances[0]:.2f} / 页脚{expected_distances[1]:.2f} cm", "统一页眉页脚距离"))

    for style_name, spec in profile.style_specs.items():
        if style_name not in document.styles:
            issues.append(AuditIssue("missing_style", "全文样式", style_name, "缺少样式", _format_spec_text(spec), f"创建并应用“{style_name}”样式"))
            continue
        style = document.styles[style_name]
        if not _style_matches_spec(style, spec):
            affected = sum(1 for paragraph in _all_body_paragraphs(document) if _style_name(paragraph) == style_name)
            issues.append(AuditIssue("style_format", "全文样式", style_name, _describe_style(style), _format_spec_text(spec), f"统一“{style_name}”样式及其所属段落", affected=max(affected, 1)))

        conflicting = [
            paragraph
            for paragraph in _all_body_paragraphs(document)
            if paragraph._p not in special_paragraphs
            and _style_name(paragraph) == style_name
            and _conflicting_direct_format(paragraph, spec)
        ]
        if conflicting:
            issues.append(AuditIssue("direct_format", "全文段落", f"“{style_name}”样式的局部格式", f"{len(conflicting)}段存在覆盖样式的局部格式", _format_spec_text(spec), "清除冲突的局部格式并按所属样式统一", affected=len(conflicting)))

    front_titles = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == FRONT_TITLE]
    wrong_front_titles = [paragraph for paragraph in front_titles if not _explicit_format_matches(paragraph, front_title_spec)]
    if wrong_front_titles:
        issues.append(AuditIssue("front_title_format", "前置说明", "前置说明总标题", f"{len(wrong_front_titles)}处格式不一致", _format_spec_text(front_title_spec), "按前置说明总标题格式统一", affected=len(wrong_front_titles)))
    wrong_front_sections = [document.paragraphs[index] for index in front_sections if not _explicit_format_matches(document.paragraphs[index], front_section_spec)]
    if wrong_front_sections:
        issues.append(AuditIssue("front_section_format", "前置说明", "前置说明分节标题", f"{len(wrong_front_sections)}处格式不一致", _format_spec_text(front_section_spec), "按前置说明分节标题格式统一", affected=len(wrong_front_sections)))

    expected_header = normalized_options["header"]["content"]
    for index, header in enumerate(_unique_headers(document), start=1):
        actual = "\n".join(p.text.strip() for p in header.paragraphs if p.text.strip())
        if actual != expected_header:
            issues.append(AuditIssue("header_content", f"页眉{index}", "页眉文字", actual or "空", expected_header or "空", "替换为用户填写的页眉内容"))
        first = next((p for p in header.paragraphs if p.text.strip()), header.paragraphs[0])
        border_ok = _has_bottom_border(first) == normalized_options["header"]["bottom_border"]
        text_format_ok = _explicit_format_matches(first, header_spec) if expected_header else first.alignment == header_spec.alignment
        if not text_format_ok or not border_ok:
            issues.append(AuditIssue("header_format", f"页眉{index}", "页眉格式", "未完全符合", _format_spec_text(header_spec) + ("、下横线" if normalized_options["header"]["bottom_border"] else "、无下横线"), "统一页眉格式"))

    if _formula_count(document) and _math_font(document) != "Cambria Math":
        issues.append(AuditIssue("math_font", "全文公式", "公式字体", _math_font(document) or "未指定", "Cambria Math", "设置Word原生公式默认字体；不改写公式内容", affected=_formula_count(document)))

    image_paragraphs = [paragraph for paragraph in _all_body_paragraphs(document) if _is_image_only(paragraph)]
    image_alignment = ALIGNMENT_VALUES[normalized_options["image"]["alignment"]]
    wrong_image_paragraphs = [paragraph for paragraph in image_paragraphs if paragraph.alignment != image_alignment or _style_name(paragraph) != "图"]
    if wrong_image_paragraphs:
        image_alignment_label = "居中" if normalized_options["image"]["alignment"] == "center" else "左对齐"
        issues.append(AuditIssue("image_paragraph", "全文图片", "独立图片段落", f"{len(wrong_image_paragraphs)}处未使用图样式及指定对齐", f"使用“图”样式、{image_alignment_label}、不缩进", "统一独立图片段落；不缩放或重绘图片", affected=len(wrong_image_paragraphs)))

    content_width = Cm(normalized_options["image"]["max_width_cm"])
    content_height = Cm(normalized_options["image"]["max_height_cm"])
    oversized = [shape for shape in document.inline_shapes if shape.width > content_width or shape.height > content_height]
    if oversized:
        issues.append(AuditIssue("image_overflow", "全文图片", "图片超出版心", f"{len(oversized)}张图片宽度或高度超过当前版心", "图片完整位于版心内并保持原始宽高比", "仅对超限图片等比缩小，不放大其他图片", affected=len(oversized)))

    table_spec = _table_spec(normalized_options, profile.style_specs.get("表", profile.style_specs["Normal"]))
    wrong_table_paragraphs = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    effective_spec = table_spec
                    if normalized_options["table"]["alignment"] == "preserve" and paragraph.alignment is not None:
                        effective_spec = replace(table_spec, alignment=paragraph.alignment)
                    if _conflicting_direct_format(paragraph, effective_spec):
                        wrong_table_paragraphs.append(paragraph)
    if wrong_table_paragraphs:
        issues.append(AuditIssue("table_format", "全文表格", "表格文字与段落格式", f"{len(wrong_table_paragraphs)}个表格段落格式不一致", _format_spec_text(table_spec), "统一表格字体字号；按所选方式处理对齐", affected=len(wrong_table_paragraphs)))

    exact_rows = _exact_height_rows(document)
    if exact_rows:
        issues.append(AuditIssue("table_row_height", "全文表格", "固定行高", f"{exact_rows}行使用固定行高", "行高自动适应内容", "移除固定行高，避免文字被截断", affected=exact_rows))

    footers = list(_unique_footers(document))
    footer_has_page = any(_has_page_field(p) for footer in footers for p in footer.paragraphs)
    if not footer_has_page:
        issues.append(AuditIssue("page_number", "全文页脚", "页码", "未找到PAGE页码域", _format_spec_text(page_spec), "在页脚增加Word自动页码"))

    wrong_page_format = 0
    wrong_footer_text = 0
    for footer in footers:
        for paragraph in footer.paragraphs:
            if _has_page_field(paragraph):
                wrong_page_format += not _explicit_format_matches(paragraph, page_spec)
            elif paragraph.text.strip():
                wrong_footer_text += not _explicit_format_matches(paragraph, footer_text_spec)
    if wrong_page_format:
        issues.append(AuditIssue("page_number_format", "全文页脚", "页码格式", f"{wrong_page_format}处格式不一致", _format_spec_text(page_spec), "统一现有自动页码格式", affected=wrong_page_format))
    if wrong_footer_text:
        issues.append(AuditIssue("footer_text_format", "全文页脚", "页脚宣传语格式", f"{wrong_footer_text}处格式不一致", _format_spec_text(footer_text_spec), "保留原文字，仅统一格式", affected=wrong_footer_text))

    known_styles = set(profile.style_specs) | {"Normal"}
    unknown_counts: dict[str, int] = {}
    for paragraph in _all_body_paragraphs(document):
        name = _style_name(paragraph)
        if paragraph.text.strip() and name and name not in known_styles:
            unknown_counts[name] = unknown_counts.get(name, 0) + 1
    if unknown_counts:
        detail = "、".join(f"{name}（{count}段）" for name, count in sorted(unknown_counts.items()))
        issues.append(AuditIssue("unknown_styles", "全文样式", "标准外样式", detail, f"使用{profile.label}标准中的已定义层级", "请人工确认这些段落的内容角色后再决定是否改样式", severity="warning", fixable=False, affected=sum(unknown_counts.values())))

    error_count = sum(issue.severity == "error" for issue in issues)
    warning_count = sum(issue.severity == "warning" for issue in issues)
    return {
        "profile": profile.key,
        "profile_label": profile.label,
        "task_options": normalized_options,
        "source_name": Path(path).name,
        "summary": {
            "issue_count": len(issues),
            "fixable_count": sum(issue.fixable for issue in issues),
            "error_count": error_count,
            "warning_count": warning_count,
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "formula_count": _formula_count(document),
            "image_count": len(document.element.body.xpath('.//w:drawing')),
        },
        "issues": [issue.to_dict() for issue in issues],
    }


def repair_docx(source: str | Path, destination: str | Path, profile_key: str, header_text: str, task_options: dict | None = None) -> dict:
    normalized_options = _resolved_options(profile_key, header_text, task_options)
    profile = _profile(profile_key, normalized_options)
    front_title_spec = _option_spec(normalized_options["special"]["front_title"], FRONT_TITLE_SPEC)
    front_section_spec = _option_spec(normalized_options["special"]["front_section"], FRONT_SECTION_SPEC)
    header_spec = _option_spec(normalized_options["header"], ParagraphSpec())
    footer_text_spec = _option_spec(normalized_options["footer"], ParagraphSpec())
    page_spec = _option_spec(normalized_options["page_number"], ParagraphSpec())
    table_spec = _table_spec(normalized_options, profile.style_specs.get("表", profile.style_specs["Normal"]))
    figure_spec = replace(profile.style_specs["图"], alignment=ALIGNMENT_VALUES[normalized_options["image"]["alignment"]])
    page_width_cm, page_height_cm = _page_dimensions(normalized_options)
    document = Document(str(source))

    for section in document.sections:
        section.page_width = Cm(page_width_cm)
        section.page_height = Cm(page_height_cm)
        margin = Cm(normalized_options["page"]["margin_cm"])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
        section.header_distance = Cm(normalized_options["page"]["header_distance_cm"])
        section.footer_distance = Cm(normalized_options["page"]["footer_distance_cm"])

    for style_name, spec in profile.style_specs.items():
        style = document.styles[style_name] if style_name in document.styles else document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        _set_style(style, spec)

    front_sections = _front_section_indexes(document, profile)
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text == FRONT_TITLE:
            _format_paragraph(paragraph, front_title_spec, force_body_bold=True)
            continue
        if index in front_sections:
            _format_paragraph(paragraph, front_section_spec, force_body_bold=True)
            continue
        style_name = _style_name(paragraph)
        if style_name in profile.style_specs:
            _format_paragraph(paragraph, profile.style_specs[style_name])
        if _is_image_only(paragraph):
            paragraph.style = document.styles["图"]
            _format_paragraph(paragraph, figure_spec)

    seen_cells: set[object] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph in cell.paragraphs:
                    if profile.key == "answer":
                        paragraph.style = document.styles["表"]
                        _format_paragraph(paragraph, table_spec)
                    else:
                        if normalized_options["table"]["alignment"] == "preserve":
                            for run in paragraph.runs:
                                if not run._r.xpath(".//w:drawing|.//w:object|.//m:oMath"):
                                    _set_run_fonts(run, table_spec, set_bold=False)
                        else:
                            _format_paragraph(paragraph, table_spec)
                    if _is_image_only(paragraph):
                        paragraph.style = document.styles["图"]
                        _format_paragraph(paragraph, figure_spec)

    for header in _unique_headers(document):
        first = header.paragraphs[0]
        first.clear()
        first.add_run(normalized_options["header"]["content"])
        _format_paragraph(first, header_spec, force_body_bold=True)
        if normalized_options["header"]["bottom_border"]:
            _set_bottom_border(first)
        else:
            _remove_bottom_border(first)
        for extra in header.paragraphs[1:]:
            extra.clear()

    for footer in _unique_footers(document):
        page_paragraph = None
        for paragraph in footer.paragraphs:
            if _has_page_field(paragraph):
                page_paragraph = paragraph
                _format_paragraph(paragraph, page_spec)
            elif paragraph.text.strip():
                _format_paragraph(paragraph, footer_text_spec)
        if page_paragraph is None:
            page_paragraph = footer.add_paragraph()
            _append_page_field(page_paragraph)
            _format_paragraph(page_paragraph, page_spec)

    _set_math_font(document)
    _remove_exact_row_heights(document)
    content_width = Cm(normalized_options["image"]["max_width_cm"])
    content_height = Cm(normalized_options["image"]["max_height_cm"])
    for shape in document.inline_shapes:
        if shape.width <= 0 or shape.height <= 0:
            continue
        scale = min(1.0, content_width / shape.width, content_height / shape.height)
        if scale < 1.0:
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)
    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return audit_docx(destination, profile_key, normalized_options["header"]["content"], normalized_options)
