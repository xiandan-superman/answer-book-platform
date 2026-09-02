from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from lxml import etree

from app.docx_audit import (
    audit_docx_v4,
    math_node_has_empty_delimiter_character,
    math_node_has_empty_delimiter_slots,
)
from app.docx_v4 import _answer_summary_formula_candidates, build_docx_from_fragments
from app.expression_normalization import normalize_control_word_boundaries
from app.expression_promotion import promote_inline_mathematical_expressions
from app.final_acceptance import build_final_acceptance_report
from app.omml import FormulaConversionError, normalize_latex, omml_from_latex


class FormulaRenderingGuardTests(unittest.TestCase):
    def test_multiline_display_latex_is_one_formula_candidate(self) -> None:
        text = "结果：$$\nx=\\frac{a}{b}\n$$。"

        candidates = _answer_summary_formula_candidates(text)

        self.assertEqual([r"x=\frac{a}{b}"], [latex for _start, _end, latex in candidates])

    def test_unicode_large_operator_with_subscript_is_promoted_to_latex(self) -> None:
        candidates = _answer_summary_formula_candidates("答案中的 ∑_B 需要保留。")

        self.assertEqual([r"\sum_{B}"], [latex for _start, _end, latex in candidates])

    def test_cases_environment_allows_intentionally_invisible_right_delimiter(self) -> None:
        latex = (
            r"F_{HKL}=\begin{cases}"
            r"2f,&H+K+L\text{为偶数}\\"
            r"0,&H+K+L\text{为奇数}"
            r"\end{cases}"
        )

        omml = omml_from_latex(latex)

        self.assertTrue(list(omml))
        self.assertFalse(math_node_has_empty_delimiter_character(omml))

    def test_primary_formula_conversion_failure_does_not_silently_degrade(self) -> None:
        with patch("app.omml.omml_from_latex_via_mathml", side_effect=RuntimeError("broken transform")), patch.dict(
            "os.environ",
            {
                "ANSWER_BOOK_DISABLE_MATHML_OMML": "0",
                "ANSWER_BOOK_ALLOW_DEGRADED_OMML_FALLBACK": "0",
            },
        ):
            with self.assertRaises(FormulaConversionError):
                omml_from_latex(r"\frac{a}{b}")

    def test_degraded_formula_fallback_requires_explicit_emergency_switch(self) -> None:
        with patch("app.omml.omml_from_latex_via_mathml", side_effect=RuntimeError("broken transform")), patch.dict(
            "os.environ",
            {"ANSWER_BOOK_ALLOW_DEGRADED_OMML_FALLBACK": "1"},
            clear=False,
        ):
            omml = omml_from_latex(r"\frac{a}{b}")

        self.assertTrue(list(omml))

    def test_vector_norm_product_does_not_create_empty_omml_slots(self) -> None:
        latex = r"\cos\varphi=\frac{\mathbf{g}_1\cdot\mathbf{g}_2}{\left|\mathbf{g}_1\right|\left|\mathbf{g}_2\right|}"

        normalized = normalize_latex(latex)
        omml = omml_from_latex(latex)
        xml = etree.tostring(omml, encoding="unicode")

        self.assertIn(r"\vert", normalized)
        self.assertNotIn("<m:e/>", xml)
        self.assertFalse(math_node_has_empty_delimiter_slots(omml))

    def test_percent_before_chemical_symbol_uses_separate_portable_omml_runs(self) -> None:
        latex = r"\mathrm{Cu}-18\,\mathrm{at.\%Ni}"

        normalized = normalize_latex(latex)
        omml = omml_from_latex(latex)
        texts = omml.xpath(
            ".//m:t/text()",
            namespaces={"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"},
        )

        self.assertIn(r"\mathrm{at.\%}\mathrm{Ni}", normalized)
        self.assertIn("%", texts)
        self.assertIn("Ni", texts)
        self.assertNotIn("%Ni", texts)

    def test_xrightarrow_annotations_use_portable_plain_arrow(self) -> None:
        self.assertEqual(
            normalize_latex(r"\gamma\xrightarrow[\,T<M_s\,]{\text{水淬}}\mathrm{M}"),
            r"\gamma\rightarrow\mathrm{M}",
        )

    def test_docx_audit_reports_empty_formula_delimiter_slots(self) -> None:
        xml = etree.fromstring(
            b"""
            <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
              <m:d>
                <m:dPr><m:begChr m:val="|"/><m:endChr m:val="|"/></m:dPr>
                <m:e/><m:e/><m:e/>
              </m:d>
            </m:oMath>
            """
        )

        self.assertTrue(math_node_has_empty_delimiter_slots(xml))

    def test_docx_audit_reports_empty_formula_delimiter_character(self) -> None:
        xml = etree.fromstring(
            b"""
            <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
              <m:d><m:dPr><m:endChr m:val=""/></m:dPr><m:e><m:r><m:t>x</m:t></m:r></m:e></m:d>
            </m:oMath>
            """
        )

        self.assertTrue(math_node_has_empty_delimiter_character(xml))

    def test_parenthesized_fraction_with_script_has_no_empty_delimiter_character(self) -> None:
        latex = r"\left(\frac{\partial x}{\partial T}\right)_p=-\frac{y}{T}"

        omml = omml_from_latex(latex)
        xml = etree.tostring(omml, encoding="unicode")

        self.assertFalse(math_node_has_empty_delimiter_character(omml), xml)
        self.assertNotIn("<m:begChr m:val=\"\"", xml)
        self.assertNotIn("<m:endChr m:val=\"\"", xml)

    def test_expression_normalization_is_idempotent_across_nested_fraction_groups(self) -> None:
        latex = (
            r"\left(\frac{\partial \Delta_{\mathrm{r}} G_m^\ominus}{\partial T}\right)_p"
            r"=-\frac{\Delta_{\mathrm{r}} H_m^\ominus}{T}"
        )

        once = normalize_latex(latex)
        twice = normalize_latex(once)

        self.assertEqual(twice, once)
        self.assertEqual(twice.count(r"\frac"), 2)
        self.assertEqual(twice.count("{"), twice.count("}"))
        self.assertIn(r"G_m^{\theta}", twice)

    def test_docx_audit_passes_generated_vector_norm_formula(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            docx_path = Path(raw_tmp) / "formula.docx"
            doc = Document()
            p = doc.add_paragraph()
            p._p.append(
                omml_from_latex(
                    r"\cos\varphi=\frac{\mathbf{g}_1\cdot\mathbf{g}_2}{\left|\mathbf{g}_1\right|\left|\mathbf{g}_2\right|}"
                )
            )
            doc.save(docx_path)

            issues = audit_docx_v4(docx_path, min_formulas=1)

        self.assertFalse(any("empty delimiter slots" in issue for issue in issues), issues)

    def test_compact_phase_reaction_has_command_boundary_in_word_math(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            docx_path = Path(raw_tmp) / "reaction.docx"
            doc = Document()
            p = doc.add_paragraph()
            p._p.append(omml_from_latex(r"L+bcc\tofcc"))
            p._p.append(omml_from_latex(r"L\tobcc+\mathrm{Fe}_{2}\mathrm{Ti}"))
            doc.save(docx_path)

            issues = audit_docx_v4(docx_path, min_formulas=2)

        self.assertFalse(any("raw latex marker" in issue for issue in issues), issues)

    def test_compact_thermodynamic_commands_render_without_raw_latex(self) -> None:
        self.assertEqual(normalize_latex(r"\Delta U=nCv\DeltaT"), r"\Delta U=nCv\Delta T")
        self.assertEqual(normalize_latex(r"K\thetaT > 1"), r"K_{T}^{\theta} > 1")
        self.assertEqual(
            normalize_latex(r"\Delta rHm\theta < 0"),
            r"\Delta_{\mathrm{r}} H_{\mathrm{m}}^{\theta} < 0",
        )

        with tempfile.TemporaryDirectory() as raw_tmp:
            docx_path = Path(raw_tmp) / "thermodynamics.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph._p.append(omml_from_latex(r"\Delta U=nCv\DeltaT"))
            paragraph._p.append(omml_from_latex(r"K\thetaT > 1"))
            doc.save(docx_path)

            issues = audit_docx_v4(docx_path, min_formulas=2)

        self.assertFalse(any("raw latex marker" in issue for issue in issues), issues)

    def test_supported_control_words_use_longest_prefix_tokenization(self) -> None:
        samples = {
            r"\DeltaTemperature": r"\Delta Temperature",
            r"\thetaCondition": r"\theta Condition",
            r"L\tophase": r"L\to phase",
            r"\gammaRate": r"\gamma Rate",
        }

        for raw, expected in samples.items():
            with self.subTest(raw=raw):
                self.assertEqual(normalize_control_word_boundaries(raw), expected)

    def test_pending_formula_review_language_is_not_exposed_in_formal_docx(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            fragments_json = root / "answer_fragments.json"
            output_docx = root / "answer_book.docx"
            fragments_json.write_text(
                json.dumps(
                    {
                        "schema_version": "answer_book.answer_fragments.v4",
                        "fragments": [
                            {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "fill_pending_formula",
                                "section": "一、填空题",
                                "question_type": "填空题",
                                "number": "3",
                                "answer": "减少",
                                "answer_summary": "减少",
                                "evidence_ids": [],
                                "blocks": [
                                    {
                                        "label": "待复核公式",
                                        "segments": [
                                            {
                                                "type": "text",
                                                "text": "以下公式未能自然融入解析，请复核其必要性与放置位置。",
                                            },
                                            {"type": "formula_ref", "formula_id": "f1"},
                                        ],
                                    }
                                ],
                                "formulas": [
                                    {
                                        "formula_id": "f1",
                                        "latex": r"\Delta U=nCv\DeltaT",
                                        "role": "relation",
                                        "display": True,
                                    }
                                ],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_docx_from_fragments(fragments_json, output_docx)
            document = Document(output_docx)
            paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            issues = audit_docx_v4(output_docx, min_formulas=1)

        self.assertIn("补充公式", paragraph_text)
        self.assertNotIn("待复核公式", paragraph_text)
        self.assertNotIn("以下公式未能自然融入解析", paragraph_text)
        self.assertFalse(any("internal review language" in issue for issue in issues), issues)

    def test_short_answer_answer_field_renders_dollar_latex_as_omml(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            fragments_json = root / "answer_fragments.json"
            output_docx = root / "answer_book.docx"
            fragments_json.write_text(
                json.dumps(
                    {
                        "schema_version": "answer_book.answer_fragments.v4",
                        "fragments": [
                            {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "q_diffusion",
                                "section": "六、简答题",
                                "question_type": "简答题",
                                "number": "1",
                                "answer": "见解析",
                                "answer_summary": r"稳态扩散流量 $J = \frac{D_1 D_2 c_0}{L(D_2 + 2D_1)}$",
                                "evidence_ids": [],
                                "blocks": [],
                                "formulas": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_docx_from_fragments(fragments_json, output_docx)
            issues = audit_docx_v4(output_docx, min_formulas=1)

        self.assertFalse(any("raw latex marker in normal text" in issue for issue in issues), issues)

    def test_structured_multiline_answer_summary_renders_without_raw_latex(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            fragments_json = root / "answer_fragments.json"
            output_docx = root / "answer_book.docx"
            fragment = promote_inline_mathematical_expressions(
                {
                    "schema_version": "answer_book.answer_fragment.v4",
                    "question_id": "qa_s01_02_02",
                    "section": "六、简答题",
                    "question_type": "简答题",
                    "number": "2",
                    "answer": "见解析",
                    "answer_summary": (
                        "异氰酸酯的物质的量为\n"
                        "$$\n"
                        r"n_{\mathrm{OCN-C_6H_4-NCO}}=\frac{m}{M}"
                        "\n$$"
                    ),
                    "evidence_ids": [],
                    "blocks": [],
                    "formulas": [],
                }
            )
            fragments_json.write_text(
                json.dumps(
                    {"schema_version": "answer_book.answer_fragments.v4", "fragments": [fragment]},
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_docx_from_fragments(fragments_json, output_docx)
            document = Document(output_docx)
            paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            issues = audit_docx_v4(output_docx, min_formulas=1)

        self.assertNotIn("$$", paragraph_text)
        self.assertFalse(any("raw latex marker" in issue for issue in issues), issues)

    def test_fill_blank_answer_field_renders_dollar_latex_as_omml(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            fragments_json = root / "answer_fragments.json"
            output_docx = root / "answer_book.docx"
            fragments_json.write_text(
                json.dumps(
                    {
                        "schema_version": "answer_book.answer_fragments.v4",
                        "fragments": [
                            {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "fill_formula",
                                "section": "二、填空题",
                                "question_type": "填空题",
                                "number": "1",
                                "answer": r"$\\Delta G < 0$",
                                "answer_summary": "",
                                "evidence_ids": [],
                                "blocks": [],
                                "formulas": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_docx_from_fragments(fragments_json, output_docx)
            issues = audit_docx_v4(output_docx, min_formulas=1)

        self.assertFalse(any("raw latex marker in normal text" in issue for issue in issues), issues)

    def test_fill_blank_answer_renders_braced_thermodynamic_subscript(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            fragments_json = root / "answer_fragments.json"
            output_docx = root / "answer_book.docx"
            fragments_json.write_text(
                json.dumps(
                    {
                        "schema_version": "answer_book.answer_fragments.v4",
                        "fragments": [
                            {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "fill_thermodynamic_formula",
                                "section": "一、填空题",
                                "question_type": "填空题",
                                "number": "2",
                                "answer": "见解析",
                                "answer_summary": "Q_V = Δ_rU；Δ_rC_{V,m}；不确定",
                                "evidence_ids": [],
                                "blocks": [],
                                "formulas": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_docx_from_fragments(fragments_json, output_docx)
            issues = audit_docx_v4(output_docx, min_formulas=3)

        self.assertFalse(any("raw latex marker in normal text" in issue for issue in issues), issues)

    def test_fill_blank_answer_renders_equation_chain_next_to_chinese_text(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            fragments_json = root / "answer_fragments.json"
            output_docx = root / "answer_book.docx"
            fragments_json.write_text(
                json.dumps(
                    {
                        "schema_version": "answer_book.answer_fragments.v4",
                        "fragments": [
                            {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "fill_steady_state",
                                "section": "一、填空题",
                                "question_type": "填空题",
                                "number": "7",
                                "answer": "见解析",
                                "answer_summary": (
                                    "稳态时dcB/dt≈0；其方程为dcB/dt = k1cA - k2cB ≈ 0。"
                                ),
                                "evidence_ids": [],
                                "blocks": [],
                                "formulas": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_docx_from_fragments(fragments_json, output_docx)
            issues = audit_docx_v4(output_docx, min_formulas=2)

        self.assertFalse(any("raw latex marker in normal text" in issue for issue in issues), issues)

    def test_final_acceptance_fails_demo_fragments_end_to_end(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            stage_dir = root / "stage_outputs"
            output_dir = root / "outputs"
            stage_dir.mkdir()
            (output_dir / "word_rendered").mkdir(parents=True)
            (output_dir / "answer_book.docx").write_bytes(b"placeholder")
            (output_dir / "word_rendered" / "answer_book.pdf").write_bytes(b"placeholder")
            for name in (
                "environment_check.json",
                "exam_structure_audit.json",
                "retrieval_audit.json",
                "answer_coverage_audit.json",
                "content_quality_audit.json",
                "docx_audit.json",
                "figure_size_audit.json",
                "render_audit.json",
            ):
                payload = {"ok": True, "issues": [], "warnings": []}
                if name == "environment_check.json":
                    payload["formula_conversion"] = {"preferred_chain_ready": True}
                (stage_dir / name).write_text(json.dumps(payload), encoding="utf-8")
            (stage_dir / "acceptance_report.json").write_text(json.dumps({"status": "passed"}), encoding="utf-8")
            (stage_dir / "pipeline_status.json").write_text(json.dumps({"stages": []}), encoding="utf-8")
            (stage_dir / "answer_fragments.json").write_text(
                json.dumps({"provider": "demo", "fragments": [{"question_id": "q1", "answer": "待复核"}]}, ensure_ascii=False),
                encoding="utf-8",
            )

            report = build_final_acceptance_report(stage_dir, output_dir)

        self.assertFalse(report["ok"])
        self.assertEqual("failed", report["status"])
        self.assertTrue(any("demo 占位流程" in issue for issue in report["issues"]))


if __name__ == "__main__":
    unittest.main()
