from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import parse_xml

from app.omml_input import clear_omml_input_caches, mixed_text_with_structured_math
from app.practice_inputs import _docx_content


def test_docx_practice_input_preserves_paragraph_table_body_order() -> None:
    with tempfile.TemporaryDirectory() as raw_tmp:
        path = Path(raw_tmp) / "ordered.docx"
        document = Document()
        document.add_paragraph("表格之前")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格正文"
        paragraph = document.add_paragraph("表格之后")
        paragraph._p.append(
            parse_xml(
                '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                "<m:r><m:t>x</m:t></m:r>"
                "</m:oMath>"
            )
        )
        document.save(path)

        with patch("app.practice_inputs.find_omml2mathml_xsl", return_value=None), patch(
            "app.omml_input.find_omml2mathml_xsl",
            return_value=None,
        ):
            text, _images, diagnostics = _docx_content(path.name, path.read_bytes())

    assert text.index("表格之前") < text.index("表格正文") < text.index("表格之后")
    assert "x⟦OMML_STRUCTURE_UNAVAILABLE⟧" in text
    assert diagnostics["omml_degraded_formula_count"] == 1
    assert diagnostics["omml_structured_formula_count"] == 0


def test_omml_uses_configured_structural_xslt_adapter() -> None:
    stylesheet = """<?xml version="1.0" encoding="UTF-8"?>
    <xsl:stylesheet version="1.0"
      xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
      xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
      xmlns="http://www.w3.org/1998/Math/MathML">
      <xsl:output method="xml" omit-xml-declaration="yes"/>
      <xsl:template match="/">
        <math><mi><xsl:value-of select="//m:t"/></mi></math>
      </xsl:template>
    </xsl:stylesheet>
    """
    with tempfile.TemporaryDirectory() as raw_tmp:
        xsl = Path(raw_tmp) / "omml2mathml.xsl"
        xsl.write_text(stylesheet, encoding="utf-8")
        node = parse_xml(
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            "<m:r><m:t>alpha</m:t></m:r>"
            "</m:oMath>"
        )
        with patch.dict("os.environ", {"OMML2MATHML_XSL": str(xsl)}, clear=False), patch(
            "app.omml_input.platform.system",
            return_value="Linux",
        ):
            clear_omml_input_caches()
            result = mixed_text_with_structured_math(node)
            clear_omml_input_caches()

    assert result.formula_count == 1
    assert result.structured_formula_count == 1
    assert result.degraded_formula_count == 0
    assert "MATHML" in result.text
    assert "alpha" in result.text
