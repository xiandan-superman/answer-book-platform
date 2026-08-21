from __future__ import annotations

import json
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from app.document_contracts import (
    HEADER_FOOTER_CONTRACT,
    NUMBERING_CONTRACT,
    PAGE_CONTRACT,
    QUESTION_STRUCTURE_CONTRACT,
    TEXT_CONTRACT,
)
from app.document_presentation import question_unit_rows
from app.docx_audit import audit_docx_v4
from app.docx_v4 import build_docx_from_fragments


def _cm(value) -> float:
    return value.cm


def test_document_contract_keeps_existing_question_structures_and_numbering() -> None:
    assert QUESTION_STRUCTURE_CONTRACT["名词解释"] == ("题号+教材依据", "答案标题", "缩进答案正文")
    assert QUESTION_STRUCTURE_CONTRACT["作图题"][-3:] == ("图示", "解析", "易错点及注意事项")
    assert QUESTION_STRUCTURE_CONTRACT["计算题"][0] == "题号+教材依据"
    assert NUMBERING_CONTRACT == {
        "question": "{number}、",
        "subquestion": "({number})",
        "requirement_1_to_20": "①、…⑳、",
    }
    rows = question_unit_rows(
        {
            "subquestions": [
                {"number": "1", "stem": "第一问"},
                {
                    "number": "2",
                    "stem": "第二问",
                    "requirements": [
                        {"number": "2.1", "stem": "第一项"},
                        {"number": "2.2", "stem": "第二项"},
                    ],
                },
            ]
        }
    )
    assert [row["heading"] for row in rows] == ["(1)第一问", "①、第一项", "②、第二项"]


def test_synthetic_empty_parent_is_flattened_to_first_level_subquestions() -> None:
    rows = question_unit_rows(
        {
            "number": "6",
            "subquestions": [
                {
                    "number": "6",
                    "stem": "写出三相反应；同时画出扩散曲线。",
                    "raw": "写出三相反应；同时画出扩散曲线。",
                    "synthetic_parent": True,
                    "requirements": [
                        {"number": "6.1", "stem": "写出三相反应"},
                        {"number": "6.2", "stem": "画出扩散曲线"},
                    ],
                }
            ],
        }
    )

    assert [row["number"] for row in rows] == ["6.1", "6.2"]
    assert [row["heading"] for row in rows] == ["(1)写出三相反应", "(2)画出扩散曲线"]
    assert all(not row["parent_number"] for row in rows)


def test_generated_docx_obeys_page_font_spacing_and_answer_indent_contract(tmp_path) -> None:
    payload = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "short_1",
                "section": "一、简答题",
                "question_type": "简答题",
                "number": "1",
                "answer": "答案正文。",
                "answer_summary": "答案正文。",
                "formulas": [],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材内容。"}]},
                    {"label": "解析", "segments": [{"type": "text", "text": "解析正文。"}]},
                    {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "注意事项。"}]},
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    doc = Document(output)
    section = doc.sections[0]
    assert _cm(section.page_width) == pytest.approx(PAGE_CONTRACT.width_cm, abs=0.02)
    assert _cm(section.page_height) == pytest.approx(PAGE_CONTRACT.height_cm, abs=0.02)
    assert all(
        _cm(getattr(section, side)) == pytest.approx(PAGE_CONTRACT.margin_cm, abs=0.02)
        for side in ("top_margin", "bottom_margin", "left_margin", "right_margin")
    )

    title = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "真题答案解析")
    assert title.runs[0].font.size.pt == TEXT_CONTRACT.title_size_pt
    body = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "答案正文。")
    assert body.paragraph_format.first_line_indent.cm == pytest.approx(
        TEXT_CONTRACT.answer_first_line_indent_cm, abs=0.02
    )
    analysis = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("解析："))
    assert analysis.paragraph_format.line_spacing == TEXT_CONTRACT.line_spacing
    note = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("易错点及注意事项："))
    assert note.paragraph_format.line_spacing == TEXT_CONTRACT.note_line_spacing

    body_run = body.runs[0]
    assert body_run.font.name == TEXT_CONTRACT.latin_font
    assert body_run._element.rPr.rFonts.get(qn("w:eastAsia")) == TEXT_CONTRACT.east_asia_font
    assert body_run.font.size.pt == TEXT_CONTRACT.body_size_pt
    header_run = section.header.paragraphs[0].runs[0]
    footer_run = section.footer.paragraphs[0].runs[0]
    assert header_run._element.rPr.rFonts.get(qn("w:eastAsia")) == HEADER_FOOTER_CONTRACT.header_font
    assert header_run.font.size.pt == HEADER_FOOTER_CONTRACT.header_size_pt
    assert header_run.bold is True
    assert footer_run._element.rPr.rFonts.get(qn("w:eastAsia")) == HEADER_FOOTER_CONTRACT.footer_font
    assert footer_run.font.size.pt == HEADER_FOOTER_CONTRACT.footer_size_pt
    with zipfile.ZipFile(output) as archive:
        font_table = etree.fromstring(archive.read("word/fontTable.xml"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    alternate_fonts = {
        node.get(qn("w:name")): node.find("w:altName", namespaces=namespace).get(qn("w:val"))
        for node in font_table.findall("w:font", namespaces=namespace)
        if node.find("w:altName", namespaces=namespace) is not None
    }
    assert alternate_fonts[TEXT_CONTRACT.east_asia_font] == TEXT_CONTRACT.east_asia_fallback_font
    assert alternate_fonts[HEADER_FOOTER_CONTRACT.header_font] == HEADER_FOOTER_CONTRACT.header_fallback_font
    assert alternate_fonts[HEADER_FOOTER_CONTRACT.footer_font] == HEADER_FOOTER_CONTRACT.footer_fallback_font
    assert audit_docx_v4(output) == []


def test_student_docx_hides_unconfirmed_evidence_process_diagnostics(tmp_path) -> None:
    payload = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "short_1",
                "section": "一、简答题",
                "question_type": "简答题",
                "number": "1",
                "answer": "答案正文。",
                "answer_summary": "答案正文。",
                "formulas": [],
                "blocks": [
                    {
                        "label": "教材依据",
                        "segments": [
                            {"type": "text", "text": "已确认知识点：课本-p10"},
                            {"type": "text", "text": "；"},
                            {"type": "text", "text": "未确认知识点：未确认到可用教材依据", "highlight": "unconfirmed_evidence"},
                        ],
                    }
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "已确认知识点：课本-p10" in text
    assert "未确认到可用教材依据" not in text


def test_docx_audit_blocks_silent_page_or_font_contract_drift(tmp_path) -> None:
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    damaged = tmp_path / "damaged.docx"
    source.write_text(
        json.dumps(
            {
                "schema_version": "answer_book.answer_fragments.v4",
                "fragments": [
                    {
                        "question_id": "q1",
                        "section": "一、判断题",
                        "question_type": "判断题",
                        "number": "1",
                        "answer": "正确",
                        "answer_summary": "",
                        "formulas": [],
                        "blocks": [],
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    build_docx_from_fragments(source, output)
    with zipfile.ZipFile(output) as archive:
        members = {name: archive.read(name) for name in archive.namelist()}
    root = etree.fromstring(members["word/document.xml"])
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root.xpath("//w:pgSz", namespaces=namespace)[-1].set(qn("w:w"), "9000")
    first_body_run = root.xpath("//w:body/w:p/w:r[w:t]", namespaces=namespace)[0]
    first_body_run.find("w:rPr/w:rFonts", namespaces=namespace).set(qn("w:eastAsia"), "Arial")
    members["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    with zipfile.ZipFile(damaged, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in members.items():
            archive.writestr(name, payload)

    issues = audit_docx_v4(damaged)

    assert "document contract page width mismatch" in issues
    assert any("document contract font mismatch" in issue for issue in issues)


def test_formal_answer_renders_ascii_chemical_indices_as_omml(tmp_path) -> None:
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(
        json.dumps(
            {
                "schema_version": "answer_book.answer_fragments.v4",
                "fragments": [
                    {
                        "question_id": "q1",
                        "section": "一、简答题",
                        "question_type": "简答题",
                        "number": "1",
                        "answer": "形成Fe2Ti、CuAl2和H2O。",
                        "answer_summary": "",
                        "formulas": [],
                        "blocks": [],
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    build_docx_from_fragments(source, output)

    with zipfile.ZipFile(output) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    namespace = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    normal_text = "".join(root.xpath("//w:t/text()", namespaces=namespace))
    math_text = ["".join(node.xpath(".//m:t/text()", namespaces=namespace)) for node in root.xpath("//m:oMath", namespaces=namespace)]
    assert "Fe2Ti" not in normal_text
    assert "CuAl2" not in normal_text
    assert {"Fe2Ti", "CuAl2", "H2O"}.issubset(set(math_text))
    assert audit_docx_v4(output) == []


def test_docx_audit_blocks_internal_review_language_in_formal_delivery(tmp_path) -> None:
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(
        json.dumps(
            {
                "schema_version": "answer_book.answer_fragments.v4",
                "fragments": [
                    {
                        "question_id": "q1",
                        "section": "一、简答题",
                        "question_type": "简答题",
                        "number": "1",
                        "answer": "该处需人工复核。",
                        "answer_summary": "",
                        "formulas": [],
                        "blocks": [],
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    build_docx_from_fragments(source, output)

    assert any("contains internal review language" in issue for issue in audit_docx_v4(output))


def test_formal_answer_keeps_part_numbers_without_repeating_source_questions(tmp_path) -> None:
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    repeated_stem = "指出金属液相结晶时的热力学条件，它一定需要过冷吗？"
    payload = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "question_id": "q1",
                "section": "二、简答题",
                "question_type": "简答题",
                "number": "1",
                "answer": "见解析",
                "answer_summary": "需要过冷。",
                "subquestions": [
                    {"number": "1", "stem": repeated_stem, "question_type": "简答题"},
                    {"number": "2", "stem": "说明原因。", "question_type": "简答题"},
                ],
                "answer_units": [
                    {"number": "1", "answer": "需要过冷。"},
                    {"number": "2", "answer": "自由能降低。"},
                ],
                "formulas": [],
                "blocks": [
                    {
                        "label": "解析",
                        "segments": [
                            {"type": "text", "text": f"(1){repeated_stem}\n"},
                            {"type": "text", "text": "答案：需要过冷。"},
                            {"type": "text", "text": "(2)说明原因。\n"},
                            {"type": "text", "text": "答案：自由能降低。"},
                        ],
                    }
                ],
            }
        ],
    }
    source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert repeated_stem not in text
    assert "说明原因" not in text
    assert "(1)" in text and "(2)" in text
    assert "答案：需要过冷" in text


def test_formal_heading_marker_does_not_treat_prose_list_comma_as_numbering(tmp_path) -> None:
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    stem = "如果 A、B 原子随机占据所有位置，对应哪种点阵？"
    source.write_text(
        json.dumps(
            {
                "fragments": [
                    {
                        "question_id": "q_ab",
                        "section": "简答题",
                        "question_type": "简答题",
                        "number": "1",
                        "answer": "简单立方点阵。",
                        "answer_summary": "简单立方点阵。",
                        "subquestions": [{"number": "1", "stem": stem, "raw": f"(1){stem}"}],
                        "answer_units": [{"number": "1", "answer": "简单立方点阵。"}],
                        "formulas": [],
                        "blocks": [
                            {
                                "label": "解析",
                                "segments": [
                                    {"type": "text", "text": f"(1){stem}\n"},
                                    {"type": "text", "text": "答案：简单立方点阵。"},
                                ],
                            }
                        ],
                    }
                ]
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    build_docx_from_fragments(source, output)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert stem not in text
    assert "(1)如果 A、" not in text
    assert "(1)答案：简单立方点阵" in text


def test_mixed_short_answer_and_calculation_units_keep_source_order(tmp_path) -> None:
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(
        json.dumps(
            {
                "fragments": [
                    {
                        "question_id": "mixed_order",
                        "section": "简答题",
                        "question_type": "简答题",
                        "number": "1",
                        "answer": "见解析",
                        "answer_summary": "(1)甲；(3)丙。",
                        "subquestions": [
                            {"number": "1", "stem": "说明甲"},
                            {"number": "2", "stem": "计算乙", "question_type": "计算题"},
                            {"number": "3", "stem": "说明丙"},
                        ],
                        "answer_units": [
                            {"number": "1", "answer": "甲"},
                            {"number": "2", "answer": "乙"},
                            {"number": "3", "answer": "丙"},
                        ],
                        "formulas": [],
                        "blocks": [
                            {
                                "label": "解析",
                                "segments": [
                                    {"type": "text", "text": "(1)说明甲"},
                                    {"type": "text", "text": "甲的解析"},
                                    {"type": "text", "text": "(3)说明丙"},
                                    {"type": "text", "text": "丙的解析"},
                                ],
                            },
                            {
                                "label": "解题步骤",
                                "segments": [
                                    {"type": "text", "text": "(2)计算乙"},
                                    {"type": "text", "text": "乙的计算"},
                                ],
                            },
                        ],
                    }
                ]
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    build_docx_from_fragments(source, output)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    analysis = text[text.index("解析：") :]
    assert analysis.index("(1)") < analysis.index("(2)") < analysis.index("(3)")
    assert "说明甲" not in analysis
    assert "计算乙" not in analysis
    assert "说明丙" not in analysis
