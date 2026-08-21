from __future__ import annotations

from docx import Document

from app.docx_v4 import add_mixed_paragraph


def test_scientific_inline_list_wraps_do_not_create_one_paragraph_per_item() -> None:
    document = Document()

    add_mixed_paragraph(
        document,
        [
            {
                "type": "text",
                "text": "晶面指数为(110)、\n(200)、\n(211)、\n(220)。",
            }
        ],
        {},
        label="解析",
    )

    non_empty = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert non_empty == ["解析：晶面指数为(110)、(200)、(211)、(220)。"]


def test_real_subquestion_break_is_preserved() -> None:
    document = Document()

    add_mixed_paragraph(
        document,
        [{"type": "text", "text": "(1)第一问。\n(2)第二问。"}],
        {},
        label="解析",
    )

    non_empty = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert non_empty == ["解析：(1)第一问。", "(2)第二问。"]
