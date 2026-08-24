from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document
from lxml import etree

from app.practice_document_contracts import (
    PRACTICE_DOCUMENT_CONTRACT_VERSION,
    PRACTICE_NUMBERING_CONTRACT,
    PRACTICE_STRUCTURE_CONTRACT,
)
from app.practice_export import (
    build_practice_question_docx,
    build_practice_solution_docx,
    resolve_practice_export_payload,
    validate_docx_output,
    validate_practice_export,
)


def _practice() -> dict:
    return {
        "source_analysis": {"subject": "材料科学", "question_type": "计算题", "difficulty": "进阶"},
        "blueprint": {"training_goal": "掌握晶格基矢与密度计算"},
        "exercises": [{
            "number": 1,
            "question_type": "计算题",
            "difficulty": "进阶",
            "target_skill": "公式应用",
            "stem": "已知基矢 $\\mathbf{a}_1$，计算 $x^2$。",
            "options": [],
            "knowledge_points": ["晶格基矢"],
        }],
    }


def _text(data: bytes) -> str:
    return "\n".join(paragraph.text for paragraph in Document(BytesIO(data)).paragraphs)


def _document_xml(data: bytes) -> str:
    with ZipFile(BytesIO(data)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _omml_texts(data: bytes) -> list[str]:
    with ZipFile(BytesIO(data)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    namespace = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    return ["".join(node.itertext()) for node in root.xpath("//m:oMath", namespaces=namespace)]


def test_practice_export_creates_question_only_document():
    questions = build_practice_question_docx(_practice())

    question_text = _text(questions)
    assert "专项练习题目卷" in question_text
    assert "参考答案" not in question_text
    assert "训练能力" not in question_text
    assert "第 1 题" in question_text
    assert "1. 已知基矢" not in question_text


def test_practice_word_contract_keeps_independent_structure_and_numbering():
    assert PRACTICE_STRUCTURE_CONTRACT["questions"] == (
        "专项练习题目卷",
        "练习题",
        "第 N 题",
        "题干",
        "A. / B. 选项",
    )
    assert PRACTICE_NUMBERING_CONTRACT["question_heading"] == "第 {number} 题"
    assert PRACTICE_NUMBERING_CONTRACT["option"] == "A. / B. / …"


def test_practice_docx_validator_reports_contract_version_and_passes_current_format():
    data = _practice()
    report = validate_docx_output(build_practice_question_docx(data), data)
    assert report["ok"] is True
    assert report["document_contract_version"] == PRACTICE_DOCUMENT_CONTRACT_VERSION


def test_practice_docx_validator_blocks_silent_style_contract_drift():
    data = _practice()
    content = build_practice_question_docx(data)
    with ZipFile(BytesIO(content)) as archive:
        members = {name: archive.read(name) for name in archive.namelist()}
    styles = etree.fromstring(members["word/styles.xml"])
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    normal_font = styles.xpath(".//w:style[@w:styleId='Normal']/w:rPr/w:rFonts", namespaces=namespace)[0]
    normal_font.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Arial")
    members["word/styles.xml"] = etree.tostring(styles, xml_declaration=True, encoding="UTF-8", standalone=True)
    damaged = BytesIO()
    with ZipFile(damaged, "w") as archive:
        for name, payload in members.items():
            archive.writestr(name, payload)

    report = validate_docx_output(damaged.getvalue(), data)

    assert report["ok"] is False
    assert "练习 Word 契约：Normal 中文字体发生变化。" in report["issues"]


def test_practice_export_converts_inline_latex_and_sets_portable_chinese_font():
    questions = build_practice_question_docx(_practice())
    xml = _document_xml(questions)
    assert "$\\mathbf{a}_1$" not in xml
    assert "m:oMath" in xml
    assert 'w:eastAsia="宋体"' in xml
    assert 'w:ascii="Cambria Math"' in xml
    with ZipFile(BytesIO(questions)) as archive:
        font_table = archive.read("word/fontTable.xml").decode("utf-8")
        settings = archive.read("word/settings.xml").decode("utf-8")
    assert "宋体" in font_table
    assert 'm:mathFont m:val="Cambria Math"' in settings
    assert "w:doNotExpandShiftReturn" not in settings


def test_practice_export_formats_choice_options_without_bold_or_blank_lines():
    data = _practice()
    data["exercises"][0]["question_type"] = "单选题"
    data["exercises"][0]["options"] = [
        {"label": "A", "text": "A. **第一项**"},
        {"label": "A", "text": "第二项；"},
    ]
    questions = build_practice_question_docx(data)
    document = Document(BytesIO(questions))
    option_paragraphs = [p for p in document.paragraphs if p.text.startswith(("A. ", "B. "))]
    assert [p.text for p in option_paragraphs] == ["A. 第一项。", "B. 第二项。"]
    assert all(not any(run.bold for run in paragraph.runs) for paragraph in option_paragraphs)
    assert all(paragraph.paragraph_format.left_indent.pt == 22 for paragraph in option_paragraphs)
    assert all(paragraph.paragraph_format.right_indent.pt == 0 for paragraph in option_paragraphs)
    assert all(paragraph.paragraph_format.first_line_indent.pt == -22 for paragraph in option_paragraphs)
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
    assert paragraph_texts.index("B. 第二项。") == paragraph_texts.index("A. 第一项。") + 1


def test_practice_export_uses_hanging_indent_for_solution_lists():
    data = _practice()
    data["exercises"][0]["answer"] = "答案正文。"
    data["exercises"][0]["solution_steps"] = ["第一步。", "第二步。"]
    document = Document(BytesIO(build_practice_solution_docx(data)))

    list_format = document.styles["List Number"].paragraph_format
    assert list_format.left_indent.pt == 22
    assert list_format.right_indent.pt == 0
    assert list_format.first_line_indent.pt == -22


def test_practice_solution_restarts_numbered_steps_for_each_question():
    data = _practice()
    data["exercises"][0]["answer"] = "第一题答案。"
    data["exercises"][0]["solution_steps"] = ["第一题步骤一。", "第一题步骤二。"]
    second = {**data["exercises"][0]}
    second["number"] = 2
    second["answer"] = "第二题答案。"
    second["solution_steps"] = ["第二题步骤一。", "第二题步骤二。"]
    data["exercises"].append(second)
    content = build_practice_solution_docx(data)
    report = validate_docx_output(content, data)
    assert report["ok"] is True
    with ZipFile(BytesIO(content)) as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
        numbering = etree.fromstring(archive.read("word/numbering.xml"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    list_ids = document.xpath(
        ".//w:p[w:pPr/w:pStyle[@w:val='ListNumber']]/w:pPr/w:numPr/w:numId/@w:val",
        namespaces=namespace,
    )
    assert list_ids[0] == list_ids[1]
    assert list_ids[2] == list_ids[3]
    assert list_ids[0] != list_ids[2]
    restarting_ids = numbering.xpath(
        ".//w:num[w:lvlOverride[@w:ilvl='0']/w:startOverride[@w:val='1']]/@w:numId",
        namespaces=namespace,
    )
    assert set(list_ids) <= set(restarting_ids)


def test_practice_export_converts_display_math_and_table_math():
    data = _practice()
    data["exercises"][0]["stem"] = r"计算 \[x^2=4\]。"
    data["exercises"][0]["tables"] = [{
        "location": "stem",
        "headers": [r"$T$", "数值"],
        "rows": [["温度", r"$263.15\,\mathrm{K}$"]],
    }]
    xml = _document_xml(build_practice_question_docx(data))
    assert "\\[" not in xml
    assert xml.count("<m:oMath") >= 3


def test_practice_export_converts_obvious_bare_equation_to_office_math():
    data = _practice()
    data["exercises"][0]["stem"] = "已知 ΔH = Q_p，判断该关系的适用条件。"
    xml = _document_xml(build_practice_question_docx(data))
    assert "ΔH = Q_p" not in xml
    assert "<m:oMath" in xml


def test_practice_export_attaches_formula_only_superscript_to_unit_base():
    data = _practice()
    data["exercises"][0]["stem"] = r"单位为 mol$^{-1}$ 和 cm$^3$。"
    xml = _document_xml(build_practice_question_docx(data))
    assert "<m:t>mol</m:t>" in xml
    assert "<m:t>cm</m:t>" in xml
    assert "<m:e><m:r" in xml


def test_practice_export_converts_chemical_and_electrode_notation_in_options():
    data = _practice()
    data["exercises"][0]["question_type"] = "单选题"
    data["exercises"][0]["options"] = [{
        "text": "铜电极（Cu²⁺|Cu）、氢电极（Pt|H₂|H⁺）、甘汞电极（Hg₂Cl₂(s)|Hg(l)|KCl(aq)）",
    }]
    content = build_practice_question_docx(data)
    formulas = _omml_texts(content)
    xml = _document_xml(content)
    assert "Cu2+|Cu" in formulas
    assert "Pt|H2|H+" in formulas
    assert "Hg2Cl2(s)|Hg(l)|KCl(aq)" in formulas
    assert "Cu²⁺|Cu" not in xml
    assert "Hg₂Cl₂(s)|Hg(l)|KCl(aq)" not in xml


def test_practice_export_converts_complete_bare_reaction_to_one_office_math_object():
    data = _practice()
    data["exercises"][0]["stem"] = "相变过程为 L+δ→γ，请说明各相的含义。"
    formulas = _omml_texts(build_practice_question_docx(data))

    assert "L+δ→γ" in formulas


def test_practice_export_preserves_shared_chemical_typography_decision():
    data = _practice()
    data["exercises"][0]["stem"] = (
        "反应式为 2H₂+O₂→2H₂O，参比电极为（Hg₂Cl₂(s)|Hg(l)|KCl(aq)）。"
    )
    content = build_practice_question_docx(data)
    with ZipFile(BytesIO(content)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    namespace = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    objects = root.xpath("//m:oMath", namespaces=namespace)
    targets = [
        node
        for node in objects
        if "2H2+O2→2H2O" in "".join(node.itertext())
        or "Hg2Cl2(s)|Hg(l)|KCl(aq)" in "".join(node.itertext())
    ]

    assert len(targets) == 2
    for target in targets:
        runs = target.xpath(".//m:r", namespaces=namespace)
        upright_runs = target.xpath(".//m:r[m:rPr/m:sty[@m:val='p']]", namespaces=namespace)
        assert len(upright_runs) == len(runs)


def test_practice_export_keeps_complete_electrode_equation_in_one_math_object():
    data = _practice()
    data["exercises"][0]["stem"] = (
        "甘汞电极（Hg₂Cl₂(s)|Hg(l)|KCl(aq)）的标准电极电势 "
        "E°(Hg₂Cl₂/Hg)=0.268 V。"
    )
    formulas = _omml_texts(build_practice_question_docx(data))
    assert "Eθ(Hg2Cl2/Hg)=0.268V" in formulas
    assert "Hg)=0" not in formulas


def test_practice_export_normalizes_standard_state_o_and_degree_to_theta():
    data = _practice()
    data["exercises"][0]["stem"] = (
        r"标准平衡常数 $K^o=\exp(-\Delta G^\circ/RT)$，且 ΔH°、E^{O} 均指标准态。"
    )
    formulas = _omml_texts(build_practice_question_docx(data))
    joined = "|".join(formulas)
    assert "Kθ" in joined
    assert "ΔGθ" in joined
    assert "ΔHθ" in joined
    assert "Eθ" in joined
    assert "Kᵒ" not in joined


def test_practice_export_uses_real_paragraphs_for_multiline_question_text():
    data = _practice()
    data["exercises"][0]["stem"] = "第 1 题\n请完成下列\n各问：\n\n（1）写出基本关系式。\n\n2. 说明适用条件。"
    content = build_practice_question_docx(data)
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "<w:br" not in _document_xml(content)
    assert "第 1 题" in paragraphs
    assert any("请完成下列 各问" in text for text in paragraphs)
    assert "(1) 写出基本关系式。" in paragraphs
    assert "(2) 说明适用条件。" in paragraphs


def test_practice_export_repairs_single_escaped_latex_theta():
    data = _practice()
    data["exercises"][0]["stem"] = "已知标准电极电势 $E^" + "\t" + "heta$，计算反应的标准态函数。"
    formulas = "|".join(_omml_texts(build_practice_question_docx(data)))

    assert "Eθ" in formulas
    assert "heta" not in formulas


def test_practice_export_preserves_formula_and_table_specific_layouts():
    data = _practice()
    data["exercises"][0]["formulas"] = [{"location": "stem", "latex": r"E = E^\theta - \frac{RT}{nF}\ln Q"}]
    data["exercises"][0]["tables"] = [{"location": "stem", "headers": ["物理量", "数值"], "rows": [["温度", "298.15 K"]]}]
    document = Document(BytesIO(build_practice_question_docx(data)))

    formula_paragraph = next(
        paragraph for paragraph in document.paragraphs
        if "oMath" in paragraph._p.xml and paragraph.alignment == 1
    )
    assert formula_paragraph.alignment == 1
    for cell in document.tables[0].rows[0].cells:
        paragraph = cell.paragraphs[0]
        assert paragraph.paragraph_format.first_line_indent.pt == 0
        assert paragraph.alignment == 0


def test_practice_export_strips_delimiters_from_structured_formula_fields():
    data = _practice()
    data["exercises"][0]["formulas"] = [
        {
            "location": "stem",
            "latex": r"$(\frac{\partial U}{\partial V})_T = 0 \quad (\text{Ideal Gas})$",
        }
    ]

    content = build_practice_question_docx(data)
    xml = _document_xml(content)
    formulas = _omml_texts(content)

    assert "$" not in xml
    normalized_formulas = [formula.replace("\xa0", " ") for formula in formulas]
    assert any("Ideal Gas" in formula and "∂U" in formula for formula in normalized_formulas)
    assert validate_docx_output(content, data)["office_math_markup_leak_count"] == 0


def test_practice_docx_validator_rejects_markup_inside_office_math_objects():
    data = _practice()
    content = build_practice_question_docx(data)
    with ZipFile(BytesIO(content)) as archive:
        members = {name: archive.read(name) for name in archive.namelist()}
    document = etree.fromstring(members["word/document.xml"])
    namespace = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    math_text = document.xpath(".//m:oMath//m:t", namespaces=namespace)[0]
    math_text.text = f"${math_text.text}$"
    members["word/document.xml"] = etree.tostring(
        document,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )
    damaged = BytesIO()
    with ZipFile(damaged, "w") as archive:
        for name, payload in members.items():
            archive.writestr(name, payload)

    report = validate_docx_output(damaged.getvalue(), data)

    assert report["ok"] is False
    assert report["office_math_markup_leak_count"] == 1
    assert "DOCX Office 公式对象中仍含提供方的 LaTeX/Markdown 定界标记。" in report["issues"]


def test_practice_export_renders_node_edge_diagram_as_media():
    data = _practice()
    data["exercises"][0]["stem"] = "根据图示说明入口与出口的关系。"
    data["exercises"][0]["figures"] = [{
        "figure_id": "g1",
        "location": "stem",
        "figure_type": "diagram",
        "nodes": [
            {"id": "a", "label": "入口", "x": 0.15, "y": 0.5, "shape": "box"},
            {"id": "b", "label": "出口", "x": 0.85, "y": 0.5, "shape": "box"},
        ],
        "edges": [{"from": "a", "to": "b", "label": "流向", "directed": True}],
    }]
    content = build_practice_question_docx(data)
    report = validate_docx_output(content, data)
    assert report["ok"]
    assert report["media_count"] >= 1


def test_practice_export_renders_unlabelled_wireframe_vertices_and_shaded_plane():
    data = _practice()
    data["exercises"][0]["stem"] = "根据图示判断晶向与晶面。"
    data["exercises"][0]["figures"] = [{
        "figure_id": "g1",
        "location": "stem",
        "figure_type": "diagram",
        "nodes": [
            {"id": "o", "label": "O", "x": 0.2, "y": 0.5, "shape": "circle"},
            {"id": "corner", "label": "", "x": 0.4, "y": 0.8, "shape": "circle"},
            {"id": "a", "label": "A", "x": 0.5, "y": 0.7, "shape": "box"},
            {"id": "b", "label": "B", "x": 0.7, "y": 0.7, "shape": "box"},
            {"id": "c", "label": "C", "x": 0.7, "y": 0.3, "shape": "box"},
            {"id": "d", "label": "D", "x": 0.5, "y": 0.3, "shape": "box"},
        ],
        "edges": [
            {"from": "o", "to": "corner", "label": "晶向箭头", "directed": True},
            {"from": "a", "to": "b", "label": "阴影晶面边界", "directed": False},
            {"from": "b", "to": "c", "label": "平行 z 轴", "directed": False},
            {"from": "c", "to": "d", "label": "阴影晶面边界", "directed": False},
            {"from": "d", "to": "a", "label": "平行 z 轴", "directed": False},
        ],
    }]

    content = build_practice_question_docx(data)
    report = validate_docx_output(content, data)

    assert report["ok"]
    assert report["media_count"] >= 1


def test_practice_figure_title_stays_with_following_image():
    data = _practice()
    data["exercises"][0]["figures"] = [{
        "figure_id": "g1",
        "location": "stem",
        "figure_type": "line",
        "title": "应力-应变曲线",
        "series": [{"name": "曲线", "points": [[0, 0], [1, 1]]}],
    }]

    document = Document(BytesIO(build_practice_question_docx(data)))
    title = next(paragraph for paragraph in document.paragraphs if paragraph.text == "应力-应变曲线")

    assert title.paragraph_format.keep_with_next is True


def test_practice_formula_caption_stays_with_formula():
    data = _practice()
    data["exercises"][0]["formulas"] = [{
        "formula_id": "f1",
        "location": "stem",
        "latex": "x=1",
        "caption": "已知关系",
        "display": True,
    }]

    document = Document(BytesIO(build_practice_question_docx(data)))
    caption = next(paragraph for paragraph in document.paragraphs if paragraph.text == "已知关系")

    assert caption.paragraph_format.keep_with_next is True


def test_practice_export_marks_unreviewed_edit_as_candidate_without_blocking_download():
    data = _practice()
    data["semantic_review"] = {
        "status": "failed",
        "review_scope": "stale_after_edit",
        "items": [{"number": 1, "status": "not_reviewed", "risks": []}],
    }

    report = validate_practice_export(data)

    assert report["ok"] is True
    assert report["release_level"] == "review_candidate"
    assert "不应视为正式发布版" in report["warning_issues"][0]


def test_practice_export_preflights_inline_math_before_building_word():
    data = _practice()
    data["exercises"][0]["stem"] = r"计算 $\left(x$。"

    report = validate_practice_export(data)

    assert report["ok"] is False
    assert report["release_level"] == "blocked"
    assert any("行内公式" in issue for issue in report["blocking_issues"])


def test_selected_export_does_not_inherit_another_questions_review_candidate_status():
    data = _practice()
    data["exercises"].append({
        **data["exercises"][0],
        "number": 2,
        "plan_item_id": "plan_item_02",
        "stem": "用户刚刚修改、尚未复核的第二题。",
    })
    data["exercises"][0]["plan_item_id"] = "plan_item_01"
    data["semantic_review"] = {
        "status": "failed",
        "items": [
            {"number": 1, "status": "passed", "risks": []},
            {"number": 2, "status": "not_reviewed", "risks": []},
        ],
    }

    selected = resolve_practice_export_payload(
        {"export_scope": "selected", "selected_exercise_ids": ["plan_item_01"]},
        data,
    )
    report = validate_practice_export(selected)

    assert len(selected["exercises"]) == 1
    assert report["ok"] is True
    assert report["release_level"] == "formal"


def test_export_resolves_stale_browser_payload_against_latest_saved_questions():
    stale = _practice()
    stale["exercises"][0]["plan_item_id"] = "plan_item_01"
    stale["exercises"][0]["stem"] = "旧页面中的题干"
    latest = _practice()
    latest["exercises"][0]["plan_item_id"] = "plan_item_01"
    latest["exercises"][0]["stem"] = "服务器已经保存的新题干"

    resolved_all = resolve_practice_export_payload(
        {**stale, "history_id": "practice_example1234", "export_scope": "all"},
        latest,
    )
    resolved_selected = resolve_practice_export_payload(
        {
            **stale,
            "history_id": "practice_example1234",
            "export_scope": "selected",
            "selected_exercise_ids": ["plan_item_01"],
        },
        latest,
    )

    assert resolved_all["exercises"][0]["stem"] == "服务器已经保存的新题干"
    assert resolved_selected["exercises"][0]["stem"] == "服务器已经保存的新题干"
