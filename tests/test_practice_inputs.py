import base64
from io import BytesIO
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

from docx import Document
from lxml import etree
from PIL import Image

from app.practice_inputs import _reference_image_data_url, parse_practice_sources


PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _payload(path: Path) -> dict:
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return {"source_files": [{"name": path.name, "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "data_url": "data:application/octet-stream;base64," + encoded}]}


class PracticeInputTests(unittest.TestCase):
    def test_large_embedded_image_is_losslessly_compacted_for_transport(self):
        source = Image.new("RGB", (320, 240))
        source.putdata([
            ((x * 7) % 256, (y * 11) % 256, ((x + y) * 5) % 256)
            for y in range(240) for x in range(320)
        ])
        raw = BytesIO()
        source.save(raw, format="PNG", compress_level=0)

        compact_url = _reference_image_data_url(raw.getvalue(), "image/png")
        header, encoded = compact_url.split(",", 1)
        compact = base64.b64decode(encoded)

        self.assertIn("image/webp", header)
        self.assertLess(len(compact), len(raw.getvalue()))
        with Image.open(BytesIO(compact)) as restored:
            self.assertEqual(list(source.getdata()), list(restored.convert("RGB").getdata()))

    def test_explicit_docx_media_directory_entry_is_not_treated_as_an_image(self):
        from app.practice_inputs import _docx_content

        with tempfile.TemporaryDirectory() as raw:
            path = self.make_docx(Path(raw), formulas=False, images=1, table_formula=False)
            with ZipFile(path, "a") as archive:
                archive.writestr("word/media/", b"")

            _text, images, diagnostics = _docx_content(path.name, path.read_bytes())

        self.assertEqual(1, diagnostics["embedded_image_count"])
        self.assertEqual(1, len(diagnostics["embedded_image_order"]))
        self.assertEqual(1, len(images))

    def make_docx(self, root: Path, *, formulas: bool, images: int, table_formula: bool = False, nested_formula: bool = False) -> Path:
        path = root / ("fixture_formula.docx" if formulas else "fixture_plain.docx")
        doc = Document()
        doc.add_paragraph("题干前")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格前"
        for index in range(images):
            image = root / f"image_{index}.png"
            image.write_bytes(PNG + bytes([index]))
            doc.add_picture(str(image))
        doc.save(path)
        if formulas:
            with ZipFile(path, "r") as source:
                members = {name: source.read(name) for name in source.namelist()}
            xml = members["word/document.xml"].decode("utf-8")
            formula = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>0.25</m:t></m:r><m:r><m:t>%</m:t></m:r></m:oMath>'
            xml = xml.replace("<w:body>", '<w:body><w:p><w:r><w:t>题干中</w:t></w:r>' + formula + '</w:p>', 1)
            if table_formula:
                table_formula_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>x</m:t></m:r></m:oMath>'
                xml = xml.replace("<w:t>表格前</w:t>", "<w:t>表格前</w:t>" + table_formula_xml, 1)
            if nested_formula:
                nested = '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t>y</m:t></m:r></m:oMath></m:oMathPara>'
                xml = xml.replace("</w:body>", nested + "</w:body>", 1)
            members["word/document.xml"] = xml.encode("utf-8")
            with ZipFile(path.with_suffix(".tmp.docx"), "w") as target:
                for name, data in members.items():
                    target.writestr(name, data)
            path.with_suffix(".tmp.docx").replace(path)
        return path

    def test_omml_order_table_formula_and_image_warning(self):
        with tempfile.TemporaryDirectory() as raw:
            path = self.make_docx(Path(raw), formulas=True, images=9, table_formula=True)
            result = parse_practice_sources(_payload(path))
            diagnostics = result["file_diagnostics"][0]
            self.assertIn("题干前", result["text"])
            self.assertIn("0.25%", result["text"])
            self.assertIn("表格前x", result["text"])
            self.assertEqual(diagnostics["omml_formula_count"], 2)
            self.assertEqual(diagnostics["table_count"], 1)
            self.assertEqual(diagnostics["image_count_included"], 8)
            self.assertEqual(diagnostics["reference_image_count_included"], 9)
            self.assertEqual(diagnostics["image_anchor_count_included"], 9)
            self.assertEqual(len(diagnostics["embedded_image_order"]), 9)
            self.assertEqual(len(diagnostics["reference_image_order"]), 9)
            self.assertEqual(result["text"].count("⟦IMAGE_REF:"), 9)
            self.assertEqual(result["reference_image_count"], 9)
            self.assertTrue(diagnostics["warnings"])

    def test_plain_docx_has_no_formula_warning(self):
        with tempfile.TemporaryDirectory() as raw:
            path = self.make_docx(Path(raw), formulas=False, images=0)
            diagnostics = parse_practice_sources(_payload(path))["file_diagnostics"][0]
            self.assertEqual(diagnostics["omml_formula_count"], 0)
            self.assertFalse(diagnostics["warnings"])

    def test_docx_with_substantial_native_text_does_not_force_vision(self):
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            path = root / "text_first.docx"
            image = root / "decorative.png"
            image.write_bytes(PNG)
            doc = Document()
            doc.add_paragraph("这是一道拥有完整题干的文档题目，已经包含足够多的可提取文字，应当优先使用文字模型进行分析。")
            doc.add_picture(str(image))
            doc.save(path)

            result = parse_practice_sources(_payload(path))

            self.assertEqual(result["analysis_mode"], "text")
            self.assertEqual(result["images"], [])
            self.assertEqual(result["reference_image_count"], 1)
            self.assertEqual(result["text"].count("⟦IMAGE_REF:"), 1)
            self.assertEqual(result["file_diagnostics"][0]["image_count_included"], 0)
            self.assertEqual(result["file_diagnostics"][0]["reference_image_count_included"], 1)

    def test_nested_omml_para_counts_once(self):
        root = etree.fromstring(
            '<w:body xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            '<m:oMathPara><m:oMath><m:r><m:t>y</m:t></m:r></m:oMath></m:oMathPara></w:body>'
        )
        count = len(root.xpath('.//m:oMathPara|.//m:oMath[not(ancestor::m:oMathPara)]', namespaces={"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}))
        self.assertEqual(count, 1)

    def test_global_image_budget_is_reconciled_across_files(self):
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            first = self.make_docx(root, formulas=False, images=5)
            second = root / "second.docx"
            first.rename(second)
            first = self.make_docx(root, formulas=False, images=5)
            payload = _payload(first)
            payload["source_files"].append(_payload(second)["source_files"][0])
            result = parse_practice_sources(payload)
            counts = [item["image_count_included"] for item in result["file_diagnostics"]]
            self.assertEqual(counts, [5, 3])
            reference_counts = [item["reference_image_count_included"] for item in result["file_diagnostics"]]
            anchor_counts = [item["image_anchor_count_included"] for item in result["file_diagnostics"]]
            self.assertEqual(reference_counts, [5, 5])
            self.assertEqual(anchor_counts, [5, 5])
            self.assertEqual(result["text"].count("⟦IMAGE_REF:"), 10)
            self.assertIn("⟦IMAGE_REF:6;", result["text"])
            self.assertEqual(result["reference_image_count"], 10)
            self.assertTrue(result["file_diagnostics"][1]["warnings"])

    def test_corrupt_docx_is_rejected(self):
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "broken.docx"
            path.write_bytes(b"not a zip")
            with self.assertRaises(Exception):
                parse_practice_sources(_payload(path))

    def test_pdf_diagnostics_report_used_and_omitted_pages(self):
        encoded = base64.b64encode(b"pdf-fixture").decode("ascii")
        payload = {"source_files": [{"name": "scan.pdf", "type": "application/pdf", "data_url": f"data:application/pdf;base64,{encoded}"}]}
        diagnostics = {
            "format": "pdf",
            "page_count_total": 12,
            "page_numbers_used": list(range(1, 9)),
            "page_numbers_omitted": list(range(9, 13)),
            "warnings": ["PDF 共 12 页，仅向模型传递前 8 页图像。"],
        }
        with patch("app.practice_inputs._pdf_content", return_value=("", ["data:image/jpeg;base64,AA=="] * 8, diagnostics)):
            result = parse_practice_sources(payload)
        row = result["file_diagnostics"][0]
        self.assertEqual(row["image_count_included"], 8)
        self.assertEqual(row["page_numbers_used"], list(range(1, 9)))
        self.assertEqual(row["page_numbers_omitted"], list(range(9, 13)))


if __name__ == "__main__":
    unittest.main()
