from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.shared import Cm
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))


class FigureRuntimeObservabilityTests(unittest.TestCase):
    def test_local_image_integrity_rejects_blank_and_accepts_visible_figure(self) -> None:
        from app.figures import audit_figure_image_integrity

        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            blank = root / "blank.png"
            visible = root / "visible.png"
            Image.new("RGB", (400, 300), "white").save(blank)
            rendered = Image.new("RGB", (400, 300), "white")
            for x in range(30, 370):
                rendered.putpixel((x, 150), (0, 0, 0))
                rendered.putpixel((x, 151), (0, 0, 0))
            rendered.save(visible)

            self.assertIn("figure_image_blank_or_nearly_uniform", audit_figure_image_integrity(blank))
            self.assertEqual([], audit_figure_image_integrity(visible))

    def test_generation_removes_stale_output_when_renderer_fails_integrity(self) -> None:
        from app.figures import generate_figures

        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            output_dir = root / "figures"
            output_dir.mkdir()
            stale = output_dir / "f1.png"
            Image.new("RGB", (400, 300), "black").save(stale)
            specs = root / "specs.json"
            specs.write_text(
                json.dumps({"figures": [{"figure_id": "f1", "question_id": "q1", "kind": "custom_diagram"}]}),
                encoding="utf-8",
            )
            with patch("app.figures.draw_custom_diagram", lambda _spec, output: Image.new("RGB", (400, 300), "white").save(output)):
                generated = generate_figures(specs, output_dir)

            self.assertEqual([], generated)
            self.assertFalse(stale.exists())
            stored = json.loads(specs.read_text(encoding="utf-8"))["figures"][0]
            self.assertIn("figure_image_blank_or_nearly_uniform", stored["validation_issues"])

    def test_visual_qa_reuses_unchanged_content_by_fingerprint(self) -> None:
        from app.figures import audit_figures_with_vision
        from app.settings import ProviderConfig

        provider = ProviderConfig(
            name="vision",
            type="openai_compatible",
            base_url="http://unused",
            api_key="key",
            default_model="text",
            model_options=("text",),
            allow_custom_model=True,
            model_hint="",
            temperature=0.2,
            max_tokens=12288,
            vision_model="vision-model",
            supports_vision=True,
        )
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            figures = root / "figures"
            figures.mkdir()
            Image.new("RGB", (200, 120), "white").save(figures / "f1.png")
            specs = root / "figure_specs.json"
            specs.write_text(
                json.dumps({"figures": [{"figure_id": "f1", "question_id": "q1", "kind": "curve", "caption": "图"}]}),
                encoding="utf-8",
            )
            report_path = root / "figure_visual_qa.json"
            worker_result = {
                "items": [
                    {
                        "question_id": "q1",
                        "figure_id": "f1",
                        "path": str(figures / "f1.png"),
                        "qa": {"ok": True, "summary": "ok"},
                    }
                ]
            }
            with patch("app.figures._audit_figures_with_vision_serial", return_value=worker_result) as audit:
                first = audit_figures_with_vision(
                    {"items": [{"question_id": "q1", "stem": "画图"}]},
                    specs,
                    figures,
                    report_path,
                    provider=provider,
                    model="vision-model",
                )
                second = audit_figures_with_vision(
                    {"items": [{"question_id": "q1", "stem": "画图"}]},
                    specs,
                    figures,
                    report_path,
                    provider=provider,
                    model="vision-model",
                )

        self.assertEqual(1, audit.call_count)
        self.assertEqual(1, first["cache"]["audited_count"])
        self.assertEqual(1, second["cache"]["reused_count"])
        self.assertTrue(second["items"][0]["reused"])

    def test_visual_qa_does_not_cache_transient_reviewer_error(self) -> None:
        from app.figures import audit_figures_with_vision
        from app.settings import ProviderConfig

        provider = ProviderConfig(
            name="vision", type="openai_compatible", base_url="http://unused", api_key="key",
            default_model="text", model_options=("text",), allow_custom_model=True, model_hint="",
            temperature=0.2, max_tokens=12288, vision_model="vision-model", supports_vision=True,
        )
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            figures = root / "figures"
            figures.mkdir()
            Image.new("RGB", (200, 120), "white").save(figures / "f1.png")
            specs = root / "figure_specs.json"
            specs.write_text(
                json.dumps({"figures": [{"figure_id": "f1", "question_id": "q1", "kind": "curve"}]}),
                encoding="utf-8",
            )
            report_path = root / "figure_visual_qa.json"
            errored = {"items": [{"question_id": "q1", "figure_id": "f1", "qa": {"ok": False, "error": "timeout"}}]}
            passed = {"items": [{"question_id": "q1", "figure_id": "f1", "qa": {"ok": True}}]}
            with patch("app.figures._audit_figures_with_vision_serial", side_effect=[errored, passed]) as audit:
                audit_figures_with_vision(
                    {"items": [{"question_id": "q1", "stem": "画图"}]}, specs, figures, report_path,
                    provider=provider, model="vision-model",
                )
                second = audit_figures_with_vision(
                    {"items": [{"question_id": "q1", "stem": "画图"}]}, specs, figures, report_path,
                    provider=provider, model="vision-model",
                )

        self.assertEqual(2, audit.call_count)
        self.assertEqual(1, second["cache"]["audited_count"])
        self.assertEqual(0, second["cache"]["reused_count"])

    def test_wide_figure_is_flagged_when_embedded_too_short(self) -> None:
        from app.figure_size_audit import audit_docx_figure_sizes

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            image_path = tmp / "wide.png"
            docx_path = tmp / "answer_book.docx"
            Image.new("RGB", (1800, 600), "white").save(image_path)
            document = Document()
            document.add_paragraph().add_run().add_picture(str(image_path), width=Cm(10.5))
            document.save(docx_path)
            report = audit_docx_figure_sizes(docx_path)

        self.assertFalse(report["ok"])
        self.assertTrue(any("too small" in issue or "too short" in issue for issue in report["issues"]))

    def test_no_figure_question_skips_empty_word_figure_warning(self) -> None:
        from app.figure_size_audit import audit_docx_figure_sizes

        with tempfile.TemporaryDirectory() as raw_tmp:
            docx_path = Path(raw_tmp) / "answer_book.docx"
            Document().save(docx_path)
            report = audit_docx_figure_sizes(
                docx_path,
                structured_exam={
                    "items": [
                        {
                            "question_id": "q_text",
                            "stem": "说明二元反应扩散中不存在两相区的原因。",
                            "answer_figure_required": False,
                            "source_image_required": False,
                        }
                    ]
                },
            )

        self.assertTrue(report["ok"])
        self.assertFalse(report["applicable"])
        self.assertEqual("no_figure_required", report["skipped_reason"])
        self.assertEqual([], report["warnings"])
        self.assertEqual([], report["required_question_ids"])

    def test_required_figure_question_keeps_empty_word_figure_warning(self) -> None:
        from app.figure_size_audit import audit_docx_figure_sizes

        with tempfile.TemporaryDirectory() as raw_tmp:
            docx_path = Path(raw_tmp) / "answer_book.docx"
            Document().save(docx_path)
            report = audit_docx_figure_sizes(
                docx_path,
                structured_exam={
                    "items": [
                        {
                            "question_id": "q_draw",
                            "stem": "请画出二元相图并标出两相区。",
                            "answer_figure_required": True,
                        }
                    ]
                },
            )

        self.assertTrue(report["ok"])
        self.assertTrue(report["applicable"])
        self.assertEqual(["q_draw"], report["required_question_ids"])
        self.assertTrue(any("no inline figure images" in item for item in report["warnings"]))

    def test_wide_figure_uses_near_page_width(self) -> None:
        from app.docx_v4 import add_figure_picture
        from app.figure_size_audit import audit_docx_figure_sizes

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            image_path = tmp / "wide.png"
            docx_path = tmp / "answer_book.docx"
            Image.new("RGB", (1800, 600), "white").save(image_path)
            document = Document()
            add_figure_picture(document.add_paragraph(), image_path)
            document.save(docx_path)
            report = audit_docx_figure_sizes(docx_path)

        self.assertTrue(report["ok"], report)
        self.assertGreaterEqual(report["figures"][0]["width_cm"], 13.7)

    def test_very_wide_figure_is_padded_for_word_minimum_height(self) -> None:
        from app.docx_v4 import add_figure_picture
        from app.figure_size_audit import MIN_FIGURE_HEIGHT_CM, audit_docx_figure_sizes

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            image_path = tmp / "very_wide.png"
            docx_path = tmp / "answer_book.docx"
            Image.new("RGB", (2400, 500), "white").save(image_path)
            document = Document()
            add_figure_picture(document.add_paragraph(), image_path)
            document.save(docx_path)
            report = audit_docx_figure_sizes(docx_path)
            fitted = tmp / "very_wide_wordfit.png"
            fit_report = json.loads((tmp / "figure_word_fit.json").read_text(encoding="utf-8"))
            self.assertTrue(fitted.exists())
            self.assertTrue(report["ok"], report)
            self.assertGreaterEqual(report["figures"][0]["height_cm"], MIN_FIGURE_HEIGHT_CM)
            self.assertEqual("white_vertical_padding", fit_report["items"][0]["method"])

    def test_figure_progress_tracker_writes_heartbeat_state(self) -> None:
        from app.pipeline import FigureProgressTracker

        with tempfile.TemporaryDirectory() as raw_tmp:
            output = Path(raw_tmp) / "figure_progress.json"
            tracker = FigureProgressTracker(output, heartbeat_seconds=60)
            tracker.emit("figure_rendering", {"figure_id": "fig_01", "question_id": "q1"})
            data = __import__("json").loads(output.read_text(encoding="utf-8"))

        self.assertEqual("figures", data["stage"])
        self.assertEqual("figure_rendering", data["active_event"])
        self.assertEqual("fig_01", data["figure_id"])

    def test_figure_progress_tracker_persists_terminal_status(self) -> None:
        from app.pipeline import FigureProgressTracker

        with tempfile.TemporaryDirectory() as raw_tmp:
            output = Path(raw_tmp) / "figure_progress.json"
            tracker = FigureProgressTracker(output, heartbeat_seconds=60)
            tracker.emit("stage_completed", {"status": "passed", "generated_count": 3})
            data = __import__("json").loads(output.read_text(encoding="utf-8"))

        self.assertEqual("passed", data["status"])
        self.assertEqual(3, data["generated_count"])

    def test_figure_generation_emits_per_figure_events(self) -> None:
        from app.figures import prepare_figures_for_fragments

        structured_exam = {
            "items": [
                {"question_id": "q1", "question_type": "作图题", "drawing_generation_mode": "figure_specs", "stem": "画出[110]带轴电子衍射花样。"},
            ]
        }
        fragments = {
            "fragments": [
                {
                    "question_id": "q1",
                    "blocks": [],
                    "_draft": {
                        "figure_specs": [
                            {
                                "kind": "zone_axis_diffraction",
                                "caption": "[110]带轴电子衍射花样",
                                "zone_axis": [1, 1, 0],
                                "lattice": "generic_cubic",
                                "max_index": 2,
                                "label_indices": [[0, 0, 0], [1, -1, 0], [0, 0, 1]],
                            }
                        ]
                    },
                }
            ]
        }
        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            fragments_json = tmp / "answer_fragments.json"
            fragments_json.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")
            events = []
            prepare_figures_for_fragments(
                structured_exam,
                fragments_json,
                tmp / "figure_specs.json",
                tmp / "figures",
                progress_callback=lambda event, detail: events.append((event, detail)),
            )

        self.assertTrue(any(event == "figure_rendering" for event, _ in events))
        self.assertTrue(any(event == "figure_rendered" for event, _ in events))

    def test_visual_audit_input_downscales_image_and_omits_code(self) -> None:
        from app.figures import _compact_figure_spec_for_visual_qa, _vision_audit_image_data_url

        with tempfile.TemporaryDirectory() as raw_tmp:
            image_path = Path(raw_tmp) / "large.png"
            Image.new("RGB", (2400, 1600), "white").save(image_path)
            data_url, metadata = _vision_audit_image_data_url(image_path)

        compact = _compact_figure_spec_for_visual_qa(
            {"figure_id": "fig_01", "kind": "model_drawing_code", "code": "print('very long code')", "prompt": "large prompt", "caption": "图示"}
        )
        self.assertTrue(data_url.startswith("data:image/jpeg;base64,"))
        self.assertTrue(metadata["downscaled"])
        self.assertEqual(1536, metadata["processed_width"])
        self.assertNotIn("code", compact)
        self.assertNotIn("prompt", compact)

    def test_retry_token_budget_honors_requested_ceiling(self) -> None:
        from app.llm_client import OpenAICompatibleClient
        from app.settings import ProviderConfig

        provider = ProviderConfig(
            name="test",
            type="openai_compatible",
            base_url="https://example.test/v1",
            api_key="test",
            default_model="model-a",
            model_options=("model-a", "model-b"),
            allow_custom_model=False,
            model_hint="",
            temperature=0.1,
            max_tokens=12288,
        )
        plans = OpenAICompatibleClient(provider)._json_retry_plans([], "model-a", 8192, 1)
        self.assertTrue(plans)
        self.assertTrue(all(plan["max_tokens"] <= 8192 for plan in plans))

    def test_generic_json_retry_never_switches_model_disables_thinking_or_compacts(self) -> None:
        import inspect

        from app.llm_client import OpenAICompatibleClient
        from app.settings import ProviderConfig

        provider = ProviderConfig(
            name="test",
            type="openai_compatible",
            base_url="https://example.test/v1",
            api_key="test",
            default_model="model-a",
            model_options=("model-a", "model-b"),
            allow_custom_model=False,
            model_hint="",
            temperature=0.1,
            max_tokens=12288,
            thinking_mode="medium",
        )
        messages = [{"role": "user", "content": "full evidence"}]
        signature = inspect.signature(OpenAICompatibleClient.chat_json_object)
        self.assertNotIn("fallback_model", signature.parameters)
        self.assertNotIn("compact_messages", signature.parameters)
        plans = OpenAICompatibleClient(provider)._json_retry_plans(
            messages,
            "model-a",
            8192,
            3,
        )

        self.assertEqual(3, len(plans))
        self.assertTrue(all(plan["model"] == "model-a" for plan in plans))
        self.assertTrue(all(plan["thinking"] == "medium" for plan in plans))
        self.assertTrue(all(plan["messages"] is messages for plan in plans))
        self.assertTrue(all(plan["strategy"].startswith("same_route_attempt_") for plan in plans))

    def test_drawing_code_request_forces_disabled_thinking_with_its_own_budget(self) -> None:
        from app.drawing_code import build_drawing_code_prompt, generate_drawing_code_spec
        from app.settings import DRAWING_CODE_MAX_TOKENS

        class FakeClient:
            def __init__(self):
                self.kwargs = {}

            def chat_json_object(self, _messages, **kwargs):
                self.kwargs = kwargs
                return {"figure_id": "fig_01", "caption": "测试图", "code": "def draw(output_path: str):\n    pass"}

        client = FakeClient()
        question = {
            "question_id": "q1",
            "stem": "画出衍射花样。",
            "question_understanding": {"images": [{"detected_labels": ["000"]}]},
            "figure_schema_plan": {"figure_semantic_contract": {"required_labels": ["000"]}},
        }
        fragment = {"question_id": "q1", "answer_summary": "标出中心斑点和衍射斑点。"}
        prompt = build_drawing_code_prompt(question, fragment)
        result = generate_drawing_code_spec(client, question, fragment, model="deepseek-v4-flash")

        self.assertEqual("fig_01", result["figure_id"])
        self.assertEqual(DRAWING_CODE_MAX_TOKENS, client.kwargs["max_tokens"])
        self.assertEqual("disabled", client.kwargs["thinking"])
        payload = __import__("json").loads(prompt[1]["content"])
        self.assertIn("monochrome_code_template", payload)
        self.assertIn("color=BLACK", payload["monochrome_code_template"])
        self.assertEqual(["000"], payload["question"]["question_understanding"]["images"][0]["detected_labels"])
        self.assertEqual(["000"], payload["question"]["figure_schema_plan"]["figure_semantic_contract"]["required_labels"])

    def test_drawing_code_file_block_protocol_keeps_code_outside_json(self) -> None:
        from app.drawing_code import generate_drawing_code_spec, parse_drawing_code_model_response
        from app.llm_client import LLMResult

        content = """<JSON>{"figure_id":"fig_01","caption":"测试图","code_ref":"fig_01.py","notes":"ok"}</JSON>
<FILE name="fig_01.py">
def draw(output_path: str) -> None:
    value = r"\\bar{2}"
    pass
</FILE>"""
        spec, notes = parse_drawing_code_model_response(content)
        self.assertEqual("fig_01", spec["figure_id"])
        self.assertEqual("ok", spec["notes"])
        self.assertIn(r"\bar{2}", spec["code"])
        self.assertEqual([], notes)

        class FakeClient:
            def __init__(self):
                self.kwargs = {}

            def chat_text(self, _messages, **kwargs):
                self.kwargs = kwargs
                return LLMResult(provider="fake", model="m", content=content, raw={})

        client = FakeClient()
        result = generate_drawing_code_spec(client, {"question_id": "q1"}, {"question_id": "q1"}, model="m")
        self.assertEqual("fig_01", result["figure_id"])
        self.assertIn(r"\bar{2}", result["code"])
        self.assertEqual("disabled", client.kwargs["thinking"])

    def test_dashscope_qwen_image_size_aliases_are_normalized(self) -> None:
        from app.llm_client import _dashscope_image_size

        self.assertEqual("2048*2048", _dashscope_image_size("qwen-image-2.0-pro", explicit_size="2K", configured_size="2K"))
        self.assertEqual("1024*1024", _dashscope_image_size("qwen-image-2.0-pro", explicit_size=None, configured_size="1K"))

    def test_crystallographic_index_whitelist_suppresses_only_physics_claims(self) -> None:
        from app.figures import _apply_crystallographic_index_whitelist

        question = {"stem": "画出体心立方[110]带轴电子衍射花样，并标出晶面指数。"}
        spec = {
            "kind": "zone_axis_diffraction",
            "caption": "BCC [110] 电子衍射花样",
            "zone_axis": [1, 1, 0],
            "lattice": "bcc",
            "max_index": 3,
            "label_indices": [[0, 0, 0], [1, -1, 0], [0, 0, 2]],
            "spot_size": 42,
            "apply_extinction": True,
        }
        physics_only = {
            "ok": False,
            "summary": "(110) 为非法反射，违反消光条件。",
            "missing_requirements": ["(110) 不应出现，违反 h+k=0。"],
            "label_issues": ["(110) 标签错误。"],
            "visual_issues": [],
        }
        filtered = _apply_crystallographic_index_whitelist(physics_only, question, spec)
        self.assertTrue(filtered["ok"])
        self.assertTrue(filtered["crystallographic_index_whitelist"]["applied"])
        self.assertFalse(filtered["missing_requirements"])
        self.assertFalse(filtered["label_issues"])

        mixed = dict(physics_only)
        mixed["visual_issues"] = ["标签与斑点重叠，难以辨认。"]
        filtered_mixed = _apply_crystallographic_index_whitelist(mixed, question, spec)
        self.assertFalse(filtered_mixed["ok"])
        self.assertEqual(["标签与斑点重叠，难以辨认。"], filtered_mixed["visual_issues"])

    def test_drawing_validator_allows_numeric_variable_reassignment(self) -> None:
        from app.drawing_code import validate_drawing_code

        code = """import matplotlib.pyplot as plt
import numpy as np

def draw(output_path: str) -> None:
    x, y = np.meshgrid(np.arange(-2, 3), np.arange(-2, 3))
    x = x.flatten()
    y = y.flatten()
    fig, ax = plt.subplots()
    ax.scatter(x, y, c='black')
    fig.savefig(output_path)
    plt.close(fig)
"""
        self.assertEqual([], validate_drawing_code(code))

    def test_drawing_validator_allows_latex_crystallographic_indices(self) -> None:
        from app.drawing_code import validate_drawing_code

        code = """import matplotlib.pyplot as plt

def draw(output_path: str) -> None:
    fig, ax = plt.subplots()
    ax.scatter([0], [0], color='black')
    ax.text(0, 0, r'$(00\\bar{2})$', color='black')
    ax.text(0, 1, r'$(1\\overline{1}0)$', color='black')
    fig.savefig(output_path)
    plt.close(fig)
"""

        self.assertEqual([], validate_drawing_code(code))

    def test_drawing_validator_still_rejects_english_mathtext_label(self) -> None:
        from app.drawing_code import validate_drawing_code

        code = """import matplotlib.pyplot as plt

def draw(output_path: str) -> None:
    fig, ax = plt.subplots()
    ax.plot([0, 1], [0, 1], color='black')
    ax.set_ylabel('$Intensity$')
    fig.savefig(output_path)
    plt.close(fig)
"""

        issues = validate_drawing_code(code)
        self.assertTrue(any("Chinese or standard notation" in issue for issue in issues), issues)

    def test_drawing_validator_allows_compact_scientific_point_and_phase_labels(self) -> None:
        from app.drawing_code import validate_drawing_code

        code = """import matplotlib.pyplot as plt

def draw(output_path: str) -> None:
    fig, ax = plt.subplots()
    ax.plot([0, 1], [0, 1], color='black')
    ax.annotate('A', xy=(0, 0), color='black')
    ax.annotate('Fe₃CⅡ', xy=(1, 1), color='black')
    fig.savefig(output_path)
    plt.close(fig)
"""

        self.assertEqual([], validate_drawing_code(code))

    def test_drawing_prompt_scope_excludes_non_drawing_siblings(self) -> None:
        from app.drawing_code import _drawing_prompt_question

        question = {
            "stem": "综合题",
            "subquestions": [
                {"number": "1", "question_type": "简答题", "stem": "解释反应。"},
                {
                    "number": "2",
                    "question_type": "计算题",
                    "stem": "作图并计算。",
                    "requirements": [
                        {"number": "2.1", "question_type": "作图题", "stem": "画出冷却曲线。"},
                        {"number": "2.2", "question_type": "计算题", "stem": "计算质量比。"},
                    ],
                },
            ],
        }

        scoped = _drawing_prompt_question(question)

        self.assertIn("画出冷却曲线", scoped["stem"])
        self.assertNotIn("解释反应", scoped["stem"])
        self.assertNotIn("计算质量比", scoped["stem"])

    def test_literal_string_resolver_stops_on_cyclic_aliases(self) -> None:
        import ast

        from app.drawing_code import _literal_strings

        tree = ast.parse("a = b\nb = a\n")
        assignments = {node.targets[0].id: node.value for node in tree.body if isinstance(node, ast.Assign)}
        self.assertEqual([], _literal_strings(ast.Name(id="a"), assignments))

    def test_visual_repair_promotes_only_a_reaudited_passing_candidate(self) -> None:
        from app.figures import repair_figures_with_model_for_visual_qa

        code = """import matplotlib.pyplot as plt
def draw(output_path: str) -> None:
    fig, ax = plt.subplots()
    ax.plot([0, 1], [0, 1], color='black')
    fig.savefig(output_path)
    plt.close(fig)
"""
        structured_exam = {"items": [{"question_id": "q1", "question_type": "作图题", "stem": "画图"}]}
        initial_spec = {"figure_id": "fig_01", "question_id": "q1", "kind": "model_drawing_code", "caption": "测试图", "code": code}
        initial_qa = {
            "enabled": True,
            "provider": "vision",
            "vision_model": "vision-model",
            "items": [{"question_id": "q1", "figure_id": "fig_01", "path": "", "qa": {"ok": False, "summary": "标签重叠"}}],
            "skipped": [],
        }
        repair_provider = SimpleNamespace(name="repair", api_key="key", default_model="repair-model")
        vision_provider = SimpleNamespace(name="vision", api_key="key", supports_vision=True, vision_model="vision-model")

        class FakeClient:
            def __init__(self, provider):
                self.provider = provider

            def chat_json_object(self, *_args, **_kwargs):
                return {
                    "drawing_code_spec": {"caption": "测试图", "code": code},
                    "repair_notes": [f"{self.provider.name} candidate"],
                }

        def fake_audit(_exam, candidate_specs_json, candidate_dir, report_json, **_kwargs):
            candidate = json.loads(candidate_specs_json.read_text(encoding="utf-8"))["figures"][0]
            passed = candidate.get("source") == "visual_qa_vision_reviewer_candidate"
            result = {
                "enabled": True,
                "items": [{
                    "question_id": "q1",
                    "figure_id": "fig_01",
                    "path": str(candidate_dir / "fig_01.png"),
                    "qa": {"ok": passed, "summary": "passed" if passed else "failed"},
                }],
                "skipped": [],
            }
            report_json.write_text(json.dumps(result), encoding="utf-8")
            return result

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            specs_json = tmp / "figure_specs.json"
            figures_dir = tmp / "figures"
            specs_json.write_text(json.dumps({"figures": [initial_spec]}), encoding="utf-8")
            with patch("app.figures.OpenAICompatibleClient", FakeClient), patch("app.figures.audit_figures_with_vision", fake_audit):
                report = repair_figures_with_model_for_visual_qa(
                    structured_exam,
                    None,
                    specs_json,
                    figures_dir,
                    tmp / "figure_visual_qa.json",
                    tmp / "repair.json",
                    qa_report=initial_qa,
                    provider=repair_provider,
                    model="repair-model",
                    vision_provider=vision_provider,
                    vision_model="vision-model",
                )
            final_spec = json.loads(specs_json.read_text(encoding="utf-8"))["figures"][0]
            final_image_exists = (figures_dir / "fig_01.png").exists()

        target = report["rounds"][0]["targets"][0]
        self.assertEqual("vision_reviewer", target["selected_strategy"])
        self.assertEqual("visual_qa_vision_reviewer_candidate", final_spec["source"])
        self.assertTrue(report["latest_visual_qa"]["items"][0]["qa"]["ok"])
        self.assertEqual(2, len(target["candidates"]))
        self.assertTrue(final_image_exists)

    def test_visual_repair_reuses_completed_identical_attempt_without_model_calls(self) -> None:
        from app.figures import repair_figures_with_model_for_visual_qa

        structured_exam = {"items": [{"question_id": "q1", "question_type": "作图题", "stem": "画图"}]}
        spec = {
            "figure_id": "fig_01",
            "question_id": "q1",
            "kind": "generic_axis_curve",
            "caption": "测试图",
            "x_label": "x",
            "y_label": "y",
            "points": [[0, 0], [1, 1]],
        }
        qa = {
            "enabled": True,
            "items": [{"question_id": "q1", "figure_id": "fig_01", "qa": {"ok": False, "summary": "标签问题"}}],
            "skipped": [],
        }
        repair_provider = SimpleNamespace(name="repair", api_key="key", default_model="repair-model")
        vision_provider = SimpleNamespace(name="vision", api_key="key", supports_vision=True, vision_model="vision-model")

        class FakeClient:
            calls = 0

            def __init__(self, _provider):
                pass

            def chat_json_object(self, *_args, **_kwargs):
                type(self).calls += 1
                return {"figure_spec": spec, "repair_notes": []}

        def fake_audit(_exam, _candidate_specs_json, candidate_dir, report_json, **_kwargs):
            result = {
                "enabled": True,
                "items": [
                    {
                        "question_id": "q1",
                        "figure_id": "fig_01",
                        "path": str(candidate_dir / "fig_01.png"),
                        "qa": {"ok": False, "summary": "仍未通过，但服务已正常完成"},
                    }
                ],
                "skipped": [],
            }
            report_json.write_text(json.dumps(result), encoding="utf-8")
            return result

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            specs_json = tmp / "figure_specs.json"
            figures_dir = tmp / "figures"
            figures_dir.mkdir()
            specs_json.write_text(json.dumps({"figures": [spec]}), encoding="utf-8")
            from PIL import Image

            Image.new("RGB", (120, 120), "white").save(figures_dir / "fig_01.png")
            kwargs = {
                "structured_exam": structured_exam,
                "fragments_json": None,
                "specs_json": specs_json,
                "output_dir": figures_dir,
                "visual_qa_json": tmp / "figure_visual_qa.json",
                "repair_report_json": tmp / "repair.json",
                "qa_report": qa,
                "provider": repair_provider,
                "model": "repair-model",
                "vision_provider": vision_provider,
                "vision_model": "vision-model",
            }
            with patch("app.figures.OpenAICompatibleClient", FakeClient), patch(
                "app.figures.audit_figures_with_vision", fake_audit
            ):
                first = repair_figures_with_model_for_visual_qa(**kwargs)
                first_calls = FakeClient.calls
                second = repair_figures_with_model_for_visual_qa(**kwargs)

        self.assertGreater(first_calls, 0)
        self.assertFalse(first["cache"]["hit"])
        self.assertTrue(second["cache"]["hit"])
        self.assertEqual(first_calls, FakeClient.calls)
        self.assertEqual(0, second["remote_model_calls_this_run"])


if __name__ == "__main__":
    unittest.main()
