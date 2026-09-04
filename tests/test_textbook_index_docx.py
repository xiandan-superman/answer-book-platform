from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document

from app import textbook_index
from app.textbook_index import build_textbook_index_for_files, discover_textbooks
from app.textbook_package import TextbookPackage


def _rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def _mineru_package(tmp_path: Path, source: Path, blocks: list[dict]) -> TextbookPackage:
    root = tmp_path / f"mineru-{source.stem}"
    root.mkdir()
    content = root / "content_list.json"
    content.write_text(json.dumps(blocks, ensure_ascii=False), encoding="utf-8")
    audit = root / "audit.json"
    audit.write_text("{}", encoding="utf-8")
    return TextbookPackage(source.stem, root, source.stem, source.stem, content, None, None, None, None, root, audit)


def test_docx_textbook_is_indexed_as_retrievable_prose_and_table_content(tmp_path: Path, monkeypatch) -> None:
    textbook = tmp_path / "ideal_gas_reference.docx"
    document = Document()
    document.add_paragraph("For an ideal gas, pressure, volume, amount, and temperature obey pV = nRT.")
    document.add_paragraph("Use absolute temperature when applying the equation of state.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Quantity"
    table.cell(0, 1).text = "Unit"
    table.cell(1, 0).text = "Pressure"
    table.cell(1, 1).text = "Pa"
    document.save(textbook)
    package = _mineru_package(
        tmp_path,
        textbook,
        [
            {"type": "text", "page_idx": 0, "text": "For an ideal gas, pressure, volume, amount, and temperature obey pV = nRT."},
            {"type": "text", "page_idx": 0, "text": "Use absolute temperature when applying the equation of state."},
            {"type": "table", "page_idx": 0, "table_body": "<table><tr><td>Pressure</td><td>Pa</td></tr></table>"},
        ],
    )
    monkeypatch.setattr(textbook_index, "parse_document", lambda _path: package)

    stage = tmp_path / "stage"
    result = build_textbook_index_for_files([textbook], stage)
    indexed = _rows(stage / "textbook_blocks.csv")

    assert result.block_count == 3
    assert all(row["text"].strip() for row in indexed)
    assert any("pV = nRT" in row["retrieval_text"] for row in indexed)
    table_row = next(row for row in indexed if row["block_type"] == "table")
    assert "Pressure Pa" in table_row["retrieval_text"]
    assert table_row["source_type"] == "table_block"
    assert _rows(stage / "textbook_page_map.csv")[0]["verified"] == "false"


def test_discover_textbooks_includes_supported_docx_and_pdf_files(tmp_path: Path) -> None:
    (tmp_path / "reference.docx").write_bytes(b"placeholder")
    (tmp_path / "reference.pdf").write_bytes(b"placeholder")

    assert [path.name for path in discover_textbooks(tmp_path)] == ["reference.docx", "reference.pdf"]


def test_pdf_textbook_uses_mineru_text_and_visual_blocks(tmp_path: Path, monkeypatch) -> None:
    textbook = tmp_path / "reference.pdf"
    textbook.write_bytes(b"pdf")
    page = tmp_path / "page-1.jpg"
    page.write_bytes(b"jpeg")
    package = _mineru_package(
        tmp_path,
        textbook,
        [
            {"type": "text", "page_idx": 0, "text": "相平衡条件与相图"},
            {"type": "image", "page_idx": 0, "img_path": str(page), "image_caption": ["相图"]},
        ],
    )
    monkeypatch.setattr(textbook_index, "parse_document", lambda _path: package)

    result = build_textbook_index_for_files([textbook], tmp_path / "stage")
    indexed = _rows(Path(result.blocks_csv))
    status = (tmp_path / "stage" / "textbook_index_status.json").read_text(encoding="utf-8")

    assert any(row["source_type"] == "text_block" and "相平衡" in row["retrieval_text"] for row in indexed)
    visual = next(row for row in indexed if row["source_type"] == "figure_block")
    assert visual["asset_path"] == str(page)
    assert "input_representations" in status
    assert '"parser": "mineru"' in status
