from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches
from PIL import Image

from standalone_word_format_reviewer import server as reviewer_server
from standalone_word_format_reviewer.format_engine import audit_docx, default_task_options, repair_docx
from standalone_word_format_reviewer.server import _safe_filename


def _fixture(path: Path) -> None:
    image_path = path.with_suffix(".png")
    Image.new("RGB", (1200, 300), "white").save(image_path)

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.header.paragraphs[0].text = "旧页眉"
    document.add_paragraph("复习规划及本书使用指南")
    document.add_paragraph("一、使用说明")
    document.add_paragraph("正文 English 123")
    document.add_picture(str(image_path))
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格 English 456"
    table.rows[0].height = Inches(0.1)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    formula_paragraph = document.add_paragraph()
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    math.append(math_run)
    formula_paragraph._p.append(math)
    document.save(path)


def test_answer_repair_fixes_supported_items_and_preserves_objects(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _fixture(source)

    before = audit_docx(source, "answer", "自定义页眉")
    codes = {issue["code"] for issue in before["issues"]}
    assert {"page_size", "header_content", "table_row_height", "page_number", "image_overflow"} <= codes
    assert before["summary"]["formula_count"] == 1
    assert before["summary"]["image_count"] == 1

    after = repair_docx(source, output, "answer", "自定义页眉")
    assert after["summary"]["fixable_count"] == 0
    assert after["summary"]["formula_count"] == 1
    assert after["summary"]["image_count"] == 1

    repaired = Document(output)
    assert repaired.sections[0].header.paragraphs[0].text == "自定义页眉"
    assert repaired.styles["Normal"].paragraph_format.first_line_indent.pt == 10
    assert repaired.tables[0].rows[0].height_rule != WD_ROW_HEIGHT_RULE.EXACTLY
    assert repaired.inline_shapes[0].width <= Cm(14.8)
    assert "PAGE" in repaired.sections[0].footer._element.xml.upper()
    assert repaired.settings.element.find(".//" + qn("m:mathFont")).get(qn("m:val")) == "Cambria Math"


def test_lecture_profile_uses_lecture_body_indent_and_styles(tmp_path: Path) -> None:
    source = tmp_path / "lecture.docx"
    output = tmp_path / "lecture_fixed.docx"
    _fixture(source)

    after = repair_docx(source, output, "lecture", "讲义页眉")
    repaired = Document(output)

    assert after["summary"]["fixable_count"] == 0
    assert repaired.styles["Normal"].paragraph_format.first_line_indent.pt == 22.1
    assert "一级标题" in repaired.styles
    assert repaired.styles["一级标题"].font.size.pt == 22
    assert repaired.styles["二级标题"].font.size.pt == 14
    assert repaired.styles["三级标题"].font.size.pt == 12
    assert repaired.styles["注"].paragraph_format.first_line_indent.pt == 5


def test_manual_line_break_uses_word_compatibility_without_disabling_justify(tmp_path: Path) -> None:
    source = tmp_path / "manual_break.docx"
    output = tmp_path / "manual_break_fixed.docx"
    document = Document()
    paragraph = document.add_paragraph("第一条短要求")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("第二条短要求")
    document.add_paragraph("这是保持普通两端对齐样式的完整正文段落。")
    document.save(source)

    repair_docx(source, output, "answer", "")

    repaired = Document(output)
    compat = repaired.settings.element.find(qn("w:compat"))
    assert compat is not None
    assert compat.find(qn("w:doNotExpandShiftReturn")) is not None
    assert repaired.styles["Normal"].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert repaired.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert repaired.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_safe_filename_removes_paths_and_reserved_characters() -> None:
    assert _safe_filename("../../有:问题?.docx") == "有问题.docx"
    assert _safe_filename("") == "document.docx"


def test_task_options_are_applied_without_changing_fixed_rules(tmp_path: Path) -> None:
    source = tmp_path / "custom.docx"
    output = tmp_path / "custom_fixed.docx"
    _fixture(source)
    options = {"chinese_font": "仿宋", "latin_font": "Arial", "body_size_pt": 12, "line_spacing": 2}

    report = repair_docx(source, output, "answer", "自定义页眉", options)
    repaired = Document(output)
    normal = repaired.styles["Normal"]
    r_fonts = normal.element.rPr.find(qn("w:rFonts"))

    assert report["task_options"]["styles"]["Normal"]["chinese_font"] == "仿宋"
    assert report["task_options"]["styles"]["Normal"]["latin_font"] == "Arial"
    assert report["task_options"]["styles"]["Normal"]["size_pt"] == 12
    assert report["task_options"]["styles"]["Normal"]["line_spacing"] == 2
    assert r_fonts.get(qn("w:eastAsia")) == "仿宋"
    assert r_fonts.get(qn("w:ascii")) == "Arial"
    assert normal.font.size.pt == 12
    assert normal.paragraph_format.line_spacing == 2
    assert repaired.styles["Heading 1"].font.size.pt == 18


def test_optional_style_bold_does_not_create_an_unrepairable_residual(tmp_path: Path) -> None:
    source = tmp_path / "explicit_not_bold.docx"
    output = tmp_path / "explicit_not_bold_fixed.docx"
    document = Document()
    document.styles["Normal"].font.bold = False
    document.add_paragraph("正文保持常规字重。")
    document.save(source)

    report = repair_docx(source, output, "answer", "")

    assert not any(
        issue["code"] == "style_format" and issue["item"] == "Normal"
        for issue in report["issues"]
    )
    assert report["summary"]["fixable_count"] == 0


def test_frontend_keeps_configuration_inside_read_only_rule_page() -> None:
    html = (Path(__file__).parents[1] / "standalone_word_format_reviewer" / "web" / "index.html").read_text(encoding="utf-8")

    assert 'name="profile" value="answer"' in html
    assert 'name="profile" value="lecture"' in html
    assert '<select id="profile"' not in html
    assert 'id="chineseFont"' not in html
    assert 'id="taskConfigSummary"' not in html
    assert '当前标准配置' not in html
    assert 'id="standardsView"' in html
    assert 'id="saveScopeModal"' in html
    assert 'id="initialSettings"' in html
    assert 'data-lucide="pencil-line"' in html
    assert 'onclick="beginRuleEdit(this)"' in html
    assert 'onclick="cancelRuleEdit()"' in html
    assert '<span>保存</span>' in html
    assert '<span>取消</span>' in html
    assert "styles.Normal" in html
    assert "header.content" in html
    assert "仅用于本次任务" in html
    assert "设为永久默认" in html


def test_server_embeds_settings_in_initial_page() -> None:
    html = reviewer_server._render_web_page().decode("utf-8")

    assert "__INITIAL_SETTINGS_JSON__" not in html
    assert '"answer":{"defaults"' in html
    assert '"lecture":{"defaults"' in html


def test_standalone_job_payload_keeps_before_after_changes_and_recovery_links(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(reviewer_server, "JOB_ROOT", tmp_path)
    job_id = "a" * 32
    job = tmp_path / job_id
    job.mkdir()
    source = job / "source.docx"
    modified = job / "modified.docx"
    _fixture(source)
    before = audit_docx(source, "answer", "新页眉")
    after = repair_docx(source, modified, "answer", "新页眉")
    reviewer_server._write_json(job / "meta.json", {
        "job_id": job_id,
        "filename": "真实题目.docx",
        "profile": "answer",
        "mode": "review",
        "status": "completed",
    })
    reviewer_server._write_json(job / "audit.json", before)
    reviewer_server._write_json(job / "final_audit.json", after)

    payload = reviewer_server._job_payload(job_id)

    assert payload["report"] == before
    assert payload["final_report"] == after
    assert payload["changes"]["resolved"]
    assert payload["suggested_filename"] == "真实题目_格式已修改.docx"
    assert payload["source_download_url"].endswith("/source")
    assert payload["source_preview_url"].endswith("preview?version=source")
    assert payload["modified_preview_url"].endswith("preview?version=modified")
    assert payload["can_restore"] is True


def test_frontend_explains_history_recovery_preview_cancel_and_download_state() -> None:
    html = (Path(__file__).parents[1] / "standalone_word_format_reviewer" / "web" / "index.html").read_text(encoding="utf-8")

    assert "修改前问题" in html
    assert "最终复查" in html
    assert "实际处理" in html
    assert "wordFormatReviewer:lastJob" in html
    assert "previewModal" in html
    assert "cancelCurrentOperation" in html
    assert "restoreOriginal" in html
    assert "重试上一步" in html
    assert "已向浏览器发起下载" in html
    assert "桌面版已保存" in html
    assert "实际路径" in html
    assert "saved?.status === 'cancelled'" in html


def test_permanent_profile_settings_are_stored_and_reloaded(tmp_path: Path, monkeypatch) -> None:
    settings_path = tmp_path / "settings.json"
    monkeypatch.setattr(reviewer_server, "SETTINGS_FILE", settings_path)
    defaults = reviewer_server.default_task_options("lecture")
    defaults["styles"]["Normal"]["chinese_font"] = "微软雅黑"
    defaults["header"]["content"] = "永久页眉"

    saved = reviewer_server._save_profile_settings("lecture", defaults)
    payload = reviewer_server._settings_payload()

    assert settings_path.exists()
    assert saved["styles"]["Normal"]["chinese_font"] == "微软雅黑"
    assert payload["profiles"]["lecture"]["effective"]["header"]["content"] == "永久页眉"
    assert payload["profiles"]["answer"]["saved"] is None


def test_detailed_standard_options_are_applied_to_docx(tmp_path: Path) -> None:
    source = tmp_path / "detailed.docx"
    output = tmp_path / "detailed_fixed.docx"
    _fixture(source)
    options = default_task_options("answer")
    options["page"].update({"size": "a4", "margin_cm": 2.0, "header_distance_cm": 1.5, "footer_distance_cm": 1.4})
    options["styles"]["Heading 1"]["size_pt"] = 20
    options["header"].update({"content": "完整配置页眉", "size_pt": 10.5, "bottom_border": False})
    options["page_number"].update({"size_pt": 10, "alignment": "left"})
    options["table"].update({"size_pt": 10.5, "alignment": "left"})
    options["image"].update({"alignment": "left", "max_width_cm": 10})

    report = repair_docx(source, output, "answer", "", options)
    repaired = Document(output)
    section = repaired.sections[0]
    header = section.header.paragraphs[0]
    page_paragraph = next(p for p in section.footer.paragraphs if "PAGE" in p._p.xml.upper())

    assert report["summary"]["fixable_count"] == 0
    assert round(section.page_width.cm, 1) == 21.0
    assert round(section.page_height.cm, 1) == 29.7
    assert round(section.left_margin.cm, 1) == 2.0
    assert repaired.styles["Heading 1"].font.size.pt == 20
    assert header.text == "完整配置页眉"
    assert header.runs[0].font.size.pt == 10.5
    assert header._p.pPr.find(qn("w:pBdr")) is None or header._p.pPr.find(qn("w:pBdr")).find(qn("w:bottom")) is None
    assert page_paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert repaired.inline_shapes[0].width <= Cm(10)
    assert repaired.tables[0].cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
