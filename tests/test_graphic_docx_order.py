from __future__ import annotations

import json
import zipfile

from docx import Document
from lxml import etree

from app.document_presentation import plan_ordered_answer_units
from app.docx_v4 import build_docx_from_fragments


WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def test_calculation_parent_with_graphic_child_keeps_calculation_steps_and_formulas(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "calc_mixed",
                "section": "九、计算题",
                "question_type": "计算题",
                "subquestions": [{"number": "1", "question_type": "作图题"}],
                "number": "1",
                "answer": "见解析",
                "answer_summary": "见解析",
                "evidence_ids": [],
                "formulas": [
                    {"formula_id": "f1", "latex": "x=1", "role": "result", "display": True}
                ],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材。"}]},
                    {"label": "图示", "segments": [{"type": "text", "text": "图示。"}]},
                    {
                        "label": "解题步骤",
                        "segments": [
                            {"type": "text", "text": "计算："},
                            {"type": "formula_ref", "formula_id": "f1"},
                        ],
                    },
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    doc = Document(output)
    assert any(paragraph.text == "答案：" for paragraph in doc.paragraphs)
    with zipfile.ZipFile(output) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    assert len(root.xpath("//m:oMath", namespaces=WORD_NS)) == 1


def test_composite_calculation_renders_each_unit_payload_in_original_order(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "calc_ordered",
                "section": "九、计算题",
                "question_type": "计算题",
                "number": "1",
                "answer": "见解析",
                "answer_summary": "见解析",
                "subquestions": [
                    {"number": "1", "stem": "说明原理", "question_type": "简答题"},
                    {"number": "2", "stem": "计算结果", "question_type": "计算题"},
                ],
                "answer_units": [
                    {"number": "1", "question_type": "简答题", "answer": "原理答案"},
                    {"number": "2", "question_type": "计算题", "answer": "x=1"},
                ],
                "evidence_ids": [],
                "formulas": [{"formula_id": "f1", "latex": "x=1", "role": "result", "display": True}],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材。"}]},
                    {
                        "label": "解析",
                        "segments": [
                            {"type": "text", "text": "(1)说明原理\n"},
                            {"type": "text", "text": "答案：原理答案\n原理解析。\n"},
                            {"type": "text", "text": "(2)计算结果\n"},
                            {"type": "text", "text": "答案：x=1\n"},
                        ],
                    },
                    {
                        "label": "解题步骤",
                        "segments": [
                            {"type": "text", "text": "(2)计算结果\n"},
                            {"type": "text", "text": "建立关系。"},
                            {"type": "formula_ref", "formula_id": "f1"},
                        ],
                    },
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs if paragraph.text.strip()]
    first_heading = paragraphs.index("(1)")
    first_answer = paragraphs.index("答案：原理答案")
    first_analysis = paragraphs.index("原理解析。")
    second_heading = paragraphs.index("(2)")
    second_answer = paragraphs.index("答案：x=1")
    steps_heading = paragraphs.index("解题步骤：")
    assert first_heading < first_answer < first_analysis < second_heading < second_answer < steps_heading


def test_ordered_projection_joins_soft_break_before_parenthesized_value() -> None:
    fragment = {
        "subquestions": [
            {"number": "1", "stem": "判断成分"},
            {"number": "2", "stem": "说明组织"},
        ],
        "answer_units": [{"number": "1"}, {"number": "2"}],
        "blocks": [
            {
                "label": "解析",
                "segments": [
                    {"type": "text", "text": "(1)判断成分"},
                    {"type": "text", "text": "合金位于共晶点\n（60）左侧。"},
                    {"type": "text", "text": "(2)说明组织"},
                    {"type": "text", "text": "组织说明。"},
                ],
            }
        ],
    }

    plan = plan_ordered_answer_units(fragment)

    assert plan["ok"] is True
    assert plan["units"][0]["analysis_segments"][0]["text"] == "合金位于共晶点（60）左侧。"


def test_ordered_projection_falls_back_when_legacy_figure_cannot_be_assigned() -> None:
    fragment = {
        "subquestions": [
            {"number": "1", "stem": "作图"},
            {"number": "2", "stem": "说明"},
        ],
        "answer_units": [{"number": "1"}, {"number": "2"}],
        "blocks": [
            {
                "label": "解析",
                "segments": [
                    {"type": "text", "text": "(1)作图"},
                    {"type": "text", "text": "作图解析。"},
                    {"type": "text", "text": "(2)说明"},
                    {"type": "text", "text": "说明解析。"},
                ],
            },
            {
                "label": "图示",
                "segments": [{"type": "image_ref", "path": "legacy.png"}],
            },
        ],
    }

    plan = plan_ordered_answer_units(fragment)

    assert plan["ok"] is False
    assert plan["unassigned"]["figures"] == [{"type": "image_ref", "path": "legacy.png"}]


def test_page_number_field_is_valid_paragraph_child(tmp_path) -> None:
    fragments = {"schema_version": "answer_book.answer_fragments.v4", "fragments": []}
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments), encoding="utf-8")

    build_docx_from_fragments(source, output)

    with zipfile.ZipFile(output) as archive:
        footer_name = next(name for name in archive.namelist() if name.startswith("word/footer"))
        footer = etree.fromstring(archive.read(footer_name))
        settings = etree.fromstring(archive.read("word/settings.xml"))
    fields = footer.xpath("//w:p/w:fldSimple[@w:instr=' PAGE ']", namespaces=WORD_NS)
    assert len(fields) == 1
    assert not footer.xpath("//w:r/w:fldSimple", namespaces=WORD_NS)
    assert settings.xpath("//w:updateFields[@w:val='true']", namespaces=WORD_NS)


def test_legacy_promoted_sentence_formula_renders_inline(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "q1",
                "section": "六、简答题",
                "question_type": "简答题",
                "number": "1",
                "answer": "见解析",
                "answer_summary": "见解析",
                "evidence_ids": [],
                "formulas": [
                    {
                        "formula_id": "f1",
                        "latex": r"\beta_{II}",
                        "role": "relation",
                        "display": True,
                        "source_note": "程序在 Word 生成前从普通文本中识别出的公式片段。",
                    }
                ],
                "blocks": [
                    {
                        "label": "解析",
                        "segments": [
                            {"type": "text", "text": "析出二次相"},
                            {"type": "formula_ref", "formula_id": "f1"},
                            {"type": "text", "text": "，继续冷却。"},
                        ],
                    }
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    analysis = next(paragraph for paragraph in Document(output).paragraphs if paragraph.text.startswith("解析：析出二次相"))
    assert "，继续冷却。" in analysis.text
    assert analysis._p.xpath(".//m:oMath")


def test_graphic_question_renders_evidence_before_answer(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "graphic_s01_01",
                "section": "一、作图题",
                "question_type": "作图题",
                "number": "1",
                "answer": "见解析",
                "answer_summary": "应画出 Frank-Read 位错源增殖过程。",
                "evidence_ids": ["ev1"],
                "formulas": [],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材中的位错源增殖机制。"}]},
                    {"label": "图示", "segments": [{"type": "text", "text": "Frank-Read 位错源示意图"}]},
                    {"label": "解析", "segments": [{"type": "text", "text": "解析说明。"}]},
                    {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "注意钉扎点。"}]},
                    {"label": "待复核公式", "segments": [{"type": "text", "text": "不应进入作图题正式格式。"}]},
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    doc = Document(output)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    assert paragraphs[1:8] == [
        "一、作图题",
        "1、教材依据：教材中的位错源增殖机制。",
        "答案：",
        "应画出 Frank-Read 位错源增殖过程。",
        "图示：Frank-Read 位错源示意图",
        "解析：解析说明。",
        "易错点及注意事项：注意钉扎点。",
    ]
    answer_body = next(p for p in doc.paragraphs if p.text == "应画出 Frank-Read 位错源增殖过程。")
    assert answer_body.paragraph_format.first_line_indent is not None
    assert "易错点及注意事项：注意钉扎点。" in paragraphs
    assert not any("待复核公式" in text for text in paragraphs)
    assert not any(text == "1、见解析" for text in paragraphs)


def test_short_answer_renders_answer_and_analysis_separately(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "qa_s01_06_02",
                "section": "六、简答题",
                "question_type": "简答题",
                "number": "2",
                "answer": "扩散系数随温度升高而增大，表面扩散最快。",
                "answer_summary": "扩散系数随温度升高而增大，表面扩散最快。",
                "evidence_ids": ["ev1"],
                "formulas": [],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材扩散章节。"}]},
                    {"label": "解析", "segments": [{"type": "text", "text": "同温下看曲线高低，表面高于晶界和晶内。"}]},
                    {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "不能把横轴直接当线性温度。"}]},
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    doc = Document(output)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    assert paragraphs[1:7] == [
        "六、简答题",
        "2、教材依据：教材扩散章节。",
        "答案：",
        "扩散系数随温度升高而增大，表面扩散最快。",
        "解析：同温下看曲线高低，表面高于晶界和晶内。",
        "易错点及注意事项：不能把横轴直接当线性温度。",
    ]
    answer_body = next(p for p in doc.paragraphs if p.text == "扩散系数随温度升高而增大，表面扩散最快。")
    assert answer_body.paragraph_format.first_line_indent is not None


def test_calculation_renders_number_and_evidence_in_same_paragraph(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "calc_s01_01",
                "section": "五、计算题",
                "question_type": "计算题",
                "number": "1",
                "answer": "计算结果见解题步骤。",
                "answer_summary": "计算结果见解题步骤。",
                "evidence_ids": ["ev1"],
                "formulas": [],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材扩散定律章节。"}]},
                    {"label": "解析", "segments": [{"type": "text", "text": "先统一单位。"}]},
                    {"label": "解题步骤", "segments": [{"type": "text", "text": "代入已知条件计算。"}]},
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    paragraphs = [p.text for p in Document(output).paragraphs if p.text.strip()]
    assert "1、教材依据：教材扩散定律章节。" in paragraphs
    assert "1、" not in paragraphs
    assert sum("教材依据：" in text for text in paragraphs) == 1
    assert "解析：先统一单位。" in paragraphs


def test_docx_normalizes_saved_subquestion_and_requirement_markers(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "nested_s01_01",
                "section": "六、简答题",
                "question_type": "简答题",
                "number": "1",
                "answer": "(1)、第一问答案。\n①、第一项要求。",
                "answer_summary": "(1)、第一问答案。\n①、第一项要求。",
                "evidence_ids": ["ev1"],
                "formulas": [],
                "blocks": [
                    {"label": "教材依据", "segments": [{"type": "text", "text": "教材相关章节。"}]},
                    {
                        "label": "解析",
                        "segments": [{"type": "text", "text": "(1)、第一小问\n(2) 第二小问\n①、第一项要求\n②、第二项要求"}],
                    },
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    paragraphs = [p.text for p in Document(output).paragraphs if p.text.strip()]
    joined = "\n".join(paragraphs)
    assert "(1)第一小问" in joined
    assert "(2)第二小问" in joined
    assert "①、第一项要求" in joined
    assert "②、第二项要求" in joined
    assert "(1)、" not in joined
