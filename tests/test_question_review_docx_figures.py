from __future__ import annotations

import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))


class QuestionReviewDocxFigureTests(unittest.TestCase):
    def test_snapshot_migrates_relative_images_and_embeds_media(self) -> None:
        from app.question_review_docx import _snapshot_question

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            stage = tmp / "stage"
            output = tmp / "output"
            (stage / "figures").mkdir(parents=True)
            Image.new("RGB", (320, 180), "white").save(stage / "figures" / "q1.png")
            fragment = {
                "question_id": "q1",
                "number": "1",
                "section": "作图题",
                "question_type": "作图题",
                "answer": "见图",
                "blocks": [{
                    "label": "图示",
                    "segments": [{"type": "image_ref", "image_id": "q1_fig", "path": "figures/q1.png"}],
                }],
            }
            with patch("app.question_review_docx.export_docx_to_pdf", lambda _docx, pdf: pdf.write_bytes(b"pdf")), patch(
                "app.question_review_docx.render_pdf_to_png", return_value=[]
            ):
                _snapshot_question(stage, output, fragment)

            snapshot = output / "question_review_snapshots" / "q1"
            self.assertTrue((snapshot / "figures" / "q1.png").is_file())
            with zipfile.ZipFile(snapshot / "q1.docx") as archive:
                self.assertTrue(any(name.startswith("word/media/") for name in archive.namelist()))

    def test_missing_snapshot_image_blocks_instead_of_writing_placeholder(self) -> None:
        from app.question_review_docx import _snapshot_question

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            fragment = {
                "question_id": "q1",
                "blocks": [{
                    "label": "图示",
                    "segments": [{"type": "image_ref", "image_id": "missing_fig", "path": "figures/missing.png"}],
                }],
            }
            with self.assertRaisesRegex(FileNotFoundError, "missing_fig"):
                _snapshot_question(tmp / "stage", tmp / "output", fragment)
            self.assertFalse((tmp / "output" / "question_review_snapshots" / "q1" / "q1.docx").exists())

    def test_prepare_figures_archives_stage_images_for_review_docx(self) -> None:
        from app.figures import prepare_figures_for_fragments
        from app.question_review_docx import collect_question_figure_review_items

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            fragments_json = tmp / "answer_fragments.json"
            fragments_json.write_text(
                json.dumps(
                    {
                        "fragments": [
                            {
                                "question_id": "q1",
                                "answer": "见图。",
                                "blocks": [],
                                "_draft": {
                                    "figure_specs": [
                                        {
                                            "kind": "line_chart",
                                            "figure_id": "q1_fig_01",
                                            "caption": "趋势图",
                                            "points": [[0, 0], [1, 1]],
                                        }
                                    ]
                                },
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            structured_exam = {
                "items": [
                    {
                        "question_id": "q1",
                        "question_type": "作图题",
                        "drawing_generation_mode": "figure_specs",
                        "stem": "画趋势图。",
                    }
                ]
            }
            (tmp / "structured_exam.json").write_text(json.dumps(structured_exam, ensure_ascii=False), encoding="utf-8")

            prepare_figures_for_fragments(structured_exam, fragments_json, tmp / "figure_specs.json", tmp / "figures")
            manifest = json.loads((tmp / "figure_stage_images.json").read_text(encoding="utf-8"))
            items = collect_question_figure_review_items(tmp)

        self.assertEqual("answer_book.figure_stage_images.v1", manifest["schema_version"])
        self.assertTrue(any(item.get("stage") == "initial_render" for item in manifest["items"]))
        self.assertTrue(
            any(
                any("阶段归档/initial_render" in stage for stage in figure["stages"])
                for figure in items[0]["figures"]
            )
        )

    def test_figure_review_docx_is_separate_from_question_review_docx(self) -> None:
        from app.question_review_docx import build_figure_review_docx, build_question_review_docx, collect_question_figure_review_items

        with tempfile.TemporaryDirectory() as raw_tmp:
            tmp = Path(raw_tmp)
            stage = tmp / "stage"
            output = tmp / "out"
            figures = stage / "figures"
            candidate = stage / "figure_visual_qa_candidates" / "q1" / "round_1" / "vision_reviewer"
            figures.mkdir(parents=True)
            candidate.mkdir(parents=True)
            Image.new("RGB", (900, 500), "white").save(figures / "q1_fig_01.png")
            Image.new("RGB", (900, 500), "white").save(candidate / "q1_fig_01.png")

            (stage / "structured_exam.json").write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "question_id": "q1",
                                "section": "一",
                                "number": "1",
                                "question_type": "作图题",
                                "stem": "画出示意图。",
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (stage / "answer_fragments.json").write_text(
                json.dumps(
                    {
                        "fragments": [
                            {
                                "question_id": "q1",
                                "answer": "见图。",
                                "blocks": [
                                    {
                                        "label": "图示",
                                        "segments": [
                                            {
                                                "type": "image_ref",
                                                "image_id": "q1_fig_01",
                                                "path": "figures/q1_fig_01.png",
                                            }
                                        ],
                                    }
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (stage / "figure_specs.json").write_text(
                json.dumps(
                    {
                        "figures": [
                            {
                                "question_id": "q1",
                                "figure_id": "q1_fig_01",
                                "kind": "model_drawing_code",
                                "source": "independent_code_generator",
                                "run_result": {"ok": True},
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (stage / "figure_visual_qa.json").write_text(
                json.dumps(
                    {
                        "enabled": True,
                        "items": [
                            {
                                "question_id": "q1",
                                "figure_id": "q1_fig_01",
                                "path": str(figures / "q1_fig_01.png"),
                                "qa": {"ok": True, "summary": "正式图通过。"},
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (stage / "figure_visual_qa_repair.json").write_text(
                json.dumps(
                    {
                        "enabled": True,
                        "rounds": [
                            {
                                "round": 1,
                                "targets": [
                                    {
                                        "question_id": "q1",
                                        "figure_id": "q1_fig_01",
                                        "candidates": [
                                            {
                                                "strategy": "vision_reviewer",
                                                "passed": False,
                                                "status": "visual_qa_failed",
                                                "path": str(candidate / "q1_fig_01.png"),
                                                "repair_notes": ["候选图仍需复核。"],
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            items = collect_question_figure_review_items(stage)
            question_review_path = build_question_review_docx(stage, output, render_snapshots=False)
            figure_review_path = build_figure_review_docx(stage, output)
            question_review = Document(question_review_path)
            figure_review = Document(figure_review_path)

        question_text = "\n".join(paragraph.text for paragraph in question_review.paragraphs)
        figure_text = "\n".join(paragraph.text for paragraph in figure_review.paragraphs)
        self.assertEqual(1, len(items))
        self.assertEqual(1, len(items[0]["figures"]))
        self.assertIn("正式采用", items[0]["figures"][0]["statuses"])
        self.assertIn("未采用（视觉QA未通过）", items[0]["figures"][0]["statuses"])
        self.assertIn("未发现需要单独审查的题目", question_text)
        self.assertNotIn("作图题全流程图片", question_text)
        self.assertIn("作图题全流程图片", figure_text)
        self.assertEqual(1, len(figure_review.inline_shapes))
        self.assertIn("正式采用", figure_text)
        self.assertIn("未采用（视觉QA未通过）", figure_text)


if __name__ == "__main__":
    unittest.main()
