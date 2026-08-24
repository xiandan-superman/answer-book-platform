from __future__ import annotations

import json

from docx import Document

from app.docx_v4 import build_docx_from_fragments
from app.question_types import (
    infer_question_type,
    is_calculation_question,
    is_short_answer_question,
    normalize_question_type,
    question_has_type,
    question_kind,
)


def test_term_explanation_type_normalization() -> None:
    assert normalize_question_type("名词解释题") == "名词解释"
    assert infer_question_type({"section": "一、名词解释题", "stem": "点阵畸变"}) == "名词解释"
    assert question_kind({"question_type": "名词解释"}) == "term_explanation"


def test_mixed_calculation_section_preserves_unanimous_short_answer_group() -> None:
    item = {
        "question_id": "calc_s04_01",
        "section": "四、计算与综合分析题",
        "section_raw": "四、计算与综合分析题",
        "question_type": "计算题",
        "confirmed_question_type": "计算题",
        "stem": "扩散是材料科学中的重要过程，请回答下列问题。",
        "subquestions": [
            {"number": "(1)", "question_type": "简答题", "stem": "说明扩散机制。"},
            {"number": "(2)", "question_type": "简答题", "stem": "讨论影响因素。"},
        ],
    }

    assert infer_question_type(item) == "简答题"
    assert question_kind(item) == "short_answer"
    assert question_has_type(item, "简答题") is True
    assert question_has_type(item, "计算题") is False
    assert is_short_answer_question(item) is True
    assert is_calculation_question(item) is False


def test_term_explanation_docx_renders_evidence_and_answer_only(tmp_path) -> None:
    fragments = {
        "schema_version": "answer_book.answer_fragments.v4",
        "fragments": [
            {
                "schema_version": "answer_book.answer_fragment.v4",
                "question_id": "term_s01_01",
                "section": "一、名词解释题",
                "question_type": "名词解释",
                "number": "1",
                "answer": "点阵畸变：晶体点阵偏离理想周期排列的局部畸变。",
                "answer_summary": "点阵畸变：晶体点阵偏离理想周期排列的局部畸变。",
                "evidence_ids": ["ev1"],
                "formulas": [],
                "blocks": [
                    {
                        "label": "教材依据",
                        "segments": [{"type": "text", "text": "材料科学基础：晶体缺陷相关定义。"}],
                    },
                    {
                        "label": "解析",
                        "segments": [{"type": "text", "text": "这段解析不应进入名词解释题最终 Word。"}],
                    },
                ],
            }
        ],
    }
    source = tmp_path / "fragments.json"
    output = tmp_path / "answer.docx"
    source.write_text(json.dumps(fragments, ensure_ascii=False), encoding="utf-8")

    build_docx_from_fragments(source, output)

    paragraphs = [p.text for p in Document(output).paragraphs if p.text.strip()]
    doc = Document(output)
    joined = "\n".join(paragraphs)
    assert "1、教材依据：材料科学基础：晶体缺陷相关定义。" in joined
    assert "答案：" in paragraphs
    assert "点阵畸变：晶体点阵偏离理想周期排列的局部畸变。" in paragraphs
    answer_index = paragraphs.index("答案：")
    answer_body_text = paragraphs[answer_index + 1]
    answer_body = next(p for p in doc.paragraphs if p.text == answer_body_text)
    assert answer_body.paragraph_format.first_line_indent is not None
    assert "解析：" not in joined
    assert "这段解析不应进入名词解释题最终 Word" not in joined
