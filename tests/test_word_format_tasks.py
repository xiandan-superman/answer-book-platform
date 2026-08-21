from __future__ import annotations

import base64
from pathlib import Path

from docx import Document

from app import word_format_tasks


def _docx_bytes(path: Path) -> bytes:
    document = Document()
    document.add_heading("测试标题", level=1)
    document.add_paragraph("这是一段需要按标准统一的正文。English 123")
    document.save(path)
    return path.read_bytes()


def _configure_storage(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(word_format_tasks, "TASKS_DIR", tmp_path / "tasks")
    monkeypatch.setattr(word_format_tasks, "SETTINGS_FILE", tmp_path / "config" / "word_format_settings.json")


def test_review_task_persists_and_can_be_applied_downloaded_and_deleted(tmp_path: Path, monkeypatch) -> None:
    _configure_storage(tmp_path, monkeypatch)
    content = _docx_bytes(tmp_path / "input.docx")

    created = word_format_tasks.create_word_format_task(
        {
            "filename": "示例答案.docx",
            "content_base64": base64.b64encode(content).decode("ascii"),
            "profile": "answer",
            "mode": "review",
            "header_text": "测试页眉",
        }
    )

    task_id = created["task_id"]
    assert created["status"] == "needs_input"
    assert created["download_url"] is None
    assert created["report"]["source_name"] == "示例答案.docx"
    assert created["task"]["task_kind"] == "format"
    assert created["task"]["capabilities"]["view_result"] is True
    assert [task["task_id"] for task in word_format_tasks.list_word_format_tasks()] == [task_id]

    applied = word_format_tasks.apply_word_format_task(task_id)
    assert applied["status"] in {"completed", "completed_with_issues"}
    assert applied["final_report"]["source_name"] == "示例答案.docx"
    assert applied["download_url"].endswith(f"/{task_id}/download")
    output, filename = word_format_tasks.word_format_download_path(task_id)
    assert output.exists()
    assert filename == "示例答案_格式已修改.docx"

    deleted = word_format_tasks.delete_word_format_task(task_id)
    assert deleted["ok"] is True
    assert word_format_tasks.list_word_format_tasks() == []


def test_auto_task_finishes_with_download_and_profile_settings_are_persistent(tmp_path: Path, monkeypatch) -> None:
    _configure_storage(tmp_path, monkeypatch)
    settings = word_format_tasks.word_format_settings_payload()["profiles"]["lecture"]["effective"]
    settings["styles"]["Normal"]["chinese_font"] = "微软雅黑"
    saved = word_format_tasks.save_word_format_profile_settings("lecture", settings)

    content = _docx_bytes(tmp_path / "lecture.docx")
    created = word_format_tasks.create_word_format_task(
        {
            "filename": "讲义.docx",
            "content_base64": base64.b64encode(content).decode("ascii"),
            "profile": "lecture",
            "mode": "auto",
        }
    )

    assert saved["styles"]["Normal"]["chinese_font"] == "微软雅黑"
    assert word_format_tasks.word_format_settings_payload()["profiles"]["lecture"]["effective"]["styles"]["Normal"]["chinese_font"] == "微软雅黑"
    assert created["status"] in {"completed", "completed_with_issues"}
    assert created["download_url"]
    assert created["task"]["capabilities"]["download"] is True


def test_invalid_upload_does_not_create_a_task(tmp_path: Path, monkeypatch) -> None:
    _configure_storage(tmp_path, monkeypatch)
    try:
        word_format_tasks.create_word_format_task(
            {
                "filename": "not-docx.docx",
                "content_base64": base64.b64encode(b"not a zip").decode("ascii"),
                "profile": "answer",
                "mode": "review",
            }
        )
    except ValueError as exc:
        assert "DOCX" in str(exc)
    else:
        raise AssertionError("无效文件应被拒绝")
    assert not (tmp_path / "tasks").exists()
