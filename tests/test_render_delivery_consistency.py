from __future__ import annotations

import json
from unittest.mock import patch

import pytest
from docx import Document
from PIL import Image

from app.capabilities.expression_rendering import render_expression_omml
from app.docx_v4 import build_docx_from_fragments
from app.practice_export import build_practice_question_docx
from app.render_audit import audit_docx_pdf_consistency

NASA_MIXED_PROSE = "Explain which quantities are connected by the ideal-gas equation pV=nRT."


def _write_docx(path, *, with_drawing: bool = False) -> str:
    text = "这是一段用于确认正式Word与渲染PDF来自同一版本的足够长文本，任何旧文件都不应通过一致性检查。"
    document = Document()
    document.add_paragraph(text)
    if with_drawing:
        image = path.with_suffix(".png")
        Image.new("RGB", (80, 60), "white").save(image)
        document.add_picture(str(image))
    document.save(path)
    return text


def test_consistency_audit_blocks_stale_pdf_text(tmp_path) -> None:
    docx = tmp_path / "answer.docx"
    pdf = tmp_path / "answer.pdf"
    _write_docx(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=("完全不同的旧版内容", 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert not report["ok"]
    assert report["compared"] is True
    assert report["comparison_status"] == "compared"
    assert any("does not match current DOCX" in issue for issue in report["issues"])


def test_consistency_audit_accepts_matching_text_and_blocks_omitted_drawing(tmp_path) -> None:
    docx = tmp_path / "answer.docx"
    pdf = tmp_path / "answer.pdf"
    text = _write_docx(docx, with_drawing=True)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=(text, 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["anchor_match_ratio"] == 1.0
    assert report["compared"] is True
    assert not report["ok"]
    assert any("image count" in issue for issue in report["issues"])


def _write_mixed_formula_docx(path) -> None:
    document = Document()
    paragraph = document.add_paragraph("Explain which quantities are connected by the ideal-gas equation ")
    paragraph._p.append(render_expression_omml("pV=nRT", display=False, location="test"))
    paragraph.add_run(".")
    document.save(path)


def _write_true_exam_docx(path, tmp_path) -> None:
    fragments = tmp_path / "fragments.json"
    fragments.write_text(
        json.dumps(
            {
                "fragments": [
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "q_nasa",
                        "display_number": "1",
                        "section": "简答题",
                        "question_type": "简答题",
                        "answer": "见解析",
                        "evidence_ids": [],
                        "formulas": [{"formula_id": "f1", "latex": "pV=nRT"}],
                        "blocks": [
                            {
                                "label": "解析",
                                "segments": [
                                    {
                                        "type": "text",
                                        "text": "Explain which quantities are connected by the ideal-gas equation ",
                                    },
                                    {"type": "formula_ref", "formula_id": "f1", "inline": True},
                                    {"type": "text", "text": "."},
                                ],
                            }
                        ],
                    }
                ]
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    build_docx_from_fragments(fragments, path)


@pytest.mark.parametrize("source_mode", ["exam", "knowledge"])
def test_practice_delivery_paths_compare_mixed_prose(source_mode, tmp_path) -> None:
    docx = tmp_path / f"{source_mode}.docx"
    pdf = tmp_path / f"{source_mode}.pdf"
    docx.write_bytes(
        build_practice_question_docx(
            {
                "source_mode": source_mode,
                "exercises": [{"number": 1, "question_type": "简答题", "stem": NASA_MIXED_PROSE}],
            }
        )
    )
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=("stale unrelated output", 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["compared"] is True
    assert report["anchor_count"] > 0
    assert not report["ok"]


def test_true_exam_delivery_path_compares_mixed_prose(tmp_path) -> None:
    docx = tmp_path / "true_exam.docx"
    pdf = tmp_path / "true_exam.pdf"
    _write_true_exam_docx(docx, tmp_path)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=("stale unrelated output", 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["compared"] is True
    assert report["anchor_count"] > 0
    assert not report["ok"]


def test_mixed_formula_paragraph_keeps_prose_comparable(tmp_path) -> None:
    docx = tmp_path / "mixed.docx"
    pdf = tmp_path / "mixed.pdf"
    _write_mixed_formula_docx(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=(NASA_MIXED_PROSE, 0)):
        matching = audit_docx_pdf_consistency(docx, pdf)
    with patch("app.render_audit._pdf_text_and_image_count", return_value=("stale unrelated output", 0)):
        stale = audit_docx_pdf_consistency(docx, pdf)

    assert matching["text_comparison_status"] == "compared"
    assert matching["anchor_match_ratio"] == 1.0
    assert matching["ok"]
    assert not stale["ok"]


def test_table_long_text_is_comparable(tmp_path) -> None:
    text = "表格中的这一段教材说明足够长，必须参与正式Word与PDF的一致性比较，不能因位于单元格而被忽略。"
    docx = tmp_path / "table.docx"
    pdf = tmp_path / "table.pdf"
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = text
    document.save(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=(text, 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["text_comparison_status"] == "compared"
    assert report["anchor_match_ratio"] == 1.0
    assert report["ok"]


def test_short_answer_is_comparable(tmp_path) -> None:
    answer = "光合作用产生氧气"
    docx = tmp_path / "short.docx"
    pdf = tmp_path / "short.pdf"
    document = Document()
    document.add_paragraph(answer)
    document.save(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=(answer, 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["anchor_count"] == 1
    assert report["text_comparison_status"] == "compared"
    assert report["ok"]


def test_pure_formula_is_comparable(tmp_path) -> None:
    docx = tmp_path / "formula.docx"
    pdf = tmp_path / "formula.pdf"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph._p.append(render_expression_omml("pV=nRT", display=True, location="test"))
    document.save(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=("𝑝𝑉 = 𝑛𝑅𝑇", 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["anchor_count"] == 1
    assert report["text_comparison_status"] == "not_comparable"
    assert report["formula_comparison_status"] == "compared"
    assert report["comparison_status"] == "compared"
    assert report["compared"] is True
    assert report["anchor_match_ratio"] == 1.0
    assert report["ok"]


def test_pure_image_uses_image_comparison_without_text_anchor(tmp_path) -> None:
    docx = tmp_path / "image.docx"
    pdf = tmp_path / "image.pdf"
    image = tmp_path / "source.png"
    Image.new("RGB", (80, 60), "white").save(image)
    document = Document()
    document.add_picture(str(image))
    document.save(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=("", 1)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["anchor_count"] == 0
    assert report["text_comparison_status"] == "not_comparable"
    assert report["image_comparison_status"] == "compared"
    assert report["comparison_status"] == "compared"
    assert report["ok"]


def test_no_comparable_content_is_explicit_but_not_blocked(tmp_path) -> None:
    docx = tmp_path / "empty.docx"
    pdf = tmp_path / "empty.pdf"
    Document().save(docx)
    pdf.write_bytes(b"placeholder")

    with patch("app.render_audit._pdf_text_and_image_count", return_value=("", 0)):
        report = audit_docx_pdf_consistency(docx, pdf)

    assert report["compared"] is False
    assert report["comparison_status"] == "not_comparable"
    assert report["text_comparison_status"] == "not_comparable"
    assert report["formula_comparison_status"] == "not_comparable"
    assert report["image_comparison_status"] == "not_comparable"
    assert report["ok"] is True
    assert report["issues"] == []
