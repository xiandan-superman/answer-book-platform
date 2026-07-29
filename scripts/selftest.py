#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
import tempfile
import time
import zipfile
import base64
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

import app.answer_generation as answer_generation_module
import app.evidence_selection as evidence_selection_module
import app.figures as figures_module
import app.knowledge_planning as knowledge_planning_module
import app.task_store as task_store_module
from app.concurrency import run_limited_concurrent
from app.formula_audit import audit_text_segments_no_formula, looks_like_formula
from app.answer_generation import answer_generation_worker_count, attach_program_evidence_block, bind_top_evidence, evidence_for_answer_generation, fragment_from_analysis_draft, has_bound_evidence, semantic_generation_issues
from app.audit_model_repair import collect_audit_issue_targets
from app.audit_review_gate import apply_allowed_to_audit_report
from app.content_quality_audit import audit_content_quality
from app.docx_audit import audit_docx_v4
from app.docx_v4 import build_docx_from_fragments
from app.docx_model_repair import repair_fragments_with_model_for_docx
from app.delivery_package import build_task_delivery_package
from app.exam_audit import audit_exam_structure
from app.exam_extract import extract_exam_structure, question_items
from app.exam_structure_review import apply_exam_structure_review_updates, build_exam_structure_review_request
from app.question_types import QUESTION_TYPES, infer_question_type, question_kind
from app.evidence_selection import citation_groups_from_selection, confirm_evidence_selection, filter_candidates_by_selection
from app.final_acceptance import build_final_acceptance_report
from app.fragment_repair import repair_answer_fragments_for_docx
from app.library_files import scan_library_files
from app.textbook_index_cache import prepare_textbook_index_cache
from app.llm_client import LLMError, LLMResult, OpenAICompatibleClient, parse_json_content
from app.omml import latex_to_plain
from app.paths import CONFIG_DIR
from app.pipeline import (
    _filter_audit_report_for_model_repair,
    _filter_docx_issues_for_model_repair,
    build_user_allowed_docx_candidate,
    build_user_allowed_docx_placeholder,
)
from app.prompts import build_answer_draft_prompt
from app.question_understanding import (
    build_question_understanding,
    is_complex_table,
    needs_vision_model,
    render_table_to_image,
)
from app.question_review_docx import build_question_review_docx, collect_question_review_items
from app.retrieval import EvidenceCandidate, build_candidates, is_invalid_evidence_row, retrieval_query_text
from app.review_notes import build_answer_review_notes
from app.settings import ProviderConfig
from app.server import _task_duration_summary
from app.task_diagnostics import build_task_diagnostics
from app.task_store import TaskRecord
from app.textbook_index import audit_page_map, build_page_map, enrich_rows_with_sections, write_csv, BLOCK_FIELDS, PAGE_MAP_FIELDS
from app.v4_schema import validate_v4_answer_fragment


def main() -> int:
    good = {
        "schema_version": "answer_book.answer_fragment.v4",
        "question_id": "demo_01",
        "answer": "A",
        "evidence_ids": ["ev_demo_01"],
        "blocks": [
            {"label": "解析", "segments": [
                {"type": "text", "text": "根据教材证据，反应的吉布斯函数与电动势关系如下。"},
                {"type": "formula_ref", "formula_id": "f_demo_01"}
            ]}
        ],
        "formulas": [
            {"formula_id": "f_demo_01", "latex": "\\Delta_r G_m=-nFE", "role": "relation", "display": True}
        ],
        "warnings": []
    }
    bad = {
        "schema_version": "answer_book.answer_fragment.v4",
        "question_id": "demo_02",
        "answer": "A",
        "evidence_ids": ["ev_demo_02"],
        "blocks": [
            {"label": "解析", "segments": [
                {"type": "text", "text": "由 ΔrGm=-nFE 可得。"}
            ]}
        ],
        "formulas": [],
        "warnings": []
    }
    good_issues = validate_v4_answer_fragment(good)
    bad_issues = validate_v4_answer_fragment(bad)
    result = {
        "formula_detector_examples": {
            "plain_text": looks_like_formula("根据教材证据可知该过程自发。"),
            "formula_text": looks_like_formula("由 ΔrGm=-nFE 可得。"),
            "chinese_formula_paraphrase": looks_like_formula("吉布斯自由能变小于零时反应自发，等于焓变减去温度与熵变的乘积。"),
        },
        "good_fragment_issues": good_issues,
        "bad_fragment_issues": bad_issues,
    }
    started: list[int] = []
    finished: list[int] = []

    def concurrent_worker(value: int) -> int:
        started.append(value)
        time.sleep(0.04 if value == 1 else 0.01)
        finished.append(value)
        return value * 10

    concurrent_start = time.perf_counter()
    result["limited_concurrency"] = {
        "values": run_limited_concurrent([1, 2, 3], concurrent_worker, max_workers=2),
        "started": started,
        "finished": finished,
        "elapsed": round(time.perf_counter() - concurrent_start, 3),
    }
    result["answer_generation_default_workers"] = answer_generation_worker_count()
    multipart_section = {
        "cn": "三",
        "major_no": 3,
        "raw_title": "三、计算题",
        "title_tail": "计算题",
        "body": [
            "1、根据相图回答下列问题：",
            "1. 画出冷却曲线。",
            "2. 用杠杆定律计算组织组成。",
            "3. 分析不同冷却条件下的性能差异。",
            "2、说明扩散机制对组织演变的影响。",
        ],
        "subject_index": 1,
        "subject": "",
    }
    multipart_items = question_items(multipart_section)
    arabic_comma_subquestion_section = {
        "cn": "二",
        "major_no": 2,
        "raw_title": "二、简答题",
        "title_tail": "简答题",
        "body": [
            "1、围绕电化学体系回答下列问题：",
            "1、判断该电池是否可逆。",
            "2、说明盐桥的作用。",
        ],
        "subject_index": 1,
        "subject": "",
    }
    arabic_comma_subquestion_items = question_items(arabic_comma_subquestion_section)
    unnumbered_intro_subquestion_section = {
        "cn": "三",
        "major_no": 3,
        "raw_title": "三、计算题（本题共15分）",
        "title_tail": "计算题",
        "body": [
            "将1mol苯在正常沸点353K和10kPa压力下向真空蒸发，后恒温可逆膨胀至50.5kPa。试求:",
            "1. 整个过程的Q、W、∆S和∆G；",
            "2. 整个过程环境的熵变∆S环；",
            "3. 根据计算结果，判断上述过程的可逆性；",
        ],
        "subject_index": 1,
        "subject": "",
    }
    unnumbered_intro_items = question_items(unnumbered_intro_subquestion_section)
    nested_requirement_section = {
        "cn": "九",
        "major_no": 9,
        "raw_title": "九、相图分析题（本题共14分）",
        "title_tail": "相图分析题",
        "body": [
            "1、题九4图为 Fe-Fe3C 合金相图。",
            "(1)相图中三个恒温转变都是什么转变？其转变产物是什么？（2分）",
            "(2)画出 3.5%C 的 A 合金平衡凝固冷却曲线和室温平衡凝固组织示意图，计算其室温组织组成和组成质量比。（8分）",
        ],
        "subject_index": 1,
        "subject": "",
    }
    nested_requirement_items = question_items(nested_requirement_section)
    result["multipart_question_extraction"] = {
        "item_count": len(multipart_items),
        "first_stem": multipart_items[0].get("stem") if multipart_items else "",
        "first_subquestions": multipart_items[0].get("subquestions") if multipart_items else [],
        "second_number": multipart_items[1].get("number") if len(multipart_items) > 1 else "",
        "comma_item_count": len(arabic_comma_subquestion_items),
        "comma_subquestions": arabic_comma_subquestion_items[0].get("subquestions") if arabic_comma_subquestion_items else [],
        "unnumbered_intro_count": len(unnumbered_intro_items),
        "unnumbered_intro_stem": unnumbered_intro_items[0].get("stem") if unnumbered_intro_items else "",
        "unnumbered_intro_subquestions": unnumbered_intro_items[0].get("subquestions") if unnumbered_intro_items else [],
        "nested_requirement_subquestions": nested_requirement_items[0].get("subquestions") if nested_requirement_items else [],
    }
    review_exam = {
        "items": [
            {
                "question_id": "calc_review_type_01",
                "major_number": "3",
                "section": "三、计算题",
                "section_raw": "三、计算题（本题共15分）",
                "number": "1",
                "stem": "说明材料组织演变规律。",
            },
            {
                "question_id": "qa_review_type_02",
                "major_number": "2",
                "section": "二、简答题",
                "section_raw": "二、简答题",
                "number": "2",
                "stem": "简述盐桥作用。",
            },
        ]
    }
    review_request = build_exam_structure_review_request("review_task", review_exam)
    reviewed_exam = apply_exam_structure_review_updates(
        review_exam,
        [
            {"question_id": "calc_review_type_01", "question_type": "简答题"},
            {"question_id": "qa_review_type_02", "question_type": "计算题"},
        ],
    )
    result["exam_structure_review"] = {
        "request_item_count": len(review_request.get("items", [])),
        "request_first_type": review_request.get("items", [{}])[0].get("question_type"),
        "first_question_type": reviewed_exam["items"][0].get("question_type"),
        "first_section": reviewed_exam["items"][0].get("section"),
        "first_section_raw": reviewed_exam["items"][0].get("section_raw"),
        "first_extracted_section": reviewed_exam["items"][0].get("extracted_section"),
        "first_reviewed": reviewed_exam["items"][0].get("type_reviewed"),
        "second_question_type": reviewed_exam["items"][1].get("question_type"),
        "second_section": reviewed_exam["items"][1].get("section"),
    }
    subquestion_review_exam = {
        "items": [
            {
                "question_id": "mixed_review_type_01",
                "major_number": "4",
                "section": "四、简答题",
                "section_raw": "四、简答题",
                "number": "1",
                "question_type": "简答题",
                "stem": "回答下列问题。",
                "subquestions": [
                    {"number": "1", "marker": "1.", "stem": "计算焓变。"},
                    {"number": "2", "marker": "2.", "stem": "画出示意图。"},
                ],
            }
        ]
    }
    subquestion_request = build_exam_structure_review_request("subquestion_review_task", subquestion_review_exam)
    subquestion_reviewed = apply_exam_structure_review_updates(
        subquestion_review_exam,
        [
            {
                "question_id": "mixed_review_type_01",
                "question_type": "简答题",
                "subquestions": [
                    {"number": "1", "question_type": "计算题"},
                    {"number": "2", "question_type": "作图题"},
                ],
            }
        ],
    )
    subquestion_structure_edited = apply_exam_structure_review_updates(
        subquestion_review_exam,
        [
            {
                "question_id": "mixed_review_type_01",
                "question_type": "计算题",
                "subquestions": [
                    {"number": "1", "stem": "计算焓变并写出单位。", "question_type": "计算题"},
                    {"number": "3", "stem": "判断过程是否可逆。", "question_type": "简答题"},
                ],
            }
        ],
    )
    result["question_type_contract"] = {
        "types": list(QUESTION_TYPES),
        "inferred_choice": infer_question_type({"section": "一、选择题", "stem": "下列说法正确的是"}),
        "normalized_drawing": infer_question_type({"question_type": "图示题", "stem": "画图"}),
        "kind_calc": question_kind({"question_type": "计算题"}),
        "kind_mixed": question_kind(subquestion_reviewed["items"][0]),
        "kind_nested_mixed": question_kind(nested_requirement_items[0]) if nested_requirement_items else "",
        "request_subquestion_type": subquestion_request["items"][0]["subquestions"][0]["question_type"],
        "subquestion_first_type": subquestion_reviewed["items"][0]["subquestions"][0]["question_type"],
        "subquestion_second_type": subquestion_reviewed["items"][0]["subquestions"][1]["question_type"],
        "subquestion_structure_numbers": [sub.get("number") for sub in subquestion_structure_edited["items"][0]["subquestions"]],
        "subquestion_structure_first_stem": subquestion_structure_edited["items"][0]["subquestions"][0].get("stem"),
    }
    with tempfile.TemporaryDirectory() as table_tmp:
        table_tmp_path = Path(table_tmp)
        table_docx = table_tmp_path / "table_exam.docx"
        table_doc = Document()
        table_doc.add_paragraph("三、计算题")
        table_doc.add_paragraph("1、已知电池：Pt|H2(p=101kPa)|NaOH(aq,稀)|Bi2O3(s)|Bi，求反应焓。")
        table_doc.add_paragraph("附表.常数及对数")
        doc_table = table_doc.add_table(rows=2, cols=6)
        table_rows = [
            ["量", "R/J·K-1·mol-1", "F/C·mol-1", "0℃", "ln10", "ln2"],
            ["值", "8.314", "96500", "273K", "2.303", "0.693"],
        ]
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                doc_table.cell(row_index, col_index).text = value
        table_doc.save(table_docx)
        table_exam = extract_exam_structure(table_docx, table_tmp_path / "structured_exam.json")
        table_item = table_exam["items"][0]
        table_understanding = build_question_understanding(table_item, table_tmp_path / "understanding")
        table_prompt_payload = json.loads(build_answer_draft_prompt(table_item, [], table_understanding)[1]["content"])
        complex_table = {
            "rows": [
                ["条件", "A\nB", "C/D"],
                ["T/K", "300", "350", "400"],
                ["k / s^-1", "1.2×10^-3", "2.4×10^-3", "5.1×10^-3"],
                ["备注", "含上下标与合并表头语义", "", ""],
            ]
        }
        complex_table_image = render_table_to_image(complex_table, table_tmp_path / "complex_table.png")
        result["docx_table_attachment_extraction"] = {
            "table_count": len((table_item.get("attachments") or {}).get("tables", [])),
            "rows": ((table_item.get("attachments") or {}).get("tables", [{}])[0].get("rows") or []),
            "stem_has_marker": "__ANSWER_BOOK_TABLE__" in table_item.get("stem", ""),
            "source_has_table_text": any("R/J·K-1·mol-1" in line and "96500" in line for line in table_exam.get("source_paragraphs", [])),
            "prompt_has_question_understanding": bool(table_prompt_payload.get("question_understanding")),
            "understanding_has_table_rows": bool((table_understanding.get("tables") or [{}])[0].get("table_rows")),
            "simple_table_needs_vision": needs_vision_model(table_item),
            "complex_table_detected": is_complex_table(complex_table),
            "complex_table_render_exists": complex_table_image.exists(),
        }
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        structured_for_selection = {
            "items": [
                {"question_id": "q_parallel_1", "number": "1", "stem": "题目1"},
                {"question_id": "q_parallel_2", "number": "2", "stem": "题目2"},
                {"question_id": "q_parallel_3", "number": "3", "stem": "题目3"},
            ]
        }
        plans_for_selection = {
            item["question_id"]: {"question_id": item["question_id"], "knowledge_points": ["考查点"]}
            for item in structured_for_selection["items"]
        }
        candidates_for_selection = [
            EvidenceCandidate(
                evidence_id=f"ev_parallel_{index}",
                question_id=item["question_id"],
                textbook="教材",
                citation_textbook="教材",
                chapter_section="",
                source_file="",
                pdf_page_idx=str(index),
                printed_page=str(index),
                score=1.0,
                evidence_text="候选依据",
                verified_page=True,
                knowledge_point="考查点",
            )
            for index, item in enumerate(structured_for_selection["items"], start=1)
        ]
        original_select_one = evidence_selection_module._select_one
        evidence_finished: list[str] = []

        def fake_select_one(client, provider_config, model_name, question, plan, candidates, expanded=False):
            qid = str(question.get("question_id"))
            time.sleep(0.04 if qid.endswith("_1") else 0.005)
            evidence_finished.append(qid)
            selected_id = candidates[0].evidence_id if candidates else ""
            return {
                "question_id": qid,
                "knowledge_points": [
                    {
                        "knowledge_point": "考查点",
                        "selected_evidence_ids": [selected_id] if selected_id else [],
                        "rejected_evidence_ids": [],
                        "reason": "测试并发证据确认。",
                        "no_suitable_evidence_reason": "",
                        "needs_expansion": False,
                    }
                ],
                "citation_groups": [],
                "_meta": {"llm_retry": {"ok": True, "attempts": []}},
            }

        try:
            evidence_selection_module._select_one = fake_select_one
            selection_provider = ProviderConfig(
                name="test",
                type="openai_compatible",
                base_url="http://example.invalid",
                api_key="",
                default_model="test-model",
                model_options=("test-model",),
                allow_custom_model=True,
                model_hint="",
                temperature=0.1,
                max_tokens=4096,
            )
            evidence_result, confirmed = confirm_evidence_selection(
                structured_for_selection,
                plans_for_selection,
                candidates_for_selection,
                selection_provider,
                "test-model",
                tmp_path / "evidence_selection.json",
                tmp_path / "blocks.csv",
                tmp_path / "page_map.csv",
                progress_json=tmp_path / "evidence_selection_progress.json",
                use_model=False,
            )
        finally:
            evidence_selection_module._select_one = original_select_one
        evidence_output = json.loads((tmp_path / "evidence_selection.json").read_text(encoding="utf-8"))
        evidence_progress = json.loads((tmp_path / "evidence_selection_progress.json").read_text(encoding="utf-8"))
    result["evidence_selection_parallel"] = {
        "finished": evidence_finished,
        "worker_count": evidence_output.get("concurrency", {}).get("max_workers"),
        "progress_worker_count": evidence_progress.get("max_workers"),
        "parallel_enabled": evidence_output.get("concurrency", {}).get("parallel_enabled"),
        "confirmed_count": len(confirmed),
        "selected_question_count": evidence_result.selected_question_count,
    }
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        original_client = knowledge_planning_module.OpenAICompatibleClient
        original_workers_env = knowledge_planning_module.os.environ.get("KNOWLEDGE_PLANNING_MAX_WORKERS")
        knowledge_finished: list[str] = []

        class FakeKnowledgeClient:
            def __init__(self, provider):
                self.last_json_retry_report = {}

            def chat_json_object(self, messages, **kwargs):
                payload = json.loads(messages[-1]["content"])
                qid = payload["question"]["question_id"]
                time.sleep(0.04 if qid.endswith("_1") else 0.005)
                knowledge_finished.append(qid)
                self.last_json_retry_report = {"ok": True, "attempts": [{"strategy": "primary", "model": kwargs.get("model")}]}
                return {
                    "question_id": qid,
                    "knowledge_points": ["考点"],
                    "formulas": [],
                    "key_terms": ["术语"],
                    "search_queries": ["考点 术语"],
                    "warnings": [],
                }

        try:
            knowledge_planning_module.OpenAICompatibleClient = FakeKnowledgeClient
            knowledge_planning_module.os.environ["KNOWLEDGE_PLANNING_MAX_WORKERS"] = "3"
            provider = ProviderConfig(
                name="test",
                type="openai_compatible",
                base_url="http://example.invalid",
                api_key="test-key",
                default_model="test-model",
                model_options=("test-model",),
                allow_custom_model=True,
                model_hint="",
                temperature=0.1,
                max_tokens=2048,
            )
            knowledge_result = knowledge_planning_module.generate_knowledge_plans(
                {
                    "items": [
                        {"question_id": "kp_1", "number": "1", "stem": "题目1"},
                        {"question_id": "kp_2", "number": "2", "stem": "题目2"},
                        {"question_id": "kp_3", "number": "3", "stem": "题目3"},
                    ]
                },
                provider,
                "test-model",
                tmp_path / "knowledge_plans.json",
                use_model=True,
                progress_json=tmp_path / "knowledge_planning_progress.json",
            )
        finally:
            knowledge_planning_module.OpenAICompatibleClient = original_client
            if original_workers_env is None:
                knowledge_planning_module.os.environ.pop("KNOWLEDGE_PLANNING_MAX_WORKERS", None)
            else:
                knowledge_planning_module.os.environ["KNOWLEDGE_PLANNING_MAX_WORKERS"] = original_workers_env
        knowledge_output = json.loads((tmp_path / "knowledge_plans.json").read_text(encoding="utf-8"))
        knowledge_progress = json.loads((tmp_path / "knowledge_planning_progress.json").read_text(encoding="utf-8"))
    result["knowledge_planning_parallel"] = {
        "finished": knowledge_finished,
        "worker_count": knowledge_output.get("concurrency", {}).get("max_workers"),
        "progress_worker_count": knowledge_progress.get("max_workers"),
        "parallel_enabled": knowledge_output.get("concurrency", {}).get("parallel_enabled"),
        "progress_status": knowledge_progress.get("status"),
        "plan_count": knowledge_result.plan_count,
    }
    with tempfile.TemporaryDirectory() as tmp:
        original_tasks_dir = task_store_module.TASKS_DIR
        task_store_module.TASKS_DIR = Path(tmp)
        try:
            running_dir = task_store_module.TASKS_DIR / "stale_running_task"
            running_dir.mkdir(parents=True)
            task_store_module.save_task(
                TaskRecord(
                    task_id="stale_running_task",
                    exam_path="exam.docx",
                    textbooks_dir="textbooks",
                    provider="mock",
                    model="mock-model",
                    status="running",
                    created_at="2026-07-04 10:00:00",
                    updated_at="2026-07-04 10:00:01",
                    current_stage="knowledge_planning",
                )
            )
            completed_dir = task_store_module.TASKS_DIR / "completed_task"
            completed_dir.mkdir(parents=True)
            task_store_module.save_task(
                TaskRecord(
                    task_id="completed_task",
                    exam_path="exam.docx",
                    textbooks_dir="textbooks",
                    provider="mock",
                    model="mock-model",
                    status="completed",
                    created_at="2026-07-04 10:00:00",
                    updated_at="2026-07-04 10:05:00",
                    current_stage="completed",
                )
            )
            try:
                recovered = task_store_module.recover_interrupted_tasks("server_startup")
            except AttributeError:
                recovered = []
            stale_after = task_store_module.load_task("stale_running_task")
            completed_after = task_store_module.load_task("completed_task")
            event_text = (running_dir / "events.jsonl").read_text(encoding="utf-8") if (running_dir / "events.jsonl").exists() else ""
            result["interrupted_task_recovery"] = {
                "recovered_ids": [item.get("task_id") for item in recovered],
                "stale_status": stale_after.status,
                "stale_stage": stale_after.current_stage,
                "stale_error": stale_after.error,
                "completed_status": completed_after.status,
                "has_interrupted_event": "server_startup" in event_text,
            }
        finally:
            task_store_module.TASKS_DIR = original_tasks_dir
    with tempfile.TemporaryDirectory() as tmp:
        original_tasks_dir = task_store_module.TASKS_DIR
        task_store_module.TASKS_DIR = Path(tmp)
        try:
            running_dir = task_store_module.TASKS_DIR / "normal_running_task"
            stage_outputs = running_dir / "stage_outputs"
            stage_outputs.mkdir(parents=True)
            task_store_module.save_task(
                TaskRecord(
                    task_id="normal_running_task",
                    exam_path="exam.docx",
                    textbooks_dir="textbooks",
                    provider="mock",
                    model="mock-model",
                    status="running",
                    created_at="2026-07-04 10:00:00",
                    updated_at="2026-07-04 10:00:03",
                    current_stage="knowledge_planning",
                )
            )
            (stage_outputs / "pipeline_status.json").write_text(
                json.dumps(
                    {
                        "stages": [
                            {"stage": "environment", "status": "passed", "detail": {}},
                            {"stage": "extract_exam", "status": "passed", "detail": {"question_count": 2}},
                            {"stage": "textbook_index", "status": "passed", "detail": {"block_count": 10}},
                        ]
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            diagnostics = build_task_diagnostics("normal_running_task")
            result["running_task_diagnostics"] = {
                "needs_attention": diagnostics.get("needs_attention"),
                "title": diagnostics.get("summary", {}).get("title"),
                "issue_count": diagnostics.get("summary", {}).get("issue_count"),
                "warning_count": diagnostics.get("summary", {}).get("warning_count"),
                "status": diagnostics.get("status"),
            }
        finally:
            task_store_module.TASKS_DIR = original_tasks_dir
    index_html = (ROOT / "web" / "index.html").read_text(encoding="utf-8")
    app_js = (ROOT / "web" / "app.js").read_text(encoding="utf-8")
    pipeline_py = (ROOT / "app" / "pipeline.py").read_text(encoding="utf-8")
    final_acceptance_py = (ROOT / "app" / "final_acceptance.py").read_text(encoding="utf-8")
    delivery_py = (ROOT / "app" / "delivery_package.py").read_text(encoding="utf-8")
    server_py = (ROOT / "app" / "server.py").read_text(encoding="utf-8")
    type_sensitive_sources = {
        "prompts": (ROOT / "app" / "prompts.py").read_text(encoding="utf-8"),
        "answer_generation": (ROOT / "app" / "answer_generation.py").read_text(encoding="utf-8"),
        "content_quality_audit": (ROOT / "app" / "content_quality_audit.py").read_text(encoding="utf-8"),
        "question_review_docx": (ROOT / "app" / "question_review_docx.py").read_text(encoding="utf-8"),
        "docx_v4": (ROOT / "app" / "docx_v4.py").read_text(encoding="utf-8"),
        "task_result_view": (ROOT / "app" / "task_result_view.py").read_text(encoding="utf-8"),
    }
    required_ui_labels = ["开工前检查", "准备教材库", "创建真题项目", "生成与复核", "验收与交付"]
    result["workflow_ui_labels"] = {label: (label in index_html) for label in required_ui_labels}
    required_guidance_labels = ["第二步只做一件事", "页码校准在生成后进行"]
    result["workflow_guidance_labels"] = {label: (label in index_html) for label in required_guidance_labels}
    required_picker_labels = ["选择真题文件", "选择教材文件", "上传真题", "上传教材", "教材可多选", "重复文件审查", "建立教材索引", "已建立索引可复用"]
    result["workflow_picker_labels"] = {label: (label in index_html) for label in required_picker_labels}
    result["frontend_answer_structure"] = {
        "hide_top_answer_helper": "function shouldHideTopAnswer" in app_js,
        "calculation_solution_as_answer": 'solution = resultBlock(question, "解题步骤")' in app_js and "<h4>答案</h4><p>${escapeHtml(solution.text)}</p>" in app_js,
        "short_answer_analysis_as_answer": 'const analysisTitle = isShortAnswerQuestion(question) ? "答案" : "解析";' in app_js,
    }
    result["task_summary_state"] = {
        "has_active_task_id": "activeTaskId" in app_js,
        "uses_active_task_lookup": "task.task_id === activeTaskId" in app_js,
    }
    result["thinking_mode_ui"] = {
        "control_exists": 'id="thinkingModeSelect"' in index_html,
        "sent_to_provider_test": "thinking_mode: selectedThinkingMode()" in app_js,
        "sent_to_task_create": "model_thinking: selectedThinkingMode()" in app_js,
    }
    result["task_duration_ui"] = {
        "server_duration_text": "duration_text" in server_py,
        "frontend_duration_function": "function taskDurationText" in app_js,
        "task_card_duration": "运行 ${escapeHtml(taskDurationText(task))}" in app_js,
    }
    result["task_duration_summary"] = _task_duration_summary(
        {
            "status": "completed",
            "created_at": "2026-07-04 10:00:00",
            "updated_at": "2026-07-04 10:03:05",
        }
    )
    result["pipeline_uses_knowledge_planning"] = {
        "has_stage": "knowledge_planning" in pipeline_py,
        "loads_plans_for_retrieval": "knowledge_plans=knowledge_plans" in pipeline_py,
    }
    result["pipeline_uses_evidence_selection"] = {
        "has_stage": "evidence_selection" in pipeline_py,
        "passes_confirmed_candidates": "confirmed_candidates" in pipeline_py,
    }
    result["ui_evidence_selection_progress"] = {
        "stage_label": "模型教材引用确认" in app_js,
        "expansion_progress": "引用证据扩建中" in app_js,
    }
    result["exam_structure_review_ui"] = {
        "pipeline_stage": "exam_structure_review" in pipeline_py and 'mark("exam_structure_review", "started"' in pipeline_py,
        "server_pending_flag": "exam_structure_review_pending" in server_py,
        "server_api": "exam-structure-review" in server_py,
        "frontend_stage_label": "确认题目与题型" in app_js,
        "frontend_pending_flag": "exam_structure_review_pending" in app_js,
        "frontend_api": "exam-structure-review" in app_js,
        "modal_exists": 'id="examStructureReviewModal"' in index_html,
    }
    result["question_type_unification"] = {
        "uses_shared_module": all("question_types" in source for source in type_sensitive_sources.values()),
        "no_local_question_kind": not any("def _question_kind" in source for source in type_sensitive_sources.values()),
        "frontend_subquestion_select": "subquestions" in app_js and "data-subquestion-number" in app_js,
        "six_types_only": list(QUESTION_TYPES) == ["选择题", "判断题", "填空题", "简答题", "计算题", "作图题"],
    }
    result["content_quality_pipeline_contract"] = {
        "pipeline_stage": "content_quality" in pipeline_py,
        "pipeline_report": "content_quality_audit.json" in pipeline_py,
        "final_acceptance_gate": "content_quality_audit.json" in final_acceptance_py,
        "server_summary": "content_quality" in server_py and "content_quality_audit.json" in server_py,
        "delivery_report": "content_quality_audit.json" in delivery_py,
        "ui_stage_label": "内容质量审计" in app_js,
        "review_decision_api": "review-decision" in server_py,
        "review_decision_ui": "checkReviewDecision" in app_js and "review-decision" in app_js,
        "figure_repair_rerun_logged": "figures_after_content_quality_model_repair" in pipeline_py,
    }
    result["late_stage_progress_contract"] = {
        "server_tracks_repair_substages": '"content_quality_model_repair"' in server_py and '"docx_repair"' in server_py,
        "frontend_labels_repair_substages": "质量审查模型回修" in app_js and "Word 文档程序修复" in app_js,
        "frontend_maps_repair_substages": 'content_quality_model_repair: "content_quality"' in app_js and 'docx_repair: "docx"' in app_js,
        "pipeline_marks_content_quality_started": 'mark("content_quality", "started"' in pipeline_py,
        "pipeline_marks_docx_started": 'mark("docx", "started"' in pipeline_py,
        "pipeline_updates_acceptance_stage": 'update_task(task_id, current_stage="acceptance")' in pipeline_py,
        "pipeline_marks_final_acceptance_started": 'mark("final_acceptance", "started"' in pipeline_py,
    }
    model_repair_filtered_quality = _filter_audit_report_for_model_repair(
        {
            "ok": False,
            "issues": [
                {"question_id": "q1", "code": "missing_required_figure", "message": "缺图"},
                {"question_id": "q2", "code": "formula_in_plain_text", "message": "普通正文公式"},
                {"question_id": "q3", "code": "choice_missing_final_conclusion", "message": "故选"},
                {"question_id": "q4", "code": "missing_analysis", "message": "缺解析"},
            ],
            "warnings": [
                {"question_id": "q5", "code": "calculation_answer_missing_unit", "message": "缺单位"},
                {"question_id": "q6", "code": "noncalculation_unintegrated_formulas", "message": "公式未融入"},
            ],
        },
        {
            "missing_required_figure",
            "calculation_missing_mistake_notes",
            "formula_absence_after_retry",
            "calculation_missing_steps",
            "calculation_answer_missing_unit",
            "calculation_missing_substitution",
            "missing_analysis",
        },
    )
    result["model_repair_filter_policy"] = {
        "content_issue_codes": [item.get("code") for item in model_repair_filtered_quality.get("issues", [])],
        "content_warning_codes": [item.get("code") for item in model_repair_filtered_quality.get("warnings", [])],
        "docx_issues": _filter_docx_issues_for_model_repair(
            [
                "Formula-like text must not be written as normal text: 示例",
                "OMML formula count 0 below expected minimum 3",
                "paragraph 1 raw radical in normal text: √(2/3)",
            ]
        ),
        "source_has_choice_missing_final_conclusion": "choice_missing_final_conclusion" in (ROOT / "app" / "content_quality_audit.py").read_text(encoding="utf-8"),
    }
    library_files = scan_library_files()
    smallest_textbook = min(library_files["textbooks"], key=lambda item: int(item.get("size", 0)))
    first_cache = prepare_textbook_index_cache([smallest_textbook["path"]])
    second_cache = prepare_textbook_index_cache([smallest_textbook["path"]])
    result["library_scan"] = {
        "exam_count": len(library_files["exams"]),
        "textbook_count": len(library_files["textbooks"]),
        "duplicate_review_exists": "duplicate_review" in library_files,
        "cache_reused": bool(second_cache.get("cached")),
        "cache_block_count": second_cache.get("block_count"),
    }
    provider_example = json.loads((CONFIG_DIR / "providers.example.json").read_text(encoding="utf-8"))
    deepseek = provider_example["providers"]["deepseek"]
    result["deepseek_models"] = {
        "default_model": deepseek.get("default_model"),
        "vision_model": deepseek.get("vision_model"),
        "supports_vision": deepseek.get("supports_vision"),
        "model_options": deepseek.get("model_options", []),
        "max_tokens": deepseek.get("max_tokens"),
    }
    ark = provider_example["providers"].get("ark", {})
    result["ark_provider"] = {
        "base_url": ark.get("base_url"),
        "api_key_env": ark.get("api_key_env"),
        "allow_custom_model": ark.get("allow_custom_model"),
        "default_model": ark.get("default_model"),
        "vision_model": ark.get("vision_model"),
        "supports_vision": ark.get("supports_vision"),
        "image_model": ark.get("image_model"),
        "max_tokens": ark.get("max_tokens"),
        "deepseek_flash_label": ark.get("model_option_labels", {}).get("deepseek-v4-flash-260425"),
    }
    zhipu = provider_example["providers"].get("zhipu", {})
    result["zhipu_provider"] = {
        "base_url": zhipu.get("base_url"),
        "api_key_env": zhipu.get("api_key_env"),
        "default_model": zhipu.get("default_model"),
        "vision_model": zhipu.get("vision_model"),
        "supports_vision": zhipu.get("supports_vision"),
        "image_model": zhipu.get("image_model"),
        "max_tokens": zhipu.get("max_tokens"),
        "model_options": zhipu.get("model_options", []),
        "glm_46v_label": zhipu.get("model_option_labels", {}).get("glm-4.6v"),
        "thinking_mode": zhipu.get("thinking_mode"),
    }
    try:
        parse_json_content("")
    except LLMError as exc:
        result["empty_model_content_error"] = str(exc)
    else:
        result["empty_model_content_error"] = ""
    try:
        parse_json_content("not-json")
    except LLMError as exc:
        result["invalid_model_content_error"] = str(exc)
    else:
        result["invalid_model_content_error"] = ""
    result["json_repair_parse"] = parse_json_content('```json\n{"ping":"pong"}\n``` trailing text')
    retry_provider = ProviderConfig(
        name="mock",
        type="openai_compatible",
        base_url="http://127.0.0.1",
        api_key="mock",
        default_model="mock-model",
        model_options=(),
        allow_custom_model=False,
        model_hint="",
        temperature=0.1,
        max_tokens=1024,
    )
    retry_client = OpenAICompatibleClient(retry_provider)
    retry_calls = {"count": 0}

    def fake_chat_json(messages, model=None, temperature=None, max_tokens=None, thinking=None, timeout=120):
        retry_calls["count"] += 1
        if retry_calls["count"] == 1:
            return LLMResult("mock", str(model or "mock-model"), "", {"choices": [{"finish_reason": "stop"}], "usage": {}})
        return LLMResult("mock", str(model or "mock-model"), '{"ok": true}', {})

    retry_client.chat_json = fake_chat_json
    try:
        retry_value = retry_client.chat_json_object([{"role": "user", "content": "{}"}], model="mock-model", max_tokens=128)
    except AttributeError:
        retry_value = {}
    result["model_json_retry"] = {
        "calls": retry_calls["count"],
        "ok": retry_value.get("ok") is True,
    }
    response_format_provider = ProviderConfig(
        name="mock_response_format",
        type="openai_compatible",
        base_url="http://127.0.0.1",
        api_key="mock",
        default_model="glm-mock",
        model_options=("glm-mock",),
        allow_custom_model=True,
        model_hint="",
        temperature=0.1,
        max_tokens=1024,
    )
    response_format_client = OpenAICompatibleClient(response_format_provider)
    response_format_requests: list[bool] = []

    def fake_urlopen_response_format(req, timeout=120):
        payload = json.loads(req.data.decode("utf-8"))
        response_format_requests.append("response_format" in payload)
        if "response_format" in payload:
            raise LLMError("Provider HTTP 400: {\"error\":{\"code\":\"InvalidParameter\",\"message\":\"The parameter `response_format.type` specified in the request are not valid: `json_object` is not supported by this model.\",\"param\":\"response_format.type\"}}")

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps({"choices": [{"message": {"content": "{\"ok\": true}"}, "finish_reason": "stop"}]}).encode("utf-8")

        return FakeResponse()

    response_format_client._urlopen = fake_urlopen_response_format
    try:
        response_format_value = response_format_client.chat_json_object([{"role": "user", "content": "{}"}], model="glm-mock", max_tokens=128, attempts=1)
    except (AttributeError, LLMError):
        response_format_value = {}
    result["response_format_fallback"] = {
        "ok": response_format_value.get("ok") is True,
        "requests_used_response_format": response_format_requests,
    }
    thinking_provider = ProviderConfig(
        name="mock_thinking",
        type="openai_compatible",
        base_url="http://127.0.0.1",
        api_key="mock",
        default_model="thinking-model",
        model_options=("thinking-model",),
        allow_custom_model=True,
        model_hint="",
        temperature=0.1,
        max_tokens=1024,
        thinking_mode="disabled",
    )
    thinking_client = OpenAICompatibleClient(thinking_provider)
    thinking_requests: list[str] = []

    def fake_urlopen_thinking(req, timeout=120):
        payload = json.loads(req.data.decode("utf-8"))
        thinking_requests.append((payload.get("thinking") or {}).get("type", "auto"))

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps({"choices": [{"message": {"content": "{\"ok\": true}"}, "finish_reason": "stop"}]}).encode("utf-8")

        return FakeResponse()

    thinking_client._urlopen = fake_urlopen_thinking
    try:
        thinking_result = thinking_client.chat_json([{"role": "user", "content": "{}"}], model="thinking-model", max_tokens=128)
        thinking_value = parse_json_content(thinking_result.content)
    except LLMError:
        thinking_value = {}
    result["thinking_mode_request"] = {
        "ok": thinking_value.get("ok") is True,
        "request_modes": thinking_requests,
    }
    zhipu_reasoning_provider = ProviderConfig(
        name="zhipu",
        type="openai_compatible",
        base_url="https://open.bigmodel.cn/api/paas/v4",
        api_key="mock",
        default_model="glm-4.6v",
        model_options=("glm-4.6v",),
        allow_custom_model=True,
        model_hint="",
        temperature=0.1,
        max_tokens=8192,
        thinking_mode="auto",
    )
    zhipu_reasoning_client = OpenAICompatibleClient(zhipu_reasoning_provider)
    zhipu_reasoning_calls = []

    def fake_zhipu_reasoning_chat_json(messages, model=None, temperature=None, max_tokens=None, timeout=120, thinking=None):
        zhipu_reasoning_calls.append({"model": model, "max_tokens": max_tokens, "thinking": thinking})
        if len(zhipu_reasoning_calls) < 3:
            return LLMResult(
                "zhipu",
                str(model or "glm-4.6v"),
                "",
                {
                    "choices": [
                        {
                            "finish_reason": "length",
                            "message": {"content": "", "reasoning_content": "推理过程" * 10},
                        }
                    ],
                    "usage": {"completion_tokens": max_tokens, "completion_tokens_details": {"reasoning_tokens": max_tokens}},
                },
            )
        return LLMResult(
            "zhipu",
            str(model or "glm-4.6v"),
            '{"ok": true}',
            {"choices": [{"finish_reason": "stop", "message": {"content": '{"ok": true}'}}], "usage": {}},
        )

    zhipu_reasoning_client.chat_json = fake_zhipu_reasoning_chat_json
    try:
        zhipu_reasoning_value = zhipu_reasoning_client.chat_json_object(
            [{"role": "user", "content": "{}"}],
            model="glm-4.6v",
            max_tokens=1024,
            attempts=1,
        )
    except LLMError:
        zhipu_reasoning_value = {}
    result["zhipu_reasoning_retry_policy"] = {
        "ok": zhipu_reasoning_value.get("ok") is True,
        "calls": zhipu_reasoning_calls,
    }
    deepseek_retry_provider = ProviderConfig(
        name="deepseek",
        type="openai_compatible",
        base_url="http://127.0.0.1",
        api_key="mock",
        default_model="deepseek-v4-flash",
        model_options=("deepseek-v4-flash", "deepseek-v4-pro"),
        allow_custom_model=False,
        model_hint="",
        temperature=0.1,
        max_tokens=4096,
    )
    deepseek_retry_client = OpenAICompatibleClient(deepseek_retry_provider)
    deepseek_retry_calls = []

    def fake_deepseek_chat_json(messages, model=None, temperature=None, max_tokens=None, timeout=120, thinking=None):
        deepseek_retry_calls.append(
            {
                "model": model,
                "max_tokens": max_tokens,
                "thinking": thinking,
                "messages": messages,
            }
        )
        if len(deepseek_retry_calls) < 5:
            return LLMResult(
                "deepseek",
                str(model or "deepseek-v4-flash"),
                "",
                {
                    "choices": [{"finish_reason": "length", "message": {"content": ""}}],
                    "usage": {"completion_tokens_details": {"reasoning_tokens": max_tokens}},
                },
            )
        return LLMResult("deepseek", str(model or "deepseek-v4-pro"), '{"ok": true}', {"choices": [{"finish_reason": "stop"}]})

    deepseek_retry_client.chat_json = fake_deepseek_chat_json
    compact_messages = [{"role": "user", "content": "{\"JSON\":\"compact\"}"}]
    try:
        deepseek_retry_value = deepseek_retry_client.chat_json_object(
            [{"role": "user", "content": "{\"JSON\":\"full\"}"}],
            model="deepseek-v4-flash",
            max_tokens=1800,
            fallback_model="deepseek-v4-pro",
            compact_messages=lambda messages: compact_messages,
        )
    except (AttributeError, TypeError):
        deepseek_retry_value = {}
    result["deepseek_json_retry_policy"] = {
        "ok": deepseek_retry_value.get("ok") is True,
        "calls": [
            {
                "model": call.get("model"),
                "max_tokens": call.get("max_tokens"),
                "thinking": call.get("thinking"),
                "compact": call.get("messages") == compact_messages,
            }
            for call in deepseek_retry_calls
        ],
    }
    unit_text = "根据水质量 1 kg、摩尔质量 18.00 g·mol⁻¹、温度 100 °C、压力 101.3 kPa 计算。"
    equation_text = "根据相律 F = C - P + 1 判断。"
    chinese_formula_text = "吉布斯自由能变小于零时反应自发，等于焓变减去温度与熵变的乘积。"
    chinese_prose_text = "恒容意味着系统体积不变，加之无非体积功，故完成的体积功为零。"
    unit_issues = audit_text_segments_no_formula([{"type": "text", "text": unit_text}])
    equation_issues = audit_text_segments_no_formula([{"type": "text", "text": equation_text}])
    chinese_formula_issues = audit_text_segments_no_formula([{"type": "text", "text": chinese_formula_text}], include_chinese_paraphrase=True)
    chinese_prose_issues = audit_text_segments_no_formula([{"type": "text", "text": chinese_prose_text}])
    result["formula_audit_precision"] = {
        "unit_issues": unit_issues,
        "equation_issues": equation_issues,
        "chinese_formula_issues": chinese_formula_issues,
        "chinese_prose_issues": chinese_prose_issues,
        "has_match_detail": any("matched expression" in issue and "F = C - P + 1" in issue for issue in equation_issues),
    }
    result["chemical_latex_plain"] = latex_to_plain(r"\ce{SO2(g)+1/2O2(g)<=>SO3(g)}")
    answer_prompt_payload = json.loads(build_answer_draft_prompt({"question_id": "quality_demo", "section": "一、选择题", "stem": "示例题干"}, [])[1]["content"])
    concise_prompt_payload = json.loads(
        build_answer_draft_prompt(
            {
                "question_id": "depth_concise_demo",
                "section": "七、问答题",
                "stem": "何谓成分过冷？简述组成过冷的条件。(2分)",
            },
            [],
        )[1]["content"]
    )
    deep_prompt_payload = json.loads(
        build_answer_draft_prompt(
            {
                "question_id": "depth_deep_demo",
                "section": "三、计算题",
                "stem": "计算并分析材料相变过程。(10分)",
            },
            [],
        )[1]["content"]
    )
    section_score_prompt_payload = json.loads(
        build_answer_draft_prompt(
            {
                "question_id": "depth_section_score_demo",
                "section": "一、选择题",
                "section_raw": "一、选择题 (本题共 4 分, 每小题 2 分)",
                "stem": "下列说法正确的是：",
            },
            [],
        )[1]["content"]
    )
    quality_text = answer_prompt_payload.get("answer_content_quality_requirements", "")
    result["answer_quality_requirements_in_prompt"] = {
        "draft_task": answer_prompt_payload.get("task") == "generate_question_analysis_draft",
        "no_page_output": any("Do not output page numbers" in rule for rule in answer_prompt_payload.get("hard_rules", [])),
        "no_evidence_id_output": any("Do not output evidence_id" in rule for rule in answer_prompt_payload.get("hard_rules", [])),
        "serves_question": "解析必须服务于本题作答" in quality_text,
        "choice_requirement": "必须说明为什么选该选项" in quality_text,
        "calculation_requirement": "必须给出关键关系式" in quality_text,
        "reasoning_chain": "题干条件 -> 使用依据 -> 判断或计算 -> 结论" in quality_text,
        "no_formula_in_plain_text": "禁止把公式写在普通正文中代替公式对象" in quality_text,
        "figure_specs_schema": "figure_specs" in json.dumps(answer_prompt_payload.get("output_schema_example", {}), ensure_ascii=False),
        "custom_diagram_rule": any("custom_diagram" in rule for rule in answer_prompt_payload.get("hard_rules", [])),
        "calculation_step_groups_schema": all(
            key in (answer_prompt_payload.get("output_schema_example", {}).get("steps") or [{}])[0]
            for key in ("relation_formula_indices", "substitution_formula_indices", "result_formula_indices")
        ),
        "analysis_segments_schema": all(
            key in (answer_prompt_payload.get("output_schema_example", {}).get("analysis_segments") or [{}])[0]
            for key in ("text", "formula_indices")
        ),
        "no_formula_dump_rule": any("Do not list all formulas before substitution" in rule for rule in answer_prompt_payload.get("hard_rules", [])),
        "noncalculation_formula_flow_rule": any("non-calculation questions" in rule and "analysis_segments" in rule for rule in answer_prompt_payload.get("hard_rules", [])),
        "multi_part_calculation_rule": any("multi-part calculation" in rule and "第1小问" in rule for rule in answer_prompt_payload.get("hard_rules", [])),
    }
    result["answer_depth_profile_prompt"] = {
        "concise_depth": (concise_prompt_payload.get("answer_depth_profile") or {}).get("depth"),
        "concise_max_analysis_sentences": (concise_prompt_payload.get("answer_depth_profile") or {}).get("max_analysis_sentences"),
        "concise_mistakes_required": (concise_prompt_payload.get("answer_depth_profile") or {}).get("require_mistake_notes"),
        "deep_depth": (deep_prompt_payload.get("answer_depth_profile") or {}).get("depth"),
        "deep_min_steps": (deep_prompt_payload.get("answer_depth_profile") or {}).get("min_steps"),
        "deep_mistakes_required": (deep_prompt_payload.get("answer_depth_profile") or {}).get("require_mistake_notes"),
        "section_score": (section_score_prompt_payload.get("answer_depth_profile") or {}).get("score"),
        "section_depth": (section_score_prompt_payload.get("answer_depth_profile") or {}).get("depth"),
        "prompt_mentions_depth": any("answer_depth_profile" in rule for rule in deep_prompt_payload.get("hard_rules", [])),
    }
    empty_evidence_fragment = {
        "schema_version": "answer_book.answer_fragment.v4",
        "question_id": "demo_bind",
        "answer": "A",
        "evidence_ids": [],
        "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "根据相关概念判断。"}]}],
        "formulas": [],
        "warnings": [],
    }
    bound = bind_top_evidence(
        empty_evidence_fragment,
        [{"evidence_id": "ev_demo_bind_01", "citation_textbook": "示例教材", "printed_page": "12"}],
        reason="模型认为候选证据只部分相关，程序补用最接近证据。",
    )
    result["evidence_binding"] = {
        "has_bound_evidence": has_bound_evidence(bound),
        "evidence_ids": bound.get("evidence_ids", []),
        "warning_count": len(bound.get("warnings", [])),
        "reason": ((bound.get("_meta") or {}).get("evidence_binding") or {}).get("reason", ""),
        "first_block": (bound.get("blocks") or [{}])[0].get("label"),
    }
    evidence_formula_fragment = {
        "schema_version": "answer_book.answer_fragment.v4",
        "question_id": "evidence_formula",
        "answer": "D",
        "evidence_ids": [],
        "blocks": [
            {"label": "教材依据", "segments": [{"type": "text", "text": "弥勒-布拉菲指数 (hkil) 的约束条件 h+k+i=0：未确认到可用教材依据"}]},
            {"label": "解析", "segments": [{"type": "text", "text": "四指数中前三个指数需要满足约束关系，据此判断错误选项。"}]},
        ],
        "formulas": [],
        "warnings": [],
    }
    result["evidence_block_formula_audit"] = validate_v4_answer_fragment(evidence_formula_fragment)
    with tempfile.TemporaryDirectory() as docx_tmp:
        docx_tmp_path = Path(docx_tmp)
        evidence_docx_json = docx_tmp_path / "evidence_formula_fragments.json"
        evidence_docx_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "evidence_docx_formula",
                            "section": "一、选择题",
                            "number": "1",
                            "answer": "D",
                            "evidence_ids": ["ev_1"],
                            "blocks": [
                                {
                                    "label": "教材依据",
                                    "segments": [
                                        {
                                            "type": "text",
                                            "text": "ΔHm和ΔVm的正负关系：未确认到可用教材依据。",
                                        }
                                    ],
                                },
                                {"label": "解析", "segments": [{"type": "text", "text": "本题根据相变关系判断选项。"}]},
                            ],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        try:
            evidence_docx_path = build_docx_from_fragments(evidence_docx_json, docx_tmp_path / "evidence_formula.docx")
            evidence_docx_audit = audit_docx_v4(evidence_docx_path)
            evidence_docx_ok = evidence_docx_path.exists() and not evidence_docx_audit
        except Exception as exc:
            evidence_docx_ok = False
            evidence_docx_error = str(exc)
        else:
            evidence_docx_error = "；".join(evidence_docx_audit)
        leaking_docx_json = docx_tmp_path / "leaking_formula_fragments.json"
        leaking_docx_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "leaking_docx_formula",
                            "section": "一、选择题",
                            "number": "2",
                            "answer": "A",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "由 ΔHm=Qp 可得。"}]}],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        try:
            leaking_docx_path = build_docx_from_fragments(leaking_docx_json, docx_tmp_path / "leaking_formula.docx")
            leaking_docx_allowed = not audit_docx_v4(leaking_docx_path)
        except ValueError:
            leaking_docx_allowed = False
        chinese_paraphrase_docx_json = docx_tmp_path / "chinese_paraphrase_fragments.json"
        chinese_paraphrase_docx_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "chinese_paraphrase_docx",
                            "section": "三、计算题",
                            "number": "2",
                            "answer": "见解析",
                            "evidence_ids": [],
                            "blocks": [
                                {
                                    "label": "解题步骤",
                                    "segments": [
                                        {
                                            "type": "text",
                                            "text": "第1步：明确环境熵变计算公式。环境可视为恒温巨大热源，与系统交换的热量可逆，因此环境熵变等于环境吸收的热量除以环境温度。",
                                        }
                                    ],
                                }
                            ],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        try:
            chinese_paraphrase_docx_path = build_docx_from_fragments(chinese_paraphrase_docx_json, docx_tmp_path / "chinese_paraphrase.docx")
            chinese_paraphrase_docx_allowed = not audit_docx_v4(chinese_paraphrase_docx_path)
        except ValueError:
            chinese_paraphrase_docx_allowed = False
        dangerous_docx_json = docx_tmp_path / "dangerous_formula_fragments.json"
        dangerous_docx_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "dangerous_docx_formula",
                            "section": "一、选择题",
                            "number": "3",
                            "answer": "A",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "晶格常数之比为√(2/3)≈0.816，包晶转变为 L_B+δ_H→γ_J。"}]}],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        dangerous_docx_path = build_docx_from_fragments(dangerous_docx_json, docx_tmp_path / "dangerous_formula.docx")
        dangerous_docx_issues = audit_docx_v4(dangerous_docx_path)
        dangerous_docx_rejected = bool(dangerous_docx_issues)
        repair_report = repair_answer_fragments_for_docx(dangerous_docx_json, docx_tmp_path / "dangerous_formula.before_repair.json")
        repaired_docx_path = build_docx_from_fragments(dangerous_docx_json, docx_tmp_path / "dangerous_formula_repaired.docx")
        repaired_docx_issues = audit_docx_v4(repaired_docx_path)
        placeholder_repair_json = docx_tmp_path / "placeholder_formula_fragments.json"
        placeholder_repair_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "placeholder_docx_repair",
                            "section": "三、计算题",
                            "number": "4",
                            "answer": "见解析",
                            "answer_summary": "最终判断为{f1}",
                            "evidence_ids": [],
                            "blocks": [{"label": "解题步骤", "segments": [{"type": "text", "text": "第1步：依据{f1}判断。"}]}],
                            "formulas": [
                                {
                                    "formula_id": "f_placeholder_docx_repair_01",
                                    "latex": "\\Delta G<0",
                                    "role": "conclusion",
                                    "display": True,
                                    "source_note": "自发判据",
                                }
                            ],
                            "warnings": ["注意不要保留{f1}。"],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        placeholder_repair_report = repair_answer_fragments_for_docx(
            placeholder_repair_json,
            docx_tmp_path / "placeholder_formula.before_repair.json",
        )
        placeholder_repaired_text = placeholder_repair_json.read_text(encoding="utf-8")
        placeholder_repaired_docx_path = build_docx_from_fragments(placeholder_repair_json, docx_tmp_path / "placeholder_formula_repaired.docx")
        placeholder_repaired_docx_issues = audit_docx_v4(placeholder_repaired_docx_path)
        model_repair_json = docx_tmp_path / "model_docx_repair_fragments.json"
        model_repair_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "provider": "test",
                    "model": "test-model",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "calc_docx_model_repair",
                            "section": "三、计算题",
                            "number": "1",
                            "answer": "见解析",
                            "answer_summary": "ΔH=10 kJ·mol^-1",
                            "evidence_ids": [],
                            "blocks": [
                                {
                                    "label": "解析",
                                    "segments": [{"type": "text", "text": "本题先判断可逆相变，再计算焓变。"}],
                                },
                                {
                                    "label": "解题步骤",
                                    "segments": [
                                        {"type": "text", "text": "第1步：由可逆相变关系，气化焓等于温度与熵变的乘积。"},
                                        {"type": "text", "text": "代入："},
                                        {"type": "formula_ref", "formula_id": "f_calc_docx_model_repair_01"},
                                    ],
                                },
                            ],
                            "formulas": [
                                {
                                    "formula_id": "f_calc_docx_model_repair_01",
                                    "latex": "\\Delta H = 10\\ \\mathrm{kJ\\cdot mol^{-1}}",
                                    "role": "result",
                                    "display": True,
                                    "source_note": "结果",
                                }
                            ],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        class FakeDocxRepairClient:
            def __init__(self):
                self.calls = []
                self.last_json_retry_report = {"ok": True, "attempts": [{"model": "test-model"}]}

            def chat_json_object(self, messages, model=None, max_tokens=None, fallback_model=None):
                self.calls.append({"messages": messages, "model": model, "max_tokens": max_tokens, "fallback_model": fallback_model})
                return {
                    "schema_version": "answer_book.answer_draft.v1",
                    "question_id": "calc_docx_model_repair",
                    "answer": "\\Delta H=10\\ \\mathrm{kJ\\cdot mol^{-1}}",
                    "analysis": "本题先判断可逆相变，再计算焓变。",
                    "steps": [
                        {
                            "title": "由可逆相变关系计算焓变。",
                            "relation_formula_indices": [1],
                            "substitution_formula_indices": [2],
                            "result_formula_indices": [3],
                        }
                    ],
                    "formulas": [
                        {"latex": "\\Delta H = T\\Delta S", "role": "relation", "meaning": "可逆相变焓熵关系"},
                        {"latex": "\\Delta H = 298.15\\times 33.54", "role": "substitution", "meaning": "代入温度和熵变"},
                        {"latex": "\\Delta H=10\\ \\mathrm{kJ\\cdot mol^{-1}}", "role": "result", "meaning": "计算结果"},
                    ],
                    "mistake_notes": [],
                    "uncertainties": [],
                }

        fake_docx_repair_client = FakeDocxRepairClient()
        model_repair_report = repair_fragments_with_model_for_docx(
            model_repair_json,
            {"items": [{"question_id": "calc_docx_model_repair", "section": "三、计算题", "number": "1", "stem": "计算可逆相变焓。"}]},
            [],
            selection_data={},
            provider=ProviderConfig(
                name="test",
                type="openai_compatible",
                base_url="http://example.invalid",
                api_key="test-key",
                default_model="test-model",
                model_options=("test-model",),
                allow_custom_model=True,
                model_hint="",
                temperature=0.1,
                max_tokens=4096,
            ),
            model="test-model",
            docx_issues=["Formula-like text must not be written as normal text: 第1步：由可逆相变关系，气化焓等于温度与熵变的乘积。"],
            client=fake_docx_repair_client,
            backup_path=docx_tmp_path / "model_docx_repair.before.json",
        )
        model_repaired_docx_path = build_docx_from_fragments(model_repair_json, docx_tmp_path / "model_docx_repair.docx")
        model_repaired_docx_issues = audit_docx_v4(model_repaired_docx_path)
        result["docx_evidence_block_formula_policy"] = {
            "evidence_docx_ok": evidence_docx_ok,
            "evidence_docx_error": evidence_docx_error,
            "leaking_docx_allowed": leaking_docx_allowed,
            "chinese_paraphrase_docx_allowed": chinese_paraphrase_docx_allowed,
            "dangerous_docx_rejected": dangerous_docx_rejected,
            "dangerous_docx_issues": dangerous_docx_issues,
            "repair_report": repair_report,
            "repaired_docx_ok": not repaired_docx_issues,
            "repaired_docx_issues": repaired_docx_issues,
            "placeholder_repair_changed": placeholder_repair_report.get("changed"),
            "placeholder_repair_removed": "{f" not in placeholder_repaired_text,
            "placeholder_repaired_docx_ok": not placeholder_repaired_docx_issues,
            "model_repair_changed": model_repair_report["changed"],
            "model_repair_question_ids": model_repair_report["repaired_question_ids"],
            "model_repair_prompt_calls": len(fake_docx_repair_client.calls),
            "model_repaired_docx_ok": not model_repaired_docx_issues,
            "model_repaired_docx_issues": model_repaired_docx_issues,
        }
        formula_layout_json = docx_tmp_path / "formula_layout_fragments.json"
        formula_layout_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "formula_layout",
                            "section": "三、计算题",
                            "number": "1",
                            "answer": "见解析",
                            "answer_summary": "n=55.56 mol；ΔH=-2261 kJ",
                            "evidence_ids": ["ev_1"],
                            "blocks": [
                                {
                                    "label": "解析",
                                    "segments": [
                                        {"type": "text", "text": "先计算物质的量。相关公式："},
                                        {"type": "formula_ref", "formula_id": "f_layout_01"},
                                        {"type": "formula_ref", "formula_id": "f_layout_02"},
                                        {"type": "text", "text": "代入题干数据即可得到结果。"},
                                    ],
                                }
                            ],
                            "formulas": [
                                {"formula_id": "f_layout_01", "latex": "n=\\frac{m}{M}", "role": "relation", "display": True},
                                {"formula_id": "f_layout_02", "latex": "\\Delta H=n\\Delta H_m", "role": "relation", "display": True},
                            ],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        formula_layout_docx = build_docx_from_fragments(formula_layout_json, docx_tmp_path / "formula_layout.docx")
        import zipfile
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(formula_layout_docx) as zf:
            root = ET.fromstring(zf.read("word/document.xml"))
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        }
        math_paras = [p for p in root.findall(".//w:body/w:p", ns) if p.findall(".//m:oMath", ns)]
        math_para_texts = ["".join(t.text or "" for t in p.findall(".//w:t", ns)) for p in math_paras]
        result["docx_formula_block_layout"] = {
            "math_paragraph_count": len(math_paras),
            "math_paras_have_body_text": any("解析" in text or "相关公式" in text for text in math_para_texts),
            "docx_audit_issues": audit_docx_v4(formula_layout_docx),
        }
        summary_formula_json = docx_tmp_path / "summary_formula_fragments.json"
        summary_formula_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "summary_formula_rendering",
                            "section": "四、填空题",
                            "number": "1",
                            "answer": "见解析",
                            "answer_summary": "包晶转变: L_B+δ_H→γ_J；晶格常数之比为√(2/3)≈0.816；dS ≥ δQ/T；ΔG<0；dU=δQ+δW；2d sinθ = nλ；否",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "见解题步骤。"}]}],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        summary_formula_docx = build_docx_from_fragments(summary_formula_json, docx_tmp_path / "summary_formula.docx")
        with zipfile.ZipFile(summary_formula_docx) as zf:
            summary_xml = zf.read("word/document.xml").decode("utf-8")
            summary_root = ET.fromstring(summary_xml)
        summary_plain_text = "".join(t.text or "" for t in summary_root.findall(".//w:t", ns))
        result["docx_answer_summary_formula_rendering"] = {
            "audit_issues": audit_docx_v4(summary_formula_docx),
            "raw_subscript_removed": "L_B" not in summary_plain_text and "γ_J" not in summary_plain_text,
            "raw_sqrt_removed": "√(2/3)" not in summary_plain_text,
            "raw_inequality_removed": "dS" not in summary_plain_text and "δQ/T" not in summary_plain_text and "ΔG<0" not in summary_plain_text,
            "raw_symbolic_equation_removed": "dU=δQ+δW" not in summary_plain_text,
            "raw_bragg_equation_removed": "2d sinθ = nλ" not in summary_plain_text,
            "has_math": len(summary_root.findall(".//m:oMath", ns)) >= 6,
            "has_radical": bool(summary_root.findall(".//m:rad", ns)),
            "has_subscript": bool(summary_root.findall(".//m:sSub", ns)),
        }
        raw_summary_doc = Document()
        raw_summary_doc.add_paragraph("答：晶格常数之比为√(2/3)≈0.816，包晶转变为 L_B+δ_H→γ_J。")
        raw_summary_docx = docx_tmp_path / "raw_summary_formula.docx"
        raw_summary_doc.save(raw_summary_docx)
        result["docx_audit_checks_answer_summary"] = bool(audit_docx_v4(raw_summary_docx))
        structure_json = docx_tmp_path / "answer_structure_fragments.json"
        structure_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "calc_structure",
                            "section": "八、计算题",
                            "number": "1",
                            "answer": "见解析",
                            "answer_summary": "最终结果摘要",
                            "evidence_ids": [],
                            "blocks": [
                                {"label": "教材依据", "segments": [{"type": "text", "text": "教材依据内容"}]},
                                {"label": "解析", "segments": [{"type": "text", "text": "解析内容"}]},
                                {
                                    "label": "解题步骤",
                                    "segments": [
                                        {"type": "text", "text": "第1小问：判断标志面变化。"},
                                        {"type": "text", "text": "第1步：判断扩散方向。"},
                                        {"type": "text", "text": "第2小问：判断固溶相。"},
                                        {"type": "text", "text": "第2步：比较间隙尺寸。"},
                                    ],
                                },
                                {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "易错点内容"}]},
                            ],
                            "formulas": [],
                            "warnings": [],
                        },
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "qa_structure",
                            "section": "七、问答题",
                            "number": "2",
                            "answer": "简答顶层答案",
                            "answer_summary": "简答摘要",
                            "evidence_ids": [],
                            "blocks": [
                                {"label": "教材依据", "segments": [{"type": "text", "text": "简答教材依据"}]},
                                {"label": "解析", "segments": [{"type": "text", "text": "简答解析内容"}]},
                                {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "简答易错点"}]},
                            ],
                            "formulas": [],
                            "warnings": [],
                        },
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        structure_docx = build_docx_from_fragments(structure_json, docx_tmp_path / "answer_structure.docx")
        structure_texts = [p.text for p in Document(structure_docx).paragraphs if p.text.strip()]
        result["docx_answer_structure"] = {
            "texts": structure_texts,
            "no_calc_top_answer": "1、见解析" not in structure_texts,
            "no_calc_summary": not any(text.startswith("答：") and "最终结果摘要" in text for text in structure_texts),
            "calc_order": [text.split("：", 1)[0] for text in structure_texts if text.split("：", 1)[0] in {"教材依据", "解析", "答案", "易错点及注意事项"}][:4],
            "calc_steps_split": all(item in structure_texts for item in ["第1小问：判断标志面变化。", "第1步：判断扩散方向。", "第2小问：判断固溶相。", "第2步：比较间隙尺寸。"]),
            "no_qa_top_answer": "2、简答顶层答案" not in structure_texts,
            "no_qa_summary": not any(text.startswith("答：") and "简答摘要" in text for text in structure_texts),
            "qa_order": [text.split("：", 1)[0] for text in structure_texts if text.split("：", 1)[0] in {"教材依据", "答案", "易错点及注意事项"}][-3:],
        }
        audit_prompt_fragment = {
            "schema_version": "answer_book.answer_fragment.v4",
            "question_id": "audit_prompt_shading",
            "section": "一、填空题",
            "number": "3",
            "answer": "见解析",
            "evidence_ids": [],
            "blocks": [
                {"label": "解析", "segments": [{"type": "text", "text": "本题存在需复核的公式。"}]},
                {
                    "label": "待复核公式",
                    "segments": [
                        {"type": "text", "text": "以下公式未能自然融入解析，请复核其必要性与放置位置。"},
                        {"type": "formula_ref", "formula_id": "f_audit_prompt_shading_01"},
                    ],
                },
            ],
            "formulas": [
                {
                    "formula_id": "f_audit_prompt_shading_01",
                    "latex": "F=C-P+2",
                    "role": "relation",
                    "display": True,
                }
            ],
            "warnings": [],
        }
        audit_prompt_fragment = attach_program_evidence_block(
            audit_prompt_fragment,
            [],
            {
                "knowledge_points": [
                    {
                        "knowledge_point": "相律",
                        "selected_evidence_ids": [],
                        "no_suitable_evidence_reason": "候选证据不足",
                    }
                ]
            },
        )
        audit_prompt_json = docx_tmp_path / "audit_prompt_shading_fragments.json"
        audit_prompt_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [audit_prompt_fragment],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        audit_prompt_docx = build_docx_from_fragments(audit_prompt_json, docx_tmp_path / "audit_prompt_shading.docx")
        with zipfile.ZipFile(audit_prompt_docx) as zf:
            audit_prompt_root = ET.fromstring(zf.read("word/document.xml"))
        shaded_runs = []
        for run in audit_prompt_root.findall(".//w:r", ns):
            text = "".join(t.text or "" for t in run.findall(".//w:t", ns))
            shd = run.find(".//w:shd", ns)
            if text and shd is not None:
                shaded_runs.append({"text": text, "fill": shd.attrib.get(f"{{{ns['w']}}}fill", "")})
        result["docx_audit_prompt_shading"] = {
            "unconfirmed_evidence_segment_marked": any(
                seg.get("highlight") == "unconfirmed_evidence"
                for block in audit_prompt_fragment.get("blocks", [])
                if block.get("label") == "教材依据"
                for seg in block.get("segments", [])
            ),
            "review_formula_label_shaded": any(item["text"] == "待复核公式：" and item["fill"] for item in shaded_runs),
            "unconfirmed_evidence_shaded": any("未确认到可用教材依据" in item["text"] and item["fill"] for item in shaded_runs),
        }
    with tempfile.TemporaryDirectory() as tmp:
        docx_tmp_path = Path(tmp)
        calc_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "calc_answer_summary",
                "answer": "ΔU=-2.09×10^3 kJ；ΔH=-2.26×10^3 kJ",
                "analysis": "由热力学第一定律和焓的定义完成计算，过程见下式。",
                "option_analysis": {},
                "steps": ["列出能量守恒关系", "代入题干数据得到结果"],
                "formulas": [
                    {
                        "formula_id": "f_calc_answer_summary_01",
                        "latex": "\\Delta U=q+w",
                        "role": "relation",
                        "display": True,
                    }
                ],
                "mistake_notes": [],
                "uncertainties": ["复核{f2}是否必要。"],
            },
            {
                "question_id": "calc_answer_summary",
                "section": "三、计算题",
                "number": "1",
                "stem": "计算反应的内能变化和焓变。",
            },
            [],
        )
        if looks_like_formula(str(calc_fragment.get("answer", ""))) and calc_fragment.get("formulas"):
            calc_fragment["answer"] = "见解析"
        calc_docx_json = docx_tmp_path / "calc_answer_summary_fragments.json"
        calc_docx_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [calc_fragment],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        calc_docx_path = build_docx_from_fragments(calc_docx_json, docx_tmp_path / "calc_answer_summary.docx")
        with zipfile.ZipFile(calc_docx_path) as zf:
            calc_docx_root = ET.fromstring(zf.read("word/document.xml"))
        calc_docx_plain_text = "".join(t.text or "" for t in calc_docx_root.findall(".//w:t", ns))
        result["calc_answer_summary_policy"] = {
            "top_answer": calc_fragment.get("answer"),
            "answer_summary": calc_fragment.get("answer_summary"),
            "docx_has_answer_prefix": "答：" in calc_docx_plain_text,
            "raw_summary_formula_removed": "ΔU=-2.09×10^3 kJ" not in calc_docx_plain_text and "ΔH=-2.26×10^3 kJ" not in calc_docx_plain_text,
            "summary_math_count": len(calc_docx_root.findall(".//m:oMath", ns)),
            "docx_audit_issues": audit_docx_v4(calc_docx_path),
        }
        multipart_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "multipart_short_answer",
                "answer": "成分过冷是指合金凝固时液相实际温度低于该处平衡液相线温度而产生的过冷现象。产生成分过冷的条件包括内因和外因：内因为合金性质，外因为温度梯度小、凝固速度大。",
                "analysis": "先回答定义，再回答条件。成分过冷是由溶质再分配导致液相线温度降低形成的。产生成分过冷的条件包括合金参数和凝固条件两方面。",
                "option_analysis": {},
                "steps": [],
                "formulas": [],
                "mistake_notes": [],
                "uncertainties": ["复核{f2}是否必要。"],
            },
            {
                "question_id": "multipart_short_answer",
                "section": "七、问答题",
                "number": "3",
                "stem": "何谓成分过冷？简述组成过冷的条件。",
            },
            [],
        )
        multipart_analysis = next(
            (
                "".join(seg.get("text", "") for seg in block.get("segments", []) if seg.get("type") == "text")
                for block in multipart_fragment.get("blocks", [])
                if block.get("label") == "解析"
            ),
            "",
        )
        result["multipart_answer_layout"] = {
            "answer_summary": multipart_fragment.get("answer_summary"),
            "answer_has_second_line": "\n产生成分过冷的条件" in str(multipart_fragment.get("answer_summary", "")),
            "analysis_has_second_line": "\n产生成分过冷的条件" in multipart_analysis,
        }
        ordinal_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "ordinal_newline_fragment",
                "answer": "(1) 第一问答案。 (2) 第二问答案。",
                "analysis": "本题按小问作答。第(1)问先判断相图特征；第(2)问绘制冷却曲线；第\n(3)问分析性能差异。",
                "option_analysis": {},
                "steps": [],
                "formulas": [],
                "mistake_notes": [],
                "uncertainties": [],
            },
            {
                "question_id": "ordinal_newline_fragment",
                "section": "综合题",
                "number": "1",
                "stem": "（1）判断相图特征。（2）绘制冷却曲线。（3）分析性能差异。",
            },
            [],
        )
        ordinal_analysis = next(
            (
                "".join(seg.get("text", "") for seg in block.get("segments", []) if seg.get("type") == "text")
                for block in ordinal_fragment.get("blocks", [])
                if block.get("label") == "解析"
            ),
            "",
        )
        result["ordinal_newline_normalization"] = {
            "analysis_keeps_second_label": "第(2)问" in ordinal_analysis,
            "analysis_repairs_third_label": "第(3)问" in ordinal_analysis,
            "analysis_has_broken_label": "第\n(" in ordinal_analysis,
            "answer_keeps_numbered_line_break": "\n(2)" in str(ordinal_fragment.get("answer_summary", "")),
        }
        depth_profile_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "depth_profile_fragment",
                "answer": "凝固速度增大。",
                "analysis": "根据成分过冷条件判断。",
                "option_analysis": {},
                "steps": [],
                "formulas": [],
                "mistake_notes": [],
                "uncertainties": [],
            },
            {
                "question_id": "depth_profile_fragment",
                "section": "七、问答题",
                "number": "4",
                "stem": "简述形成成分过冷的条件。(2分)",
            },
            [],
        )
        result["answer_depth_profile_fragment"] = depth_profile_fragment.get("_draft", {}).get("answer_depth_profile")
        multipart_calc_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "multipart_calc_layout",
                "answer": "见解析",
                "analysis": "按小问分别作答。",
                "option_analysis": {},
                "steps": [
                    {"text": "判断标志面变化。", "result_text": "间距减小。"},
                    {"text": "判断氮原子主要固溶相。", "result_text": "主要固溶于 β 相。"},
                    {"text": "计算晶格常数之比。", "result_text": "约 0.816。"},
                ],
                "formulas": [],
                "mistake_notes": [],
                "uncertainties": [],
            },
            {
                "question_id": "multipart_calc_layout",
                "section": "八、计算题",
                "number": "1",
                "stem": "分析标志面位置变化。(2分)\n氮原子主要固溶于哪个相中？解释原因（2分）\n计算α和β相的晶格常数之比。（2分）",
            },
            [],
        )
        multipart_calc_steps = "".join(
            seg.get("text", "")
            for block in multipart_calc_fragment.get("blocks", [])
            if block.get("label") == "解题步骤"
            for seg in block.get("segments", [])
            if seg.get("type") == "text"
        )
        result["multipart_calculation_step_grouping"] = {
            "has_first_subquestion": "第1小问" in multipart_calc_steps,
            "has_second_subquestion": "第2小问" in multipart_calc_steps,
            "has_third_subquestion": "第3小问" in multipart_calc_steps,
        }
        structured_subquestion_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "structured_subquestion_calc",
                "answer": "见解析",
                "analysis": "按原题小问分别计算。",
                "option_analysis": {},
                "steps": [
                    {
                        "subquestion_number": "1",
                        "text": "计算系统熵变。",
                        "relation_formula_indices": [1],
                        "substitution_formula_indices": [2],
                        "result_formula_indices": [3],
                        "result_text": "系统熵变为 73.71 J·K^-1。",
                    },
                    {
                        "text": "第2小问：计算环境熵变。",
                        "relation_formula_indices": [4],
                        "substitution_formula_indices": [5],
                        "result_formula_indices": [6],
                        "result_text": "环境熵变为 -65.42 J·K^-1。",
                    },
                ],
                "formulas": [
                    {"latex": "\\Delta S_{sys}=\\frac{Q_{rev}}{T}", "meaning": "系统熵变关系式", "display": False, "role": "relation"},
                    {"latex": "\\Delta S_{sys}=26020/353", "meaning": "系统熵变代入", "display": False, "role": "substitution"},
                    {"latex": "\\Delta S_{sys}=73.71\\ \\mathrm{J\\cdot K^{-1}}", "meaning": "系统熵变结果", "display": False, "role": "result"},
                    {"latex": "\\Delta S_{amb}=-Q/T", "meaning": "环境熵变关系式", "display": False, "role": "relation"},
                    {"latex": "\\Delta S_{amb}=-23100/353", "meaning": "环境熵变代入", "display": False, "role": "substitution"},
                    {"latex": "\\Delta S_{amb}=-65.42\\ \\mathrm{J\\cdot K^{-1}}", "meaning": "环境熵变结果", "display": False, "role": "result"},
                ],
                "mistake_notes": ["注意系统和环境热量符号相反。"],
                "uncertainties": [],
            },
            {
                "question_id": "structured_subquestion_calc",
                "section": "三、计算题",
                "number": "1",
                "stem": "试求：(1) 整个过程的熵变；(2) 环境熵变。",
                "subquestions": [
                    {"number": "1", "marker": "(1)", "stem": "整个过程的熵变。"},
                    {"number": "2", "marker": "(2)", "stem": "环境熵变。"},
                ],
            },
            [],
        )
        structured_subquestion_steps = [
            seg
            for block in structured_subquestion_fragment.get("blocks", [])
            if block.get("label") == "解题步骤"
            for seg in block.get("segments", [])
        ]
        structured_subquestion_text = "".join(seg.get("text", "") for seg in structured_subquestion_steps if seg.get("type") == "text")
        result["structured_subquestion_calculation"] = {
            "program_title_uses_first_stem": "第1小问：整个过程的熵变" in structured_subquestion_text,
            "program_title_uses_second_stem": "第2小问：环境熵变" in structured_subquestion_text,
            "legacy_heading_payload_preserved": "计算环境熵变" in structured_subquestion_text
            and sum(1 for seg in structured_subquestion_steps if seg.get("type") == "formula_ref") == 6,
            "draft_subquestion_numbers": [
                step.get("subquestion_number")
                for step in structured_subquestion_fragment.get("_draft", {}).get("steps", [])
                if isinstance(step, dict) and not step.get("_subquestion_heading")
            ],
        }
        calc_layout_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "calc_solution_layout",
                "answer": "V = nRT/p",
                "analysis": "本题先明确状态方程，再代入题干条件计算。",
                "option_analysis": {},
                "steps": [
                    {
                        "text": "第1步：计算物质的量。",
                        "relation_formula_indices": [1],
                        "substitution_formula_indices": [2],
                        "result_formula_indices": [3],
                    },
                    {
                        "text": "第2步：由理想气体状态方程求体积。",
                        "relation_formula_indices": [4],
                        "substitution_formula_indices": [5],
                        "result_formula_indices": [6],
                    },
                ],
                "formulas": [
                    {"latex": "n = \\\\frac{m}{M}", "meaning": "物质的量计算式", "display": False, "role": "relation"},
                    {"latex": "n = \\\\frac{1.00 \\\\times 10^3}{18.00} = 55.56\\\\ \\\\mathrm{mol}", "meaning": "代入水的质量和摩尔质量", "display": False, "role": "substitution"},
                    {"latex": "n = 55.56\\\\ \\\\mathrm{mol}", "meaning": "物质的量结果", "display": False, "role": "result"},
                    {"latex": "V_{\\\\text{g}} = \\\\frac{nRT}{p}", "meaning": "理想气体状态方程", "display": False, "role": "relation"},
                    {"latex": "V_{\\\\text{g}} = \\\\frac{55.56 \\\\times 8.314 \\\\times 373.15}{101.3 \\\\times 10^3}", "meaning": "代入状态方程", "display": False, "role": "substitution"},
                    {"latex": "V_{\\\\text{g}} = 1.70\\\\ \\\\mathrm{m^3}", "meaning": "体积结果", "display": False, "role": "result"},
                ],
                "mistake_notes": ["温度需换算为 K。"],
                "uncertainties": [],
            },
            {
                "question_id": "calc_solution_layout",
                "section": "三、计算题",
                "number": "1",
                "stem": "计算理想气体体积。",
            },
            [],
        )
        calc_layout_blocks = {block.get("label"): block.get("segments", []) for block in calc_layout_fragment.get("blocks", [])}
        result["calc_solution_layout_policy"] = {
            "normalized_latex": [formula.get("latex") for formula in calc_layout_fragment.get("formulas", [])],
            "all_display": all(bool(formula.get("display")) for formula in calc_layout_fragment.get("formulas", [])),
            "analysis_has_formula_refs": any(seg.get("type") == "formula_ref" for seg in calc_layout_blocks.get("解析", [])),
            "steps_formula_ref_count": sum(1 for seg in calc_layout_blocks.get("解题步骤", []) if seg.get("type") == "formula_ref"),
            "steps_text": "".join(seg.get("text", "") for seg in calc_layout_blocks.get("解题步骤", []) if seg.get("type") == "text"),
            "has_formula_dump_heading": any(
                seg.get("type") == "text" and seg.get("text") in {"关系式与代入：", "补充关系式："}
                for seg in calc_layout_blocks.get("解题步骤", [])
            ),
            "step_order": [
                seg.get("text", "") if seg.get("type") == "text" else seg.get("formula_id", "")
                for seg in calc_layout_blocks.get("解题步骤", [])
            ],
        }
    with tempfile.TemporaryDirectory() as tmp:
        figure_tmp_path = Path(tmp)
        fragments_json = figure_tmp_path / "answer_fragments.json"
        fragments_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "figure_question",
                            "section": "五、简答题",
                            "number": "1",
                            "answer": "见解析",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "面心立方晶胞中原子位于顶点和面心。"}]}],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        prepare_figures = getattr(figures_module, "prepare_figures_for_fragments", None)
        if callable(prepare_figures):
            generated_paths = prepare_figures(
                {
                    "items": [
                        {
                            "question_id": "figure_question",
                            "section": "五、简答题",
                            "number": "1",
                            "stem": "画出面心立方晶胞，并标出原子位置。",
                            "image_refs": ["需要作图"],
                        }
                    ]
                },
                fragments_json,
                figure_tmp_path / "figure_specs.json",
                figure_tmp_path / "figures",
            )
            updated_fragments = json.loads(fragments_json.read_text(encoding="utf-8"))
        else:
            generated_paths = []
            updated_fragments = json.loads(fragments_json.read_text(encoding="utf-8"))
        figure_segments = [
            seg
            for fragment in updated_fragments.get("fragments", [])
            for block in fragment.get("blocks", [])
            for seg in block.get("segments", [])
        ]
        figure_specs = json.loads((figure_tmp_path / "figure_specs.json").read_text(encoding="utf-8")) if (figure_tmp_path / "figure_specs.json").exists() else {}
        result["figure_generation_pipeline"] = {
            "generated_count": len(generated_paths),
            "spec_exists": (figure_tmp_path / "figure_specs.json").exists(),
            "has_image_ref": any(seg.get("type") == "image_ref" for seg in figure_segments if isinstance(seg, dict)),
            "png_exists": bool(generated_paths and generated_paths[0].exists()),
            "kind": (figure_specs.get("figures") or [{}])[0].get("kind"),
        }
    with tempfile.TemporaryDirectory() as tmp:
        figure_tmp_path = Path(tmp)
        fragments_json = figure_tmp_path / "answer_fragments.json"
        fragments_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "curved_surface_question",
                            "section": "三、问答题",
                            "number": "1",
                            "answer": "见解析",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "弯曲液面附加压力由表面张力产生。"}]}],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        prepare_figures = getattr(figures_module, "prepare_figures_for_fragments", None)
        if callable(prepare_figures):
            generated_paths = prepare_figures(
                {
                    "items": [
                        {
                            "question_id": "curved_surface_question",
                            "section": "三、问答题",
                            "number": "1",
                            "stem": "作图并描述弯曲液面附加压力的产生，标注曲率半径和表面张力方向。",
                            "image_refs": ["需要作图"],
                        }
                    ]
                },
                fragments_json,
                figure_tmp_path / "figure_specs.json",
                figure_tmp_path / "figures",
            )
            figure_specs = json.loads((figure_tmp_path / "figure_specs.json").read_text(encoding="utf-8"))
        else:
            generated_paths = []
            figure_specs = {}
        result["curved_surface_figure_generation"] = {
            "generated_count": len(generated_paths),
            "kind": (figure_specs.get("figures") or [{}])[0].get("kind"),
            "png_exists": bool(generated_paths and generated_paths[0].exists()),
        }
    with tempfile.TemporaryDirectory() as tmp:
        figure_tmp_path = Path(tmp)
        fragments_json = figure_tmp_path / "answer_fragments.json"
        fragments_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "custom_diagram_question",
                            "section": "三、问答题",
                            "number": "1",
                            "answer": "见解析",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "按图示说明力的方向。"}]}],
                            "formulas": [],
                            "warnings": [],
                            "_draft": {
                                "figure_specs": [
                                    {
                                        "figure_id": "custom_diagram_question_fig_01",
                                        "kind": "custom_diagram",
                                        "caption": "结构化作图指令示意图",
                                        "required_labels": ["凸液面", "曲率半径", "附加压力"],
                                        "elements": [
                                            {"type": "circle", "center": [0, 0], "radius": 0.8, "label": "凸液面"},
                                            {"type": "arrow", "start": [0, 0.9], "end": [0, 0.1], "label": "附加压力", "color": "#dc2626"},
                                            {"type": "line", "start": [0, 0], "end": [0.8, 0], "label": "曲率半径", "style": "dashed"},
                                            {"type": "text", "xy": [-0.4, -0.15], "text": "曲率中心"},
                                        ],
                                    }
                                ]
                            },
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        prepare_figures = getattr(figures_module, "prepare_figures_for_fragments", None)
        if callable(prepare_figures):
            generated_paths = prepare_figures(
                {
                    "items": [
                        {
                            "question_id": "custom_diagram_question",
                            "section": "三、问答题",
                            "number": "1",
                            "stem": "作图说明弯曲液面附加压力方向。",
                            "image_refs": ["需要作图"],
                        }
                    ]
                },
                fragments_json,
                figure_tmp_path / "figure_specs.json",
                figure_tmp_path / "figures",
            )
            figure_specs = json.loads((figure_tmp_path / "figure_specs.json").read_text(encoding="utf-8"))
            updated_fragments = json.loads(fragments_json.read_text(encoding="utf-8"))
        else:
            generated_paths = []
            figure_specs = {}
            updated_fragments = {"fragments": []}
        result["custom_diagram_generation"] = {
            "generated_count": len(generated_paths),
            "kind": (figure_specs.get("figures") or [{}])[0].get("kind"),
            "png_exists": bool(generated_paths and generated_paths[0].exists()),
            "has_image_ref": any(
                segment.get("type") == "image_ref"
                for fragment in updated_fragments.get("fragments", [])
                for block in fragment.get("blocks", [])
                for segment in block.get("segments", [])
                if isinstance(segment, dict)
            ),
        }
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        provider = ProviderConfig(
            name="mock",
            type="openai_compatible",
            base_url="http://127.0.0.1",
            api_key="mock",
            default_model="mock-model",
            model_options=(),
            allow_custom_model=False,
            model_hint="",
            temperature=0.1,
            max_tokens=1024,
        )
        original_generate_one = answer_generation_module.generate_one_fragment
        try:
            def fake_generate_one(client, provider_config, question, evidence, model, evidence_selection=None, retries=2):
                qid = str(question.get("question_id", ""))
                if qid == "q_fail":
                    return None, ["simulated validation failure"]
                return {
                    "schema_version": "answer_book.answer_fragment.v4",
                    "question_id": qid,
                    "section": question.get("section", ""),
                    "number": question.get("number", ""),
                    "answer": "A",
                    "evidence_ids": [],
                    "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "本题已生成可用解析。"}]}],
                    "formulas": [],
                    "warnings": [],
                }, []

            answer_generation_module.generate_one_fragment = fake_generate_one
            continuation_result = answer_generation_module.generate_answer_fragments(
                {
                    "items": [
                        {"question_id": "q_ok_1", "section": "一、选择题", "number": "1", "stem": "第一题"},
                        {"question_id": "q_fail", "section": "一、选择题", "number": "2", "stem": "失败题"},
                        {"question_id": "q_ok_2", "section": "一、选择题", "number": "3", "stem": "第三题"},
                    ]
                },
                [],
                provider,
                "mock-model",
                tmp_path / "answer_fragments_continue.json",
            )
        finally:
            answer_generation_module.generate_one_fragment = original_generate_one
        continuation_data = json.loads((tmp_path / "answer_fragments_continue.json").read_text(encoding="utf-8"))
        result["answer_generation_continues_after_failure"] = {
            "ok": continuation_result.ok,
            "fragment_ids": [fragment.get("question_id") for fragment in continuation_data.get("fragments", [])],
            "issue_qids": [item.get("question_id") for item in continuation_data.get("issues", [])],
            "fallback_count": continuation_data.get("fallback_count"),
            "failed_flags": [
                flag.get("code")
                for fragment in continuation_data.get("fragments", [])
                if fragment.get("question_id") == "q_fail"
                for flag in fragment.get("_review_flags", [])
            ],
        }
        model_retry_notes = build_answer_review_notes(
            {
                "fragments": [
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "q_model_retry",
                        "section": "一、选择题",
                        "number": "1",
                        "answer": "A",
                        "evidence_ids": ["ev_retry"],
                        "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "切换备用模型后已生成合格解析。"}]}],
                        "formulas": [],
                        "warnings": ["自动切换模型 deepseek-v4-pro 完成结构化生成。"],
                        "_meta": {"provider": "deepseek", "model": "deepseek-v4-pro", "recovered_by": "model_retry"},
                    }
                ]
            }
        )
        result["model_retry_not_review_required"] = {
            "note_count": model_retry_notes.get("note_count"),
            "rows": model_retry_notes.get("rows", []),
        }
        exam_docx = tmp_path / "multi_subject_exam.docx"
        doc = Document()
        for para in [
            "物理化学",
            "一、选择题 (本题共 4 分, 每小题 2 分)",
            "1、物化选择题一（ ）",
            "A. 甲 B. 乙 C. 丙 D. 丁",
            "2、物化选择题二（ ）",
            "三、计算题 (本题共 12 分)",
            "已知水的摩尔气化焓，计算 ΔU 和 ΔH。",
            "计算题 (本题共 15 分)",
            "一级反应 A→Y+Z，计算速率常数和活化能。",
            "“材料现代研究” 部分",
            "五、选择题 (本题共 4 分, 每小题 2 分)",
            "1、X 射线是（ ）",
            "A. 可见光 B. 电磁波 C. 声波 D. 超声波",
            "2、晶体中的旋转对称轴共有（ ）",
        ]:
            doc.add_paragraph(para)
        doc.save(exam_docx)
        extracted_exam = extract_exam_structure(exam_docx, tmp_path / "multi_subject_exam.json")
        extracted_items = extracted_exam.get("items", [])
        result["multi_subject_exam_extract"] = {
            "ids": [item.get("question_id") for item in extracted_items],
            "count": len(extracted_items),
            "calc_stems": [item.get("stem", "") for item in extracted_items if str(item.get("question_id", "")).startswith("calc_")],
            "audit_issues": audit_exam_structure(extracted_exam, tmp_path / "multi_subject_exam_audit.json"),
        }

        source_image = tmp_path / "source_question.png"
        source_image.write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="))
        image_docx = tmp_path / "image_exam.docx"
        image_doc = Document()
        image_doc.add_paragraph("一、计算题")
        image_doc.add_paragraph("1、观察下图相图，画出冷却曲线。")
        image_doc.add_picture(str(source_image))
        image_doc.save(image_docx)
        image_exam = extract_exam_structure(image_docx, tmp_path / "image_exam.json")
        image_item = (image_exam.get("items") or [{}])[0]
        image_refs = image_item.get("image_refs") or []
        snapshot_refs = image_item.get("question_snapshot_refs") or []
        image_review_request = build_exam_structure_review_request("image_review_task", image_exam)
        review_snapshot_refs = (image_review_request.get("items") or [{}])[0].get("question_snapshot_refs") or []
        image_understanding = build_question_understanding(image_item, tmp_path / "image_understanding")
        multimodal_messages = build_answer_draft_prompt(image_item, [], image_understanding)
        multimodal_user_content = multimodal_messages[1].get("content")
        multimodal_payload = json.loads(multimodal_user_content if isinstance(multimodal_user_content, str) else next((part.get("text", "{}") for part in multimodal_user_content if isinstance(part, dict) and part.get("type") == "text"), "{}"))
        result["docx_image_multimodal_input"] = {
            "image_ref_count": len(image_refs),
            "image_exists": bool(image_refs and Path(str(image_refs[0])).exists()),
            "question_snapshot_count": len(snapshot_refs),
            "question_snapshot_exists": bool(snapshot_refs and Path(str(snapshot_refs[0])).exists()),
            "review_request_snapshot_count": len(review_snapshot_refs),
            "stem_has_marker": "__ANSWER_BOOK_IMAGE__" in str(image_item.get("stem", "")),
            "understanding_needs_vision": image_understanding.get("needs_vision_model"),
            "understanding_has_image": bool(image_understanding.get("images")),
            "prompt_uses_question_understanding": bool(multimodal_payload.get("question_understanding")),
            "prompt_has_raw_image_url": isinstance(multimodal_user_content, list)
            and any(isinstance(part, dict) and part.get("type") == "image_url" for part in multimodal_user_content),
        }

        exam_audit_path = tmp_path / "exam_structure_audit.json"
        audit_exam_structure(
            {
                "source_paragraphs": ["一、选择题", "1、相变焓随温度如何变化？", "2、缺失的原文题目？"],
                "items": [
                    {
                        "question_id": "choice_01",
                        "section": "一、选择题",
                        "number": "1",
                        "stem": "相变焓随温度如何变化？",
                    }
                ],
            },
            exam_audit_path,
        )
        result["exam_source_coverage"] = json.loads(exam_audit_path.read_text(encoding="utf-8")).get("source_coverage", {})

        section_rows = enrich_rows_with_sections(
            [
                {
                    "textbook": "示例教材",
                    "source_file": "demo.json",
                    "page_idx": 1,
                    "block_index": 1,
                    "block_type": "text",
                    "chapter_section": "",
                    "text": "92.5 表中编号，不是章节标题。",
                    "char_count": 26,
                },
                {
                    "textbook": "示例教材",
                    "source_file": "demo.json",
                    "page_idx": 2,
                    "block_index": 2,
                    "block_type": "title",
                    "chapter_section": "",
                    "text": "2.5 相变焓与温度的关系",
                    "char_count": 14,
                },
                {
                    "textbook": "示例教材",
                    "source_file": "demo.json",
                    "page_idx": 2,
                    "block_index": 3,
                    "block_type": "text",
                    "chapter_section": "",
                    "text": "正文内容应继承标题章节。",
                    "char_count": 11,
                },
            ]
        )
        result["title_only_sections"] = [row.get("chapter_section", "") for row in section_rows]
        page_map_from_page_number = build_page_map(
            [
                {
                    "textbook": "页码教材",
                    "source_file": "page.json",
                    "page_idx": 10,
                    "block_index": 1,
                    "block_type": "text",
                    "chapter_section": "",
                    "text": "此外，详见405页脚注。",
                    "char_count": 12,
                },
                {
                    "textbook": "页码教材",
                    "source_file": "page.json",
                    "page_idx": 10,
                    "block_index": 2,
                    "block_type": "page_number",
                    "chapter_section": "",
                    "text": "40",
                    "char_count": 2,
                },
            ]
        )
        page_map_from_reference_only = build_page_map(
            [
                {
                    "textbook": "页码教材",
                    "source_file": "page.json",
                    "page_idx": 89,
                    "block_index": 1,
                    "block_type": "text",
                    "chapter_section": "",
                    "text": "角动量平方算符与投影算符可同时精确测量，详见405页脚注。",
                    "char_count": 34,
                }
            ]
        )
        page_audit_ok, page_audit_issues = audit_page_map(
            [
                {"textbook": "缺页码教材", "pdf_page_idx": str(i), "printed_page": ""}
                for i in range(8)
            ]
        )
        result["page_number_mapping"] = {
            "page_number_block": page_map_from_page_number[0].get("printed_page"),
            "reference_only": page_map_from_reference_only[0].get("printed_page"),
            "audit_ok": page_audit_ok,
            "audit_issue_count": len(page_audit_issues),
        }

        blocks_csv = tmp_path / "blocks.csv"
        page_map_csv = tmp_path / "pages.csv"
        candidates_csv = tmp_path / "candidates.csv"
        block_rows = [
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "1",
                "block_index": "1",
                "block_type": "text",
                "chapter_section": "2.5",
                "text": "相变焓随温度变化可由热容差积分判断。",
                "char_count": "20",
            },
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "2",
                "block_index": "2",
                "block_type": "title",
                "chapter_section": "2.5",
                "text": "目录 第二章 相变焓 温度 关系 56",
                "char_count": "20",
            },
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "3",
                "block_index": "3",
                "block_type": "text",
                "chapter_section": "2.5",
                "text": "习题 相变焓 温度 关系",
                "char_count": "12",
            },
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "4",
                "block_index": "4",
                "block_type": "table",
                "chapter_section": "2.5",
                "text": "<table><tr><td>相变焓</td><td>温度</td></tr></table>",
                "char_count": "40",
            },
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "5",
                "block_index": "5",
                "block_type": "text",
                "chapter_section": "2.5",
                "text": "相变焓 温度 关系，但页码为空。",
                "char_count": "16",
            },
        ]
        page_rows = [
            {"textbook": "示例教材", "citation_textbook": "示例教材", "source_file": "demo.json", "pdf_page_idx": "1", "printed_page": "56", "page_source": "manual", "verified": "true", "confidence": "high", "notes": ""},
            {"textbook": "示例教材", "citation_textbook": "示例教材", "source_file": "demo.json", "pdf_page_idx": "2", "printed_page": "57", "page_source": "manual", "verified": "true", "confidence": "high", "notes": ""},
            {"textbook": "示例教材", "citation_textbook": "示例教材", "source_file": "demo.json", "pdf_page_idx": "3", "printed_page": "58", "page_source": "manual", "verified": "true", "confidence": "high", "notes": ""},
            {"textbook": "示例教材", "citation_textbook": "示例教材", "source_file": "demo.json", "pdf_page_idx": "4", "printed_page": "59", "page_source": "manual", "verified": "true", "confidence": "high", "notes": ""},
            {"textbook": "示例教材", "citation_textbook": "示例教材", "source_file": "demo.json", "pdf_page_idx": "5", "printed_page": "", "page_source": "unmapped", "verified": "false", "confidence": "none", "notes": ""},
        ]
        write_csv(blocks_csv, BLOCK_FIELDS, block_rows)
        write_csv(page_map_csv, PAGE_MAP_FIELDS, page_rows)
        filtered_candidates = build_candidates(
            {
                "items": [
                    {
                        "question_id": "choice_01",
                        "section": "一、选择题",
                        "stem": "相变焓和温度的关系",
                    }
                ]
            },
            blocks_csv,
            page_map_csv,
            candidates_csv,
            top_k=5,
        )
        result["retrieval_filtering"] = [candidate.evidence_text for candidate in filtered_candidates]
        result["directory_evidence_filtering"] = {
            "catalog_page": is_invalid_evidence_row(
                {
                    "block_type": "text",
                    "text": "章 电化学 …… 309 § 7.1 电极过程、电解质溶液及法拉第定律 …… 309 1. 电解池和原电池 …… 309 2. 电解质溶液和法拉第定律 …… 311",
                },
                {"printed_page": "319"},
            ),
            "real_page": is_invalid_evidence_row(
                {
                    "block_type": "text",
                    "text": "原则上讲，对于 ΔG<0 的反应都可设计成原电池，设计的方法是将给定反应拆分为两个电极反应。",
                },
                {"printed_page": "375"},
            ),
        }
        planned_candidates_csv = tmp_path / "planned_candidates.csv"
        plan_driven_candidates = build_candidates(
            {
                "items": [
                    {
                        "question_id": "choice_02",
                        "section": "一、选择题",
                        "stem": "这道题表面词只提到冰和溶液，不能直接拿题干检索。",
                    }
                ]
            },
            blocks_csv,
            page_map_csv,
            planned_candidates_csv,
            top_k=1,
            knowledge_plans={
                "choice_02": {
                    "knowledge_points": ["相变焓与温度的关系"],
                    "formulas": ["热容差积分"],
                    "key_terms": ["相变焓", "温度", "热容差"],
                    "search_queries": ["相变焓 温度 热容差 积分"],
                }
            },
        )
        result["retrieval_query_uses_plan"] = retrieval_query_text(
            {
                "question_id": "choice_02",
                "section": "一、选择题",
                "stem": "这道题表面词只提到冰和溶液，不能直接拿题干检索。",
            },
            {
                "knowledge_points": ["相变焓与温度的关系"],
                "formulas": ["热容差积分"],
                "key_terms": ["相变焓", "温度", "热容差"],
                "search_queries": ["相变焓 温度 热容差 积分"],
            },
        )
        result["plan_driven_retrieval"] = [candidate.evidence_text for candidate in plan_driven_candidates]
        point_blocks_csv = tmp_path / "point_blocks.csv"
        point_candidates_csv = tmp_path / "point_candidates.csv"
        point_block_rows = [
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "1",
                "block_index": "1",
                "block_type": "text",
                "chapter_section": "1.1",
                "text": "甲知识点 甲依据 A。",
                "char_count": "20",
            },
            {
                "textbook": "示例教材",
                "source_file": "demo.json",
                "page_idx": "2",
                "block_index": "2",
                "block_type": "text",
                "chapter_section": "1.2",
                "text": "乙知识点 乙依据 B。",
                "char_count": "20",
            },
        ]
        write_csv(point_blocks_csv, BLOCK_FIELDS, point_block_rows)
        point_candidates = build_candidates(
            {"items": [{"question_id": "q_point", "section": "一、选择题", "stem": "题干不用于截断"}]},
            point_blocks_csv,
            page_map_csv,
            point_candidates_csv,
            top_k=1,
            knowledge_plans={
                "q_point": {
                    "knowledge_points": ["甲知识点", "乙知识点"],
                    "key_terms": ["甲知识点", "乙知识点"],
                    "search_queries": ["甲知识点", "乙知识点"],
                }
            },
        )
        result["retrieval_topk_per_knowledge_point"] = {
            "candidate_count": len(point_candidates),
            "knowledge_points": [candidate.knowledge_point for candidate in point_candidates],
        }
        pool = [
            EvidenceCandidate("ev_multi_01", "q_multi", "物理化学第6版上1", "物理化学", "2.5", "demo", "10", "350", 9.0, "相律定义第一页。", True),
            EvidenceCandidate("ev_multi_02", "q_multi", "物理化学第6版上1", "物理化学", "2.5", "demo", "11", "351", 8.0, "相律跨页续述。", True),
            EvidenceCandidate("ev_multi_03", "q_multi", "物理化学第6版上1", "物理化学", "2.5", "demo", "12", "352", 7.0, "自由度公式。", True),
            EvidenceCandidate("ev_multi_04", "q_multi", "物理化学第6版上1", "物理化学", "3.1", "demo", "20", "400", 6.0, "无关内容。", True),
            EvidenceCandidate("ev_multi_05", "q_multi", "物理化学第6版上1", "物理化学", "2.5", "demo", "13", "352", 6.5, "自由度公式同页补充。", True),
            EvidenceCandidate("ev_multi_06", "q_multi", "物理化学第6版上1", "物理化学", "2.6", "demo", "14", "353", 6.2, "相数补充说明。", True),
            EvidenceCandidate("ev_multi_07", "q_multi", "物理化学第6版上1", "物理化学", "2.6", "demo", "15", "354", 6.1, "相数跨页说明。", True),
        ]
        selection = {
            "question_id": "q_multi",
            "knowledge_points": [
                {
                    "knowledge_point": "相律",
                    "selected_evidence_ids": ["ev_multi_01", "ev_multi_02"],
                    "rejected_evidence_ids": [],
                    "reason": "跨页共同说明相律。",
                },
                {
                    "knowledge_point": "自由度公式",
                    "selected_evidence_ids": ["ev_multi_03", "ev_multi_05", "ev_multi_06", "ev_multi_07"],
                    "rejected_evidence_ids": ["ev_multi_04"],
                    "reason": "公式证据独立支撑。",
                },
            ],
        }
        result["multi_point_confirmed_ids"] = [candidate.evidence_id for candidate in filter_candidates_by_selection(pool, {"q_multi": selection})]
        result["answer_generation_evidence_ids"] = [row["evidence_id"] for row in evidence_for_answer_generation(pool, "q_multi", selection)]
        result["multi_point_citations"] = citation_groups_from_selection(selection, pool)
        grouped_fragment = attach_program_evidence_block(
            {
                "question_id": "q_multi",
                "evidence_ids": ["ev_multi_01", "ev_multi_02", "ev_multi_03"],
                "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "示例解析。"}]}],
            },
            [candidate.__dict__ for candidate in pool],
            selection,
        )
        result["grouped_evidence_block"] = grouped_fragment["blocks"][0]["segments"][0]["text"]
        draft_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "q_multi",
                "answer": "C",
                "analysis": "本题根据相律和自由度公式判断，C 项正确。",
                "analysis_segments": [
                    {"text": "本题根据 {f1} 判断，C 项正确。", "formula_indices": [1]}
                ],
                "option_analysis": {
                    "A": "A 项忽略了相数限制。",
                    "B": "B 项条件不完整。",
                    "C": "C 项符合相律判断。",
                    "D": "D 项自由度判断错误。",
                },
                "steps": ["确定相数和组分数。", "根据自由度判断共存条件。"],
                "formulas": [
                    {"latex": "F=C-P+1", "meaning": "凝聚体系相律表达式", "role": "relation", "display": True}
                ],
                "figure_specs": [
                    {
                        "kind": "custom_diagram",
                        "required_labels": ["相律"],
                        "elements": [
                            {"type": "text", "xy": [0, 0], "text": "相律", "label": "相律"}
                        ],
                    }
                ],
                "mistake_notes": ["容易误把水合物种类数直接当作相数。"],
                "uncertainties": [],
            },
            {"question_id": "q_multi", "section": "一、选择题", "number": "5"},
            [candidate.__dict__ for candidate in pool],
            selection,
        )
        result["draft_to_fragment"] = {
            "schema_version": draft_fragment.get("schema_version"),
            "answer": draft_fragment.get("answer"),
            "evidence_ids": draft_fragment.get("evidence_ids"),
            "evidence_block": draft_fragment["blocks"][0]["segments"][0]["text"],
            "block_labels": [block.get("label") for block in draft_fragment.get("blocks", [])],
            "formula_count": len(draft_fragment.get("formulas", [])),
            "figure_specs_preserved": bool(draft_fragment.get("_draft", {}).get("figure_specs")),
            "figure_specs_top_level": bool(draft_fragment.get("figure_specs")),
            "analysis_starts_with_formula_list": next(
                (
                    block.get("segments", [{}])[0].get("text") == "关键公式："
                    for block in draft_fragment.get("blocks", [])
                    if block.get("label") == "解析" and block.get("segments")
                ),
                False,
            ),
            "issues": validate_v4_answer_fragment(draft_fragment),
        }
        integrated_formula_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "integrated_formula",
                "answer": "2",
                "analysis": "备用解析不应追加全部公式。",
                "analysis_segments": [
                    {"text": "合成氨反应达到平衡时，按 {f1} 判断自由度。", "formula_indices": [1]},
                    {"text": "体系为单一气相，代入组分数和相数即可得到结果。"},
                ],
                "option_analysis": {},
                "steps": [],
                "formulas": [
                    {"latex": "F=C-P+2", "meaning": "相律", "role": "relation", "display": True},
                    {"latex": "C=S-R-R'", "meaning": "独立组分数", "role": "relation", "display": True},
                ],
                "mistake_notes": [],
                "uncertainties": [],
            },
            {"question_id": "integrated_formula", "section": "一、填空题", "number": "2", "stem": "求自由度。"},
            [],
        )
        integrated_blocks = {block.get("label"): block.get("segments", []) for block in integrated_formula_fragment.get("blocks", [])}
        result["noncalculation_formula_integration"] = {
            "analysis_formula_refs": [seg.get("formula_id") for seg in integrated_blocks.get("解析", []) if seg.get("type") == "formula_ref"],
            "analysis_formula_refs_inline": [bool(seg.get("inline")) for seg in integrated_blocks.get("解析", []) if seg.get("type") == "formula_ref"],
            "review_formula_refs": [seg.get("formula_id") for seg in integrated_blocks.get("待复核公式", []) if seg.get("type") == "formula_ref"],
            "has_review_warning": any("未融入解析正文" in str(warning) for warning in integrated_formula_fragment.get("warnings", [])),
            "issues": validate_v4_answer_fragment(integrated_formula_fragment),
        }
        placeholder_choice_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "placeholder_choice",
                "answer": "B",
                "analysis": "备用解析按 {f1} 判断。",
                "analysis_segments": [],
                "option_analysis": {
                    "A": "A 项不满足 {f2}。",
                    "B": "B 项符合 {f1}。",
                },
                "steps": ["先判断系统熵变符号 {f1}。"],
                "formulas": [
                    {"latex": "\\Delta S_\\text{sys}<0", "meaning": "系统熵变为负", "role": "conclusion", "display": True},
                    {"latex": "\\Delta S_\\text{sur}<0", "meaning": "环境熵变为负", "role": "distractor", "display": True},
                ],
                "mistake_notes": ["注意不要把 {f2} 当作正确判断。"],
                "uncertainties": [],
            },
            {"question_id": "placeholder_choice", "section": "一、选择题", "number": "9", "stem": "判断熵变符号。"},
            [],
        )
        placeholder_choice_blocks = {block.get("label"): block.get("segments", []) for block in placeholder_choice_fragment.get("blocks", [])}
        placeholder_choice_text = json.dumps(placeholder_choice_fragment.get("blocks", []), ensure_ascii=False)
        result["noncalculation_placeholder_parsing"] = {
            "no_raw_placeholders": "{f" not in placeholder_choice_text,
            "analysis_formula_refs": [seg.get("formula_id") for seg in placeholder_choice_blocks.get("解析", []) if seg.get("type") == "formula_ref"],
            "option_formula_refs": [seg.get("formula_id") for seg in placeholder_choice_blocks.get("选项分析", []) if seg.get("type") == "formula_ref"],
            "steps_formula_refs": [seg.get("formula_id") for seg in placeholder_choice_blocks.get("解题步骤", []) if seg.get("type") == "formula_ref"],
            "mistake_formula_refs": [seg.get("formula_id") for seg in placeholder_choice_blocks.get("易错点及注意事项", []) if seg.get("type") == "formula_ref"],
            "review_formula_refs": [seg.get("formula_id") for seg in placeholder_choice_blocks.get("待复核公式", []) if seg.get("type") == "formula_ref"],
            "issues": validate_v4_answer_fragment(placeholder_choice_fragment),
        }
        citation_leak_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "citation_leak",
                "answer": "不能。",
                "analysis": "",
                "analysis_segments": [
                    {"text": "教材依据：相律：课本-p238、p240"},
                    {"text": "该条件不满足可逆过程要求，因此不能构成可逆电池。"},
                ],
                "option_analysis": {},
                "steps": [],
                "formulas": [],
                "mistake_notes": [],
                "uncertainties": [],
            },
            {"question_id": "citation_leak", "section": "二、简答题", "number": "1", "stem": "判断可逆性。"},
            [],
        )
        placeholder_calc_fragment = fragment_from_analysis_draft(
            {
                "schema_version": "answer_book.answer_draft.v1",
                "question_id": "placeholder_calc",
                "answer": "最终为{f3}",
                "analysis": "先列{f1}关系，再代入。",
                "analysis_segments": [],
                "option_analysis": {},
                "steps": [
                    {"text": "第1小问：先确认{f1}。"},
                    {
                        "text": "第1步：根据定义，{f1}。代入数据，{f2}。得到{f3}。",
                        "relation_formula_indices": [1],
                        "substitution_formula_indices": [2],
                        "result_formula_indices": [3],
                        "result_text": "得到最终结果{f3}。",
                    }
                ],
                "formulas": [
                    {"latex": "n=\\frac{m}{M}", "role": "relation", "display": True},
                    {"latex": "n=\\frac{18}{18}=1\\ \\mathrm{mol}", "role": "substitution", "display": True},
                    {"latex": "n=1\\ \\mathrm{mol}", "role": "result", "display": True},
                ],
                "mistake_notes": ["注意不要把{f3}误写成质量。"],
                "uncertainties": [],
            },
            {"question_id": "placeholder_calc", "section": "三、计算题", "number": "1", "stem": "计算物质的量。"},
            [],
        )
        placeholder_text = json.dumps(placeholder_calc_fragment.get("blocks", []), ensure_ascii=False)
        citation_text = json.dumps([block for block in citation_leak_fragment.get("blocks", []) if block.get("label") != "教材依据"], ensure_ascii=False)
        result["repair_context_sanitizers"] = {
            "citation_removed": "教材依据" not in citation_text and "课本-p" not in citation_text,
            "placeholder_removed": "{f" not in placeholder_text,
            "answer_summary_placeholder_removed": "{f" not in str(placeholder_calc_fragment.get("answer_summary", "")),
            "draft_analysis_placeholder_removed": "{f" not in str(placeholder_calc_fragment.get("_draft", {}).get("analysis", "")),
            "warnings_placeholder_removed": "{f" not in json.dumps(placeholder_calc_fragment.get("warnings", []), ensure_ascii=False),
            "calc_formula_refs": [
                seg.get("formula_id")
                for block in placeholder_calc_fragment.get("blocks", [])
                if block.get("label") == "解题步骤"
                for seg in block.get("segments", [])
                if seg.get("type") == "formula_ref"
            ],
            "mistake_inline_refs": [
                bool(seg.get("inline"))
                for block in placeholder_calc_fragment.get("blocks", [])
                if block.get("label") == "易错点及注意事项"
                for seg in block.get("segments", [])
                if seg.get("type") == "formula_ref"
            ],
        }
        quality_report = audit_content_quality(
            {
                "items": [
                    {"question_id": "q_multi", "section": "一、选择题", "number": "5", "stem": "选择题示例"},
                    {"question_id": "q_calc", "section": "三、计算题", "number": "1", "stem": "计算题示例"},
                ]
            },
            {
                "fragments": [
                    draft_fragment,
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "q_calc",
                        "answer": "待复核",
                        "evidence_ids": ["ev_calc_01"],
                        "blocks": [
                            {"label": "教材依据", "segments": [{"type": "text", "text": "计算依据：示例教材 2.1-p20"}]},
                            {"label": "解析", "segments": [{"type": "text", "text": "根据教材可知，吉布斯自由能变小于零时反应自发。本题需要人工复核。"}]},
                        ],
                        "formulas": [],
                        "warnings": [],
                    },
                ]
            },
            {
                "drafts": [
                    {
                        "schema_version": "answer_book.answer_draft.v1",
                        "question_id": "q_multi",
                        "answer": "C",
                        "analysis": "本题根据相律和自由度公式判断，故选 C。",
                        "option_analysis": {"A": "错误", "B": "错误", "C": "正确", "D": "错误"},
                        "steps": ["确定相数和组分数。"],
                        "formulas": [{"latex": "F=C-P+1"}],
                        "mistake_notes": ["容易误把水合物种类数直接当作相数。"],
                    },
                    {
                        "schema_version": "answer_book.answer_draft.v1",
                        "question_id": "q_calc",
                        "answer": "",
                        "analysis": "根据教材可知，代入公式得。",
                        "option_analysis": {},
                        "steps": [],
                        "formulas": [],
                        "mistake_notes": [],
                    },
                ]
            },
            {
                "selections": [
                    {"question_id": "q_multi", "knowledge_points": [{"selected_evidence_ids": ["ev_multi_01", "ev_multi_02", "ev_multi_03"], "rejected_evidence_ids": ["ev_multi_04"]}]},
                    {"question_id": "q_calc", "knowledge_points": [{"selected_evidence_ids": ["ev_calc_01"], "rejected_evidence_ids": []}]},
                ]
            },
            tmp_path / "content_quality_audit.json",
        )
        result["content_quality_audit"] = {
            "ok": quality_report.get("ok"),
            "issue_count": quality_report.get("issue_count"),
            "warning_count": quality_report.get("warning_count"),
            "issue_codes": [issue.get("code") for issue in quality_report.get("issues", [])],
            "warning_codes": [warning.get("code") for warning in quality_report.get("warnings", [])],
            "wrote_file": (tmp_path / "content_quality_audit.json").exists(),
        }
        overlap_quality_report = audit_content_quality(
            {"items": [{"question_id": "q_overlap", "section": "二、简答题", "stem": "证据交叉选择示例"}]},
            {
                "fragments": [
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "q_overlap",
                        "answer": "见解析",
                        "evidence_ids": ["ev_a", "ev_b"],
                        "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "本题围绕两个考查点分别说明，证据 A 和证据 B 分别支撑不同部分。"}]}],
                        "formulas": [],
                        "warnings": [],
                    }
                ]
            },
            {
                "drafts": [
                    {
                        "schema_version": "answer_book.answer_draft.v1",
                        "question_id": "q_overlap",
                        "answer": "见解析",
                        "analysis": "本题围绕两个考查点分别说明，证据 A 和证据 B 分别支撑不同部分。",
                        "steps": [],
                        "formulas": [],
                        "mistake_notes": [],
                    }
                ]
            },
            {
                "selections": [
                    {
                        "question_id": "q_overlap",
                        "knowledge_points": [
                            {"knowledge_point": "考查点A", "selected_evidence_ids": ["ev_a"], "rejected_evidence_ids": ["ev_b"]},
                            {"knowledge_point": "考查点B", "selected_evidence_ids": ["ev_b"], "rejected_evidence_ids": ["ev_a"]},
                        ],
                    }
                ]
            },
        )
        result["overlap_evidence_quality"] = [issue.get("code") for issue in overlap_quality_report.get("issues", [])]
        short_answer_quality_report = audit_content_quality(
            {
                "items": [
                    {
                        "question_id": "qa_choose_basis",
                        "section": "五、简答题",
                        "number": "2",
                        "stem": "简述 X 射线滤波片的滤波原理，以及选择依据。",
                    }
                ]
            },
            {
                "fragments": [
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "qa_choose_basis",
                        "answer": "见解析",
                        "evidence_ids": ["ev_qa"],
                        "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "滤波片利用吸收限差异削弱不需要的特征辐射，选择依据是吸收边位置与靶材特征线的匹配关系。"}]}],
                        "formulas": [],
                        "warnings": [],
                    }
                ]
            },
            {
                "drafts": [
                    {
                        "schema_version": "answer_book.answer_draft.v1",
                        "question_id": "qa_choose_basis",
                        "answer": "见解析",
                        "analysis": "滤波片利用吸收限差异削弱不需要的特征辐射，选择依据是吸收边位置与靶材特征线的匹配关系。",
                        "steps": [],
                        "formulas": [],
                        "mistake_notes": [],
                    }
                ]
            },
            {"selections": [{"question_id": "qa_choose_basis", "knowledge_points": [{"selected_evidence_ids": ["ev_qa"], "rejected_evidence_ids": []}]}]},
        )
        result["short_answer_with_choose_text"] = {
            "issue_codes": [issue.get("code") for issue in short_answer_quality_report.get("issues", [])],
            "warning_codes": [warning.get("code") for warning in short_answer_quality_report.get("warnings", [])],
        }
        calc_missing_substitution_report = audit_content_quality(
            {
                "items": [
                    {
                        "question_id": "calc_missing_substitution",
                        "section": "三、计算题",
                        "number": "1",
                        "stem": "计算理想气体体积。",
                    }
                ]
            },
            {
                "fragments": [
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "calc_missing_substitution",
                        "answer": "见解析",
                        "answer_summary": "V=1.70 m^3",
                        "evidence_ids": ["ev_calc"],
                        "blocks": [
                            {"label": "解析", "segments": [{"type": "text", "text": "先求物质的量，再用状态方程求体积。"}]},
                            {
                                "label": "解题步骤",
                                "segments": [
                                    {"type": "text", "text": "第1步：计算物质的量。"},
                                    {"type": "formula_ref", "formula_id": "f_calc_missing_substitution_01"},
                                    {"type": "text", "text": "55.56 mol"},
                                ],
                            },
                            {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "温度要换算为 K。"}]},
                        ],
                        "formulas": [
                            {"formula_id": "f_calc_missing_substitution_01", "latex": "n=\\frac{m}{M}", "role": "relation", "display": True}
                        ],
                        "warnings": [],
                    }
                ]
            },
            {
                "drafts": [
                    {
                        "schema_version": "answer_book.answer_draft.v1",
                        "question_id": "calc_missing_substitution",
                        "answer": "V=1.70 m^3",
                        "analysis": "先求物质的量，再用状态方程求体积。",
                        "steps": [
                            {
                                "text": "第1步：计算物质的量。",
                                "relation_formula_indices": [1],
                                "substitution_formula_indices": [],
                                "result_formula_indices": [],
                                "result_text": "55.56 mol",
                            }
                        ],
                        "formulas": [{"latex": "n=\\frac{m}{M}", "role": "relation", "display": True}],
                        "mistake_notes": ["温度要换算为 K。"],
                    }
                ]
            },
            {"selections": [{"question_id": "calc_missing_substitution", "knowledge_points": [{"selected_evidence_ids": ["ev_calc"], "rejected_evidence_ids": []}]}]},
        )
        result["calculation_missing_substitution_quality"] = [issue.get("code") for issue in calc_missing_substitution_report.get("issues", [])]
        nested_requirement_quality = audit_content_quality(
            {
                "items": [
                    {
                        "question_id": "nested_requirement_quality",
                        "section": "九、简答题",
                        "number": "1",
                        "stem": "相图分析题。",
                        "subquestions": [
                            {
                                "number": "2",
                                "stem": "画图并计算。",
                                "question_type": "简答题",
                                "requirements": [
                                    {"number": "2.1", "stem": "画出冷却曲线。", "question_type": "作图题"},
                                    {"number": "2.2", "stem": "画出组织示意图。", "question_type": "作图题"},
                                    {"number": "2.3", "stem": "计算组织质量比。", "question_type": "计算题"},
                                ],
                            }
                        ],
                    }
                ]
            },
            {
                "fragments": [
                    {
                        "schema_version": "answer_book.answer_fragment.v4",
                        "question_id": "nested_requirement_quality",
                        "answer": "见解析",
                        "answer_summary": "质量比为 1:1",
                        "evidence_ids": [],
                        "blocks": [
                            {"label": "解析", "segments": [{"type": "text", "text": "先按相图确定组织，再对计算部分使用杠杆定律。"}]},
                            {
                                "label": "解题步骤",
                                "segments": [
                                    {"type": "text", "text": "2.3 计算组织质量比。"},
                                    {"type": "formula_ref", "formula_id": "f_nested_req_01"},
                                    {"type": "text", "text": "带入数值："},
                                    {"type": "formula_ref", "formula_id": "f_nested_req_02"},
                                    {"type": "text", "text": "求得："},
                                    {"type": "formula_ref", "formula_id": "f_nested_req_03"},
                                ],
                            },
                            {"label": "易错点及注意事项", "segments": [{"type": "text", "text": "注意只对计算要求使用杠杆定律。"}]},
                        ],
                        "formulas": [
                            {"formula_id": "f_nested_req_01", "latex": "w=\\frac{x_1}{x_2}", "role": "relation", "display": True},
                            {"formula_id": "f_nested_req_02", "latex": "w=1/1", "role": "substitution", "display": True},
                            {"formula_id": "f_nested_req_03", "latex": "w=1", "role": "result", "display": True},
                        ],
                        "warnings": [],
                    }
                ]
            },
            {
                "drafts": [
                    {
                        "schema_version": "answer_book.answer_draft.v1",
                        "question_id": "nested_requirement_quality",
                        "answer": "见解析",
                        "analysis": "先按相图确定组织，再对计算部分使用杠杆定律。",
                        "steps": [
                            {
                                "subquestion_number": "2.3",
                                "text": "计算组织质量比。",
                                "relation_formula_indices": [1],
                                "substitution_formula_indices": [2],
                                "result_formula_indices": [3],
                                "result_text": "质量比为 1:1。",
                            }
                        ],
                        "formulas": [
                            {"latex": "w=\\frac{x_1}{x_2}", "role": "relation", "display": True},
                            {"latex": "w=1/1", "role": "substitution", "display": True},
                            {"latex": "w=1", "role": "result", "display": True},
                        ],
                        "mistake_notes": ["注意只对计算要求使用杠杆定律。"],
                    }
                ]
            },
        )
        result["nested_requirement_quality"] = [issue.get("code") for issue in nested_requirement_quality.get("issues", [])]
        review_gate_stage = tmp_path / "review_gate_stage"
        review_gate_stage.mkdir(parents=True, exist_ok=True)
        gate_report = {
            "ok": False,
            "issues": [{"question_id": "q_gate", "code": "missing_analysis", "message": "缺少解析。", "severity": "issue"}],
            "warnings": [],
            "issue_count": 1,
            "warning_count": 0,
        }
        gate_request = {
            "request_id": "gate_demo",
            "stage": "content_quality",
            "title": "质量审查仍有问题",
            "items": [{"question_id": "q_gate", "code": "missing_analysis", "message": "缺少解析。"}],
        }
        gate_response = {"decision": "allow", "note": "测试允许", "updated_at": "2026-07-04 18:00:00"}
        gate_updated = apply_allowed_to_audit_report(review_gate_stage, "content_quality", gate_report, gate_request, gate_response, review_gate_stage / "content_quality_audit.json")
        gate_targets = collect_audit_issue_targets(gate_report, {"q_gate"})
        result["audit_review_gate"] = {
            "ok_after_allow": gate_updated.get("ok"),
            "issue_count_after_allow": gate_updated.get("issue_count"),
            "warning_count_after_allow": gate_updated.get("warning_count"),
            "allowed_file_exists": (review_gate_stage / "user_allowed_audit_issues.json").exists(),
            "target_qids": sorted(gate_targets.keys()),
        }
        allowed_docx_stage = tmp_path / "allowed_docx_stage"
        allowed_docx_output = tmp_path / "allowed_docx_output"
        allowed_docx_stage.mkdir(parents=True, exist_ok=True)
        allowed_docx_json = allowed_docx_stage / "answer_fragments.json"
        allowed_docx_path = allowed_docx_output / "answer_book.docx"
        allowed_docx_json.write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "q_allowed_docx",
                            "section": "四、填空题",
                            "number": "1",
                            "answer": "见解析",
                            "answer_summary": "\\frac{x}{y}；需要人工复核，但不应阻断正式候选版生成。",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "原始解析存在 DOCX 审计问题。"}]}],
                            "formulas": [],
                            "warnings": ["用户允许 DOCX 审计问题继续。"],
                            "_review_flags": [{"code": "docx_user_allowed", "message": "DOCX 审计由用户允许继续。"}],
                            "_review_candidate_fragment": {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "q_allowed_docx",
                                "section": "四、填空题",
                                "number": "1",
                                "answer": "候选答案",
                                "answer_summary": "候选摘要：需要人工复核，但应作为正式总版 Word 优先内容。",
                                "evidence_ids": [],
                                "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "候选解析：用户允许后应优先使用待复核前的模型解析。"}]}],
                                "formulas": [],
                                "warnings": [],
                            },
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        allowed_docx_candidate = build_user_allowed_docx_candidate(
            allowed_docx_json,
            allowed_docx_path,
            allowed_docx_stage,
            "测试用户允许后按当前完整解析生成正式总版 Word。",
        )
        allowed_docx_text = "\n".join(p.text for p in Document(allowed_docx_path).paragraphs) if allowed_docx_path.exists() else ""
        fallback_docx_path = allowed_docx_output / "fallback_answer_book.docx"
        allowed_docx_placeholder = build_user_allowed_docx_placeholder(
            allowed_docx_json,
            fallback_docx_path,
            allowed_docx_stage,
            "测试正式总版仍失败后的兜底生成。",
        )
        fallback_docx_text = "\n".join(p.text for p in Document(fallback_docx_path).paragraphs) if fallback_docx_path.exists() else ""
        result["user_allowed_docx_candidate"] = {
            "ok": allowed_docx_candidate.get("ok"),
            "docx_exists": allowed_docx_path.exists(),
            "report_exists": (allowed_docx_stage / "docx_user_allowed_candidate.json").exists(),
            "contains_full_answer": "候选解析：用户允许后应优先使用待复核前的模型解析。" in allowed_docx_text,
            "keeps_allowed_summary": "需要人工复核" in allowed_docx_text,
            "uses_review_candidate": allowed_docx_candidate.get("candidate_count") == 1 and allowed_docx_candidate.get("source_mode") == "review_candidate_fragment",
            "original_not_used": "原始解析存在 DOCX 审计问题。" not in allowed_docx_text,
            "is_not_placeholder": "待复核版" not in allowed_docx_text,
            "fallback_still_available": bool(allowed_docx_placeholder.get("ok")) and "待复核版" in fallback_docx_text,
        }
        nonblocking_stage = tmp_path / "nonblocking_quality_stage"
        nonblocking_output = tmp_path / "nonblocking_quality_output"
        nonblocking_stage.mkdir(parents=True, exist_ok=True)
        nonblocking_output.mkdir(parents=True, exist_ok=True)
        review_doc = Document()
        review_doc.add_heading("待复核题目", level=1)
        review_doc.add_paragraph("q_review 需要人工评估。")
        review_doc.save(nonblocking_output / "question_review.docx")
        (nonblocking_output / "word_rendered").mkdir(parents=True, exist_ok=True)
        (nonblocking_output / "word_rendered" / "answer_book.pdf").write_bytes(b"placeholder")
        for filename in ("exam_structure_audit.json", "retrieval_audit.json", "answer_coverage_audit.json", "docx_audit.json", "render_audit.json"):
            (nonblocking_stage / filename).write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        (nonblocking_stage / "environment_check.json").write_text(
            json.dumps({"formula_conversion": {"preferred_chain_ready": True}}, ensure_ascii=False),
            encoding="utf-8",
        )
        (nonblocking_stage / "content_quality_audit.json").write_text(
            json.dumps(
                {
                    "ok": False,
                    "issues": [{"question_id": "q_review", "code": "missing_analysis", "message": "缺少【解析】内容。"}],
                    "warnings": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (nonblocking_stage / "answer_fragments.json").write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "q_review",
                            "section": "一、填空题",
                            "number": "1",
                            "answer": "待复核",
                            "evidence_ids": [],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "需要人工评估。"}]}],
                            "formulas": [],
                            "warnings": ["模型结构化解析生成失败；已进入存疑题目审查文档。"],
                            "_review_candidate_fragment": {
                                "schema_version": "answer_book.answer_fragment.v4",
                                "question_id": "q_review",
                                "section": "一、填空题",
                                "number": "1",
                                "answer": "候选答案",
                                "evidence_ids": [],
                                "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "候选解析：本题按模型候选内容生成。"}]}],
                                "formulas": [],
                                "warnings": [],
                            },
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        build_docx_from_fragments(nonblocking_stage / "answer_fragments.json", nonblocking_output / "answer_book.docx")
        (nonblocking_stage / "question_review_docx.json").write_text(
            json.dumps({"ok": True, "review_question_count": 1, "docx": str(nonblocking_output / "question_review.docx")}, ensure_ascii=False),
            encoding="utf-8",
        )
        (nonblocking_stage / "acceptance_report.json").write_text(json.dumps({"status": "passed", "rendered": True}, ensure_ascii=False), encoding="utf-8")
        (nonblocking_stage / "pipeline_status.json").write_text(
            json.dumps(
                {
                    "stages": [
                        {"stage": "content_quality", "status": "failed"},
                        {"stage": "docx_user_allowed_candidate", "status": "failed"},
                        {"stage": "docx_placeholder", "status": "applied"},
                        {"stage": "docx", "status": "failed"},
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        nonblocking_acceptance = build_final_acceptance_report(nonblocking_stage, nonblocking_output, require_render=True)
        result["content_quality_nonblocking_acceptance"] = {
            "ok": nonblocking_acceptance.get("ok"),
            "status": nonblocking_acceptance.get("status"),
            "issue_count": nonblocking_acceptance.get("issue_count"),
            "warning_count": nonblocking_acceptance.get("warning_count"),
            "warnings": nonblocking_acceptance.get("warnings", []),
            "review_ack_required": (nonblocking_acceptance.get("review_acknowledgement") or {}).get("required"),
        }
        result["recoverable_docx_candidate_failure_acceptance"] = {
            "ok": nonblocking_acceptance.get("ok"),
            "status": nonblocking_acceptance.get("status"),
            "has_pipeline_failed_issue": any("pipeline_status.json contains failed stage" in issue for issue in nonblocking_acceptance.get("issues", [])),
        }
        decision_delivery = build_task_delivery_package("review_ack_demo", nonblocking_stage, nonblocking_output)
        pending_delivery = build_task_delivery_package("review_ack_demo", nonblocking_stage, nonblocking_output, review_policy="keep_pending")
        pending_manifest = {}
        pending_docx_text = ""
        if pending_delivery.get("zip") and Path(pending_delivery["zip"]).exists():
            with zipfile.ZipFile(pending_delivery["zip"]) as zf:
                pending_manifest = json.loads(zf.read("manifest.json").decode("utf-8"))
                pending_docx = tmp_path / "pending_answer_book.docx"
                pending_docx.write_bytes(zf.read("answer_book.docx"))
                pending_docx_text = "\n".join(p.text for p in Document(pending_docx).paragraphs)
        candidate_delivery = build_task_delivery_package("review_ack_demo", nonblocking_stage, nonblocking_output, review_policy="use_candidate")
        candidate_manifest = {}
        candidate_docx_text = ""
        if candidate_delivery.get("zip") and Path(candidate_delivery["zip"]).exists():
            with zipfile.ZipFile(candidate_delivery["zip"]) as zf:
                candidate_manifest = json.loads(zf.read("manifest.json").decode("utf-8"))
                candidate_docx = tmp_path / "candidate_answer_book.docx"
                candidate_docx.write_bytes(zf.read("answer_book.docx"))
                candidate_docx_text = "\n".join(p.text for p in Document(candidate_docx).paragraphs)
        result["delivery_review_acknowledgement"] = {
            "decision_status": decision_delivery.get("status"),
            "decision_ok": decision_delivery.get("ok"),
            "decision_required": (decision_delivery.get("review_acknowledgement") or {}).get("required"),
            "pending_ok": pending_delivery.get("ok"),
            "pending_zip_exists": bool(pending_delivery.get("zip") and Path(pending_delivery["zip"]).exists()),
            "pending_docx_contains_placeholder": "待复核" in pending_docx_text or "需要人工评估" in pending_docx_text,
            "pending_docx_uses_candidate": "候选解析" in pending_docx_text,
            "pending_manifest_user_allowed": (pending_manifest.get("review_acknowledgement") or {}).get("user_allowed"),
            "pending_manifest_policy": (pending_manifest.get("review_acknowledgement") or {}).get("review_policy"),
            "candidate_ok": candidate_delivery.get("ok"),
            "candidate_zip_exists": bool(candidate_delivery.get("zip") and Path(candidate_delivery["zip"]).exists()),
            "candidate_docx_uses_candidate": "候选解析" in candidate_docx_text,
            "candidate_docx_contains_placeholder": "待复核" in candidate_docx_text,
            "review_docx_included": "question_review.docx" in (candidate_delivery.get("files") or []),
            "manifest_user_allowed": (candidate_manifest.get("review_acknowledgement") or {}).get("user_allowed"),
            "manifest_review_file_included": (candidate_manifest.get("review_acknowledgement") or {}).get("review_file_included"),
            "manifest_review_candidate_count": (candidate_manifest.get("review_acknowledgement") or {}).get("review_candidate_count"),
            "manifest_review_policy": (candidate_manifest.get("review_acknowledgement") or {}).get("review_policy"),
        }
        blocking_stage = tmp_path / "blocking_quality_stage"
        blocking_output = tmp_path / "blocking_quality_output"
        blocking_stage.mkdir(parents=True, exist_ok=True)
        blocking_output.mkdir(parents=True, exist_ok=True)
        (blocking_output / "answer_book.docx").write_bytes(b"placeholder")
        (blocking_output / "word_rendered").mkdir(parents=True, exist_ok=True)
        (blocking_output / "word_rendered" / "answer_book.pdf").write_bytes(b"placeholder")
        for filename in ("exam_structure_audit.json", "retrieval_audit.json", "answer_coverage_audit.json", "docx_audit.json", "render_audit.json"):
            (blocking_stage / filename).write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        (blocking_stage / "environment_check.json").write_text(
            json.dumps({"formula_conversion": {"preferred_chain_ready": True}}, ensure_ascii=False),
            encoding="utf-8",
        )
        (blocking_stage / "content_quality_audit.json").write_text(
            json.dumps(
                {
                    "ok": False,
                    "issues": [{"question_id": "q_figure", "code": "missing_required_figure", "message": "题目存在作图或图示需求，但最终解析未插入图片。"}],
                    "warnings": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (blocking_stage / "acceptance_report.json").write_text(json.dumps({"status": "passed", "rendered": True}, ensure_ascii=False), encoding="utf-8")
        (blocking_stage / "pipeline_status.json").write_text(
            json.dumps({"stages": [{"stage": "content_quality", "status": "review_required"}]}, ensure_ascii=False),
            encoding="utf-8",
        )
        blocking_acceptance = build_final_acceptance_report(blocking_stage, blocking_output, require_render=True)
        result["content_quality_blocking_acceptance"] = {
            "ok": blocking_acceptance.get("ok"),
            "status": blocking_acceptance.get("status"),
            "issues": blocking_acceptance.get("issues", []),
        }
        result["semantic_generation_gate"] = {
            "calc_without_formula": semantic_generation_issues(
                {"question_id": "calc_no_formula", "section": "三、计算题", "stem": "计算晶面指数"},
                {
                    "schema_version": "answer_book.answer_fragment.v4",
                    "question_id": "calc_no_formula",
                    "answer": "(100)",
                    "evidence_ids": ["ev_calc"],
                    "blocks": [
                        {"label": "解析", "segments": [{"type": "text", "text": "根据平方和比例判断晶面指数。"}]},
                        {"label": "解题步骤", "segments": [{"type": "text", "text": "比较前三个峰的比例并归属。"}]},
                    ],
                    "formulas": [],
                    "warnings": [],
                },
            ),
            "calc_without_formula_after_retry": semantic_generation_issues(
                {"question_id": "calc_no_formula", "section": "三、计算题", "stem": "计算晶面指数"},
                {
                    "schema_version": "answer_book.answer_fragment.v4",
                    "question_id": "calc_no_formula",
                    "answer": "(100)",
                    "evidence_ids": ["ev_calc"],
                    "blocks": [
                        {"label": "解析", "segments": [{"type": "text", "text": "根据平方和比例判断晶面指数。"}]},
                        {"label": "解题步骤", "segments": [{"type": "text", "text": "比较前三个峰的比例并归属。"}]},
                    ],
                    "formulas": [],
                    "warnings": ["模型重生成后仍未给出公式，按无需公式放行；请在存疑题目审查文档复核。"],
                    "_review_flags": [{"code": "formula_absence_after_retry", "message": "模型重生成后仍未给出公式。"}],
                },
                allow_formula_absence_after_retry=True,
            ),
            "choice_without_formula": semantic_generation_issues(
                {"question_id": "choice_no_formula", "section": "一、选择题", "stem": "概念选择"},
                {
                    "schema_version": "answer_book.answer_fragment.v4",
                    "question_id": "choice_no_formula",
                    "answer": "A",
                    "evidence_ids": ["ev_choice"],
                    "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "概念判断无需公式。"}]}],
                    "formulas": [],
                    "warnings": [],
                },
            ),
        }
        review_stage = tmp_path / "review_stage"
        review_output = tmp_path / "review_output"
        review_stage.mkdir()
        (review_stage / "structured_exam.json").write_text(
            json.dumps(
                {
                    "items": [
                        {
                            "question_id": "calc_review",
                            "section": "三、计算题",
                            "number": "2",
                            "stem": "列出三种立方晶体的前四个衍射峰指数。",
                        }
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (review_stage / "answer_fragments.json").write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "calc_review",
                            "section": "三、计算题",
                            "number": "2",
                            "answer": "简单立方：(100)、(110)、(111)、(200)",
                            "evidence_ids": ["ev_calc_review_01"],
                            "blocks": [
                                {"label": "教材依据", "segments": [{"type": "text", "text": "衍射峰指数：课本 3.2-p142"}]},
                                {"label": "解析", "segments": [{"type": "text", "text": "该题按消光规律和指数平方和顺序列出结果。"}]},
                            ],
                            "formulas": [],
                            "warnings": ["模型重生成后仍未给出公式，按无需公式放行；请在存疑题目审查文档复核。"],
                            "_review_flags": [{"code": "formula_absence_after_retry", "message": "模型重生成后仍未给出公式。"}],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (review_stage / "answer_review_notes.json").write_text(json.dumps({"ok": True, "note_count": 0, "rows": []}, ensure_ascii=False), encoding="utf-8")
        (review_stage / "answer_coverage_audit.json").write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        (review_stage / "content_quality_audit.json").write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        review_items = collect_question_review_items(review_stage)
        review_docx = build_question_review_docx(review_stage, review_output, render_snapshots=False)
        result["question_review_docx"] = {
            "item_count": len(review_items),
            "first_qid": review_items[0].get("question_id") if review_items else "",
            "reason_text": " | ".join(review_items[0].get("notes", [])) if review_items else "",
            "docx_exists": review_docx.exists(),
        }

        nonreview_stage = tmp_path / "nonreview_stage"
        nonreview_stage.mkdir()
        (nonreview_stage / "structured_exam.json").write_text(
            json.dumps(
                {
                    "items": [
                        {"question_id": "choice_hint", "section": "一、选择题", "number": "1", "stem": "选择题示例。"},
                        {"question_id": "calc_unit_hint", "section": "三、计算题", "number": "2", "stem": "计算题示例。"},
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (nonreview_stage / "answer_fragments.json").write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "choice_hint",
                            "answer": "A",
                            "evidence_ids": ["ev_choice_hint"],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "该选项符合题干条件。"}]}],
                            "formulas": [],
                            "warnings": [],
                        },
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "calc_unit_hint",
                            "answer": "12.5",
                            "evidence_ids": ["ev_calc_hint"],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "按题目关系计算得到结果。"}]}],
                            "formulas": [{"formula_id": "f_calc_hint", "latex": "x=12.5", "role": "result", "display": True}],
                            "warnings": [],
                        },
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (nonreview_stage / "answer_review_notes.json").write_text(json.dumps({"ok": True, "note_count": 0, "rows": []}, ensure_ascii=False), encoding="utf-8")
        (nonreview_stage / "answer_coverage_audit.json").write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        (nonreview_stage / "content_quality_audit.json").write_text(
            json.dumps(
                {
                    "ok": True,
                    "issues": [],
                    "warnings": [
                        {
                            "question_id": "calc_unit_hint",
                            "code": "calculation_answer_missing_unit",
                            "message": "计算题数值答案可能缺少单位。",
                            "severity": "warning",
                        },
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        result["nonreview_quality_warnings"] = {
            "item_count": len(collect_question_review_items(nonreview_stage)),
        }

        model_retry_stage = tmp_path / "model_retry_stage"
        model_retry_stage.mkdir()
        (model_retry_stage / "structured_exam.json").write_text(
            json.dumps(
                {"items": [{"question_id": "q_model_retry", "section": "一、选择题", "number": "1", "stem": "示例选择题。"}]},
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (model_retry_stage / "answer_fragments.json").write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "model_token_feedback": [
                        {
                            "question_id": "q_model_retry",
                            "stage": "answer_generation",
                            "ok": True,
                            "attempts": [
                                {"strategy": "primary", "model": "deepseek-v4-flash", "max_tokens": 4096, "error": "Model returned empty JSON content"},
                                {"strategy": "disable_thinking", "model": "deepseek-v4-pro", "max_tokens": 8192},
                            ],
                        }
                    ],
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "q_model_retry",
                            "answer": "A",
                            "evidence_ids": ["ev_retry"],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "切换备用模型后已生成合格解析。"}]}],
                            "formulas": [],
                            "warnings": ["自动切换模型 deepseek-v4-pro 完成结构化生成。"],
                            "_meta": {
                                "provider": "deepseek",
                                "model": "deepseek-v4-pro",
                                "recovered_by": "model_retry",
                                "llm_retry": {
                                    "ok": True,
                                    "attempts": [
                                        {"strategy": "primary", "model": "deepseek-v4-flash", "max_tokens": 4096, "error": "Model returned empty JSON content"},
                                        {"strategy": "disable_thinking", "model": "deepseek-v4-pro", "max_tokens": 8192},
                                    ],
                                },
                            },
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (model_retry_stage / "answer_review_notes.json").write_text(json.dumps({"ok": True, "note_count": 0, "rows": []}, ensure_ascii=False), encoding="utf-8")
        (model_retry_stage / "answer_coverage_audit.json").write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        (model_retry_stage / "content_quality_audit.json").write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        model_retry_docx = build_question_review_docx(model_retry_stage, tmp_path / "model_retry_output", render_snapshots=False)
        model_retry_docx_text = "\n".join(p.text for p in Document(model_retry_docx).paragraphs) if model_retry_docx.exists() else ""
        result["model_retry_not_question_review"] = {
            "item_count": len(collect_question_review_items(model_retry_stage)),
            "docx_has_retry_section": "模型重试策略" in model_retry_docx_text,
            "docx_has_retry_question": "q_model_retry" in model_retry_docx_text,
            "docx_has_retry_strategy": "disable_thinking" in model_retry_docx_text,
        }

        stale_stage = tmp_path / "stale_stage"
        stale_stage.mkdir()
        (stale_stage / "structured_exam.json").write_text(
            json.dumps(
                {
                    "items": [
                        {
                            "question_id": "qa_stale_choice",
                            "section": "五、简答题",
                            "number": "2",
                            "stem": "简述滤波片的滤波原理，以及选择依据。",
                        }
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (stale_stage / "answer_fragments.json").write_text(
            json.dumps(
                {
                    "schema_version": "answer_book.answer_fragments.v4",
                    "fragments": [
                        {
                            "schema_version": "answer_book.answer_fragment.v4",
                            "question_id": "qa_stale_choice",
                            "answer": "见解析",
                            "evidence_ids": ["ev_qa"],
                            "blocks": [{"label": "解析", "segments": [{"type": "text", "text": "滤波片依据吸收边位置实现滤波。"}]}],
                            "formulas": [],
                            "warnings": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (stale_stage / "answer_review_notes.json").write_text(json.dumps({"ok": True, "note_count": 0, "rows": []}, ensure_ascii=False), encoding="utf-8")
        (stale_stage / "answer_coverage_audit.json").write_text(json.dumps({"ok": True, "issues": [], "warnings": []}, ensure_ascii=False), encoding="utf-8")
        (stale_stage / "content_quality_audit.json").write_text(
            json.dumps(
                {
                    "ok": False,
                    "issues": [
                        {
                            "question_id": "qa_stale_choice",
                            "code": "choice_missing_option_analysis",
                            "message": "选择题缺少选项辨析。",
                            "severity": "issue",
                        }
                    ],
                    "warnings": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        result["stale_wrong_kind_review"] = {
            "item_count": len(collect_question_review_items(stale_stage)),
        }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if good_issues:
        return 1
    if not bad_issues:
        return 1
    if not all(result["workflow_ui_labels"].values()):
        return 1
    if not all(result["workflow_guidance_labels"].values()):
        return 1
    if not all(result["workflow_picker_labels"].values()):
        return 1
    if not all(result["frontend_answer_structure"].values()):
        return 1
    if not all(result["task_summary_state"].values()):
        return 1
    if not all(result["pipeline_uses_knowledge_planning"].values()):
        return 1
    if not all(result["pipeline_uses_evidence_selection"].values()):
        return 1
    if not all(result["ui_evidence_selection_progress"].values()):
        return 1
    if not all(result["exam_structure_review_ui"].values()):
        return 1
    if not all(result["question_type_unification"].values()):
        return 1
    if not all(result["content_quality_pipeline_contract"].values()):
        return 1
    if not all(result["late_stage_progress_contract"].values()):
        return 1
    if result["model_repair_filter_policy"] != {
        "content_issue_codes": ["missing_required_figure", "missing_analysis"],
        "content_warning_codes": ["calculation_answer_missing_unit"],
        "docx_issues": ["Formula-like text must not be written as normal text: 示例"],
        "source_has_choice_missing_final_conclusion": False,
    }:
        return 1
    if result["library_scan"]["exam_count"] < 1 or result["library_scan"]["textbook_count"] < 1:
        return 1
    if not result["library_scan"]["duplicate_review_exists"]:
        return 1
    if not result["library_scan"]["cache_reused"] or int(result["library_scan"]["cache_block_count"] or 0) < 1:
        return 1
    if deepseek.get("default_model") != "deepseek-v4-flash":
        return 1
    if result["deepseek_models"] != {
        "default_model": "deepseek-v4-flash",
        "vision_model": "",
        "supports_vision": False,
        "model_options": ["deepseek-v4-flash", "deepseek-v4-pro"],
        "max_tokens": 8192,
    }:
        return 1
    if result["ark_provider"] != {
        "base_url": "https://ark.cn-beijing.volces.com/api/v3",
        "api_key_env": "ARK_API_KEY",
        "allow_custom_model": True,
        "default_model": "doubao-seed-2-1-pro-260628",
        "vision_model": "doubao-seed-2-1-pro-260628",
        "supports_vision": True,
        "image_model": "doubao-seedream-5-0-260128",
        "max_tokens": 8192,
        "deepseek_flash_label": "DeepSeek-V4-flash",
    }:
        return 1
    if result["zhipu_provider"] != {
        "base_url": "https://open.bigmodel.cn/api/paas/v4",
        "api_key_env": "ZHIPU_API_KEY",
        "default_model": "glm-5v-turbo",
        "vision_model": "glm-5v-turbo",
        "supports_vision": True,
        "image_model": "glm-image",
        "max_tokens": 8192,
        "model_options": ["glm-5v-turbo", "glm-4.6v"],
        "glm_46v_label": "GLM-4.6V",
        "thinking_mode": "disabled",
    }:
        return 1
    if "empty" not in result["empty_model_content_error"].lower():
        return 1
    if "not-json" not in result["invalid_model_content_error"]:
        return 1
    if result["json_repair_parse"].get("ping") != "pong":
        return 1
    if result["model_json_retry"] != {"calls": 2, "ok": True}:
        return 1
    if result["formula_detector_examples"]["plain_text"] or not result["formula_detector_examples"]["formula_text"]:
        return 1
    if not result["formula_detector_examples"]["chinese_formula_paraphrase"]:
        return 1
    expected_deepseek_retry = [
        {"model": "deepseek-v4-flash", "max_tokens": 1800, "thinking": None, "compact": False},
        {"model": "deepseek-v4-flash", "max_tokens": 4096, "thinking": None, "compact": False},
        {"model": "deepseek-v4-flash", "max_tokens": 4096, "thinking": "disabled", "compact": False},
        {"model": "deepseek-v4-pro", "max_tokens": 4096, "thinking": None, "compact": False},
        {"model": "deepseek-v4-pro", "max_tokens": 4096, "thinking": "disabled", "compact": True},
    ]
    if result["deepseek_json_retry_policy"] != {"ok": True, "calls": expected_deepseek_retry}:
        return 1
    if result["formula_audit_precision"]["unit_issues"]:
        return 1
    if not result["formula_audit_precision"]["equation_issues"] or not result["formula_audit_precision"]["has_match_detail"]:
        return 1
    if not result["formula_audit_precision"]["chinese_formula_issues"]:
        return 1
    if result["formula_audit_precision"]["chinese_prose_issues"]:
        return 1
    if "\\ce" in result["chemical_latex_plain"] or "⇌" not in result["chemical_latex_plain"]:
        return 1
    if not all(result["answer_quality_requirements_in_prompt"].values()):
        return 1
    if result["answer_depth_profile_prompt"] != {
        "concise_depth": "concise",
        "concise_max_analysis_sentences": 4,
        "concise_mistakes_required": False,
        "deep_depth": "deep",
        "deep_min_steps": 4,
        "deep_mistakes_required": True,
        "section_score": 2,
        "section_depth": "concise",
        "prompt_mentions_depth": True,
    }:
        return 1
    if not result["evidence_binding"]["has_bound_evidence"] or result["evidence_binding"]["first_block"] != "教材依据":
        return 1
    if "部分相关" not in result["evidence_binding"]["reason"]:
        return 1
    if result["evidence_block_formula_audit"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["evidence_docx_ok"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["leaking_docx_allowed"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["chinese_paraphrase_docx_allowed"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["dangerous_docx_rejected"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["repair_report"]["changed"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["repaired_docx_ok"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["placeholder_repair_changed"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["placeholder_repair_removed"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["placeholder_repaired_docx_ok"]:
        return 1
    if not result["docx_evidence_block_formula_policy"]["model_repair_changed"]:
        return 1
    if result["docx_evidence_block_formula_policy"]["model_repair_question_ids"] != ["calc_docx_model_repair"]:
        return 1
    if result["docx_evidence_block_formula_policy"]["model_repair_prompt_calls"] != 1:
        return 1
    if not result["docx_evidence_block_formula_policy"]["model_repaired_docx_ok"]:
        return 1
    if result["docx_formula_block_layout"]["math_paragraph_count"] < 2:
        return 1
    if result["docx_formula_block_layout"]["math_paras_have_body_text"]:
        return 1
    if result["docx_formula_block_layout"]["docx_audit_issues"]:
        return 1
    if result["docx_answer_summary_formula_rendering"]["audit_issues"]:
        return 1
    if not all(
        result["docx_answer_summary_formula_rendering"][key]
        for key in (
            "raw_subscript_removed",
            "raw_sqrt_removed",
            "raw_inequality_removed",
            "raw_symbolic_equation_removed",
            "raw_bragg_equation_removed",
            "has_math",
            "has_radical",
            "has_subscript",
        )
    ):
        return 1
    if not result["docx_audit_checks_answer_summary"]:
        return 1
    if not result["docx_answer_structure"]["no_calc_top_answer"]:
        return 1
    if not result["docx_answer_structure"]["no_calc_summary"]:
        return 1
    if result["docx_answer_structure"]["calc_order"] != ["教材依据", "解析", "答案", "易错点及注意事项"]:
        return 1
    if not result["docx_answer_structure"]["calc_steps_split"]:
        return 1
    if not result["docx_answer_structure"]["no_qa_top_answer"]:
        return 1
    if not result["docx_answer_structure"]["no_qa_summary"]:
        return 1
    if result["docx_answer_structure"]["qa_order"] != ["教材依据", "答案", "易错点及注意事项"]:
        return 1
    if result["docx_audit_prompt_shading"] != {
        "unconfirmed_evidence_segment_marked": True,
        "review_formula_label_shaded": True,
        "unconfirmed_evidence_shaded": True,
    }:
        return 1
    if not result["ordinal_newline_normalization"]["analysis_keeps_second_label"]:
        return 1
    if not result["ordinal_newline_normalization"]["analysis_repairs_third_label"]:
        return 1
    if result["ordinal_newline_normalization"]["analysis_has_broken_label"]:
        return 1
    if not result["ordinal_newline_normalization"]["answer_keeps_numbered_line_break"]:
        return 1
    if result["answer_generation_continues_after_failure"]["fragment_ids"] != ["q_ok_1", "q_fail", "q_ok_2"]:
        return 1
    if result["answer_generation_continues_after_failure"]["issue_qids"] != ["q_fail"]:
        return 1
    if result["answer_generation_continues_after_failure"]["fallback_count"] != 1:
        return 1
    if "answer_generation_failed" not in result["answer_generation_continues_after_failure"]["failed_flags"]:
        return 1
    if result["model_retry_not_review_required"]["note_count"] != 0 or result["model_retry_not_review_required"]["rows"]:
        return 1
    if result["calc_answer_summary_policy"]["top_answer"] != "见解析":
        return 1
    if result["calc_answer_summary_policy"]["answer_summary"] != "ΔU=-2.09×10^3 kJ；ΔH=-2.26×10^3 kJ":
        return 1
    if result["calc_answer_summary_policy"]["docx_has_answer_prefix"]:
        return 1
    if not result["calc_answer_summary_policy"]["raw_summary_formula_removed"]:
        return 1
    if result["calc_answer_summary_policy"]["summary_math_count"] != 0:
        return 1
    if result["calc_answer_summary_policy"]["docx_audit_issues"]:
        return 1
    if not result["multipart_answer_layout"]["answer_has_second_line"]:
        return 1
    if not result["multipart_answer_layout"]["analysis_has_second_line"]:
        return 1
    if (result["answer_depth_profile_fragment"] or {}).get("depth") != "concise":
        return 1
    if (result["answer_depth_profile_fragment"] or {}).get("score") != 2:
        return 1
    if not all(result["multipart_calculation_step_grouping"].values()):
        return 1
    if not all(result["structured_subquestion_calculation"].get(key) for key in ("program_title_uses_first_stem", "program_title_uses_second_stem", "legacy_heading_payload_preserved")):
        return 1
    if result["structured_subquestion_calculation"]["draft_subquestion_numbers"] != ["1", "2"]:
        return 1
    if any("\\\\" in latex for latex in result["calc_solution_layout_policy"]["normalized_latex"]):
        return 1
    if not result["calc_solution_layout_policy"]["all_display"]:
        return 1
    if result["calc_solution_layout_policy"]["analysis_has_formula_refs"]:
        return 1
    if result["calc_solution_layout_policy"]["steps_formula_ref_count"] != 6:
        return 1
    if "计算物质的量" not in result["calc_solution_layout_policy"]["steps_text"] or "由理想气体状态方程求体积" not in result["calc_solution_layout_policy"]["steps_text"]:
        return 1
    if "带入数值：" not in result["calc_solution_layout_policy"]["steps_text"] or "求得：" not in result["calc_solution_layout_policy"]["steps_text"]:
        return 1
    if result["calc_solution_layout_policy"]["has_formula_dump_heading"]:
        return 1
    step_order = result["calc_solution_layout_policy"]["step_order"]
    if not (
        step_order.index("计算物质的量。") < step_order.index("带入数值：") < step_order.index("求得：")
        < step_order.index("由理想气体状态方程求体积。")
    ):
        return 1
    if result["figure_generation_pipeline"]["generated_count"] != 1:
        return 1
    if result["figure_generation_pipeline"]["kind"] != "fcc_cell":
        return 1
    if not result["figure_generation_pipeline"]["spec_exists"]:
        return 1
    if not result["figure_generation_pipeline"]["has_image_ref"]:
        return 1
    if not result["figure_generation_pipeline"]["png_exists"]:
        return 1
    if result["curved_surface_figure_generation"]["generated_count"] != 1:
        return 1
    if result["curved_surface_figure_generation"]["kind"] != "curved_liquid_surface":
        return 1
    if not result["curved_surface_figure_generation"]["png_exists"]:
        return 1
    if result["custom_diagram_generation"]["generated_count"] != 1:
        return 1
    if result["custom_diagram_generation"]["kind"] != "custom_diagram":
        return 1
    if not result["custom_diagram_generation"]["png_exists"]:
        return 1
    if not result["custom_diagram_generation"]["has_image_ref"]:
        return 1
    if result["exam_source_coverage"].get("missing_item_like_count") != 1:
        return 1
    if result["multi_subject_exam_extract"]["ids"] != [
        "choice_s01_01_01",
        "choice_s01_01_02",
        "calc_s01_03_01",
        "calc_s01_04_01",
        "choice_s02_05_01",
        "choice_s02_05_02",
    ]:
        return 1
    if result["multi_subject_exam_extract"]["audit_issues"]:
        return 1
    if any("材料现代研究" in stem or "计算题 (本题共 15 分)" in stem for stem in result["multi_subject_exam_extract"]["calc_stems"]):
        return 1
    if result["docx_image_multimodal_input"] != {
        "image_ref_count": 1,
        "image_exists": True,
        "question_snapshot_count": 1,
        "question_snapshot_exists": True,
        "review_request_snapshot_count": 1,
        "stem_has_marker": False,
        "understanding_needs_vision": True,
        "understanding_has_image": True,
        "prompt_uses_question_understanding": True,
        "prompt_has_raw_image_url": False,
    }:
        return 1
    if result["title_only_sections"] != ["", "2.5", "2.5"]:
        return 1
    if result["page_number_mapping"]["page_number_block"] != "40":
        return 1
    if result["page_number_mapping"]["reference_only"] != "":
        return 1
    if result["page_number_mapping"]["audit_ok"] or result["page_number_mapping"]["audit_issue_count"] != 1:
        return 1
    if result["retrieval_filtering"] != ["相变焓随温度变化可由热容差积分判断。"]:
        return 1
    if not result["directory_evidence_filtering"]["catalog_page"] or result["directory_evidence_filtering"]["real_page"]:
        return 1
    if "冰" in result["retrieval_query_uses_plan"] or "溶液" in result["retrieval_query_uses_plan"]:
        return 1
    if result["plan_driven_retrieval"] != ["相变焓随温度变化可由热容差积分判断。"]:
        return 1
    if result["multi_point_confirmed_ids"] != [
        "ev_multi_01",
        "ev_multi_02",
        "ev_multi_03",
        "ev_multi_05",
        "ev_multi_06",
        "ev_multi_07",
    ]:
        return 1
    if result["answer_generation_evidence_ids"] != result["multi_point_confirmed_ids"]:
        return 1
    citation_texts = [x.get("citation", "") for x in result["multi_point_citations"]]
    if not any("相律" in text and "课本-p350-p351" in text for text in citation_texts):
        return 1
    if not any("自由度公式" in text and "课本-p352-p354" in text for text in citation_texts):
        return 1
    if "相律：课本-p350-p351" not in result["grouped_evidence_block"]:
        return 1
    if "自由度公式：课本-p352-p354" not in result["grouped_evidence_block"]:
        return 1
    if result["draft_to_fragment"]["issues"]:
        return 1
    if result["draft_to_fragment"]["block_labels"] != ["教材依据", "解析", "选项分析", "解题步骤", "易错点及注意事项"]:
        return 1
    if result["draft_to_fragment"]["evidence_ids"] != [
        "ev_multi_01",
        "ev_multi_02",
        "ev_multi_03",
        "ev_multi_05",
        "ev_multi_06",
        "ev_multi_07",
    ]:
        return 1
    if not result["draft_to_fragment"]["figure_specs_preserved"]:
        return 1
    if not result["draft_to_fragment"]["figure_specs_top_level"]:
        return 1
    if result["draft_to_fragment"]["analysis_starts_with_formula_list"]:
        return 1
    if result["noncalculation_formula_integration"]["analysis_formula_refs"] != ["f_integrated_formula_01"]:
        return 1
    if result["noncalculation_formula_integration"]["analysis_formula_refs_inline"] != [True]:
        return 1
    if result["noncalculation_formula_integration"]["review_formula_refs"] != ["f_integrated_formula_02"]:
        return 1
    if not result["noncalculation_formula_integration"]["has_review_warning"]:
        return 1
    if result["noncalculation_formula_integration"]["issues"]:
        return 1
    if not result["noncalculation_placeholder_parsing"]["no_raw_placeholders"]:
        return 1
    if result["noncalculation_placeholder_parsing"]["analysis_formula_refs"] != ["f_placeholder_choice_01"]:
        return 1
    if result["noncalculation_placeholder_parsing"]["option_formula_refs"] != [
        "f_placeholder_choice_02",
        "f_placeholder_choice_01",
    ]:
        return 1
    if result["noncalculation_placeholder_parsing"]["steps_formula_refs"] != ["f_placeholder_choice_01"]:
        return 1
    if result["noncalculation_placeholder_parsing"]["mistake_formula_refs"] != ["f_placeholder_choice_02"]:
        return 1
    if result["noncalculation_placeholder_parsing"]["review_formula_refs"]:
        return 1
    if result["noncalculation_placeholder_parsing"]["issues"]:
        return 1
    if not result["repair_context_sanitizers"]["citation_removed"]:
        return 1
    if not result["repair_context_sanitizers"]["placeholder_removed"]:
        return 1
    if not result["repair_context_sanitizers"]["answer_summary_placeholder_removed"]:
        return 1
    if not result["repair_context_sanitizers"]["draft_analysis_placeholder_removed"]:
        return 1
    if not result["repair_context_sanitizers"]["warnings_placeholder_removed"]:
        return 1
    if len(result["repair_context_sanitizers"]["calc_formula_refs"]) != 3:
        return 1
    if result["repair_context_sanitizers"]["mistake_inline_refs"] != [True]:
        return 1
    if result["content_quality_audit"]["ok"]:
        return 1
    expected_issue_codes = {"missing_answer", "calculation_missing_formula", "calculation_missing_steps", "calculation_missing_mistake_notes", "forbidden_process_text"}
    if not expected_issue_codes.issubset(set(result["content_quality_audit"]["issue_codes"])):
        return 1
    if "generic_analysis_phrase" not in result["content_quality_audit"]["warning_codes"]:
        return 1
    if not result["content_quality_audit"]["wrote_file"]:
        return 1
    if "calculation_missing_substitution" not in result["calculation_missing_substitution_quality"]:
        return 1
    if "missing_required_figure" not in result["nested_requirement_quality"]:
        return 1
    if "calculation_missing_subquestion_steps" in result["nested_requirement_quality"]:
        return 1
    if result["audit_review_gate"] != {
        "ok_after_allow": True,
        "issue_count_after_allow": 0,
        "warning_count_after_allow": 1,
        "allowed_file_exists": True,
        "target_qids": ["q_gate"],
    }:
        return 1
    if result["user_allowed_docx_candidate"] != {
        "ok": True,
        "docx_exists": True,
        "report_exists": True,
        "contains_full_answer": True,
        "keeps_allowed_summary": True,
        "uses_review_candidate": True,
        "original_not_used": True,
        "is_not_placeholder": True,
        "fallback_still_available": True,
    }:
        return 1
    if not result["content_quality_blocking_acceptance"]["ok"]:
        return 1
    if result["content_quality_blocking_acceptance"]["status"] != "passed_with_warnings":
        return 1
    if "uses_rejected_evidence" in result["overlap_evidence_quality"]:
        return 1
    if "choice_missing_option_analysis" in result["short_answer_with_choose_text"]["issue_codes"]:
        return 1
    if not result["content_quality_nonblocking_acceptance"]["ok"]:
        return 1
    if result["content_quality_nonblocking_acceptance"]["issue_count"] != 0:
        return 1
    if not any("content_quality" in warning for warning in result["content_quality_nonblocking_acceptance"]["warnings"]):
        return 1
    if not result["content_quality_nonblocking_acceptance"]["review_ack_required"]:
        return 1
    if result["recoverable_docx_candidate_failure_acceptance"] != {
        "ok": True,
        "status": "passed_with_warnings",
        "has_pipeline_failed_issue": False,
    }:
        return 1
    if result["delivery_review_acknowledgement"]["decision_ok"]:
        return 1
    if result["delivery_review_acknowledgement"]["decision_status"] != "review_ack_required":
        return 1
    if not result["delivery_review_acknowledgement"]["decision_required"]:
        return 1
    if not result["delivery_review_acknowledgement"]["pending_ok"] or not result["delivery_review_acknowledgement"]["pending_zip_exists"]:
        return 1
    if not result["delivery_review_acknowledgement"]["pending_docx_contains_placeholder"]:
        return 1
    if result["delivery_review_acknowledgement"]["pending_docx_uses_candidate"]:
        return 1
    if result["delivery_review_acknowledgement"]["pending_manifest_user_allowed"]:
        return 1
    if result["delivery_review_acknowledgement"]["pending_manifest_policy"] != "keep_pending":
        return 1
    if not result["delivery_review_acknowledgement"]["candidate_ok"] or not result["delivery_review_acknowledgement"]["candidate_zip_exists"]:
        return 1
    if not result["delivery_review_acknowledgement"]["candidate_docx_uses_candidate"]:
        return 1
    if result["delivery_review_acknowledgement"]["candidate_docx_contains_placeholder"]:
        return 1
    if not result["delivery_review_acknowledgement"]["review_docx_included"]:
        return 1
    if not result["delivery_review_acknowledgement"]["manifest_user_allowed"]:
        return 1
    if not result["delivery_review_acknowledgement"]["manifest_review_file_included"]:
        return 1
    if result["delivery_review_acknowledgement"]["manifest_review_candidate_count"] != 1:
        return 1
    if result["delivery_review_acknowledgement"]["manifest_review_policy"] != "use_candidate":
        return 1
    if "calculation_missing_formula" not in result["semantic_generation_gate"]["calc_without_formula"]:
        return 1
    if result["semantic_generation_gate"]["calc_without_formula_after_retry"]:
        return 1
    if result["semantic_generation_gate"]["choice_without_formula"]:
        return 1
    if result["question_review_docx"]["item_count"] != 1 or result["question_review_docx"]["first_qid"] != "calc_review":
        return 1
    if "模型重生成后仍未给出公式" not in result["question_review_docx"]["reason_text"]:
        return 1
    if not result["question_review_docx"]["docx_exists"]:
        return 1
    if result["nonreview_quality_warnings"]["item_count"] != 1:
        return 1
    if result["model_retry_not_question_review"]["item_count"] != 0:
        return 1
    if not result["model_retry_not_question_review"]["docx_has_retry_section"]:
        return 1
    if not result["model_retry_not_question_review"]["docx_has_retry_question"]:
        return 1
    if not result["model_retry_not_question_review"]["docx_has_retry_strategy"]:
        return 1
    if result["stale_wrong_kind_review"]["item_count"] != 0:
        return 1
    if result["limited_concurrency"]["values"] != [10, 20, 30]:
        return 1
    if result["limited_concurrency"]["finished"] == [1, 2, 3]:
        return 1
    if not result["response_format_fallback"]["ok"]:
        return 1
    if result["response_format_fallback"]["requests_used_response_format"] != [True, False]:
        return 1
    if not result["thinking_mode_request"]["ok"]:
        return 1
    if result["thinking_mode_request"]["request_modes"] != ["disabled"]:
        return 1
    if result["zhipu_reasoning_retry_policy"] != {
        "ok": True,
        "calls": [
            {"model": "glm-4.6v", "max_tokens": 8192, "thinking": None},
            {"model": "glm-4.6v", "max_tokens": 16384, "thinking": None},
            {"model": "glm-4.6v", "max_tokens": 16384, "thinking": "disabled"},
        ],
    }:
        return 1
    if not all(result["thinking_mode_ui"].values()):
        return 1
    if not all(result["task_duration_ui"].values()):
        return 1
    if result["task_duration_summary"]["duration_text"] != "3分5秒":
        return 1
    if result["answer_generation_default_workers"] != 5:
        return 1
    if result["multipart_question_extraction"]["item_count"] != 2:
        return 1
    if "2. 用杠杆定律计算组织组成。" not in result["multipart_question_extraction"]["first_stem"]:
        return 1
    if result["multipart_question_extraction"]["second_number"] != "2":
        return 1
    if len(result["multipart_question_extraction"]["first_subquestions"]) != 3:
        return 1
    if result["multipart_question_extraction"]["comma_item_count"] != 1:
        return 1
    if len(result["multipart_question_extraction"]["comma_subquestions"]) != 2:
        return 1
    if result["multipart_question_extraction"]["unnumbered_intro_count"] != 1:
        return 1
    if "正常沸点353K" not in result["multipart_question_extraction"]["unnumbered_intro_stem"]:
        return 1
    if len(result["multipart_question_extraction"]["unnumbered_intro_subquestions"]) != 3:
        return 1
    nested_subquestions = result["multipart_question_extraction"]["nested_requirement_subquestions"]
    if len(nested_subquestions) != 2:
        return 1
    nested_second_requirements = nested_subquestions[1].get("requirements", []) if len(nested_subquestions) > 1 else []
    if [req.get("number") for req in nested_second_requirements] != ["2.1", "2.2", "2.3"]:
        return 1
    if [req.get("question_type") for req in nested_second_requirements] != ["作图题", "作图题", "计算题"]:
        return 1
    if result["exam_structure_review"] != {
        "request_item_count": 2,
        "request_first_type": "计算题",
        "first_question_type": "简答题",
        "first_section": "三、简答题",
        "first_section_raw": "三、简答题",
        "first_extracted_section": "三、计算题",
        "first_reviewed": True,
        "second_question_type": "计算题",
        "second_section": "二、计算题",
    }:
        return 1
    if result["question_type_contract"] != {
        "types": ["选择题", "判断题", "填空题", "简答题", "计算题", "作图题"],
        "inferred_choice": "选择题",
        "normalized_drawing": "作图题",
        "kind_calc": "calculation",
        "kind_mixed": "mixed",
        "kind_nested_mixed": "mixed",
        "request_subquestion_type": "简答题",
        "subquestion_first_type": "计算题",
        "subquestion_second_type": "作图题",
        "subquestion_structure_numbers": ["1", "3"],
        "subquestion_structure_first_stem": "计算焓变并写出单位。",
    }:
        return 1
    if result["docx_table_attachment_extraction"] != {
        "table_count": 1,
        "rows": [
            ["量", "R/J·K-1·mol-1", "F/C·mol-1", "0℃", "ln10", "ln2"],
            ["值", "8.314", "96500", "273K", "2.303", "0.693"],
        ],
        "stem_has_marker": False,
        "source_has_table_text": True,
        "prompt_has_question_understanding": True,
        "understanding_has_table_rows": True,
        "simple_table_needs_vision": True,
        "complex_table_detected": True,
        "complex_table_render_exists": True,
    }:
        return 1
    if result["evidence_selection_parallel"]["worker_count"] != 5:
        return 1
    if result["evidence_selection_parallel"]["progress_worker_count"] != 5:
        return 1
    if not result["evidence_selection_parallel"]["parallel_enabled"]:
        return 1
    if result["evidence_selection_parallel"]["finished"] == ["q_parallel_1", "q_parallel_2", "q_parallel_3"]:
        return 1
    if result["evidence_selection_parallel"]["confirmed_count"] != 3:
        return 1
    if result["evidence_selection_parallel"]["selected_question_count"] != 3:
        return 1
    if result["knowledge_planning_parallel"]["worker_count"] != 3:
        return 1
    if result["knowledge_planning_parallel"]["progress_worker_count"] != 3:
        return 1
    if not result["knowledge_planning_parallel"]["parallel_enabled"]:
        return 1
    if result["knowledge_planning_parallel"]["finished"] == ["kp_1", "kp_2", "kp_3"]:
        return 1
    if result["knowledge_planning_parallel"]["progress_status"] != "completed":
        return 1
    if result["knowledge_planning_parallel"]["plan_count"] != 3:
        return 1
    if result["interrupted_task_recovery"]["recovered_ids"] != ["stale_running_task"]:
        return 1
    if result["interrupted_task_recovery"]["stale_status"] != "failed":
        return 1
    if result["interrupted_task_recovery"]["stale_stage"] != "interrupted":
        return 1
    if "服务重启" not in result["interrupted_task_recovery"]["stale_error"]:
        return 1
    if result["interrupted_task_recovery"]["completed_status"] != "completed":
        return 1
    if not result["interrupted_task_recovery"]["has_interrupted_event"]:
        return 1
    if result["running_task_diagnostics"]["needs_attention"]:
        return 1
    if result["running_task_diagnostics"]["title"] != "任务日志摘要":
        return 1
    if result["running_task_diagnostics"]["issue_count"] != 0:
        return 1
    if result["running_task_diagnostics"]["warning_count"] != 0:
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
