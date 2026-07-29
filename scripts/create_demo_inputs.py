#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document

from app.paths import EXAMS_DIR, TEXTBOOKS_DIR, ensure_project_dirs


def main() -> int:
    ensure_project_dirs()
    exam = EXAMS_DIR / "demo_物理化学真题.docx"
    doc = Document()
    doc.add_paragraph("一、选择题")
    doc.add_paragraph("1、理想气体等温可逆膨胀时，下列判断正确的是（ ）。")
    doc.add_paragraph("A. 熵变为正 B. 熵变为零 C. 熵变为负 D. 无法判断")
    doc.add_paragraph("二、计算题")
    doc.add_paragraph("1、某可逆电池电动势为 E，反应电子数为 n，写出反应吉布斯函数与电动势的关系。")
    doc.save(exam)

    textbook = TEXTBOOKS_DIR / "物理化学_demo.json"
    data = {
        "pdf_info": [
            {
                "page_idx": 1,
                "blocks": [
                    {"text": "第 100 页"},
                    {"text": "理想气体等温可逆膨胀时，系统熵变可由体积比计算，膨胀过程熵变为正。"},
                ],
            },
            {
                "page_idx": 2,
                "blocks": [
                    {"text": "第 210 页"},
                    {"text": "可逆电池反应的摩尔吉布斯函数变与电动势的关系为 ΔrGm=-nFE。"},
                ],
            },
        ]
    }
    textbook.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"exam": str(exam), "textbooks": str(TEXTBOOKS_DIR)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

